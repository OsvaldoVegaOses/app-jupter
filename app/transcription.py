"""
Transcripción de audio con diarización usando Azure OpenAI.

Este módulo proporciona funciones para:
1. Transcribir archivos de audio usando gpt-4o-transcribe o gpt-4o-transcribe-diarize
2. Parsear respuestas con segmentos de múltiples speakers (N speakers soportados)
3. Convertir transcripciones a ParagraphRecord para integración con el pipeline de ingesta

Modelos soportados:
    - gpt-4o-transcribe: Transcripción simple sin diarización
    - gpt-4o-transcribe-diarize: Transcripción con identificación de speakers

Formatos de audio soportados:
    - MP3, MP4, M4A, WAV, WebM, MPEG, MPGA, FLAC, OGG

Configuración actual:
    - Idioma: Español ("es") - fijado en frontend
    - Formato respuesta: "json" (NO verbose_json - incompatible con diarize)
    - Chunking: "auto" para archivos grandes
    - Max chunk: 5 minutos (300 segundos)
    
Uso:
    from app.transcription import transcribe_audio, audio_to_fragments
    from app.settings import load_settings
    
    settings = load_settings()
    
    # Transcripción simple
    result = transcribe_audio("entrevista.mp3", settings)
    
    # Transcripción con diarización
    result = transcribe_audio("entrevista.mp3", settings, diarize=True)
    
    # Pipeline completo: audio → fragmentos listos para ingesta
    fragments = audio_to_fragments("entrevista.mp3", settings)

Límites:
    - Tamaño máximo por archivo: 25 MB (API directa)
    - Para archivos más grandes, se usa chunking automático

Changelog:
    - 2025-12-17: Corregido response_format de "verbose_json" a "json"
                  (Azure cambió la API de gpt-4o-transcribe-diarize)
    - 2025-12-17: Documentada configuración de idioma español
"""

from __future__ import annotations

import io
import json
import os
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence, Tuple

import httpx
import structlog

from .documents import ParagraphRecord, FragmentRecord, coalesce_paragraph_records
from .settings import AppSettings

logger = structlog.get_logger(__name__)

# Formatos de audio soportados por la API
SUPPORTED_AUDIO_FORMATS = {
    ".mp3", ".mp4", ".m4a", ".wav", ".webm", 
    ".mpeg", ".mpga", ".flac", ".ogg"
}

# Tamaño máximo de archivo (25 MB)
MAX_FILE_SIZE_MB = 25
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

# Duración máxima por chunk para dividir audio (5 minutos)
MAX_CHUNK_DURATION_SEC = 300


def optimize_audio_for_transcription(
    input_path: str | Path,
    output_path: str | Path | None = None,
) -> Path:
    """
    Optimiza audio para transcripción: convierte a MP3 mono 16kHz.
    
    Beneficios:
    - Reduce tamaño en ~70% (menos tiempo de upload)
    - 16kHz es suficiente para reconocimiento de voz
    - Mono reduce tamaño a la mitad vs stereo
    - MP3 tiene buena compresión para voz
    
    Args:
        input_path: Ruta al archivo de audio original
        output_path: Ruta destino (opcional, crea archivo temporal si no se especifica)
        
    Returns:
        Path al archivo optimizado
    """
    import subprocess
    import tempfile
    
    input_path = Path(input_path)
    
    if output_path is None:
        # Crear archivo temporal con extensión .mp3
        fd, tmp = tempfile.mkstemp(suffix=".mp3")
        os.close(fd)
        output_path = Path(tmp)
    else:
        output_path = Path(output_path)
    
    log = logger.bind(
        input=str(input_path),
        output=str(output_path),
    )
    
    # Obtener tamaño original
    original_size = input_path.stat().st_size
    
    try:
        # FFmpeg: convertir a MP3 mono 16kHz con bitrate optimizado para voz
        result = subprocess.run(
            [
                "ffmpeg",
                "-y",                  # Sobrescribir si existe
                "-i", str(input_path), # Input
                "-vn",                 # No video
                "-ac", "1",            # Mono (1 canal)
                "-ar", "16000",        # 16kHz sample rate (suficiente para voz)
                "-b:a", "64k",         # 64kbps bitrate (bueno para voz)
                "-f", "mp3",           # Formato MP3
                str(output_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,  # 2 minutos timeout
        )
        
        if result.returncode != 0:
            log.error("audio.optimize.ffmpeg_failed", stderr=result.stderr)
            raise RuntimeError(f"FFmpeg failed: {result.stderr}")
        
        # Calcular reducción
        optimized_size = output_path.stat().st_size
        reduction_pct = (1 - optimized_size / original_size) * 100
        
        log.info(
            "audio.optimize.success",
            original_mb=round(original_size / 1024 / 1024, 2),
            optimized_mb=round(optimized_size / 1024 / 1024, 2),
            reduction_pct=round(reduction_pct, 1),
        )
        
        return output_path
        
    except subprocess.TimeoutExpired:
        log.error("audio.optimize.timeout")
        raise RuntimeError("Audio optimization timed out")
    except FileNotFoundError:
        log.warning("audio.optimize.ffmpeg_not_found")
        # Si ffmpeg no está disponible, retornar archivo original
        return input_path


def get_audio_duration(file_path: str | Path) -> float:
    """Obtiene la duración del audio en segundos usando ffprobe."""
    import subprocess
    path = Path(file_path)
    try:
        result = subprocess.run(
            ["ffprobe", "-v", "quiet", "-show_entries", "format=duration",
             "-of", "csv=p=0", str(path)],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return float(result.stdout.strip())
    except Exception as e:
        logger.warning("ffprobe.error", file=str(path), error=str(e))
    return 0.0


def split_audio_ffmpeg(
    file_path: str | Path,
    output_dir: str | Path | None = None,
    chunk_duration_sec: int = MAX_CHUNK_DURATION_SEC,
) -> List[Path]:
    """
    Divide un archivo de audio en chunks más pequeños usando ffmpeg.
    
    Args:
        file_path: Ruta al archivo de audio original
        output_dir: Directorio donde guardar los chunks (default: temp dir)
        chunk_duration_sec: Duración de cada chunk en segundos (default: 300 = 5 min)
        
    Returns:
        Lista de rutas a los archivos de chunks creados
    """
    import subprocess
    import tempfile
    
    path = Path(file_path)
    if output_dir is None:
        output_dir = Path(tempfile.mkdtemp(prefix="audio_chunks_"))
    else:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
    
    total_duration = get_audio_duration(path)
    if total_duration <= 0 or total_duration <= chunk_duration_sec:
        logger.info("split_audio.no_split_needed", duration=total_duration)
        return [path]
    
    num_chunks = int(total_duration / chunk_duration_sec) + 1
    chunks: List[Path] = []
    
    logger.info("split_audio.start", file=path.name, duration=total_duration, chunks=num_chunks)
    
    for i in range(num_chunks):
        start_time = i * chunk_duration_sec
        output_file = output_dir / f"{path.stem}_chunk{i:03d}{path.suffix}"
        
        try:
            result = subprocess.run(
                ["ffmpeg", "-y", "-i", str(path), "-ss", str(start_time),
                 "-t", str(chunk_duration_sec), "-c", "copy", str(output_file)],
                capture_output=True, text=True, timeout=120,
            )
            if result.returncode == 0 and output_file.exists():
                chunks.append(output_file)
        except Exception as e:
            logger.error("ffmpeg.error", chunk=i, error=str(e))
    
    logger.info("split_audio.complete", chunks=len(chunks))
    return chunks


@dataclass
class TranscriptSegment:
    """
    Segmento individual de transcripción con información de speaker.
    
    Attributes:
        speaker: Identificador del speaker (SPEAKER_00, SPEAKER_01, etc.)
        text: Texto transcrito del segmento
        start: Tiempo de inicio en segundos
        end: Tiempo de fin en segundos
        segment_index: Índice correlativo global del segmento (para ordenamiento)
    """
    speaker: str
    text: str
    start: float = 0.0
    end: float = 0.0
    segment_index: int = 0


@dataclass
class TranscriptionResult:
    """
    Resultado completo de una transcripción de audio.
    
    Attributes:
        text: Transcripción completa concatenada
        segments: Lista de segmentos con información de speaker/timing
        speaker_count: Número de speakers únicos detectados
        duration_seconds: Duración total del audio en segundos
        model: Modelo usado para la transcripción
        metadata: Metadata adicional del procesamiento
    """
    text: str
    segments: List[TranscriptSegment] = field(default_factory=list)
    speaker_count: int = 0
    duration_seconds: float = 0.0
    model: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)



def validate_audio_file(file_path: str | Path, skip_size_check: bool = False) -> Tuple[bool, str]:
    """
    Valida que el archivo de audio sea procesable.
    
    Args:
        file_path: Ruta al archivo de audio
        skip_size_check: Si True, no valida el tamaño máximo (para chunks)
        
    Returns:
        Tuple[bool, str]: (es_válido, mensaje_de_error)
    """
    path = Path(file_path)
    
    if not path.exists():
        return False, f"Archivo no encontrado: {path}"
    
    if not path.is_file():
        return False, f"No es un archivo: {path}"
    
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_AUDIO_FORMATS:
        return False, f"Formato no soportado: {suffix}. Formatos válidos: {', '.join(SUPPORTED_AUDIO_FORMATS)}"
    
    file_size = path.stat().st_size
    if not skip_size_check and file_size > MAX_FILE_SIZE_BYTES:
        size_mb = file_size / (1024 * 1024)
        return False, f"Archivo demasiado grande: {size_mb:.1f} MB (máximo: {MAX_FILE_SIZE_MB} MB)"
    
    if file_size == 0:
        return False, "Archivo vacío"
    
    return True, ""


def transcribe_audio(
    file_path: str | Path,
    settings: AppSettings,
    diarize: bool = True,
    language: str = "es",
    speaker_refs: Optional[List[Tuple[str, str | Path]]] = None,
    _skip_size_check: bool = False,  # Internal: bypass size check for chunks
) -> TranscriptionResult:
    """
    Transcribe un archivo de audio usando Azure OpenAI.
    
    Args:
        file_path: Ruta al archivo de audio (MP3, WAV, etc.)
        settings: Configuración de la aplicación con credenciales Azure
        diarize: Si True, usa gpt-4o-transcribe-diarize para separar speakers
        language: Código de idioma (default: "es" para español)
        speaker_refs: Lista opcional de (nombre, ruta_audio) para identificación por voz.
                      Máximo 4 speakers, cada clip debe ser 2-10 segundos.
                      Ejemplo: [("Entrevistador", "ref_interviewer.mp3"), ("Entrevistado", "ref_subject.mp3")]
        
    Returns:
        TranscriptionResult con texto, segmentos y metadata
        
    Raises:
        ValueError: Si el archivo no es válido o no hay deployment configurado
        httpx.HTTPError: Si la llamada a la API falla
    """
    path = Path(file_path)
    
    # Validar archivo (skip size check if called from chunked transcription)
    is_valid, error_msg = validate_audio_file(path, skip_size_check=_skip_size_check)
    if not is_valid:
        raise ValueError(error_msg)
    
    # Seleccionar deployment
    if diarize:
        deployment = settings.azure.deployment_transcribe_diarize
        if not deployment:
            raise ValueError(
                "Deployment de transcripción con diarización no configurado. "
                "Configura AZURE_OPENAI_DEPLOYMENT_TRANSCRIBE_DIARIZE en .env"
            )
    else:
        deployment = settings.azure.deployment_transcribe
        if not deployment:
            raise ValueError(
                "Deployment de transcripción no configurado. "
                "Configura AZURE_OPENAI_DEPLOYMENT_TRANSCRIBE en .env"
            )
    
    # Construir URL de la API
    endpoint = settings.azure.endpoint.rstrip("/")
    api_version = settings.azure.transcribe_api_version
    url = f"{endpoint}/openai/deployments/{deployment}/audio/transcriptions?api-version={api_version}"
    
    logger.info(
        "transcription.start",
        file=path.name,
        diarize=diarize,
        deployment=deployment,
        size_mb=path.stat().st_size / (1024 * 1024),
        speaker_refs=len(speaker_refs) if speaker_refs else 0,
    )
    
    # Preparar request multipart
    with open(path, "rb") as audio_file:
        files = {
            "file": (path.name, audio_file, _get_mime_type(path.suffix)),
        }
        # Azure docs: model ES requerido en el body
        # Para gpt-4o-transcribe-diarize, chunking_strategy ES REQUERIDO
        data = {
            "model": deployment,
            "language": language,
        }
        
        # chunking_strategy es requerido para el modelo diarize
        if diarize:
            data["chunking_strategy"] = "auto"
            # gpt-4o-transcribe-diarize solo soporta 'json' o 'text', NO 'verbose_json'
            data["response_format"] = "json"
        
        # Agregar speaker reference clips si se proporcionaron
        if speaker_refs and diarize:
            speaker_refs = speaker_refs[:4]  # Máximo 4 speakers
            speaker_names = []
            speaker_data_urls = []
            
            for name, ref_path in speaker_refs:
                ref_path = Path(ref_path)
                if ref_path.exists():
                    # Convertir a data URL
                    data_url = _audio_to_data_url(ref_path)
                    if data_url:
                        speaker_names.append(name)
                        speaker_data_urls.append(data_url)
                        logger.debug("transcription.speaker_ref", name=name, path=str(ref_path))
            
            if speaker_names:
                # Agregar como arrays de formulario
                for i, name in enumerate(speaker_names):
                    data[f"known_speaker_names[{i}]"] = name
                    data[f"known_speaker_references[{i}]"] = speaker_data_urls[i]
                logger.info("transcription.speaker_refs_added", count=len(speaker_names))
        
        # Azure OpenAI with API key authenticates via `api-key`.
        headers: Dict[str, str] = {}
        if settings.azure.api_key:
            headers["api-key"] = settings.azure.api_key
        else:
            access_token = (os.getenv("AZURE_OPENAI_ACCESS_TOKEN") or "").strip()
            if access_token:
                headers["Authorization"] = f"Bearer {access_token}"
            else:
                raise ValueError(
                    "Missing Azure OpenAI credentials for transcription "
                    "(AZURE_OPENAI_API_KEY or AZURE_OPENAI_ACCESS_TOKEN)."
                )
        
        # Llamar a la API con reintentos y backoff para uploads grandes
        timeout_config = httpx.Timeout(900.0, connect=60.0, read=600.0)
        last_error: Optional[str] = None
        for attempt in range(3):
            try:
                with httpx.Client(timeout=timeout_config) as client:
                    response = client.post(url, files=files, data=data, headers=headers)
                request_id = response.headers.get("x-request-id") or response.headers.get("x-ms-request-id")

                if response.status_code == 200:
                    result_json = response.json()
                    break

                error_detail = response.text[:500] if response.text else "Unknown error"
                last_error = f"{response.status_code}: {error_detail}"
                logger.error(
                    "transcription.api_error",
                    status=response.status_code,
                    detail=error_detail,
                    request_id=request_id,
                    attempt=attempt + 1,
                )
            except (httpx.TimeoutException, httpx.HTTPError) as exc:
                last_error = str(exc)
                logger.warning(
                    "transcription.api_retry",
                    error=str(exc),
                    attempt=attempt + 1,
                )

            if attempt < 2:
                backoff = 2 ** attempt
                import time
                time.sleep(backoff)
        else:
            raise ValueError(f"API error after retries: {last_error}")


    
    # Parsear respuesta
    transcription = _parse_transcription_response(result_json, deployment, diarize)
    
    logger.info(
        "transcription.complete",
        file=path.name,
        speakers=transcription.speaker_count,
        segments=len(transcription.segments),
        duration=transcription.duration_seconds,
    )
    
    return transcription


def _audio_to_data_url(audio_path: Path) -> Optional[str]:
    """
    Convierte un archivo de audio a data URL para speaker reference.
    
    Args:
        audio_path: Ruta al archivo de audio (2-10 segundos)
        
    Returns:
        Data URL en formato data:audio/mp3;base64,... o None si falla
    """
    import base64
    
    if not audio_path.exists():
        return None
    
    # Verificar tamaño razonable (10 seg @ 64kbps ≈ 80KB max)
    file_size = audio_path.stat().st_size
    if file_size > 500 * 1024:  # 500KB max
        logger.warning("speaker_ref.too_large", path=str(audio_path), size_kb=file_size/1024)
        return None
    
    mime_type = _get_mime_type(audio_path.suffix)
    
    try:
        with open(audio_path, "rb") as f:
            audio_bytes = f.read()
        b64_content = base64.b64encode(audio_bytes).decode("utf-8")
        return f"data:{mime_type};base64,{b64_content}"
    except Exception as e:
        logger.error("speaker_ref.encode_failed", path=str(audio_path), error=str(e))
        return None


def transcribe_audio_chunked(
    file_path: str | Path,
    settings: AppSettings,
    diarize: bool = True,
    language: str = "es",
    chunk_duration_sec: int = MAX_CHUNK_DURATION_SEC,
    speaker_refs: Optional[List[Tuple[str, str | Path]]] = None,
) -> TranscriptionResult:
    """
    Transcribe archivos de audio grandes dividiéndolos en chunks con ffmpeg.
    
    Para archivos > 5 min o > 10 MB, divide el audio en chunks más pequeños,
    transcribe cada uno, y combina los resultados.
    
    Args:
        file_path: Ruta al archivo de audio
        settings: Configuración de la aplicación
        diarize: Si usar diarización
        language: Código de idioma
        chunk_duration_sec: Duración de cada chunk (default: 300s = 5 min)
        speaker_refs: Lista opcional de (nombre, ruta_audio) para identificación por voz.
        
    Returns:
        TranscriptionResult con todos los segmentos combinados
    """
    import shutil
    
    path = Path(file_path)
    
    # Verificar que el archivo existe
    if not path.exists():
        raise ValueError(f"Archivo no encontrado: {path}")
    
    # OPTIMIZACIÓN: Convertir a MP3 mono 16kHz para reducir tamaño de upload (~70% reducción)
    # Esto acelera significativamente el proceso al reducir datos a transmitir
    optimized_path = None
    try:
        original_size = path.stat().st_size / (1024 * 1024)
        optimized_path = optimize_audio_for_transcription(path)
        optimized_size = optimized_path.stat().st_size / (1024 * 1024)
        
        # Usar archivo optimizado si es significativamente más pequeño
        if optimized_path != path and optimized_size < original_size * 0.9:
            logger.info(
                "transcribe_chunked.optimized",
                original_mb=round(original_size, 2),
                optimized_mb=round(optimized_size, 2),
                reduction_pct=round((1 - optimized_size / original_size) * 100, 1),
            )
            path = optimized_path
        else:
            # Optimización no fue efectiva, limpiar y usar original
            if optimized_path != Path(file_path):
                optimized_path.unlink(missing_ok=True)
            optimized_path = None
    except Exception as e:
        logger.warning("transcribe_chunked.optimization_failed", error=str(e))
        # Continuar con archivo original si falla la optimización
    
    # Obtener duración
    duration = get_audio_duration(path)
    file_size_mb = path.stat().st_size / (1024 * 1024)
    
    logger.info(
        "transcribe_chunked.start",
        file=path.name,
        duration=duration,
        size_mb=file_size_mb,
    )
    
    # Si es pequeño, usar transcripción directa
    if duration <= chunk_duration_sec and file_size_mb <= MAX_FILE_SIZE_MB:
        logger.info("transcribe_chunked.direct", reason="small_file")
        result = transcribe_audio(path, settings, diarize=diarize, language=language, _skip_size_check=True)
        return result
    
    # Dividir el audio
    chunks = split_audio_ffmpeg(path, chunk_duration_sec=chunk_duration_sec)
    
    if len(chunks) == 1 and chunks[0] == path:
        # No se pudo dividir, intentar transcripción directa
        result = transcribe_audio(path, settings, diarize=diarize, language=language, _skip_size_check=True)
        return result
    
    # Transcribir cada chunk y combinar
    all_segments: List[TranscriptSegment] = []
    all_text_parts: List[str] = []
    speakers_seen: set = set()
    total_duration = 0.0
    model_name = "gpt-4o-transcribe-diarize" if diarize else "gpt-4o-transcribe"
    successful_chunks = 0
    
    for i, chunk_path in enumerate(chunks):
        offset_seconds = i * chunk_duration_sec
        
        logger.info("transcribe_chunked.chunk", index=i, total=len(chunks), offset=offset_seconds)
        
        try:
            chunk_result = transcribe_audio(chunk_path, settings, diarize=diarize, language=language, _skip_size_check=True)
            
            # Segmentos ajustados de este chunk
            chunk_adjusted_segments: List[TranscriptSegment] = []
            
            # Ajustar timestamps de los segmentos
            for seg in chunk_result.segments:
                adjusted_seg = TranscriptSegment(
                    speaker=seg.speaker,
                    text=seg.text,
                    start=seg.start + offset_seconds,
                    end=seg.end + offset_seconds,
                )
                all_segments.append(adjusted_seg)
                chunk_adjusted_segments.append(adjusted_seg)
                speakers_seen.add(seg.speaker)
            
            all_text_parts.append(chunk_result.text)
            total_duration = max(total_duration, chunk_result.duration_seconds + offset_seconds)
            successful_chunks += 1
            
        except Exception as e:
            logger.error("transcribe_chunked.chunk_error", chunk=i, error=str(e))
            # Continuar con los otros chunks
    
    # Limpiar chunks temporales
    if chunks[0] != path:
        try:
            shutil.rmtree(chunks[0].parent)
        except Exception as e:
            logger.warning(
                "transcribe_chunked.cleanup_failed",
                path=str(chunks[0].parent),
                error=str(e),
            )
    
    # Combinar resultados
    combined_text = "\n".join(all_text_parts)
    
    # POST-PROCESAMIENTO: Dividir segmentos largos para mejorar granularidad
    # Esto es especialmente importante cuando hay 1 sola voz detectada por chunk
    if all_segments:
        original_count = len(all_segments)
        all_segments = split_large_segments(all_segments, max_chars=1000, min_chars=200)
        if len(all_segments) > original_count:
            logger.info(
                "transcribe_chunked.segments_expanded",
                original=original_count,
                expanded=len(all_segments),
            )
    
    result = TranscriptionResult(
        text=combined_text,
        segments=all_segments,
        speaker_count=len(speakers_seen),
        duration_seconds=total_duration,
        model=model_name,
        metadata={
            "chunked": True,
            "chunk_count": len(chunks),
            "chunks_completed": successful_chunks,
            "language": language,
            "segments_expanded": len(all_segments) > len(chunks),
        },
    )
    
    logger.info(
        "transcribe_chunked.complete",
        chunks=len(chunks),
        segments=len(all_segments),
        speakers=len(speakers_seen),
    )
    
    # Limpiar archivo optimizado temporal
    if optimized_path and optimized_path != Path(file_path):
        try:
            optimized_path.unlink(missing_ok=True)
        except Exception as e:
            logger.warning(
                "transcribe_chunked.cleanup_optimized_failed",
                path=str(optimized_path),
                error=str(e),
            )
    
    return result



def _get_mime_type(suffix: str) -> str:
    """Retorna el MIME type para una extensión de archivo."""
    mime_types = {
        ".mp3": "audio/mpeg",
        ".mp4": "audio/mp4",
        ".m4a": "audio/mp4",
        ".wav": "audio/wav",
        ".webm": "audio/webm",
        ".mpeg": "audio/mpeg",
        ".mpga": "audio/mpeg",
        ".flac": "audio/flac",
        ".ogg": "audio/ogg",
    }
    return mime_types.get(suffix.lower(), "audio/mpeg")


def _parse_transcription_response(
    response: Dict[str, Any],
    model: str,
    has_diarization: bool,
) -> TranscriptionResult:
    """
    Parsea la respuesta JSON de la API de transcripción.
    
    La respuesta puede variar según el modelo:
    - Sin diarización: segments con start/end pero sin speaker
    - Con diarización: segments con speaker, start, end
    """
    text = response.get("text", "")
    duration = response.get("duration", 0.0)
    
    segments: List[TranscriptSegment] = []
    speakers_seen: set = set()
    
    # Parsear segments si existen
    raw_segments = response.get("segments", [])
    
    for seg in raw_segments:
        speaker = seg.get("speaker", "SPEAKER_00")
        seg_text = seg.get("text", "").strip()
        start = seg.get("start", 0.0)
        end = seg.get("end", 0.0)
        
        if seg_text:
            segments.append(TranscriptSegment(
                speaker=speaker,
                text=seg_text,
                start=start,
                end=end,
            ))
            speakers_seen.add(speaker)
    
    # Si no hay segments pero hay texto, crear uno solo
    if not segments and text:
        segments.append(TranscriptSegment(
            speaker="SPEAKER_00",
            text=text,
            start=0.0,
            end=duration,
        ))
        speakers_seen.add("SPEAKER_00")
    
    return TranscriptionResult(
        text=text,
        segments=segments,
        speaker_count=len(speakers_seen),
        duration_seconds=duration,
        model=model,
        metadata={
            "language": response.get("language", ""),
            "task": response.get("task", "transcribe"),
            "has_diarization": has_diarization,
        },
    )


def split_large_segments(
    segments: List[TranscriptSegment],
    max_chars: int = 1000,
    min_chars: int = 200,
) -> List[TranscriptSegment]:
    """
    Divide segmentos largos en segmentos más pequeños basados en oraciones.
    
    Esto mejora la granularidad cuando la diarización retorna pocos segmentos
    (típicamente cuando hay 1 solo speaker por chunk de audio).
    
    Args:
        segments: Lista de segmentos originales
        max_chars: Máximo de caracteres por segmento resultante
        min_chars: Mínimo de caracteres para crear un nuevo segmento
        
    Returns:
        Lista de segmentos, posiblemente más granular que la original
    """
    if not segments:
        return []
    
    result: List[TranscriptSegment] = []
    global_index = 0
    
    # Patrones para dividir oraciones
    sentence_endings = re.compile(r'(?<=[.!?])\s+')
    
    for seg in segments:
        text = seg.text.strip()
        
        # Si el segmento es pequeño, mantenerlo como está
        if len(text) <= max_chars:
            result.append(TranscriptSegment(
                speaker=seg.speaker,
                text=text,
                start=seg.start,
                end=seg.end,
                segment_index=global_index,
            ))
            global_index += 1
            continue
        
        # Dividir por oraciones
        sentences = sentence_endings.split(text)
        if len(sentences) <= 1:
            # No hay puntos, intentar dividir por comas o punto y coma
            sentences = re.split(r'(?<=[,;])\s+', text)
        
        # Agrupar oraciones hasta alcanzar max_chars
        current_text = ""
        current_start = seg.start
        duration_per_char = (seg.end - seg.start) / max(len(text), 1)
        char_position = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Si agregar esta oración excede el máximo, cerrar segmento actual
            if current_text and len(current_text) + len(sentence) + 1 > max_chars:
                if len(current_text) >= min_chars:
                    end_time = current_start + (len(current_text) * duration_per_char)
                    result.append(TranscriptSegment(
                        speaker=seg.speaker,
                        text=current_text,
                        start=round(current_start, 2),
                        end=round(end_time, 2),
                        segment_index=global_index,
                    ))
                    global_index += 1
                    current_start = end_time
                    current_text = sentence
                    char_position += len(sentence)
                else:
                    # Muy corto, agregar aunque exceda
                    current_text = f"{current_text} {sentence}".strip()
                    char_position += len(sentence) + 1
            else:
                if current_text:
                    current_text = f"{current_text} {sentence}"
                    char_position += len(sentence) + 1
                else:
                    current_text = sentence
                    char_position += len(sentence)
        
        # Último fragmento
        if current_text:
            result.append(TranscriptSegment(
                speaker=seg.speaker,
                text=current_text,
                start=round(current_start, 2),
                end=round(seg.end, 2),
                segment_index=global_index,
            ))
            global_index += 1
    
    logger.info(
        "split_large_segments.complete",
        original=len(segments),
        result=len(result),
        expansion_ratio=round(len(result) / max(len(segments), 1), 2),
    )
    
    return result


def assign_segment_indices(segments: List[TranscriptSegment]) -> List[TranscriptSegment]:
    """
    Asigna índices correlativos a los segmentos para mantener orden claro.
    
    Args:
        segments: Lista de segmentos
        
    Returns:
        Lista con segment_index asignado secuencialmente
    """
    for i, seg in enumerate(segments):
        seg.segment_index = i
    return segments


def transcription_to_paragraphs(
    result: TranscriptionResult,
) -> List[ParagraphRecord]:
    """
    Convierte TranscriptionResult a lista de ParagraphRecord.
    
    Cada segmento se convierte en un ParagraphRecord con el speaker
    mapeado a un formato consistente.
    
    Args:
        result: Resultado de la transcripción
        
    Returns:
        Lista de ParagraphRecord listos para coalescencia
    """
    paragraphs: List[ParagraphRecord] = []
    
    for segment in result.segments:
        # Normalizar speaker a formato lowercase con underscore
        # SPEAKER_00 -> speaker_00
        speaker = segment.speaker.lower().replace(" ", "_")
        
        # Texto con prefijo de timestamp para referencia
        text = segment.text.strip()
        if not text:
            continue
        
        paragraphs.append(ParagraphRecord(
            text=text,
            speaker=speaker,
        ))
    
    return paragraphs


def audio_to_fragments(
    file_path: str | Path,
    settings: AppSettings,
    diarize: bool = True,
    min_chars: int = 200,
    max_chars: int = 1200,
    min_interviewee_tokens: int = 10,
) -> List[FragmentRecord]:
    """
    Pipeline completo: audio → fragmentos listos para ingesta.
    
    Combina transcripción, conversión a paragraphs y coalescencia
    en un solo flujo.
    
    Args:
        file_path: Ruta al archivo de audio
        settings: Configuración de la aplicación
        diarize: Si usar diarización
        min_chars: Mínimo de caracteres por fragmento
        max_chars: Máximo de caracteres por fragmento
        min_interviewee_tokens: Mínimo de tokens del entrevistado
        
    Returns:
        Lista de FragmentRecord listos para ingestión
    """
    # 1. Transcribir
    transcription = transcribe_audio(file_path, settings, diarize=diarize)
    
    # 2. Convertir a paragraphs
    paragraphs = transcription_to_paragraphs(transcription)
    
    if not paragraphs:
        logger.warning("audio_to_fragments.empty", file=str(file_path))
        return []
    
    # 3. Coalescencia (agrupar en fragmentos óptimos)
    fragments = coalesce_paragraph_records(
        paragraphs,
        min_chars=min_chars,
        max_chars=max_chars,
        min_interviewee_tokens=min_interviewee_tokens,
    )
    
    logger.info(
        "audio_to_fragments.complete",
        file=Path(file_path).name,
        paragraphs=len(paragraphs),
        fragments=len(fragments),
        speakers=transcription.speaker_count,
    )
    
    return fragments


def save_transcription_docx(
    result: TranscriptionResult,
    output_path: str | Path,
) -> Path:
    """
    Guarda la transcripción como archivo DOCX con marcas de speaker.
    
    Formato:
        [SPEAKER_00] 00:00:05
        Texto del primer segmento...
        
        [SPEAKER_01] 00:00:15
        Texto del segundo segmento...
    
    Args:
        result: Resultado de la transcripción
        output_path: Ruta donde guardar el DOCX
        
    Returns:
        Path al archivo creado
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise ImportError("python-docx es requerido. Instala con: pip install python-docx")
    
    doc = Document()
    path = Path(output_path)
    
    # Título
    title = doc.add_heading(f"Transcripción: {path.stem}", level=1)
    
    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Speakers detectados: {result.speaker_count}\n").bold = True
    meta_para.add_run(f"Duración: {_format_duration(result.duration_seconds)}\n")
    meta_para.add_run(f"Modelo: {result.model}\n")
    meta_para.add_run(f"Total segmentos: {len(result.segments)}\n")
    doc.add_paragraph()  # Espacio
    
    # Segmentos con numeración correlativa
    current_speaker = None
    for i, segment in enumerate(result.segments):
        # Usar segment_index si está disponible, sino usar el índice del loop
        seg_num = segment.segment_index if hasattr(segment, 'segment_index') and segment.segment_index > 0 else i
        
        # Cabecera con número de segmento, speaker y timestamp
        timestamp = _format_timestamp(segment.start)
        
        header = doc.add_paragraph()
        # Formato: [#003] [SPEAKER_00] 00:00:15
        run = header.add_run(f"[#{seg_num + 1:03d}] [{segment.speaker}] {timestamp}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)  # Azul
        
        # Texto del segmento
        doc.add_paragraph(segment.text)
    
    # Guardar
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    
    logger.info("transcription.saved_docx", path=str(path))
    
    return path


def _format_timestamp(seconds: float) -> str:
    """Formatea segundos como HH:MM:SS o MM:SS."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _format_duration(seconds: float) -> str:
    """Formatea duración total en formato legible."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    
    if hours > 0:
        return f"{hours}h {minutes}m {secs}s"
    if minutes > 0:
        return f"{minutes}m {secs}s"
    return f"{secs}s"
