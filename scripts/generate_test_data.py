"""
Generate synthetic interview dataset for load testing.

Usage:
    python scripts/generate_test_data.py --count 100 --output data/test_interviews
"""

import argparse
import random
from pathlib import Path
from docx import Document


# Sample interview content patterns
INTERVIEW_TOPICS = [
    "comunidad", "educación", "salud", "vivienda", "trabajo", 
    "familia", "cultura", "medio ambiente", "transporte", "seguridad"
]

INTERVIEWER_QUESTIONS = [
    "¿Podría contarme sobre su experiencia con {topic}?",
    "¿Cómo ha sido su relación con la {topic} en los últimos años?",
    "¿Qué desafíos ha enfrentado en relación a {topic}?",
    "¿Qué cambios ha observado en {topic} recientemente?",
    "¿Cómo cree que se podría mejorar la situación de {topic}?",
]

RESPONSE_TEMPLATES = [
    "Bueno, respecto a {topic}, yo diría que ha sido un proceso bastante interesante. "
    "En mi experiencia, hemos visto muchos cambios en los últimos años. "
    "Por ejemplo, antes las cosas eran muy diferentes, pero ahora hay más oportunidades. "
    "Aunque también existen desafíos, como la falta de recursos y apoyo institucional.",
    
    "Mire, el tema de {topic} es algo que nos afecta a todos en la comunidad. "
    "Yo personalmente he trabajado en esto por varios años y puedo decir que hay avances. "
    "Sin embargo, todavía queda mucho por hacer. Las autoridades podrían hacer más, "
    "y nosotros como vecinos también tenemos que organizarnos mejor.",
    
    "La verdad es que {topic} ha sido central en nuestras vidas. "
    "Recuerdo cuando éramos jóvenes, las cosas eran muy distintas. "
    "Ahora los jóvenes tienen otras prioridades, pero igual es importante mantener las tradiciones. "
    "Creo que el balance entre lo nuevo y lo tradicional es clave.",
    
    "Sobre {topic}, tengo bastante que decir. He estado involucrado en varios proyectos "
    "y he visto cómo la situación ha evolucionado. Hay cosas positivas, como el aumento "
    "de la participación ciudadana, pero también hay problemas estructurales que persisten. "
    "Necesitamos políticas públicas más efectivas y mejor coordinación entre actores.",
]


def generate_paragraph(topic: str, is_interviewer: bool) -> str:
    """Generate a single interview paragraph."""
    if is_interviewer:
        template = random.choice(INTERVIEWER_QUESTIONS)
        return f"Entrevistador: {template.format(topic=topic)}"
    else:
        template = random.choice(RESPONSE_TEMPLATES)
        # Add some variation
        prefix = random.choice(["", "Sí, ", "Claro, ", "Bueno, "])
        return f"Entrevistado: {prefix}{template.format(topic=topic)}"


def generate_interview(num_exchanges: int = 15) -> Document:
    """Generate a complete interview document."""
    doc = Document()
    
    # Header
    doc.add_heading("Transcripción de Entrevista", level=1)
    doc.add_paragraph(f"Tema principal: {random.choice(INTERVIEW_TOPICS)}")
    doc.add_paragraph("---")
    
    for i in range(num_exchanges):
        topic = random.choice(INTERVIEW_TOPICS)
        
        # Interviewer question
        question = generate_paragraph(topic, is_interviewer=True)
        doc.add_paragraph(question)
        
        # Interviewee response (sometimes multiple paragraphs)
        num_responses = random.randint(1, 3)
        for _ in range(num_responses):
            response = generate_paragraph(topic, is_interviewer=False)
            doc.add_paragraph(response)
    
    return doc


def main():
    parser = argparse.ArgumentParser(description="Generate synthetic interview data")
    parser.add_argument("--count", type=int, default=10, help="Number of interviews to generate")
    parser.add_argument("--output", type=str, default="data/test_interviews", help="Output directory")
    parser.add_argument("--min-exchanges", type=int, default=10, help="Minimum Q&A exchanges per interview")
    parser.add_argument("--max-exchanges", type=int, default=25, help="Maximum Q&A exchanges per interview")
    
    args = parser.parse_args()
    
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"Generating {args.count} synthetic interviews...")
    
    for i in range(args.count):
        num_exchanges = random.randint(args.min_exchanges, args.max_exchanges)
        doc = generate_interview(num_exchanges)
        
        filename = f"Entrevista_Sintetica_{i+1:03d}.docx"
        filepath = output_dir / filename
        doc.save(str(filepath))
        
        if (i + 1) % 10 == 0:
            print(f"  Generated {i + 1}/{args.count} interviews...")
    
    print(f"Done! Files saved to: {output_dir}")
    print(f"Total files: {args.count}")


if __name__ == "__main__":
    main()
