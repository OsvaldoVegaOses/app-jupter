"""
Aplicación principal FastAPI - API REST para análisis cualitativo.

Este es el punto de entrada del servidor HTTP que expone ~54 endpoints
para interactuar con el sistema de análisis cualitativo GraphRAG.

Grupos de endpoints:
    
    /token
        - POST: Obtener JWT token (login)
    
    /health
        - GET: Health check del servicio
    
    /api/projects
        - GET: Listar proyectos
        - POST: Crear nuevo proyecto
    
    /api/status/{project}
        - GET: Estado del proyecto (etapas completadas)
    
    /api/ingest
        - POST: Ingestar documentos DOCX
    
    /api/analyze
        - POST: Ejecutar análisis LLM
        - POST /persist: Persistir resultados de análisis
        - GET /task/{id}: Estado de tarea Celery
    
    /api/neo4j
        - POST /query: Ejecutar consulta Cypher
        - POST /export: Exportar resultados (CSV/JSON)

    /api/axial
        - POST /gds: Ejecutar algoritmos GDS (Louvain/PageRank/Betweenness)
    
    /api/coding
        - POST /assign: Asignar código a fragmento
        - POST /suggest: Sugerir fragmentos similares
        - GET /stats: Estadísticas de codificación
        - GET /codes: Lista de códigos
        - GET /fragments: Fragmentos por archivo
        - GET /citations: Citas por código

    /api/interviews
        - GET: Lista de entrevistas (archivos) por proyecto
    
    /api/search
        - POST /discover: Búsqueda conceptual con ejemplos
        - POST /similar: Fragmentos similares
    
Autenticación:
    Todos los endpoints (excepto /health y /token) requieren:
    - Authorization: Bearer <jwt_token>
    - O: X-API-Key: <api_key>

Dependencias FastAPI:
    - get_settings(): Carga configuración (cacheada)
    - get_clients(): Construye ServiceClients
    - get_neo4j_clients(): Clientes solo para Neo4j
    - require_auth(): Valida autenticación

Configuración CORS:
    Permite todos los orígenes (*) para desarrollo.
    En producción, restringir a dominios específicos.

Logging:
    - Estructlog para logging JSON
    - Eventos: api.request, api.error, etc.

Ejecución:
    uvicorn backend.app:app --reload --port 8000
"""

from __future__ import annotations

import os
import json
import re
import unicodedata
import uuid
import asyncio
from collections import Counter
from collections.abc import AsyncGenerator
from dataclasses import dataclass
from functools import lru_cache
from io import StringIO
from pathlib import Path
import tempfile
from time import perf_counter
from typing import Any, Dict, List, Optional, Tuple, Union, cast, LiteralString

import structlog
from rapidfuzz.distance import Levenshtein
from pydantic import BaseModel, ConfigDict, Field
import openai

try:
    from fastapi import Depends, FastAPI, HTTPException, Header, Query, Body, status, BackgroundTasks  # type: ignore[import]
    from fastapi.middleware.cors import CORSMiddleware  # type: ignore[import]
    from fastapi.responses import JSONResponse, PlainTextResponse, Response, FileResponse  # type: ignore[import]
except ImportError as exc:  # pragma: no cover - surface a clearer message when FastAPI is missing
    raise RuntimeError(
        "FastAPI is required to run the backend service. Install it with `pip install fastapi`."
    ) from exc

from fastapi.security import OAuth2PasswordRequestForm
from jose import jwt
from datetime import datetime, timedelta
from backend.auth import SECRET_KEY, ALGORITHM, ACCESS_TOKEN_EXPIRE_MINUTES, User, get_current_user, require_role
from backend.celery_worker import task_analyze_interview, celery_app
from backend.routers.ingest import IngestRequest
from celery.result import AsyncResult

from neo4j import GraphDatabase, Driver  # type: ignore[import]
from neo4j import Query as Neo4jQuery  # Renamed to avoid collision with fastapi.Query

from app.analysis import analyze_interview_text, persist_analysis, matriz_etapa3, matriz_etapa4, modelo_ascii
from app.clients import (
    ServiceClients,
    build_service_clients,
    get_pg_connection,
    return_pg_connection,
    PGConnection,
)
from app.postgres_block import stage0_require_ready_or_override
from app.coding import (
    assign_open_code,
    citations_for_code,
    coding_statistics,
    export_codes_maxqda_csv,
    export_codes_refi_qda,
    fetch_code_history,
    fetch_fragment_context,
    find_similar_codes,
    get_saturation_data,
    list_available_interviews,
    list_available_interviews_with_ranking_debug,
    list_interview_fragments,
    list_open_codes,
    suggest_similar_fragments,
    unassign_open_code,
    CodingError,
)
from app.axial import run_gds_analysis, AxialError, AxialNotReadyError
from app.documents import load_fragments
from app.ingestion import ingest_documents
from app.project_state import (
    DEFAULT_PROJECT,
    create_project,
    detect_stage_status,
    list_projects,
    mark_stage,
    resolve_project,
    get_project_config,
    update_project_config,
    update_project_details,
)
from app.queries import run_cypher, sample_postgres
from app.settings import AppSettings, load_settings
from app.logging_config import configure_logging
from app.qdrant_block import discover_search, search_similar
from app.embeddings import embed_batch
from app.postgres_block import (
    ensure_candidate_codes_table,
    ensure_open_coding_table,
    ensure_axial_table,
    coding_stats,
    insert_candidate_codes,
    list_candidate_codes,
    validate_candidate,
    reject_candidate,
    merge_candidates,
    merge_candidates_by_code,
    promote_to_definitive,
    get_candidate_stats_by_source,
    get_canonical_examples,
    get_backlog_health,
    get_dashboard_counts,
    add_project_member,
)

from qdrant_client.models import ContextExamplePair, FieldCondition, Filter, FilterSelector, MatchValue
from qdrant_client import models
from main import STAGE_DEFINITIONS, STAGE_ORDER

# Backend routers - ALL 7 ROUTERS
from backend.routers.admin import router as admin_router, health_router
from backend.routers.auth import router as auth_router, oauth_router, legacy_router as auth_legacy_router
from backend.routers.neo4j import router as neo4j_router
from backend.routers.discovery import router as discovery_router, qdrant_router
from backend.routers.graphrag import graphrag_router, axial_router
from backend.routers.coding import router as coding_router, codes_router
from backend.routers.dashboard import router as dashboard_router
from backend.routers.ingest import router as ingest_router
from backend.routers.interviews import router as interviews_router
from backend.routers.familiarization import router as familiarization_router
from backend.routers.agent import router as agent_router  # Autonomous Agent
from backend.routers.stage0 import router as stage0_router
from backend.routers.nucleus import router as nucleus_router


# Initialize Logging (honor env override for verbosity)
configure_logging(os.getenv("LOG_LEVEL", "INFO").upper())


ALLOWED_FORMATS = {"raw", "table", "graph"}
DEFAULT_FORMATS = ["raw"]
ENV_FILE_VAR = "APP_ENV_FILE"
API_KEY_HEADER = "X-API-Key"

logger = structlog.get_logger("neo4j.api")
api_logger = structlog.get_logger("app.api")

# Load .env early so that any import-time code that reads os.getenv() sees values.
# Honor APP_ENV_FILE if set; otherwise try to locate a .env in the cwd, then fall
# back to the repo root (so reloads don't lose config due to cwd changes).
try:
    from dotenv import load_dotenv, find_dotenv

    _env_file = os.getenv(ENV_FILE_VAR) or find_dotenv(usecwd=True)
    if not _env_file:
        try:
            from pathlib import Path

            candidate = Path(__file__).resolve().parents[1] / ".env"
            if candidate.exists():
                _env_file = str(candidate)
        except Exception:
            _env_file = None
    if _env_file:
        try:
            load_dotenv(_env_file)
        except Exception:
            # best-effort: failures here shouldn't crash the API import
            logger.debug("failed_to_load_env_file", env_file=str(_env_file))
except Exception:
    # dotenv not installed or other error — proceed; load_settings will still attempt to load later
    pass


class CypherRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    cypher: str = Field(..., description="Consulta Cypher a ejecutar en Neo4j.")
    project: Optional[str] = Field(
        default=None, description="Proyecto requerido para la consulta."
    )
    params: Optional[Dict[str, Any]] = Field(
        default=None, description="Diccionario de parámetros (clave->valor)."
    )
    database: Optional[str] = Field(
        default=None, description="Base de datos Neo4j opcional."
    )
    formats: Optional[List[str]] = Field(
        default=None, description="Formatos de respuesta requeridos (raw, table, graph)."
    )


class MaintenanceDeleteFileRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="Proyecto (nombre lógico).")
    file: str = Field(..., min_length=1, description="Nombre de archivo de entrevista (e.g. .docx).")





class CodeSuggestionRequest(BaseModel):
    fragment_text: str = Field(..., description="Texto del fragmento a analizar")
    limit: int = Field(5, ge=1, le=20)

    database: Optional[str] = Field(
        default=None, description="Base de datos Neo4j opcional."
    )


class AnalyzeHiddenRelationshipsRequest(BaseModel):
    """Request para análisis IA de relaciones ocultas."""
    model_config = ConfigDict(extra="forbid")
    project: str = Field(default="default", description="ID del proyecto")
    suggestions: List[Dict[str, Any]] = Field(..., description="Relaciones ocultas sugeridas")
    persist: bool = Field(
        default=True,
        description="Si True, persiste el análisis como artefacto auditable (axial_ai_analyses).",
    )


class HiddenRelationshipsMetricsRequest(BaseModel):
    """Request para métricas de diversidad/overlap en relaciones ocultas."""
    model_config = ConfigDict(extra="forbid")
    project: str = Field(default="default", description="ID del proyecto")
    suggestions: List[Dict[str, Any]] = Field(..., description="Relaciones ocultas sugeridas")



class CypherResponse(BaseModel):
    model_config = ConfigDict(extra="forbid")

    raw: Optional[List[Dict[str, Any]]] = Field(
        default=None, description="Filas en formato raw (dict por registro)."
    )
    table: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Vista tabular con columnas y filas.",
    )
    graph: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Nodos y relaciones en formato de grafo.",
    )


app = FastAPI(title="Neo4j Query API", version="1.0.0")

# CORS NOTE:
# - Browsers will *reject* responses when `Access-Control-Allow-Origin` is missing.
# - Using `allow_origins=["*"]` together with `allow_credentials=True` is an invalid combination
#   per the CORS spec and will often cause the middleware to omit the header.
#
# Configure via `CORS_ALLOW_ORIGINS` (comma-separated). Default targets local dev frontends.
_cors_allow_origins_raw = os.getenv("CORS_ALLOW_ORIGINS", "").strip()
if _cors_allow_origins_raw:
    _cors_allow_origins = [o.strip().rstrip("/") for o in _cors_allow_origins_raw.split(",") if o.strip()]
    _cors_allow_origin_regex = os.getenv("CORS_ALLOW_ORIGIN_REGEX", "").strip() or None
else:
    _cors_allow_origins = [
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        "http://localhost:5174",
        "http://127.0.0.1:5174",
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ]
    # Dev-friendly: allow any localhost/127.0.0.1 port (covers Vite auto-fallback 5174, 5175, etc.)
    _cors_allow_origin_regex = r"^https?://(localhost|127\.0\.0\.1)(:\\d+)?$"

_cors_allow_credentials = os.getenv("CORS_ALLOW_CREDENTIALS", "true").strip().lower() in {
    "1",
    "true",
    "yes",
}
if "*" in _cors_allow_origins and _cors_allow_credentials:
    # Prevent a misconfiguration that results in missing CORS headers.
    _cors_allow_credentials = False

app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_allow_origins,
    allow_origin_regex=_cors_allow_origin_regex,
    allow_credentials=_cors_allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Include ALL routers - REFACTORING 100% COMPLETE
app.include_router(health_router)       # /healthz endpoint
app.include_router(admin_router)        # /api/* admin + insights
app.include_router(dashboard_router)    # /api/status, /api/dashboard/*
app.include_router(oauth_router)        # /token OAuth2 endpoint
app.include_router(auth_router)         # /api/auth/* endpoints
app.include_router(auth_legacy_router)  # /register, /refresh legacy endpoints
app.include_router(neo4j_router)        # /api/neo4j/* endpoints
app.include_router(qdrant_router)       # /api/qdrant/* endpoints  
app.include_router(discovery_router)    # /api/discovery/* endpoints
app.include_router(graphrag_router)     # /api/graphrag/* endpoints
app.include_router(axial_router)        # /api/axial/* endpoints
app.include_router(coding_router)       # /api/coding/* endpoints
app.include_router(codes_router)        # /api/codes/* endpoints
app.include_router(ingest_router)       # /api/ingest, /api/upload-and-ingest, /api/transcribe/*
app.include_router(interviews_router)   # /api/interviews/* (decoupled pipeline)
app.include_router(familiarization_router)  # /api/familiarization/reviews
app.include_router(agent_router)        # /api/agent/* (autonomous agent)
app.include_router(stage0_router)       # /api/stage0/* (Etapa 0: Preparación)
app.include_router(nucleus_router)      # /api/nucleus/* (lightweight nucleus endpoints)


# =============================================================================
# RATE LIMITING con Redis
# =============================================================================

from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from slowapi.middleware import SlowAPIMiddleware
from fastapi import Request

# Configurar limiter con Redis (usa el mismo que Celery)
REDIS_URL = os.getenv("CELERY_BROKER_URL", "redis://localhost:6379/0")
limiter = Limiter(
    key_func=get_remote_address,
    storage_uri=REDIS_URL,
    default_limits=["100/minute"],  # Límite general
)
app.state.limiter = limiter


async def rate_limit_exceeded_handler(request: Request, exc: Exception) -> Response:
    """Wrapper que satisface FastAPI typing y delega en slowapi."""
    if isinstance(exc, RateLimitExceeded):
        return _rate_limit_exceeded_handler(request, exc)
    raise exc


app.add_exception_handler(RateLimitExceeded, rate_limit_exceeded_handler)


# =============================================================================
# REQUEST ID MIDDLEWARE (Observability)
# =============================================================================

from starlette.middleware.base import BaseHTTPMiddleware

LOG_SCHEMA_VERSION = "1.0"
BUILD_VERSION = os.getenv("BUILD_VERSION") or os.getenv("GIT_SHA") or "unknown"

class RequestIdMiddleware(BaseHTTPMiddleware):
    """Adds unique request_id to each request for tracing/debugging."""
    
    async def dispatch(self, request: Request, call_next):
        request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())
        session_id = request.headers.get("X-Session-ID") or request.query_params.get("session_id")
        project_id = (
            request.headers.get("X-Project-ID")
            or request.query_params.get("project")
            or request.query_params.get("project_id")
        )
        test_run_id = request.headers.get("X-Test-Run-Id") or request.query_params.get("test_run_id")
        route = request.scope.get("route")
        route_template = getattr(route, "path", None) if route else None
        sample_rate = 1.0
        start = perf_counter()
        
        # Bind request_id to structlog context
        structlog.contextvars.clear_contextvars()
        bind_payload = {
            "request_id": request_id,
            "schema_version": LOG_SCHEMA_VERSION,
            "build_version": BUILD_VERSION,
        }
        if session_id:
            bind_payload["session_id"] = session_id
        if project_id:
            bind_payload["project_id"] = project_id
        if test_run_id:
            bind_payload["test_run_id"] = test_run_id
        structlog.contextvars.bind_contextvars(**bind_payload)
        
        # Store in request state for endpoint access
        request.state.request_id = request_id
        if session_id:
            request.state.session_id = session_id
        if project_id:
            request.state.project_id = project_id
        if test_run_id:
            request.state.test_run_id = test_run_id
        request.state.route_template = route_template
        
        # Log request start
        api_logger.info(
            "request.start",
            schema_version=LOG_SCHEMA_VERSION,
            build_version=BUILD_VERSION,
            method=request.method,
            path=str(request.url.path),
            route=route_template or str(request.url.path),
            request_id=request_id,
            test_run_id=test_run_id,
            sample_rate=sample_rate,
        )
        
        response = await call_next(request)
        
        # Add request_id to response headers
        response.headers["X-Request-ID"] = request_id
        
        duration_ms = (perf_counter() - start) * 1000

        # Log request end
        api_logger.info(
            "request.end",
            schema_version=LOG_SCHEMA_VERSION,
            build_version=BUILD_VERSION,
            method=request.method,
            path=str(request.url.path),
            route=route_template or str(request.url.path),
            status_code=response.status_code,
            request_id=request_id,
            test_run_id=test_run_id,
            duration_ms=round(duration_ms, 2),
            sample_rate=sample_rate,
        )

        if duration_ms >= 5000:
            api_logger.warning(
                "request.slow",
                schema_version=LOG_SCHEMA_VERSION,
                build_version=BUILD_VERSION,
                method=request.method,
                path=str(request.url.path),
                route=route_template or str(request.url.path),
                request_id=request_id,
                test_run_id=test_run_id,
                duration_ms=round(duration_ms, 2),
                sample_rate=sample_rate,
            )
        
        return response

app.add_middleware(RequestIdMiddleware)


# =============================================================================
# NOTA: Los endpoints de autenticación están definidos después de get_settings
# en las líneas 700+. Ver:
# - POST /api/auth/login  (JSON)
# - POST /api/auth/register  (JSON)
# - POST /api/auth/refresh
# - POST /token (OAuth2 form-data)
# - POST /register
# =============================================================================


@lru_cache(maxsize=1)
def get_settings() -> AppSettings:
    env_file = os.getenv(ENV_FILE_VAR)
    if not env_file:
        try:
            from pathlib import Path

            candidate = Path(__file__).resolve().parents[1] / ".env"
            if candidate.exists():
                env_file = str(candidate)
        except Exception:
            env_file = None
    return load_settings(env_file)


@dataclass
class Neo4jOnlyClients:
    neo4j: Driver

    def close(self) -> None:
        try:
            self.neo4j.close()
        except Exception:
            pass


def build_neo4j_only(settings: AppSettings) -> Neo4jOnlyClients:
    driver = GraphDatabase.driver(
        settings.neo4j.uri,
        auth=(settings.neo4j.username, settings.neo4j.password or ""),
    )
    return Neo4jOnlyClients(neo4j=driver)


@dataclass
class PgOnlyClients:
    postgres: PGConnection

    def close(self) -> None:
        try:
            return_pg_connection(self.postgres)
        except Exception:
            pass


def build_pg_only(settings: AppSettings) -> PgOnlyClients:
    return PgOnlyClients(postgres=get_pg_connection(settings))


def build_clients_or_error(settings: AppSettings) -> ServiceClients:
    try:
        return build_service_clients(settings)
    except Exception as exc:  # noqa: BLE001
        from app.error_handling import api_error, ErrorCode
        raise api_error(
            status_code=502,
            code=ErrorCode.SERVICE_UNAVAILABLE,
            message="Error conectando con servicios externos. Intente más tarde.",
            exc=exc,
        ) from exc


def _normalize_formats(values: Optional[List[str]]) -> List[str]:
    if not values:
        return DEFAULT_FORMATS.copy()
    normalized: List[str] = []
    for entry in values:
        fmt = (entry or "").strip().lower()
        if fmt == "all":
            return list(ALLOWED_FORMATS)
        if fmt not in ALLOWED_FORMATS:
            raise ValueError(
                f"Formato '{entry}' no soportado. Usa un formato válido: raw, table, graph o all."
            )
        if fmt not in normalized:
            normalized.append(fmt)
    if not normalized:
        raise ValueError("Debe especificar al menos un formato válido (raw, table, graph, all).")
    return normalized


async def get_clients(
    settings: AppSettings = Depends(get_settings),
) -> AsyncGenerator[ServiceClients, None]:
    clients = build_clients_or_error(settings)
    try:
        yield clients
    finally:
        clients.close()



async def require_auth(
    user: User = Depends(get_current_user),
) -> User:
    return user


@app.post("/api/axial/analyze-hidden-relationships")
async def api_analyze_hidden_relationships(
    payload: Optional[AnalyzeHiddenRelationshipsRequest] = Body(default=None),
    payload_q: Optional[str] = Query(default=None, alias="payload"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Analiza relaciones ocultas con IA.

    Devuelve memo estructurado con estatus epistemológico.
    """
    from app.clients import build_service_clients

    if payload is None and payload_q:
        try:
            payload = AnalyzeHiddenRelationshipsRequest.model_validate_json(payload_q)
        except Exception as exc:
            raise HTTPException(status_code=400, detail="Invalid payload JSON") from exc
    if payload is None:
        raise HTTPException(status_code=422, detail="Missing request body")

    clients = build_service_clients(settings)
    try:
        # Multi-tenant: resolver project_id canónico + validar acceso.
        try:
            project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        # Epistemic mode → prompts versionados (audit trail)
        from app.postgres_block import get_project_epistemic_mode
        from app.prompts.loader import get_system_prompt

        epistemic_mode = get_project_epistemic_mode(clients.postgres, project_id)
        system_prompt, prompt_version = get_system_prompt(epistemic_mode, "axial_coding")

        api_logger.info(
            "api.analyze_hidden_relationships.start",
            project=project_id,
            epistemic_mode=epistemic_mode.value,
            prompt_version=prompt_version,
            suggestions_count=len(payload.suggestions or []),
        )

        MAX_SUGGESTIONS = 10
        POS_TOTAL = 24
        NEG_TOTAL = 12
        EXCERPT_CHARS = 240
        POS_IN_PROMPT = 3
        NEG_IN_PROMPT = 2

        suggestions_for_prompt = payload.suggestions[:MAX_SUGGESTIONS]

        # Evidence pack (positivo/negativo) para grounding + auditoria reproducible.
        from datetime import datetime, timezone

        evidence_pack: Dict[str, Any]
        try:
            from app.axial_evidence import build_link_prediction_evidence_pack

            evidence_pack = build_link_prediction_evidence_pack(
                clients,
                settings,
                project_id=project_id,
                suggestions=suggestions_for_prompt,
                positive_total=POS_TOTAL,
                negative_total=NEG_TOTAL,
                excerpt_chars=EXCERPT_CHARS,
            )
        except Exception as exc:  # noqa: BLE001
            api_logger.warning(
                "api.analyze_hidden_relationships.evidence_failed",
                project=project_id,
                error=str(exc)[:200],
            )
            now = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
            evidence_pack = {
                "schema_version": 1,
                "generated_at": now,
                "limits": {
                    "positive_total": POS_TOTAL,
                    "negative_total": NEG_TOTAL,
                    "excerpt_chars": EXCERPT_CHARS,
                },
                "suggestions": [],
                "totals": {"positive": 0, "negative": 0, "by_method": {}},
                "notes": {"reason": "build_failed"},
            }

        evidence_by_id: Dict[int, Dict[str, Any]] = {}
        try:
            if isinstance(evidence_pack, dict) and isinstance(evidence_pack.get("suggestions"), list):
                for item in evidence_pack["suggestions"]:
                    if not isinstance(item, dict):
                        continue
                    try:
                        sid = int(item.get("id") or 0)
                    except Exception:
                        continue
                    if sid > 0:
                        evidence_by_id[sid] = item
        except Exception:
            evidence_by_id = {}

        def _one_line(value: Any) -> str:
            return str(value or "").replace("\r", " ").replace("\n", " ").strip()

        # Preparar contexto de sugerencias (+ evidencia resumida) para el prompt.
        suggestions_text: List[str] = []
        for i, sug in enumerate(suggestions_for_prompt, 1):
            source = sug.get("source", "?")
            target = sug.get("target", "?")
            score = sug.get("score", 0)
            reason = _one_line(sug.get("reason", ""))
            try:
                score_f = float(score or 0.0)
            except Exception:
                score_f = 0.0

            block_lines = [f"{i}. {source} ↔ {target} (score: {score_f:.3f})"]
            if reason:
                block_lines.append(f"   RAZON: {reason}")

            ev = evidence_by_id.get(i) or {}
            pos_items = ev.get("positive") if isinstance(ev.get("positive"), list) else []
            neg_items = ev.get("negative") if isinstance(ev.get("negative"), list) else []

            if pos_items:
                block_lines.append("   EVIDENCIA_POSITIVA:")
                for e in pos_items[:POS_IN_PROMPT]:
                    if not isinstance(e, dict):
                        continue
                    fid = _one_line(e.get("fragmento_id"))
                    archivo = _one_line(e.get("archivo"))
                    par_idx = e.get("par_idx")
                    frag = _one_line(e.get("fragmento"))
                    method = _one_line(e.get("method"))
                    block_lines.append(
                        f"   - [{fid}] {archivo}:{par_idx} \"{frag}\" ({method})".strip()
                    )
            if neg_items:
                block_lines.append("   EVIDENCIA_NEGATIVA (tension):")
                for e in neg_items[:NEG_IN_PROMPT]:
                    if not isinstance(e, dict):
                        continue
                    fid = _one_line(e.get("fragmento_id"))
                    archivo = _one_line(e.get("archivo"))
                    par_idx = e.get("par_idx")
                    frag = _one_line(e.get("fragmento"))
                    method = _one_line(e.get("method"))
                    ps = e.get("present_source")
                    pt = e.get("present_target")
                    present = ""
                    if ps is not None or pt is not None:
                        present = f" present(source={int(bool(ps))},target={int(bool(pt))})"
                    block_lines.append(
                        f"   - [{fid}] {archivo}:{par_idx}{present} \"{frag}\" ({method})".strip()
                    )

            suggestions_text.append("\n".join(block_lines))

        suggestions_block = "\n\n".join(suggestions_text)

        prompt = f"""Analiza las siguientes relaciones ocultas detectadas en un grafo de analisis cualitativo (Teoria Fundamentada).

Para cada sugerencia se incluye evidencia en fragmentos (EVIDENCIA_POSITIVA / EVIDENCIA_NEGATIVA).
Cita la evidencia usando los `fragmento_id` proporcionados. No inventes IDs.

Sugerencias de relaciones ocultas (IDs 1..{len(suggestions_for_prompt)}):
{suggestions_block}

Contexto: estas relaciones son hipotesis y requieren validacion humana con evidencia.

IMPORTANTE: Responde UNICAMENTE con un JSON valido (sin markdown, sin ```).
La sintesis debe distinguir explicitamente el estatus epistemologico de cada afirmacion.
El JSON debe tener esta estructura exacta:

{{
    "memo_sintesis": [
        {{"type": "OBSERVATION", "text": "...", "evidence_ids": [1, 3], "evidence_fragment_ids": ["123", "456"]}},
        {{"type": "INTERPRETATION", "text": "...", "evidence_ids": [1, 3], "evidence_fragment_ids": ["123"]}},
        {{"type": "HYPOTHESIS", "text": "...", "evidence_ids": [2], "evidence_fragment_ids": ["789"]}},
        {{"type": "NORMATIVE_INFERENCE", "text": "...", "evidence_ids": [4]}}
    ]
}}

REGLAS:
1. memo_sintesis: lista de 3-6 statements.
2. type: uno de OBSERVATION | INTERPRETATION | HYPOTHESIS | NORMATIVE_INFERENCE.
3. text: una oracion clara (espanol).
4. evidence_ids: lista de enteros referidos a la numeracion de las sugerencias (1..{len(suggestions_for_prompt)}).
5. evidence_fragment_ids: lista de strings `fragmento_id` (de la evidencia provista). No inventes IDs.
6. PROHIBIDO: OBSERVATION sin evidence_fragment_ids no vacios.
7. Se conciso: evita listas largas o parrafos extensos."""

        response = clients.aoai.chat.completions.create(
            model=settings.azure.deployment_chat,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            max_completion_tokens=600,
        )

        choice = response.choices[0] if response.choices else None
        message = getattr(choice, "message", None)
        message_content = getattr(message, "content", None)
        raw_content = message_content or ""

        import json
        import re

        parsed: Optional[Dict[str, Any]] = None
        try:
            clean_content = re.sub(r"^```json?\s*", "", raw_content.strip())
            clean_content = re.sub(r"\s*```$", "", clean_content)
            parsed = json.loads(clean_content)
        except Exception:
            parsed = None

        def _normalize_memo(value: Any) -> tuple[str, List[Dict[str, Any]]]:
            allowed = {"OBSERVATION", "INTERPRETATION", "HYPOTHESIS", "NORMATIVE_INFERENCE"}
            if isinstance(value, str):
                return value.strip(), []
            if not isinstance(value, list):
                return "", []
            out: List[Dict[str, Any]] = []
            lines: List[str] = []
            for item in value:
                if not isinstance(item, dict):
                    continue
                stype = str(item.get("type") or "").strip().upper()
                text = str(item.get("text") or "").strip()
                if not text:
                    continue
                if stype not in allowed:
                    stype = "INTERPRETATION"
                evidence_ids: List[int] = []
                raw_ids = item.get("evidence_ids")
                if isinstance(raw_ids, list):
                    for v in raw_ids:
                        try:
                            evidence_ids.append(int(v))
                        except Exception:
                            continue
                evidence_fragment_ids: List[str] = []
                raw_frag_ids = item.get("evidence_fragment_ids")
                if isinstance(raw_frag_ids, list):
                    for v in raw_frag_ids:
                        fid = str(v or "").strip()
                        if fid:
                            evidence_fragment_ids.append(fid)
                # Guardrail: OBSERVATION requiere evidencia en fragmentos.
                if stype == "OBSERVATION" and not evidence_fragment_ids:
                    stype = "INTERPRETATION"
                entry: Dict[str, Any] = {"type": stype, "text": text}
                if evidence_ids:
                    entry["evidence_ids"] = evidence_ids
                if evidence_fragment_ids:
                    entry["evidence_fragment_ids"] = evidence_fragment_ids

                if evidence_fragment_ids:
                    preview = ", ".join(evidence_fragment_ids[:6])
                    suffix = "…" if len(evidence_fragment_ids) > 6 else ""
                    lines.append(f"[{stype}] {text} (frags: {preview}{suffix})")
                elif evidence_ids:
                    lines.append(f"[{stype}] {text} (links: {', '.join(str(i) for i in evidence_ids)})")
                else:
                    lines.append(f"[{stype}] {text}")
                out.append(entry)
            return "\n".join(lines).strip(), out

        if isinstance(parsed, dict) and parsed.get("memo_sintesis") is not None:
            memo_text, memo_statements = _normalize_memo(parsed.get("memo_sintesis"))
            analysis = memo_text
            structured = True
        else:
            analysis = raw_content
            structured = False
            memo_statements = []

        analysis_id = 0
        persisted = False
        evidence_schema_version = (
            int(evidence_pack.get("schema_version", 1) or 1) if isinstance(evidence_pack, dict) else 1
        )
        if payload.persist:
            from app.postgres_block import insert_axial_ai_analysis

            analysis_id = insert_axial_ai_analysis(
                clients.postgres,
                project_id=project_id,
                source_type="analyze_hidden_relationships",
                algorithm="hidden_relationships",
                algorithm_description="Relaciones ocultas (descubrimiento): evidencia positiva/negativa",
                suggestions_json=payload.suggestions[:10],
                analysis_text=analysis,
                memo_statements=memo_statements,
                structured=structured,
                estado="pendiente",
                created_by=user.user_id,
                epistemic_mode=epistemic_mode.value,
                prompt_version=prompt_version,
                llm_deployment=settings.azure.deployment_chat,
                llm_api_version=settings.azure.api_version,
                evidence_schema_version=evidence_schema_version,
                evidence_json=evidence_pack if isinstance(evidence_pack, dict) else None,
            )
            persisted = bool(analysis_id)

        return {
            "project": project_id,
            "analysis": analysis,
            "structured": structured,
            "memo_statements": memo_statements,
            "suggestions_analyzed": len(payload.suggestions[:10]),
            "epistemic_mode": epistemic_mode.value,
            "prompt_version": prompt_version,
            "llm_deployment": settings.azure.deployment_chat,
            "llm_api_version": settings.azure.api_version,
            "analysis_id": analysis_id,
            "persisted": persisted,
            "evidence_schema_version": evidence_schema_version,
            "evidence_totals": evidence_pack.get("totals") if isinstance(evidence_pack, dict) else None,
        }
    except Exception as exc:
        api_logger.error("api.analyze_hidden_relationships.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


@app.post("/api/axial/hidden-relationships/metrics")
async def api_hidden_relationships_metrics(
    payload: Optional[HiddenRelationshipsMetricsRequest] = Body(default=None),
    payload_q: Optional[str] = Query(default=None, alias="payload"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Calcula métricas de diversidad/overlap para relaciones ocultas."""
    try:
        if payload is None and payload_q:
            try:
                payload = HiddenRelationshipsMetricsRequest.model_validate_json(payload_q)
            except Exception as exc:
                raise HTTPException(status_code=400, detail="Invalid payload JSON") from exc
        if payload is None:
            raise HTTPException(status_code=422, detail="Missing request body")
        from backend.routers.agent import _compute_evidence_diversity_metrics

        codes_with_fragments = []
        suggestions_total = 0
        suggestions_with_direct_evidence = 0
        suggestions_with_any_evidence = 0
        for sug in payload.suggestions:
            suggestions_total += 1
            source = str(sug.get("source") or "?")
            target = str(sug.get("target") or "?")
            key = f"{source}↔{target}"
            evidence_ids = sug.get("evidence_ids") or []
            supporting_ids = sug.get("supporting_evidence_ids") or []
            has_direct = isinstance(evidence_ids, list) and len([v for v in evidence_ids if v]) > 0
            has_supporting = isinstance(supporting_ids, list) and len([v for v in supporting_ids if v]) > 0
            if has_direct:
                suggestions_with_direct_evidence += 1
            if has_direct or has_supporting:
                suggestions_with_any_evidence += 1
            fragments = [{"fragmento_id": fid} for fid in evidence_ids if fid]
            codes_with_fragments.append({"code": key, "fragments": fragments})

        metrics = _compute_evidence_diversity_metrics(codes_with_fragments)
        metrics["suggestions_total"] = suggestions_total
        metrics["suggestions_with_direct_evidence"] = suggestions_with_direct_evidence
        metrics["direct_evidence_ratio"] = (
            (suggestions_with_direct_evidence / suggestions_total) if suggestions_total else 0.0
        )
        metrics["suggestions_with_any_evidence"] = suggestions_with_any_evidence
        metrics["any_evidence_ratio"] = (
            (suggestions_with_any_evidence / suggestions_total) if suggestions_total else 0.0
        )
        return {
            "project": payload.project,
            "metrics": metrics,
        }
    except Exception as exc:
        api_logger.error("api.hidden_relationships.metrics_error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc


async def get_service_clients(
    settings: AppSettings = Depends(get_settings),
) -> AsyncGenerator[ServiceClients, None]:
    """
    Dependency that provides ServiceClients with automatic cleanup.
    
    Usage:
        @app.get("/api/endpoint")
        async def handler(clients: ServiceClients = Depends(get_service_clients)):
            # Use clients.postgres, clients.qdrant, etc.
            # No need to call clients.close() - FastAPI handles it!
    """
    import structlog
    _dep_logger = structlog.get_logger("app.deps")
    
    clients = build_service_clients(settings)
    _dep_logger.info("deps.get_service_clients.yield_start")
    try:
        yield clients
    finally:
        _dep_logger.info("deps.get_service_clients.finally_executing")
        clients.close()


async def get_pg_clients(
    settings: AppSettings = Depends(get_settings),
) -> AsyncGenerator[PgOnlyClients, None]:
    clients = build_pg_only(settings)
    try:
        yield clients
    finally:
        clients.close()


async def run_pg_query_with_timeout(func, *args, timeout: float = 12.0, **kwargs):
    """Execute a blocking PG helper in a worker thread with a hard timeout.

    Prevents the UI from hanging when Postgres is slow or unreachable by
    returning a 504 instead of letting the client-side 30s timeout fire.
    """
    try:
        return await asyncio.wait_for(asyncio.to_thread(func, *args, **kwargs), timeout=timeout)
    except asyncio.TimeoutError as exc:  # pragma: no cover - network timing
        raise HTTPException(
            status_code=status.HTTP_504_GATEWAY_TIMEOUT,
            detail="PostgreSQL query timed out",
        ) from exc


async def get_neo4j_clients(
    settings: AppSettings = Depends(get_settings),
) -> AsyncGenerator[Neo4jOnlyClients, None]:
    clients = build_neo4j_only(settings)
    try:
        yield clients
    finally:
        clients.close()



# =============================================================================
# DEPRECATED: Authentication endpoints moved to backend.routers.auth
# =============================================================================
# OLD ENDPOINTS - COMMENTED OUT
# - POST /token → oauth_router
# - POST /api/auth/login → auth_router
# - POST /api/auth/register → auth_router
# - POST /refresh → auth_legacy_router
# - POST /register → auth_legacy_router

# @app.post("/token")
# async def api_login_oauth(...
# [Lines 398-565 commented out - All auth endpoints moved to backend.routers.auth]

# async def api_login_json(
#     request: Dict[str, Any] = Body(...),
#     clients: ServiceClients = Depends(get_service_clients),
# ):
#     """
#     Login de usuario con JSON - genera access token y refresh token.
    
#     El frontend envía JSON en lugar de form-data.
#     """
#     from backend.auth_service import authenticate_user, create_tokens_for_user
    
#     email = request.get("email", "")
#     password = request.get("password", "")
    
#     if not email or not password:
#         raise HTTPException(
#             status_code=status.HTTP_400_BAD_REQUEST,
#             detail="Email y password son requeridos",
#         )
    
#     user, error = authenticate_user(clients.postgres, email, password)
    
#     if error:
#         raise HTTPException(
#             status_code=status.HTTP_401_UNAUTHORIZED,
#             detail=error,
#             headers={"WWW-Authenticate": "Bearer"},
#         )
    
#     try:
#         tokens = create_tokens_for_user(clients.postgres, user)
#     except Exception as e:
#         api_logger.error("auth.login.token_error", error=str(e))
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail="Error generando tokens",
#         )
    
#     api_logger.info("auth.login.json.success", user_id=user["id"], email=user["email"])
    
#     return {
#         "access_token": tokens.access_token,
#         "refresh_token": tokens.refresh_token,
#         "token_type": tokens.token_type,
#         "expires_in": tokens.expires_in,
#         "user": {
#             "id": user["id"],
#             "email": user["email"],
#             "full_name": user.get("full_name"),
#             "organization_id": user.get("organization_id", "default_org"),
#             "role": user.get("role", "analyst"),
#         }
#     }


# @app.post("/register")
# @app.post("/api/auth/register")
# async def api_register(
#     request: Dict[str, Any] = Body(...),
#     clients: ServiceClients = Depends(get_service_clients),
# ):
#     """
#     Registro de nuevo usuario.
    
#     Campos requeridos:
#     - email: Email válido
#     - password: Mínimo 8 caracteres, debe incluir mayúscula, minúscula, número y símbolo
    
#     Campos opcionales:
#     - full_name: Nombre completo
#     - organization_id: ID de organización (default: "default_org")
#     """
#     from backend.auth_service import RegisterRequest, register_user, create_tokens_for_user
    
#     try:
#         reg_request = RegisterRequest(
#             email=request.get("email", ""),
#             password=request.get("password", ""),
#             full_name=request.get("full_name"),
#             organization_id=request.get("organization_id", "default_org"),
#         )
#     except Exception as e:
#         raise HTTPException(
#             status_code=status.HTTP_400_BAD_REQUEST,
#             detail=str(e),
#         )
    
#     user, error = register_user(clients.postgres, reg_request)
    
#     if error:
#         raise HTTPException(
#             status_code=status.HTTP_400_BAD_REQUEST,
#             detail=error,
#         )
    
#     # Generar tokens automáticamente después del registro
#     try:
#         tokens = create_tokens_for_user(clients.postgres, user)
#     except Exception as e:
#         # Usuario creado pero error en tokens - retornar solo confirmación
#         api_logger.warning("auth.register.token_error", user_id=user["id"], error=str(e))
#         return {
#             "message": "Usuario registrado. Por favor inicia sesión.",
#             "user_id": user["id"],
#         }
    
#     api_logger.info("auth.register.success", user_id=user["id"], email=user["email"])
    
#     return {
#         "message": "Usuario registrado exitosamente",
#         "user_id": user["id"],
#         "access_token": tokens.access_token,
#         "refresh_token": tokens.refresh_token,
#         "token_type": tokens.token_type,
#         "user": {
#             "id": user["id"],
#             "email": user["email"],
#             "full_name": user.get("full_name"),
#             "organization_id": user.get("organization_id", "default_org"),
#             "role": user.get("role", "analyst"),
#         }
#     }


# @app.post("/refresh")
# @app.post("/api/auth/refresh")
# async def api_refresh_token(
#     request: Dict[str, Any] = Body(...),
#     clients: ServiceClients = Depends(get_service_clients),
# ):
#     """
#     Renueva access token usando refresh token.
#     """
#     from backend.auth_service import refresh_access_token
    
#     refresh_token = request.get("refresh_token")
#     if not refresh_token:
#         raise HTTPException(
#             status_code=status.HTTP_400_BAD_REQUEST,
#             detail="refresh_token es requerido",
#         )
    
#     tokens, error = refresh_access_token(clients.postgres, refresh_token)
    
#     if error:
#         raise HTTPException(
#             status_code=status.HTTP_401_UNAUTHORIZED,
#             detail=error,
#         )
    
#     return {
#         "access_token": tokens["access_token"],
#         "refresh_token": tokens["refresh_token"],
#         "token_type": "bearer",
#     }

# End of deprecated auth endpoints
# =============================================================================


# =============================================================================
# DEPRECATED: Neo4j endpoints moved to backend.routers.neo4j
# =============================================================================
# OLD ENDPOINTS - COMMENTED OUT
# - POST /api/neo4j/query → neo4j_router
# - POST /api/neo4j/export → neo4j_router
# - Helper functions: _execute_cypher, _table_to_csv

# [Lines 575-740 commented out - All Neo4j endpoints and helpers moved to backend.routers.neo4j]

# End of deprecated Neo4j endpoints
# =============================================================================
def _validate_api_key(provided_api_key: Optional[str], expected_api_key: Optional[str]) -> None:
    """Backward-compatible API-key validator used by lightweight tests/helpers."""
    if not expected_api_key:
        raise HTTPException(status_code=500, detail="API key del servidor no configurada.")
    if provided_api_key != expected_api_key:
        raise HTTPException(status_code=401, detail="X-API-Key inválida.")


def _execute_cypher(
    payload: CypherRequest,
    settings: AppSettings,
    run_callable,
) -> Tuple[Dict[str, Any], float]:
    if not payload.project:
        raise HTTPException(status_code=400, detail="El proyecto es obligatorio para consultas Neo4j.")
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    try:
        formats = _normalize_formats(payload.formats)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    database = payload.database or settings.neo4j.database
    start = perf_counter()
    try:
        result = run_callable(
            payload.cypher,
            params=payload.params or {},
            database=database,
        )
    except HTTPException as exc:
        duration_ms = (perf_counter() - start) * 1000
        logger.warning(
            "neo4j.query.http_error",
            duration_ms=duration_ms,
            database=database,
            error=exc.detail,
        )
        raise
    except ValueError as exc:
        duration_ms = (perf_counter() - start) * 1000
        logger.warning(
            "neo4j.query.validation_error",
            duration_ms=duration_ms,
            database=database,
            error=str(exc),
        )
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:  # pragma: no cover - unexpected errors
        duration_ms = (perf_counter() - start) * 1000
        logger.error(
            "neo4j.query.failure",
            duration_ms=duration_ms,
            database=database,
            error=str(exc),
        )
        raise HTTPException(status_code=502, detail="Fallo al ejecutar la consulta en Neo4j.") from exc

    filtered = {fmt: result.get(fmt) for fmt in formats if fmt in result}
    duration_ms = (perf_counter() - start) * 1000
    logger.info(
        "neo4j.query.success",
        duration_ms=duration_ms,
        database=database,
        formats=formats,
        rows=len(filtered.get("raw") or []),
    )
    return filtered, duration_ms


def _table_to_csv(table: Dict[str, Any]) -> str:
    columns = table.get("columns") or []
    rows = table.get("rows") or []
    buffer = StringIO()
    if columns:
        buffer.write(",".join(str(column) for column in columns))
        buffer.write("\n")
        for row in rows:
            buffer.write(
                ",".join(
                    '"' + str(value).replace('"', '""') + '"'
                    if isinstance(value, str)
                    else ("" if value is None else str(value))
                    for value in row
                )
            )
            buffer.write("\n")
    return buffer.getvalue()


class CypherExportRequest(CypherRequest):
    export_format: str = Field(
        default="csv",
        description="Formato de exporte (csv o json).",
        pattern="^(csv|json)$",
    )


@app.post("/api/neo4j/query", response_model=CypherResponse)
async def neo4j_query(
    payload: CypherRequest,
    settings: AppSettings = Depends(get_settings),
    clients: Neo4jOnlyClients = Depends(get_neo4j_clients),
    user: User = Depends(require_auth),
) -> Any:
    result, duration_ms = _execute_cypher(
        payload,
        settings,
        lambda cypher, params, database: run_cypher(
            cast(ServiceClients, clients),
            cypher,
            params=params,
            database=database,
        ),
    )
    response = JSONResponse(result)
    response.headers["X-Query-Duration"] = f"{duration_ms:.2f}"
    return response


@app.post("/api/neo4j/export")
async def neo4j_export(
    payload: CypherExportRequest,
    settings: AppSettings = Depends(get_settings),
    clients: Neo4jOnlyClients = Depends(get_neo4j_clients),
    user: User = Depends(require_auth),
):
    export_format = payload.export_format.lower()
    requested_formats = payload.formats
    if export_format == "csv":
        requested_formats = ["table"]
    query_payload = CypherRequest(
        cypher=payload.cypher,
        params=payload.params,
        formats=requested_formats,
        database=payload.database,
    )
    result, duration_ms = _execute_cypher(
        query_payload,
        settings,
        lambda cypher, params, database: run_cypher(
            cast(ServiceClients, clients),
            cypher,
            params=params,
            database=database,
        ),
    )
    database_name = query_payload.database or settings.neo4j.database
    if export_format == "csv":
        table = result.get("table")
        if not table:
            raise HTTPException(status_code=400, detail="La consulta no generó datos tabulares para exportar.")
        csv_body = _table_to_csv(table)
        response = PlainTextResponse(
            csv_body,
            media_type="text/csv",
            headers={"Content-Disposition": 'attachment; filename="neo4j_export.csv"'},
        )
        logger.info(
            "neo4j.export.success",
            duration_ms=duration_ms,
            database=database_name,
            export_format="csv",
        )
    else:
        response = JSONResponse(result)
        logger.info(
            "neo4j.export.success",
            duration_ms=duration_ms,
            database=database_name,
            export_format="json",
        )
    response.headers["X-Query-Duration"] = f"{duration_ms:.2f}"
    return response


# DEPRECATED: Moved to backend.routers.admin.health_router
# @app.get("/healthz")
# async def healthcheck() -> Dict[str, str]:
#     return {"status": "ok"}


# ---------- Project & Status ----------


class ProjectCreateRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    name: str
    description: Optional[str] = None
    epistemic_mode: Optional[str] = "constructivist"  # constructivist | post_positivist


# =============================================================================
# Registration Endpoints
# =============================================================================

# NOTA: El endpoint /register antiguo fue removido.
# El registro de usuarios se hace via /api/auth/register que usa PostgreSQL.



@app.get("/api/organizations")
async def api_list_organizations(
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista organizaciones (admin only)."""
    organizations: List[Dict[str, Any]] = []
    with clients.postgres.cursor() as cur:
        cur.execute(
            """
            SELECT DISTINCT organization_id
            FROM app_users
            WHERE organization_id IS NOT NULL AND organization_id <> ''
            ORDER BY organization_id
            """
        )
        organizations = [
            {"id": row[0], "name": row[0]}
            for row in cur.fetchall()
            if row and row[0]
        ]
    return {"organizations": organizations}


# =============================================================================
# Project Endpoints (filtered by org_id)
# =============================================================================

@app.get("/api/projects")
async def api_projects(
    user: User = Depends(require_auth),
    settings: AppSettings = Depends(get_settings),
) -> Dict[str, Any]:
    """
    Lista proyectos filtrados por usuario/organización.
    
    - Admin: Ve todos los proyectos
    - Otros: Ve proyectos propios + proyectos de su org + proyectos legacy (sin owner)
    """
    from app.project_state import list_projects_for_user

    clients = build_clients_or_error(settings)
    
    try:
        # Usar función de filtrado que respeta owner_id y org_id
        filtered = list_projects_for_user(
            clients.postgres,
            user_id=user.user_id,
            org_id=user.organization_id,
            role=user.roles[0] if user.roles else "analyst",
        )
        return {"projects": filtered}
    finally:
        clients.close()



@app.post("/api/projects")
async def api_create_project(
    payload: ProjectCreateRequest,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Crea un proyecto asociado al usuario y su organización.
    
    El owner_id se asigna automáticamente al usuario que lo crea.
    """
    api_logger.info(
        "project.create.start",
        project_name=payload.name,
        user=user.user_id,
        org_id=user.organization_id,
    )
    try:
        entry = create_project(
            clients.postgres,
            payload.name,
            payload.description,
            org_id=user.organization_id,
            owner_id=user.user_id,  # Asignar al usuario creador
            epistemic_mode=payload.epistemic_mode or "constructivist",
        )
        api_logger.info(
            "project.create.db_success",
            project_id=entry.get("id"),
            project_name=payload.name,
        )
        # En multi-tenant estricto, el creador debe quedar como admin del proyecto.
        try:
            add_project_member(
                clients.postgres,
                entry.get("id"),
                user.user_id,
                "admin",
                added_by=user.user_id,
            )
        except Exception as member_exc:
            # Log pero no fallar - el proyecto ya fue creado
            api_logger.warning(
                "project.create.member_add_warning",
                project_id=entry.get("id"),
                error=str(member_exc),
            )
    except ValueError as exc:
        api_logger.warning(
            "project.create.validation_error",
            project_name=payload.name,
            error=str(exc),
        )
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        api_logger.error(
            "project.create.unexpected_error",
            project_name=payload.name,
            error=str(exc),
            exc_info=True,
        )
        raise HTTPException(status_code=500, detail=f"Error interno: {exc}") from exc
    api_logger.info(
        "project.created",
        project_id=entry.get("id"),
        project_name=payload.name,
        user=user.user_id,
        org_id=user.organization_id,
    )
    return entry



@app.get("/api/projects/{project_id}/export")
async def api_export_project(
    project_id: str,
    settings: AppSettings = Depends(get_settings),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
):
    """
    Exporta un proyecto completo como archivo ZIP para backup.
    Incluye: config, fragmentos, códigos, reportes, notas.
    """
    import io
    import zipfile
    from fastapi.responses import StreamingResponse
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        export_data: Dict[str, Any] = {
            "project_id": project_id,
            "exported_at": datetime.now().isoformat(),
        }
        
        # 1. Project config (cloud-only: PostgreSQL)
        try:
            export_data["config"] = get_project_config(clients.postgres, project_id)
            zip_file.writestr(
                "project_config.json",
                json.dumps(export_data["config"], ensure_ascii=False, indent=2),
            )
        except Exception:
            pass
        
        # 2. PostgreSQL data
        try:
            with clients.postgres.cursor() as cur:
                # Fragments
                cur.execute(
                    "SELECT * FROM entrevista_fragmentos WHERE project_id = %s",
                    (project_id,)
                )
                columns = [desc[0] for desc in (cur.description or [])]
                fragments = [dict(zip(columns, row)) for row in cur.fetchall()]
                zip_file.writestr("fragments.json", json.dumps(fragments, default=str, ensure_ascii=False, indent=2))
                export_data["fragments_count"] = len(fragments)
                
                # Open codes
                cur.execute(
                    "SELECT * FROM analisis_codigos_abiertos WHERE project_id = %s",
                    (project_id,)
                )
                columns = [desc[0] for desc in (cur.description or [])]
                codes = [dict(zip(columns, row)) for row in cur.fetchall()]
                zip_file.writestr("open_codes.json", json.dumps(codes, default=str, ensure_ascii=False, indent=2))
                export_data["codes_count"] = len(codes)
                
                # Candidate codes
                try:
                    cur.execute(
                        "SELECT * FROM codigos_candidatos WHERE project_id = %s",
                        (project_id,)
                    )
                    columns = [desc[0] for desc in (cur.description or [])]
                    candidates = [dict(zip(columns, row)) for row in cur.fetchall()]
                    zip_file.writestr("candidate_codes.json", json.dumps(candidates, default=str, ensure_ascii=False, indent=2))
                except Exception:
                    pass
                
                # Interview reports
                try:
                    cur.execute(
                        "SELECT * FROM interview_reports WHERE project_id = %s",
                        (project_id,)
                    )
                    columns = [desc[0] for desc in (cur.description or [])]
                    reports = [dict(zip(columns, row)) for row in cur.fetchall()]
                    zip_file.writestr("interview_reports.json", json.dumps(reports, default=str, ensure_ascii=False, indent=2))
                except Exception:
                    pass
        except Exception as e:
            export_data["postgres_error"] = str(e)
        
        # 3. Neo4j graph data
        try:
            with clients.neo4j.session(database=settings.neo4j.database) as session:
                query = Neo4jQuery(
                    """
                    MATCH (n {project_id: $pid})
                    OPTIONAL MATCH (n)-[r]->(m {project_id: $pid})
                    RETURN labels(n) as labels, properties(n) as props,
                           type(r) as rel_type, properties(m) as target_props
                    """
                )
                result = session.run(query, pid=project_id)
                graph_data = [dict(record) for record in result]
                zip_file.writestr("neo4j_graph.json", json.dumps(graph_data, default=str, ensure_ascii=False, indent=2))
                export_data["neo4j_nodes"] = len(graph_data)
        except Exception as e:
            export_data["neo4j_error"] = str(e)
        
        # 4. Notes directory
        notes_dir = Path(f"notes/{project_id}")
        if notes_dir.exists():
            for file_path in notes_dir.rglob("*"):
                if file_path.is_file():
                    arcname = f"notes/{file_path.relative_to(notes_dir)}"
                    zip_file.write(file_path, arcname)
        
        # 5. Reports directory
        reports_dir = Path(f"reports/{project_id}")
        if reports_dir.exists():
            for file_path in reports_dir.rglob("*"):
                if file_path.is_file():
                    arcname = f"reports/{file_path.relative_to(reports_dir)}"
                    zip_file.write(file_path, arcname)
        
        # Write export manifest
        zip_file.writestr("_export_manifest.json", json.dumps(export_data, default=str, ensure_ascii=False, indent=2))
    
    zip_buffer.seek(0)
    
    api_logger.info("project.exported", project_id=project_id, user=user.user_id)
    
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{project_id}_backup.zip"'}
    )


@app.delete("/api/projects/{project_id}")
async def api_delete_project(
    project_id: str,
    settings: AppSettings = Depends(get_settings),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Elimina un proyecto y todos sus datos asociados:
    - Neo4j: nodos Fragmento, Entrevista, Codigo, Categoria
    - Qdrant: puntos con project_id
    - PostgreSQL: filas en tablas de análisis
    - Archivos: directorio del proyecto en data/projects/
    - Registry: (deshabilitado en modo cloud-only)
    """
    from qdrant_client.models import Filter, FieldCondition, MatchValue
    import shutil
    
    results = {"project_id": project_id, "deleted": {}}
    
    try:
        # 1. Neo4j cleanup
        try:
            with clients.neo4j.session(database=settings.neo4j.database) as session:
                delete_jobs = [
                    (
                        "fragmento",
                        Neo4jQuery(
                            cast(
                                LiteralString,
                                "MATCH (n:Fragmento {project_id: $pid}) DETACH DELETE n RETURN count(n) as deleted",
                            )
                        ),
                    ),
                    (
                        "entrevista",
                        Neo4jQuery(
                            cast(
                                LiteralString,
                                "MATCH (n:Entrevista {project_id: $pid}) DETACH DELETE n RETURN count(n) as deleted",
                            )
                        ),
                    ),
                    (
                        "codigo",
                        Neo4jQuery(
                            cast(
                                LiteralString,
                                "MATCH (n:Codigo {project_id: $pid}) DETACH DELETE n RETURN count(n) as deleted",
                            )
                        ),
                    ),
                    (
                        "categoria",
                        Neo4jQuery(
                            cast(
                                LiteralString,
                                "MATCH (n:Categoria {project_id: $pid}) DETACH DELETE n RETURN count(n) as deleted",
                            )
                        ),
                    ),
                ]
                for label_key, delete_query in delete_jobs:
                    result = session.run(delete_query, pid=project_id)
                    record = result.single()
                    deleted_count = record["deleted"] if record is not None else 0
                    results["deleted"][f"neo4j_{label_key}"] = deleted_count
        except Exception as e:
            results["deleted"]["neo4j_error"] = str(e)
        
        # 2. Qdrant cleanup
        try:
            clients.qdrant.delete(
                collection_name=settings.qdrant.collection,
                points_selector=Filter(
                    must=[FieldCondition(key="project_id", match=MatchValue(value=project_id))]
                ),
            )
            results["deleted"]["qdrant"] = "success"
        except Exception as e:
            results["deleted"]["qdrant_error"] = str(e)
        
        # 3. PostgreSQL cleanup
        try:
            tables = [
                "entrevista_fragmentos",
                "analisis_codigos_abiertos", 
                "analisis_axial",
                "analisis_comparacion_constante",
                "analisis_nucleo_notas",
                "interview_reports",
                "interview_files",  # New decoupled ingestion table
                "codigos_candidatos",  # Candidate codes validation table
            ]
            with clients.postgres.cursor() as cur:
                for table in tables:
                    try:
                        cur.execute(
                            "SELECT EXISTS (SELECT FROM information_schema.tables WHERE table_name = %s)",
                            (table,)
                        )
                        row = cur.fetchone()
                        if not row or not row[0]:
                            continue
                        cur.execute(f"DELETE FROM {table} WHERE project_id = %s", (project_id,))
                        results["deleted"][f"pg_{table}"] = cur.rowcount
                    except Exception:
                        clients.postgres.rollback()
                
                # IMPORTANT: Delete from proyectos table itself (this is why projects stay in list!)
                try:
                    cur.execute("DELETE FROM proyectos WHERE id = %s", (project_id,))
                    results["deleted"]["pg_proyectos"] = cur.rowcount
                except Exception as e:
                    results["deleted"]["pg_proyectos_error"] = str(e)
                    
            clients.postgres.commit()
        except Exception as e:
            results["deleted"]["postgres_error"] = str(e)
        
        # 4. Delete project files
        project_dir = Path(f"data/projects/{project_id}")
        if project_dir.exists():
            try:
                shutil.rmtree(project_dir)
                results["deleted"]["files"] = "success"
            except Exception as e:
                results["deleted"]["files_error"] = str(e)
        else:
            results["deleted"]["files"] = "not_found"
        
        # 5-6. Registry/metadata config: deshabilitado en modo cloud-only
        
        # 7. Delete notes/{project}/ and reports/{project}/ directories
        for dir_name in ["notes", "reports"]:
            dir_path = Path(dir_name) / project_id
            try:
                if dir_path.exists():
                    shutil.rmtree(dir_path)
                    results["deleted"][dir_name] = "success"
                else:
                    results["deleted"][dir_name] = "not_found"
            except Exception as e:
                results["deleted"][f"{dir_name}_error"] = str(e)

        # 7b. Delete tenant-scoped Blob artifacts (reports container).
        try:
            from app.blob_storage import CONTAINER_REPORTS, delete_prefix, tenant_prefix

            blob_prefix = tenant_prefix(org_id=str(getattr(user, "organization_id", None) or ""), project_id=project_id).rstrip("/") + "/"
            results["deleted"]["blob_reports"] = delete_prefix(container=CONTAINER_REPORTS, prefix=blob_prefix)
        except Exception as e:
            results["deleted"]["blob_reports_error"] = str(e)
        
        # 8. Backups locales: deshabilitado en modo cloud-only
        
        # 9. Update metadata/entrevistas.json (remove project entries)
        try:
            entrevistas_path = Path("metadata/entrevistas.json")
            if entrevistas_path.exists():
                with open(entrevistas_path, "r", encoding="utf-8") as f:
                    entrevistas = json.load(f)
                if isinstance(entrevistas, list):
                    original_len = len(entrevistas)
                    entrevistas = [e for e in entrevistas if e.get("project_id") != project_id]
                    with open(entrevistas_path, "w", encoding="utf-8") as f:
                        json.dump(entrevistas, f, indent=2, ensure_ascii=False)
                    results["deleted"]["entrevistas_json"] = original_len - len(entrevistas)
        except Exception as e:
            results["deleted"]["entrevistas_json_error"] = str(e)
        
        api_logger.info("project.deleted", project_id=project_id, user=user.user_id, results=results)
        return {"status": "deleted", **results}
        
    except Exception as e:
        api_logger.exception("project.delete.error", project_id=project_id)
        raise HTTPException(status_code=500, detail=f"Error eliminando proyecto: {str(e)}")


@app.get("/healthz")
async def healthz() -> Dict[str, Any]:
    """
    Health check endpoint for frontend BackendStatus component.
    Does not require authentication to allow simple connectivity checks.
    """
    return {
        "status": "ok",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0"
    }


@app.get("/api/storage/check")
async def api_storage_check(
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Safe diagnostic endpoint for Azure Blob Storage.

    Returns whether storage is configured and reachable, without returning secrets.
    """
    try:
        from app.blob_storage import check_blob_storage_health

        storage = check_blob_storage_health()
        return {
            "timestamp": datetime.now().isoformat(),
            "storage": storage,
        }
    except Exception as e:
        return {
            "timestamp": datetime.now().isoformat(),
            "storage": {
                "configured": False,
                "status": "error",
                "message": str(e)[:200],
            },
        }


# Track server start time for uptime calculation
_SERVER_START_TIME = datetime.now()


@app.get("/api/health/full")
async def api_health_full(
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Comprehensive health check for all services.
    Returns status, latency, and details for each service.
    """
    import time
    
    services = []
    overall_healthy = True
    
    # 1. Check PostgreSQL
    pg_status = {"name": "PostgreSQL", "status": "checking"}
    try:
        start = time.perf_counter()
        clients = build_service_clients(settings)
        if clients.postgres:
            with clients.postgres.cursor() as cur:
                cur.execute("SELECT 1")
                cur.fetchone()
            latency = int((time.perf_counter() - start) * 1000)
            pg_status = {
                "name": "PostgreSQL",
                "status": "online",
                "latency_ms": latency,
                "message": f"Conectado a {settings.postgres.host}:{settings.postgres.port}",
                "details": {
                    "database": settings.postgres.database,
                    "host": settings.postgres.host,
                }
            }
            clients.close()
        else:
            pg_status = {
                "name": "PostgreSQL",
                "status": "offline",
                "message": "No configurado"
            }
            overall_healthy = False
    except Exception as e:
        pg_status = {
            "name": "PostgreSQL",
            "status": "offline",
            "message": str(e)[:100]
        }
        overall_healthy = False
    services.append(pg_status)
    
    # 2. Check Neo4j
    neo4j_status = {"name": "Neo4j", "status": "checking"}
    try:
        start = time.perf_counter()
        clients = build_service_clients(settings)
        if clients.neo4j:
            with clients.neo4j.session() as session:
                result = session.run("RETURN 1 AS n")
                result.single()
            latency = int((time.perf_counter() - start) * 1000)
            neo4j_uri = settings.neo4j.uri
            neo4j_status = {
                "name": "Neo4j",
                "status": "online",
                "latency_ms": latency,
                "message": "Conexión establecida",
                "details": {
                    "uri": neo4j_uri[:30] + "..." if len(neo4j_uri) > 30 else neo4j_uri,
                    "database": settings.neo4j.database or "neo4j",
                }
            }
            clients.close()
        else:
            neo4j_status = {
                "name": "Neo4j",
                "status": "offline",
                "message": "No configurado"
            }
    except Exception as e:
        neo4j_status = {
            "name": "Neo4j",
            "status": "degraded",
            "message": str(e)[:100]
        }
    services.append(neo4j_status)
    
    # 3. Check Qdrant
    qdrant_status = {"name": "Qdrant", "status": "checking"}
    try:
        start = time.perf_counter()
        clients = build_service_clients(settings)
        if clients.qdrant:
            # Try to get collections list
            collections = clients.qdrant.get_collections()
            latency = int((time.perf_counter() - start) * 1000)
            collection_names = [c.name for c in collections.collections]
            qdrant_status = {
                "name": "Qdrant",
                "status": "online",
                "latency_ms": latency,
                "message": f"{len(collection_names)} colecciones",
                "details": {
                    "collections": collection_names[:5],  # Limit to 5
                    "collection_target": settings.qdrant.collection,
                }
            }
            clients.close()
        else:
            qdrant_status = {
                "name": "Qdrant",
                "status": "offline",
                "message": "No configurado"
            }
    except Exception as e:
        qdrant_status = {
            "name": "Qdrant",
            "status": "degraded",
            "message": str(e)[:100]
        }
    services.append(qdrant_status)
    
    # 4. Check Azure OpenAI
    azure_status = {"name": "Azure OpenAI", "status": "checking"}
    if settings.azure.api_key:
        azure_endpoint = settings.azure.endpoint or ""
        azure_status = {
            "name": "Azure OpenAI",
            "status": "online",
            "message": "API Key configurada",
            "details": {
                "endpoint": azure_endpoint[:40] + "..." if len(azure_endpoint) > 40 else azure_endpoint,
                "chat_deployment": settings.azure.deployment_chat,
                "embed_deployment": settings.azure.deployment_embed,
            }
        }
    else:
        azure_status = {
            "name": "Azure OpenAI",
            "status": "degraded",
            "message": "API Key no configurada"
        }
    services.append(azure_status)

    # 5. Check Azure Blob Storage (optional)
    storage_status = {"name": "Azure Blob Storage", "status": "checking"}
    try:
        from app.blob_storage import check_blob_storage_health

        storage = check_blob_storage_health()
        if storage.get("configured") and storage.get("status") == "ok":
            storage_status = {
                "name": "Azure Blob Storage",
                "status": "online",
                "latency_ms": storage.get("latency_ms"),
                "message": "Conectado",
                "details": {
                    "account_name": storage.get("account_name"),
                    "interviews_container": storage.get("interviews_container"),
                },
            }
        else:
            storage_status = {
                "name": "Azure Blob Storage",
                "status": "degraded",
                "message": storage.get("message") or "No configurado",
                "details": {"configured": bool(storage.get("configured"))},
            }
    except Exception as e:
        storage_status = {
            "name": "Azure Blob Storage",
            "status": "degraded",
            "message": str(e)[:100],
        }
    services.append(storage_status)
    
    # Calculate overall status
    statuses = [s["status"] for s in services]
    if all(s == "online" for s in statuses):
        overall = "healthy"
    elif "offline" in statuses:
        overall = "unhealthy"
    else:
        overall = "degraded"
    
    # Calculate uptime
    uptime = (datetime.now() - _SERVER_START_TIME).total_seconds()
    
    return {
        "timestamp": datetime.now().isoformat(),
        "services": services,
        "overall_status": overall,
        "uptime_seconds": int(uptime),
    }


@app.get("/api/status")
async def api_status(
    project: str = Query(..., description="Proyecto requerido"),
    update_state: bool = False,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    if update_state:
        api_logger.info("api.status.update_state_flag", project=project_id)
    snapshot = detect_stage_status(
        clients.postgres,
        project_id,
        stages=STAGE_DEFINITIONS,
        stage_order=STAGE_ORDER,
    )
    return snapshot


@app.get("/api/dashboard/counts")
async def api_dashboard_counts(
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene conteos en tiempo real para el dashboard.
    
    Este endpoint resuelve el Bug E1.1: "0 fragmentos" en Etapa 2.
    A diferencia de /api/status que lee el state guardado, este endpoint
    consulta directamente la base de datos para obtener conteos actualizados.
    
    Returns:
        Dict con conteos por etapa:
        - ingesta: archivos, fragmentos, speaker distribution
        - codificacion: códigos, citas, cobertura
        - axial: relaciones, categorías
        - candidatos: pendientes, validados, rechazados
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        try:
            counts = get_dashboard_counts(clients.postgres, project_id)
            return counts
        except Exception as exc:
            from app.error_handling import api_error, ErrorCode
            raise api_error(
                status_code=500,
                code=ErrorCode.DATABASE_ERROR,
                message="Error calculando conteos del dashboard.",
                exc=exc,
            ) from exc
    finally:
        clients.close()


@app.get("/api/research/overview")
async def api_research_overview(
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Resumen unificado del proyecto para UI.

    Retorna dos capas explícitas:
    - `validated`: checklist de etapas (lo que fue marcado/validado por workflow)
    - `observed`: métricas observadas directamente en BD (lo que existe en datos)

    Si PostgreSQL no está disponible, devuelve `availability.postgres=false` y
    degrada a un payload parcial con `warnings` (sin inventar ceros).
    """
    from datetime import datetime, timezone

    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        overview: Dict[str, Any] = {
            "project": project_id,
            "generated_at": datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
            "availability": {"postgres": True},
            "validated": None,
            "observed": None,
            "discovery": None,
            "warnings": [],
        }

        # Layer 1: validated checklist
        try:
            overview["validated"] = detect_stage_status(
                clients.postgres,
                project_id,
                stages=STAGE_DEFINITIONS,
                stage_order=STAGE_ORDER,
            )
        except Exception as exc:
            api_logger.warning(
                "api.research.overview.validated_failed",
                project=project_id,
                error=str(exc),
            )
            overview["availability"]["postgres"] = False
            overview["warnings"].append(
                "No fue posible cargar el checklist de etapas (PostgreSQL no disponible o error interno)."
            )

        # Layer 2: observed counts
        try:
            overview["observed"] = get_dashboard_counts(clients.postgres, project_id)
        except Exception as exc:
            api_logger.warning(
                "api.research.overview.observed_failed",
                project=project_id,
                error=str(exc),
            )
            overview["availability"]["postgres"] = False
            overview["warnings"].append(
                "No fue posible calcular conteos observados (PostgreSQL no disponible o error interno)."
            )

        # Layer 3: small Discovery summary (best-effort)
        try:
            from app.postgres_block import get_discovery_runs

            runs = get_discovery_runs(clients.postgres, project=project_id, limit=30)
            landing_rates = [r["landing_rate"] for r in runs if r.get("landing_rate") is not None]
            avg_landing = (sum(landing_rates) / len(landing_rates)) if landing_rates else None
            overview["discovery"] = {
                "recent_runs": len(runs),
                "avg_landing_rate": round(float(avg_landing), 2) if avg_landing is not None else None,
                "latest": runs[0] if runs else None,
            }
        except Exception as exc:
            api_logger.info(
                "api.research.overview.discovery_summary_skipped",
                project=project_id,
                error=str(exc),
            )

        # Cross-layer warnings (only if both are present)
        try:
            validated = overview.get("validated") or {}
            observed = overview.get("observed") or {}
            stages = (validated.get("stages") or {}) if isinstance(validated, dict) else {}

            def _stage_completed(key: str) -> bool:
                entry = stages.get(key) or {}
                return bool(entry.get("completed"))

            ingesta_frag = (((observed.get("ingesta") or {}).get("fragmentos")) if isinstance(observed, dict) else None)
            open_codes = (((observed.get("codificacion") or {}).get("codigos")) if isinstance(observed, dict) else None)
            axial_rel = (((observed.get("axial") or {}).get("relaciones")) if isinstance(observed, dict) else None)

            if isinstance(ingesta_frag, int) and ingesta_frag > 0 and not _stage_completed("ingesta"):
                overview["warnings"].append(
                    "Hay fragmentos ingeridos, pero la Etapa 'Ingesta' no está marcada como completada (checklist)."
                )
            if isinstance(open_codes, int) and open_codes > 0 and not _stage_completed("codificacion"):
                overview["warnings"].append(
                    "Hay códigos definitivos, pero la Etapa 'Codificación abierta' no está marcada como completada (checklist)."
                )
            if isinstance(axial_rel, int) and axial_rel > 0 and not _stage_completed("axial"):
                overview["warnings"].append(
                    "Hay relaciones axiales, pero la Etapa 'Codificación axial' no está marcada como completada (checklist)."
                )
        except Exception:
            # Never fail the endpoint because of warning computation.
            pass

        # Layer 4: Panorama (CTA + gating), best-effort
        try:
            from app.home_panorama import build_panorama
            from app.project_state import get_project_config
            from app.postgres_block import select_next_uncoded_fragment, get_open_coding_pending_counts
            from app.coding import get_saturation_data

            if overview.get("availability", {}).get("postgres") is True:
                config = get_project_config(clients.postgres, project_id)
                window = int(config.get("axial_min_saturation_window") or 3)
                threshold = int(config.get("axial_min_saturation_threshold") or 2)

                saturation = None
                try:
                    saturation = get_saturation_data(clients, project_id, window=window, threshold=threshold)
                except Exception:
                    saturation = None

                total_counts = get_open_coding_pending_counts(clients.postgres, project_id=project_id, archivo=None)
                pending_total = int(total_counts.get("pending_total") or 0)

                frag = select_next_uncoded_fragment(
                    clients.postgres,
                    project_id=project_id,
                    archivo=None,
                    exclude_fragment_ids=[],
                    strategy="recent",
                )
                recommended_archivo = None
                if frag and frag.get("archivo"):
                    recommended_archivo = str(frag.get("archivo") or "").strip() or None

                pending_in_recommended = None
                if recommended_archivo:
                    c = get_open_coding_pending_counts(
                        clients.postgres,
                        project_id=project_id,
                        archivo=recommended_archivo,
                    )
                    pending_in_recommended = c.get("pending_in_archivo")

                overview["panorama"] = build_panorama(
                    project=project_id,
                    validated=overview.get("validated"),
                    observed=overview.get("observed"),
                    stage_order=STAGE_ORDER,
                    pending_total=pending_total,
                    recommended_archivo=recommended_archivo,
                    pending_in_recommended=pending_in_recommended,
                    saturation=saturation,
                    config=config,
                )
        except Exception as exc:
            api_logger.info(
                "api.research.overview.panorama_skipped",
                project=project_id,
                error=str(exc),
            )

        return overview
    finally:
        clients.close()


@app.post("/api/projects/{project_id}/stages/{stage}/complete")
async def api_complete_stage(
    project_id: str,
    stage: str,
    run_id: Optional[str] = None,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    if stage not in STAGE_DEFINITIONS:
        raise HTTPException(status_code=400, detail=f"Etapa desconocida: {stage}")
    try:
        resolved_project = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    payload = mark_stage(
        clients.postgres,
        resolved_project,
        stage,
        run_id=run_id or "-",
        command="workflow",
        subcommand="complete",
        extras=None,
    )
    return payload


# ---------- Ingestion ----------
# NOTE: All ingestion and transcription endpoints have been moved to:
#       backend/routers/ingest.py
#
# Endpoints moved:
#   - POST /api/upload-and-ingest (NEW)
#   - POST /api/ingest
#   - POST /api/transcribe
#   - POST /api/transcribe/stream
#   - POST /api/transcribe/batch
#   - POST /api/transcribe/merge
#   - GET  /api/jobs/{task_id}/status
#   - GET  /api/jobs/batch/{batch_id}/status


def _expand_inputs(raw_inputs: List[str]) -> List[str]:
    """Expande patrones glob y normaliza rutas para compatibilidad Docker."""
    expanded: List[str] = []
    seen = set()
    for item in raw_inputs:
        if not item:
            continue
        # Normalizar backslashes de Windows a forward slashes para Linux
        item = item.replace("\\", "/")
        if any(ch in item for ch in ("*", "?", "[", "]")):
            for match in sorted(Path().glob(item)):
                if match.exists():
                    path = str(match)
                    if path not in seen:
                        expanded.append(path)
                        seen.add(path)
        else:
            if item not in seen:
                expanded.append(item)
                seen.add(item)
    return expanded


@app.post("/api/ingest")
async def api_ingest(
    payload: IngestRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    inputs = _expand_inputs(payload.inputs)
    if not inputs:
        raise HTTPException(status_code=400, detail="Debe especificar al menos un archivo o patrón.")

    meta = None
    if payload.meta_json:
        meta_path = Path(payload.meta_json)
        if not meta_path.exists():
            raise HTTPException(status_code=400, detail="Archivo meta_json no encontrado.")
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(status_code=400, detail=f"No se pudo leer meta_json: {exc}") from exc

    clients = build_clients_or_error(settings)
    log = api_logger.bind(endpoint="ingest", project=project_id)
    try:
        result = ingest_documents(
            clients,
            settings,
            inputs,
            batch_size=payload.batch_size,
            min_chars=payload.min_chars,
            max_chars=payload.max_chars,
            metadata=meta,
            run_id=payload.run_id,
            logger=log,
            project=project_id,
            org_id=str(getattr(user, "organization_id", None) or ""),
        )
    except Exception as exc:  # noqa: BLE001
        log.error("api.ingest.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()

    return {
        "project": project_id,
        "files": inputs,
        "result": result,
    }


# ---------- Transcription ----------


class TranscribeRequest(BaseModel):
    """Request para transcripción de audio con diarización."""
    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="Proyecto donde se ingesta (si aplica)")
    audio_base64: str = Field(..., description="Archivo de audio codificado en base64")
    filename: str = Field(..., description="Nombre original del archivo (con extensión)")
    diarize: bool = Field(True, description="Usar diarización para separar speakers")
    language: str = Field("es", description="Código de idioma")
    ingest: bool = Field(False, description="Ingestar directamente al pipeline")
    save_to_folder: bool = Field(True, description="Guardar DOCX en carpeta del proyecto")
    min_chars: int = Field(200, description="Mínimo de caracteres por fragmento")
    max_chars: int = Field(1200, description="Máximo de caracteres por fragmento")


class TranscribeSegmentResponse(BaseModel):
    speaker: str
    text: str
    start: float
    end: float


class TranscribeResponse(BaseModel):
    """Respuesta de transcripción con segmentos y metadata."""
    text: str = Field(..., description="Transcripción completa")
    segments: List[TranscribeSegmentResponse] = Field(default_factory=list)
    speaker_count: int = Field(0, description="Número de speakers detectados")
    duration_seconds: float = Field(0.0, description="Duración del audio")
    fragments_ingested: Optional[int] = Field(None, description="Fragmentos creados si ingest=True")
    saved_path: Optional[str] = Field(None, description="Ruta donde se guardó el DOCX")


# ---------- Batch Transcription (Parallel Processing) ----------

class BatchTranscribeFileItem(BaseModel):
    """Un archivo en el batch de transcripción."""
    audio_base64: str = Field(..., description="Audio codificado en base64")
    filename: str = Field(..., description="Nombre del archivo")

class BatchTranscribeRequest(BaseModel):
    """Request para transcripción paralela de múltiples archivos."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto destino")
    files: List[BatchTranscribeFileItem] = Field(..., description="Lista de archivos (max 20)")
    diarize: bool = Field(True, description="Usar diarización")
    language: str = Field("es", description="Código de idioma")
    ingest: bool = Field(True, description="Ingestar al pipeline")
    min_chars: int = Field(200)
    max_chars: int = Field(1200)

class BatchJobInfo(BaseModel):
    """Info de un job en el batch."""
    task_id: str
    filename: str

class BatchTranscribeResponse(BaseModel):
    """Respuesta de batch: IDs de jobs para tracking."""
    batch_id: str
    jobs: List[BatchJobInfo]
    message: str

class JobStatusResponse(BaseModel):
    """Estado de un job individual con soporte para progreso incremental."""
    task_id: str
    status: str  # PENDING, PROCESSING, SUCCESS, FAILURE
    filename: Optional[str] = None
    stage: Optional[str] = None  # transcribing, ingesting
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    # Campos para progreso de transcripción incremental
    chunk_index: Optional[int] = None
    total_chunks: Optional[int] = None
    chunks_completed: Optional[int] = None
    text_preview: Optional[str] = None
    segments_count: Optional[int] = None
    speakers_count: Optional[int] = None



@app.post("/api/transcribe", response_model=TranscribeResponse)
async def api_transcribe(
    payload: TranscribeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Transcribe archivo de audio con diarización automática.
    
    Soporta formatos: MP3, MP4, M4A, WAV, WebM, MPEG, FLAC, OGG
    """
    import base64
    import tempfile
    from app.transcription import (
        transcribe_audio_chunked,
        audio_to_fragments,
        save_transcription_docx,
        SUPPORTED_AUDIO_FORMATS,
    )
    
    # Validar proyecto
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    # Validar extensión
    suffix = Path(payload.filename).suffix.lower()
    if suffix not in SUPPORTED_AUDIO_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado: {suffix}. Formatos válidos: {', '.join(SUPPORTED_AUDIO_FORMATS)}"
        )
    
    log = api_logger.bind(endpoint="transcribe", project=project_id)
    
    # Decodificar audio
    try:
        audio_bytes = base64.b64decode(payload.audio_base64)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="No se pudo decodificar audio_base64") from exc
    
    # Guardar archivo temporal
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(audio_bytes)
        tmp_path = Path(tmp.name)
    
    try:
        # Transcribir usando chunked (divide automáticamente archivos grandes con ffmpeg)
        result = transcribe_audio_chunked(
            tmp_path,
            settings,
            diarize=payload.diarize,
            language=payload.language,
        )
        
        log.info(
            "api.transcribe.complete",
            speakers=result.speaker_count,
            segments=len(result.segments),
            duration=result.duration_seconds,
        )
        
        response_data = {
            "text": result.text,
            "segments": [
                {
                    "speaker": seg.speaker,
                    "text": seg.text,
                    "start": seg.start,
                    "end": seg.end,
                }
                for seg in result.segments
            ],
            "speaker_count": result.speaker_count,
            "duration_seconds": result.duration_seconds,
            "fragments_ingested": None,
            "saved_path": None,
        }
        
        # PASO 1: Guardar DOCX a temp, subir a Blob Storage tenant-scoped, y establecer saved_path al URL del blob
        from datetime import datetime
        import tempfile
        from hashlib import sha256
        from app.blob_storage import upload_local_path, logical_path_to_blob_name, CONTAINER_INTERVIEWS, CONTAINER_AUDIO

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmpdoc:
            tmp_doc_path = Path(tmpdoc.name)
        save_transcription_docx(result, tmp_doc_path)

        # Logical path under tenant/project
        logical = f"interviews/{project_id}/audio/transcriptions/{tmp_doc_path.name}"
        try:
            docx_blob_name = logical_path_to_blob_name(org_id=(getattr(user, "organization_id", None) or None), project_id=project_id, logical_path=logical)
        except ValueError as exc:
            # Tenant guard: do not allow legacy local-style blobs in strict mode.
            msg = str(exc or "")
            if "org_id is required" in msg or "org_id is required" in msg.lower():
                log.warning("api.transcribe.missing_org", project=project_id, filename=tmp_doc_path.name)
                raise HTTPException(status_code=409, detail="Organization-scoped storage required") from exc
            # propagate other validation errors
            raise
        except Exception:
            # Fallback for unexpected errors: keep legacy behavior
            docx_blob_name = f"{project_id}/{tmp_doc_path.name}"

        docx_url = upload_local_path(container=CONTAINER_INTERVIEWS, blob_name=docx_blob_name, file_path=str(tmp_doc_path), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        # compute hash for docx
        hdoc = sha256()
        with open(tmp_doc_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                hdoc.update(chunk)
        docx_hash = hdoc.hexdigest()

        response_data["saved_path"] = docx_url
        response_data["docx_blob"] = {"container": CONTAINER_INTERVIEWS, "name": docx_blob_name, "url": docx_url, "sha256": docx_hash}

        log.info(
            "api.transcribe.saved",
            path=docx_url,
            project=project_id,
            filename=tmp_doc_path.name,
        )

        # PASO 2: Ingestar al pipeline si se solicita (usando el archivo temp local)
        if payload.ingest:
            from app.ingestion import ingest_documents

            clients = build_clients_or_error(settings)
            try:
                ingest_result = ingest_documents(
                    clients,
                    settings,
                    files=[str(tmp_doc_path)],
                    batch_size=20,
                    min_chars=payload.min_chars,
                    max_chars=payload.max_chars,
                    logger=log,
                    project=project_id,
                    org_id=str(getattr(user, "organization_id", None) or ""),
                )
                totals = ingest_result.get("totals", {})
                response_data["fragments_ingested"] = totals.get("fragments", 0)
                log.info(
                    "api.transcribe.ingested",
                    fragments=response_data["fragments_ingested"],
                    docx_path=docx_url,
                    project=project_id,
                )
            finally:
                clients.close()
        # Clean up local temp
        try:
            tmp_doc_path.unlink(missing_ok=True)
        except Exception:
            pass
        
        return response_data
        
    except ValueError as exc:
        log.error("api.transcribe.error", error=str(exc))
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        log.exception("api.transcribe.error")
        raise HTTPException(status_code=500, detail=f"Error en transcripción: {exc}") from exc
    finally:
        tmp_path.unlink(missing_ok=True)


# =============================================================================
# BATCH TRANSCRIPTION (PARALLEL PROCESSING)
# =============================================================================

from backend.celery_worker import task_transcribe_audio
import uuid

MAX_BATCH_SIZE = 20


class StreamTranscribeRequest(BaseModel):
    """Request para transcripción asíncrona con entregas incrementales."""
    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="Proyecto destino")
    audio_base64: str = Field(..., description="Audio codificado en base64")
    filename: str = Field(..., description="Nombre del archivo")
    diarize: bool = Field(True, description="Usar diarización")
    language: str = Field("es", description="Código de idioma")
    ingest: bool = Field(True, description="Ingestar al pipeline")
    min_chars: int = Field(200)
    max_chars: int = Field(1200)


class StreamTranscribeResponse(BaseModel):
    """Respuesta de transcripción stream: task_id para polling."""
    task_id: str
    filename: str
    message: str


@app.post("/api/transcribe/stream", response_model=StreamTranscribeResponse)
async def api_transcribe_stream(
    payload: StreamTranscribeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Inicia transcripción asíncrona con entregas incrementales.
    
    A diferencia de /api/transcribe (síncrono), este endpoint:
    1. Retorna inmediatamente con un task_id
    2. Procesa en background con Celery worker
    3. Guarda DOCX parcial después de cada chunk (5 min de audio)
    4. Ingesta fragmentos incrementalmente si ingest=True
    
    Usa GET /api/jobs/{task_id}/status para consultar progreso:
    - chunk_index: Chunk actual en proceso
    - total_chunks: Total de chunks
    - chunks_completed: Chunks ya completados
    - text_preview: Preview del texto transcrito hasta ahora
    
    Cuando status="SUCCESS", el resultado completo está en result.
    """
    log = api_logger.bind(endpoint="transcribe_stream", project=payload.project)
    
    # Validar proyecto
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    # Validar extensión
    suffix = Path(payload.filename).suffix.lower()
    from app.transcription import SUPPORTED_AUDIO_FORMATS
    if suffix not in SUPPORTED_AUDIO_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado: {suffix}. Formatos válidos: {', '.join(SUPPORTED_AUDIO_FORMATS)}"
        )
    
    # Encolar tarea con incremental=True
    task = cast(Any, task_transcribe_audio).delay(
        org_id=getattr(user, "organization_id", None),
        audio_base64=payload.audio_base64,
        filename=payload.filename,
        project_id=project_id,
        diarize=payload.diarize,
        language=payload.language,
        ingest=payload.ingest,
        min_chars=payload.min_chars,
        max_chars=payload.max_chars,
        incremental=True,  # Habilitar procesamiento incremental
    )
    
    log.info(
        "api.transcribe_stream.started",
        task_id=task.id,
        filename=payload.filename,
    )
    
    return {
        "task_id": task.id,
        "filename": payload.filename,
        "message": "Transcripción iniciada. Usa GET /api/jobs/{task_id}/status para consultar progreso.",
    }


@app.post("/api/transcribe/batch", response_model=BatchTranscribeResponse)
async def api_transcribe_batch(
    payload: BatchTranscribeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Inicia transcripción paralela de múltiples archivos.
    
    Los archivos se procesan en background por workers Celery.
    Retorna IDs de jobs para tracking del progreso.
    
    Límite: 20 archivos por batch.
    """
    log = api_logger.bind(endpoint="transcribe_batch", project=payload.project)
    
    # Validar proyecto
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    # Validar tamaño del batch
    if len(payload.files) > MAX_BATCH_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Máximo {MAX_BATCH_SIZE} archivos por batch. Recibidos: {len(payload.files)}"
        )
    
    if len(payload.files) == 0:
        raise HTTPException(status_code=400, detail="Debe incluir al menos un archivo")
    
    # Generar batch ID
    batch_id = f"batch_{uuid.uuid4().hex[:12]}"
    
    # Encolar cada archivo como job independiente
    jobs = []
    for file_item in payload.files:
        task = cast(Any, task_transcribe_audio).delay(
            org_id=getattr(user, "organization_id", None),
            audio_base64=file_item.audio_base64,
            filename=file_item.filename,
            project_id=project_id,
            diarize=payload.diarize,
            language=payload.language,
            ingest=payload.ingest,
            min_chars=payload.min_chars,
            max_chars=payload.max_chars,
        )
        jobs.append({
            "task_id": task.id,
            "filename": file_item.filename,
        })
    
    log.info(
        "api.transcribe_batch.started",
        batch_id=batch_id,
        job_count=len(jobs),
        project=project_id,
    )
    
    return {
        "batch_id": batch_id,
        "jobs": jobs,
        "message": f"Iniciados {len(jobs)} jobs de transcripción en paralelo",
    }


@app.get("/api/jobs/{task_id}/status", response_model=JobStatusResponse)
async def api_job_status(
    task_id: str,
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Consulta el estado de un job de transcripción.
    
    Estados posibles:
    - PENDING: En cola, esperando worker
    - PROCESSING: En ejecución (transcribiendo o ingesting)
    - SUCCESS: Completado exitosamente
    - FAILURE: Falló
    
    Para transcripciones con incremental=True, el meta incluye:
    - chunk_index: Chunk actual siendo procesado
    - total_chunks: Total de chunks a procesar
    - chunks_completed: Número de chunks ya completados
    - text_preview: Preview del texto transcrito hasta ahora (500 chars)
    - segments_count: Número de segmentos transcritos
    - speakers_count: Número de speakers detectados
    """
    result = AsyncResult(task_id, app=celery_app)
    
    status = result.status
    response = {
        "task_id": task_id,
        "status": status,
        "filename": None,
        "stage": None,
        "result": None,
        "error": None,
    }
    
    if status == "PROCESSING":
        # Get meta info (stage, filename, chunk progress)
        meta = result.info or {}
        response["stage"] = meta.get("stage")
        response["filename"] = meta.get("filename")
        # Agregar info de chunks para transcripción incremental
        response["chunk_index"] = meta.get("chunk_index")
        response["total_chunks"] = meta.get("total_chunks")
        response["chunks_completed"] = meta.get("chunks_completed", 0)
        response["text_preview"] = meta.get("text_preview")
        response["segments_count"] = meta.get("segments_count")
        response["speakers_count"] = meta.get("speakers_count")
    
    elif status == "SUCCESS":
        task_result = result.result or {}
        response["result"] = task_result
        response["filename"] = task_result.get("filename")
    
    elif status == "FAILURE":
        response["error"] = str(result.result) if result.result else "Error desconocido"
    
    return response



@app.get("/api/jobs/batch/{batch_id}/status")
async def api_batch_status(
    batch_id: str,
    task_ids: str,  # Comma-separated list of task IDs
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Consulta el estado de múltiples jobs de un batch.
    """
    ids = [tid.strip() for tid in task_ids.split(",") if tid.strip()]
    
    statuses = []
    completed = 0
    failed = 0
    processing = 0
    pending = 0
    
    for tid in ids:
        result = AsyncResult(tid, app=celery_app)
        status = result.status
        
        if status == "SUCCESS":
            completed += 1
        elif status == "FAILURE":
            failed += 1
        elif status == "PROCESSING":
            processing += 1
        else:
            pending += 1
        
        statuses.append({
            "task_id": tid,
            "status": status,
        })
    
    return {
        "batch_id": batch_id,
        "total": len(ids),
        "completed": completed,
        "failed": failed,
        "processing": processing,
        "pending": pending,
        "all_done": (completed + failed) == len(ids),
        "jobs": statuses,
    }


class TranscribeMergeSegment(BaseModel):
    speaker: str
    text: str
    start: float
    end: float


class TranscribeMergeItem(BaseModel):
    filename: str
    text: str
    segments: List[TranscribeMergeSegment] = []


class TranscribeMergeRequest(BaseModel):
    """Request para combinar múltiples transcripciones en un DOCX."""
    model_config = ConfigDict(extra="forbid")
    
    project: str
    transcriptions: List[TranscribeMergeItem]


class TranscribeMergeResponse(BaseModel):
    docx_base64: str
    filename: str


@app.post("/api/transcribe/merge", response_model=TranscribeMergeResponse)
async def api_transcribe_merge(
    payload: TranscribeMergeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Combina múltiples transcripciones en un único archivo DOCX.
    Retorna el archivo como base64 para descarga directa.
    """
    import base64
    import tempfile
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Validar proyecto
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    log = api_logger.bind(endpoint="transcribe_merge", project=project_id)
    
    if not payload.transcriptions:
        raise HTTPException(status_code=400, detail="No hay transcripciones para combinar")
    
    # Validar que al menos una transcripción tenga contenido real
    valid_count = sum(
        1 for t in payload.transcriptions
        if (t.text and t.text.strip()) or (t.segments and len(t.segments) > 0)
    )
    
    if valid_count == 0:
        log.warning(
            "api.transcribe_merge.empty_content",
            total=len(payload.transcriptions),
            filenames=[t.filename for t in payload.transcriptions],
        )
        raise HTTPException(
            status_code=400,
            detail="Las transcripciones no tienen contenido. Espera a que las transcripciones terminen de procesarse."
        )
    
    try:
        # Crear documento
        doc = Document()
        
        # Título
        title = doc.add_heading(f"Transcripción Combinada - {project_id}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Archivos: {len(payload.transcriptions)}")
        doc.add_paragraph("")
        
        # Agregar cada transcripción
        for item in payload.transcriptions:
            # Header del archivo
            doc.add_heading(f"📁 {item.filename}", level=1)
            
            if item.segments:
                # Con segmentos/speakers
                for seg in item.segments:
                    para = doc.add_paragraph()
                    # Speaker en bold
                    speaker_run = para.add_run(f"[{seg.speaker}] ")
                    speaker_run.bold = True
                    speaker_run.font.color.rgb = RGBColor(0, 100, 180)
                    # Texto
                    para.add_run(seg.text)
            else:
                # Solo texto
                doc.add_paragraph(item.text)
            
            doc.add_paragraph("")  # Separador
        
        # Guardar a temp file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        
        doc.save(str(tmp_path))
        
        # Upload combined DOCX to Blob Storage tenant-scoped instead of saving locally
        import tempfile
        from hashlib import sha256
        from app.blob_storage import upload_local_path, logical_path_to_blob_name, CONTAINER_INTERVIEWS

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        doc.save(str(tmp_path))

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        logical = f"interviews/{project_id}/audio/transcriptions/combined_{timestamp}.docx"
        try:
            blob_name = logical_path_to_blob_name(org_id=str(getattr(user, "organization_id", None) or ""), project_id=project_id, logical_path=logical)
        except Exception:
            blob_name = f"{project_id}/combined_{timestamp}.docx"

        blob_url = upload_local_path(container=CONTAINER_INTERVIEWS, blob_name=blob_name, file_path=str(tmp_path), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        # compute sha256
        h = sha256()
        with open(tmp_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        doc_hash = h.hexdigest()

        log.info("api.transcribe_merge.saved_blob", blob=blob_name, url=blob_url)
        # Clean temp
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass
        
        # Leer como base64
        with open(tmp_path, "rb") as f:
            docx_bytes = f.read()
        docx_base64 = base64.b64encode(docx_bytes).decode("utf-8")
        
        # Limpiar temp
        tmp_path.unlink(missing_ok=True)
        
        return {
            "docx_base64": docx_base64,
            "filename": f"transcripciones_{project_id}_{timestamp}.docx",
        }
        
    except Exception as exc:
        log.exception("api.transcribe_merge.error")
        raise HTTPException(status_code=500, detail=f"Error generando DOCX: {exc}") from exc


# ---------- Coding ----------


class CodingAssignRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    project: str
    fragment_id: str
    codigo: str
    cita: str
    fuente: Optional[str] = None
    memo: Optional[str] = None


@app.post("/api/coding/assign")
async def api_coding_assign(
    payload: CodingAssignRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    log = api_logger.bind(endpoint="coding.assign", project=project_id)
    try:
        result = assign_open_code(
            clients,
            settings,
            fragment_id=payload.fragment_id,
            codigo=payload.codigo,
            cita=payload.cita,
            fuente=payload.fuente,
            memo=payload.memo,
            project=project_id,
            logger=log,
        )
    except CodingError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    finally:
        clients.close()
    return result


class CodingUnassignRequest(BaseModel):
    """Request para desvincular un código de un fragmento."""
    model_config = ConfigDict(extra="forbid")

    project: str
    fragment_id: str
    codigo: str


@app.delete("/api/coding/unassign")
async def api_coding_unassign(
    payload: CodingUnassignRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Desvincula un código de un fragmento.
    
    - Elimina la relación código-fragmento en PostgreSQL y Neo4j
    - NO elimina el fragmento ni el código (pueden tener otras asociaciones)
    - Útil para corregir asignaciones erróneas sin perder datos
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    log = api_logger.bind(endpoint="coding.unassign", project=project_id)
    try:
        result = unassign_open_code(
            clients,
            settings,
            fragment_id=payload.fragment_id,
            codigo=payload.codigo,
            project=project_id,
            changed_by=user.user_id if user else None,
            logger=log,
        )
    except CodingError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    finally:
        clients.close()
    return result


class CodingSuggestRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    project: str
    fragment_id: str
    top_k: int = 5
    archivo: Optional[str] = None
    area_tematica: Optional[str] = None
    actor_principal: Optional[str] = None
    requiere_protocolo_lluvia: Optional[bool] = None
    include_coded: bool = False
    run_id: Optional[str] = None
    persist: bool = False
    llm_model: Optional[str] = None


@app.post("/api/coding/suggest")
async def api_coding_suggest(
    payload: CodingSuggestRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    filters = {
        "archivo": payload.archivo,
        "area_tematica": payload.area_tematica,
        "actor_principal": payload.actor_principal,
        "requiere_protocolo_lluvia": payload.requiere_protocolo_lluvia,
    }
    filters = {k: v for k, v in filters.items() if v is not None}
    try:
        result = suggest_similar_fragments(
            clients,
            settings,
            fragment_id=payload.fragment_id,
            top_k=payload.top_k,
            filters=filters,
            exclude_coded=not payload.include_coded,
            run_id=payload.run_id,
            project=project_id,
            persist=payload.persist,
            llm_model=payload.llm_model,
            logger=api_logger.bind(endpoint="coding.suggest", project=project_id),
        )
    except CodingError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    finally:
        clients.close()
    return result


@app.get("/api/coding/stats")
async def api_coding_stats(
    project: str = Query(..., description="Proyecto requerido"),
    clients: PgOnlyClients = Depends(get_pg_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Get coding stats with a hard PG timeout to avoid UI 30s hangs."""
    import structlog
    _logger = structlog.get_logger("app.api.coding")

    try:
        project_id = await run_pg_query_with_timeout(
            resolve_project, project, allow_create=False, pg=clients.postgres
        )

        def _coding_stats_pg_only() -> Dict[str, Any]:
            # Keep everything inside the worker thread to honor the timeout.
            ensure_open_coding_table(clients.postgres)
            ensure_axial_table(clients.postgres)
            return coding_stats(clients.postgres, project_id)

        return await run_pg_query_with_timeout(_coding_stats_pg_only)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except HTTPException:
        # run_pg_query_with_timeout already returns 504 on timeout
        raise
    except Exception as exc:
        _logger.warning("api.coding.stats.timeout_or_error", error=str(exc), project=project)
        return {
            "total_codes": 0,
            "total_fragments": 0,
            "coded_fragments": 0,
            "coverage_percent": 0.0,
            "codes_per_fragment_avg": 0.0,
            "error": "Stats temporarily unavailable",
        }


@app.get("/api/coding/gate")
async def api_coding_gate(
    project: str = Query(..., description="Proyecto requerido"),
    threshold_count: int = Query(50, description="Máximo de pendientes permitidos"),
    threshold_days: int = Query(3, description="Días máximos sin resolver"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Verifica si es seguro ejecutar un nuevo análisis LLM.
    
    El gate bloquea nuevos análisis si:
    - Hay demasiados candidatos pendientes (> threshold_count)
    - Hay candidatos sin resolver por muchos días (> threshold_days)
    
    Esto evita que el backlog crezca sin control.
    
    Returns:
        - can_proceed: True si puede ejecutar análisis, False si debe validar primero
        - reason: Explicación si está bloqueado
        - health: Métricas completas del backlog
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        health = get_backlog_health(
            clients.postgres,
            project_id,
            threshold_days=threshold_days,
            threshold_count=threshold_count,
        )
        
        can_proceed = health["is_healthy"]
        reason = None
        
        if not can_proceed:
            if health["pending_count"] >= threshold_count:
                reason = f"Backlog saturado: {health['pending_count']} pendientes (máx: {threshold_count})"
            elif health["oldest_pending_days"] >= threshold_days:
                reason = f"Hay candidatos sin validar desde hace {health['oldest_pending_days']} días"
        
        return {
            "can_proceed": can_proceed,
            "reason": reason,
            "health": health,
            "recommendation": (
                "Valide los candidatos pendientes antes de ejecutar nuevo análisis"
                if not can_proceed else None
            ),
        }
    finally:
        clients.close()


@app.get("/api/coding/fragment-context")
async def api_coding_fragment_context(
    project: str = Query(..., description="Proyecto requerido"),
    fragment_id: str = Query(..., description="ID del fragmento"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene el contexto completo de un fragmento para visualización en modal.
    
    Retorna:
        - fragment: Datos completos del fragmento (texto, metadatos)
        - codes: Lista de códigos asignados a este fragmento
        - codes_count: Total de códigos
        - adjacent_fragments: Fragmentos anteriores/posteriores para contexto
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    try:
        context = fetch_fragment_context(clients, project_id, fragment_id)
        if not context:
            raise HTTPException(status_code=404, detail=f"Fragmento '{fragment_id}' no encontrado")
    finally:
        clients.close()
    return context


@app.get("/api/interviews")
async def api_interviews(
    limit: int = 25,
    order: str = Query(
        default="ingest-desc",
        description="Orden de entrevistas: ingest-desc|ingest-asc|alpha|fragments-desc|fragments-asc|max-variation|theoretical-sampling",
    ),
    include_analyzed: bool = Query(
        default=False,
        description="Solo aplica a order=theoretical-sampling. Si true, incluye entrevistas ya analizadas al final del listado.",
    ),
    focus_codes: Optional[str] = Query(
        default=None,
        description="Solo aplica a order=theoretical-sampling. CSV de códigos foco (ej: 'Dificultad de pago,Acceso a subsidio').",
    ),
    recent_window: int = Query(
        default=3,
        ge=1,
        le=20,
        description="Solo aplica a order=theoretical-sampling. Ventana de últimos N informes para detectar saturación.",
    ),
    saturation_new_codes_threshold: int = Query(
        default=2,
        ge=0,
        le=50,
        description="Solo aplica a order=theoretical-sampling. Umbral de suma de codigos_nuevos en recent_window.",
    ),
    project: str = Query(..., description="Proyecto requerido"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
        normalized_order = (order or "").strip().lower()
        if normalized_order == "theoretical-sampling":
            interviews, ranking_debug = list_available_interviews_with_ranking_debug(
                clients,
                project_id,
                limit=limit,
                order=order,
                include_analyzed=include_analyzed,
                focus_codes=focus_codes,
                recent_window=recent_window,
                saturation_new_codes_threshold=saturation_new_codes_threshold,
            )
            return {
                "interviews": interviews,
                "order": order,
                "ranking_debug": ranking_debug,
            }

        return {
            "interviews": list_available_interviews(
                clients,
                project_id,
                limit=limit,
                order=order,
                include_analyzed=include_analyzed,
                focus_codes=focus_codes,
                recent_window=recent_window,
                saturation_new_codes_threshold=saturation_new_codes_threshold,
            ),
            "order": order,
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

# =============================================================================
# CÓDIGOS CANDIDATOS - Sistema de Consolidación
# =============================================================================

class CandidateCodeRequest(BaseModel):
    """Request para crear un código candidato."""
    model_config = ConfigDict(extra="forbid")

    project: str
    codigo: str
    cita: Optional[str] = None
    fragmento_id: Optional[str] = None
    archivo: Optional[str] = None
    fuente_origen: str = "manual"  # 'llm', 'manual', 'discovery', 'semantic_suggestion'
    fuente_detalle: Optional[str] = None
    score_confianza: Optional[float] = None
    memo: Optional[str] = None


class ValidateCandidateRequest(BaseModel):
    """Request para validar/rechazar un candidato."""
    model_config = ConfigDict(extra="forbid")

    project: str
    memo: Optional[str] = None


class MergeCandidatesRequest(BaseModel):
    """Request para fusionar códigos candidatos."""
    model_config = ConfigDict(extra="forbid")

    project: str
    source_ids: List[int]
    target_codigo: str
    memo: Optional[str] = None
    dry_run: bool = False
    idempotency_key: Optional[str] = None


class PromoteCandidatesRequest(BaseModel):
    """Request para promover candidatos a definitivos."""
    model_config = ConfigDict(extra="forbid")

    project: str
    candidate_ids: Optional[List[int]] = None
    promote_all_validated: bool = False


class RevertValidatedCandidatesRequest(BaseModel):
    """Request para revertir candidatos validados a pendientes."""
    model_config = ConfigDict(extra="forbid")

    project: str
    memo: Optional[str] = None
    dry_run: bool = False


@app.get("/api/codes/candidates")
async def api_list_candidates(
    project: Optional[str] = Query(None, description="Proyecto requerido"),
    project_id: Optional[str] = Query(None, description="Alias de 'project' (compatibilidad)"),
    estado: Optional[str] = Query(None, description="Filtrar por estado: pendiente, validado, rechazado, fusionado"),
    fuente_origen: Optional[str] = Query(None, description="Filtrar por origen: llm, manual, discovery, semantic_suggestion"),
    archivo: Optional[str] = Query(None, description="Filtrar por archivo"),
    promovido: Optional[bool] = Query(None, description="Filtrar por promoción: true (ya promovidos) / false (pendientes por promover)"),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    sort_order: str = Query("desc", description="Ordenamiento: 'desc' (recientes primero) o 'asc' (antiguos primero)"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista códigos candidatos con filtros opcionales."""
    try:
        effective_project = project or project_id
        if not effective_project:
            raise HTTPException(status_code=400, detail="Missing required query param: project")
        project_id = resolve_project(effective_project, allow_create=False, pg=clients.postgres)
        candidates = list_candidate_codes(
            clients.postgres,
            project=project_id,
            estado=estado,
            fuente_origen=fuente_origen,
            archivo=archivo,
            promovido=promovido,
            limit=limit,
            offset=offset,
            sort_order=sort_order,
        )
        return {"candidates": candidates, "count": len(candidates)}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/codes/candidates/pending_count")
async def api_candidate_pending_count(
    project: Optional[str] = Query(None, description="Proyecto requerido"),
    project_id: Optional[str] = Query(None, description="Alias de 'project' (compatibilidad)"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Retorna el total de candidatos pendientes (sin depender de limit/offset)."""
    try:
        effective_project = project or project_id
        if not effective_project:
            raise HTTPException(status_code=400, detail="Missing required query param: project")
        resolved = resolve_project(effective_project, allow_create=False, pg=clients.postgres)
        from app.postgres_block import count_pending_candidates

        pending_count = count_pending_candidates(clients.postgres, resolved)
        return {"project": resolved, "pending_count": pending_count}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/notes/{project}/download")
async def api_download_note(
    project: str,
    rel: str = Query(..., description="Ruta relativa bajo notes/<project>/"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Response:
    """Descarga segura de un memo dentro de notes/<project>/ (Blob Storage, multi-tenant)."""
    try:
        project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    if not rel or not rel.strip():
        raise HTTPException(status_code=400, detail="Missing required query param: rel")

    rel_norm = rel.replace("\\", "/").lstrip("/")
    target_rel = Path(rel_norm)
    if target_rel.is_absolute() or ".." in target_rel.parts:
        raise HTTPException(status_code=400, detail="Invalid rel path")

    logical_path = f"notes/{project_id}/{rel_norm}"
    try:
        from app.blob_storage import CONTAINER_REPORTS, download_file, logical_path_to_blob_name

        blob_name = logical_path_to_blob_name(
            org_id=str(getattr(user, "organization_id", None) or ""),
            project_id=project_id,
            logical_path=logical_path,
        )
        data = download_file(CONTAINER_REPORTS, blob_name)
    except Exception:
        # Backward-compat: fallback to local filesystem if present.
        base_dir = (Path("notes") / project_id).resolve()
        target = (base_dir / rel_norm).resolve()
        if base_dir not in target.parents and target != base_dir:
            raise HTTPException(status_code=400, detail="Invalid rel path")
        if not target.exists() or not target.is_file():
            raise HTTPException(status_code=404, detail="Memo not found")
        data = target.read_bytes()

    suffix = target_rel.suffix.lower()
    media_type = "text/markdown" if suffix in {".md", ".markdown"} else "text/plain"
    headers = {"Content-Disposition": f'attachment; filename="{target_rel.name}"'}
    return Response(content=data, media_type=media_type, headers=headers)


@app.post("/api/codes/candidates")
async def api_create_candidate(
    payload: CandidateCodeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Crea un nuevo código candidato."""
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        count = insert_candidate_codes(
            clients.postgres,
            candidates=[{
                "project_id": project_id,
                "codigo": payload.codigo,
                "cita": payload.cita,
                "fragmento_id": payload.fragmento_id,
                "archivo": payload.archivo,
                "fuente_origen": payload.fuente_origen,
                "fuente_detalle": payload.fuente_detalle,
                "score_confianza": payload.score_confianza,
                "estado": "pendiente",
                "memo": payload.memo,
            }],
        )
    finally:
        clients.close()
    
    return {"success": count > 0, "inserted": count}


@app.post("/api/codes/candidates/revert-validated")
async def api_revert_validated_candidates(
    payload: RevertValidatedCandidatesRequest,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Revierte todos los candidatos validados a pendientes para un proyecto.

    Operación acotada a `codigos_candidatos` (no retrocede promociones a tablas definitivas).
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    from app.postgres_block import revert_validated_candidates_to_pending

    result = revert_validated_candidates_to_pending(
        clients.postgres,
        project_id=project_id,
        reverted_by=(user.user_id if user else None),
        memo=payload.memo,
        dry_run=bool(payload.dry_run),
    )

    api_logger.info(
        "codes.candidates.revert_validated",
        project_id=project_id,
        dry_run=bool(payload.dry_run),
        reverted_count=int(result.get("reverted_count", 0) or 0),
        would_revert=int(result.get("would_revert", 0) or 0),
        user_id=(user.user_id if user else None),
    )

    return {"success": True, **result}


@app.put("/api/codes/candidates/{candidate_id}/validate")
async def api_validate_candidate(
    candidate_id: int,
    payload: ValidateCandidateRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Marca un código candidato como validado."""
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Pre-validación robusta: evitar falsos positivos antes de validar
        with clients.postgres.cursor() as cur:
            cur.execute(
                "SELECT codigo, estado FROM codigos_candidatos WHERE project_id = %s AND id = %s",
                (project_id, candidate_id),
            )
            row = cur.fetchone()

        if row:
            candidate_code, candidate_state = row[0], row[1]
            if candidate_code and candidate_state != "validado":
                norm_candidate = _normalize_text(candidate_code)
                threshold_pre = float(os.getenv("PREVALIDATE_DUPLICATE_THRESHOLD", "0.9"))
                with clients.postgres.cursor() as cur:
                    cur.execute(
                        """
                        SELECT DISTINCT codigo
                        FROM codigos_candidatos
                        WHERE project_id = %s
                          AND estado IN ('pendiente','validado')
                          AND id <> %s
                        """,
                        (project_id, candidate_id),
                    )
                    rows = cur.fetchall()

                for (other_code,) in rows:
                    if not other_code:
                        continue
                    norm_other = _normalize_text(other_code)
                    if norm_other == norm_candidate and other_code.strip() != candidate_code.strip():
                        raise HTTPException(
                            status_code=409,
                            detail=f"Código duplicado exacto detectado: '{other_code}'.",
                        )
                    max_len = max(len(norm_candidate), len(norm_other))
                    if max_len == 0:
                        continue
                    if abs(len(norm_candidate) - len(norm_other)) > int((1 - threshold_pre) * max_len):
                        continue
                    distance = _levenshtein_distance(norm_candidate, norm_other)
                    similarity = 1 - (distance / max_len)
                    if similarity >= threshold_pre and distance > 0:
                        if not _token_overlap_ok(candidate_code, other_code):
                            continue
                        raise HTTPException(
                            status_code=409,
                            detail=(
                                "Código muy similar detectado antes de validar: "
                                f"'{other_code}' ({similarity:.0%})."
                            ),
                        )

        success = validate_candidate(
            clients.postgres,
            project_id=project_id,
            candidate_id=candidate_id,
            validated_by=user.user_id if user else None,
            memo=payload.memo,
        )

        # UX fix: once validated, make it visible in Etapa 3 immediately.
        # The Coding panel reads from the definitive open-codes table via /api/coding/codes.
        promoted_count = (
            promote_to_definitive(
                clients.postgres,
                project_id=project_id,
                candidate_ids=[candidate_id],
                promoted_by=user.user_id if user else None,
            )
            if success
            else 0
        )
    finally:
        clients.close()
    
    if not success:
        raise HTTPException(status_code=404, detail="Candidato no encontrado o ya procesado")

    api_logger.info(
        "codes.candidates.validated",
        project_id=project_id,
        candidate_id=candidate_id,
        promoted_count=promoted_count,
    )

    return {
        "success": True,
        "candidate_id": candidate_id,
        "estado": "validado",
        "promoted_count": promoted_count,
    }


@app.put("/api/codes/candidates/{candidate_id}/reject")
async def api_reject_candidate(
    candidate_id: int,
    payload: ValidateCandidateRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Marca un código candidato como rechazado."""
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        success = reject_candidate(
            clients.postgres,
            project_id=project_id,
            candidate_id=candidate_id,
            rejected_by=user.user_id if user else None,
            memo=payload.memo,
        )
    finally:
        clients.close()
    
    if not success:
        raise HTTPException(status_code=404, detail="Candidato no encontrado o ya procesado")
    
    return {"success": True, "candidate_id": candidate_id, "estado": "rechazado"}


@app.post("/api/codes/candidates/merge")
async def api_merge_candidates(
    payload: MergeCandidatesRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Fusiona múltiples códigos candidatos en uno."""
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    if not payload.source_ids:
        raise HTTPException(status_code=400, detail="source_ids no puede estar vacío")

    if not payload.target_codigo or not payload.target_codigo.strip():
        raise HTTPException(status_code=400, detail="target_codigo no puede estar vacío")
    
    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import (  # local import to avoid module init cycles
            ensure_api_idempotency_table,
            get_idempotency_response,
            store_idempotency_response,
            preview_merge_candidates_by_ids,
        )

        endpoint_key = "POST /api/codes/candidates/merge"
        ensure_api_idempotency_table(clients.postgres)

        if payload.idempotency_key:
            cached = get_idempotency_response(
                clients.postgres,
                project_id=project_id,
                endpoint=endpoint_key,
                idempotency_key=payload.idempotency_key,
            )
            if cached is not None:
                return cached

        if payload.dry_run:
            preview = preview_merge_candidates_by_ids(
                clients.postgres,
                project_id=project_id,
                source_ids=payload.source_ids,
                target_codigo=payload.target_codigo,
            )
            resp = {
                "success": (preview.get("would_merge", 0) > 0),
                "dry_run": True,
                "merged_count": int(preview.get("would_merge", 0)),
                "details": preview,
                "target_codigo": payload.target_codigo,
            }
        else:
            count = merge_candidates(
                clients.postgres,
                project_id=project_id,
                source_ids=payload.source_ids,
                target_codigo=payload.target_codigo,
                merged_by=user.user_id if user else None,
                memo=payload.memo,
            )
            resp = {
                "success": count > 0,
                "dry_run": False,
                "merged_count": count,
                "target_codigo": payload.target_codigo,
            }

        if payload.idempotency_key:
            store_idempotency_response(
                clients.postgres,
                project_id=project_id,
                endpoint=endpoint_key,
                idempotency_key=payload.idempotency_key,
                response=resp,
            )
    finally:
        clients.close()

    return resp


@app.post("/api/coding/candidates/merge")
async def api_merge_candidates_legacy(
    payload: MergeCandidatesRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Alias legacy (deprecado): usar /api/codes/candidates/merge."""
    return await api_merge_candidates(payload=payload, settings=settings, user=user)


class AutoMergePair(BaseModel):
    """Par de códigos a fusionar."""
    source_codigo: str = Field(..., description="Código fuente a fusionar")
    target_codigo: str = Field(..., description="Código destino")


class AutoMergeRequest(BaseModel):
    """Request para auto-fusión masiva de duplicados."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="ID del proyecto")
    pairs: List[AutoMergePair] = Field(..., description="Lista de pares a fusionar")
    memo: Optional[str] = Field(None, description="Memo común para auditoría (opcional)")
    dry_run: bool = Field(False, description="Si true, no persiste cambios")
    idempotency_key: Optional[str] = Field(None, description="Clave de idempotencia (opcional)")


@app.post("/api/codes/candidates/auto-merge")
async def api_auto_merge_candidates(
    payload: AutoMergeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Fusiona masivamente pares de códigos duplicados por nombre.
    
    Más robusto que merge por IDs porque busca directamente por nombre de código.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    if not payload.pairs:
        api_logger.info(
            "api.auto_merge_candidates.empty",
            project=payload.project,
        )
        return {"success": True, "total_merged": 0, "details": []}
    
    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import (
            ensure_api_idempotency_table,
            get_idempotency_response,
            store_idempotency_response,
            preview_merge_candidates_by_code,
        )

        endpoint_key = "POST /api/codes/candidates/auto-merge"
        ensure_api_idempotency_table(clients.postgres)

        if payload.idempotency_key:
            cached = get_idempotency_response(
                clients.postgres,
                project_id=project_id,
                endpoint=endpoint_key,
                idempotency_key=payload.idempotency_key,
            )
            if cached is not None:
                return cached

        details: List[Dict[str, Any]] = []
        total_merged = 0
        api_logger.info(
            "api.auto_merge_candidates.start",
            project=payload.project,
            pairs=len(payload.pairs),
            user_id=user.user_id if user else None,
        )
        
        for pair in payload.pairs:
            if not pair.source_codigo or not pair.target_codigo:
                api_logger.warning(
                    "api.auto_merge_candidates.invalid_pair",
                    project=payload.project,
                    source=pair.source_codigo,
                    target=pair.target_codigo,
                )
                details.append(
                    {
                        "source": pair.source_codigo,
                        "target": pair.target_codigo,
                        "merged_count": 0,
                        "skipped": "invalid_pair",
                    }
                )
                continue
            if pair.source_codigo.strip().lower() == pair.target_codigo.strip().lower():
                details.append({
                    "source": pair.source_codigo,
                    "target": pair.target_codigo,
                    "merged_count": 0,
                    "skipped": "self",
                })
                continue

            if payload.dry_run:
                preview = preview_merge_candidates_by_code(
                    clients.postgres,
                    project_id=project_id,
                    source_codigo=pair.source_codigo,
                    target_codigo=pair.target_codigo,
                )
                count = int(preview.get("would_merge", 0))
                details.append({
                    "source": pair.source_codigo,
                    "target": pair.target_codigo,
                    "merged_count": count,
                    "dry_run": True,
                    "details": preview,
                })
            else:
                count = merge_candidates_by_code(
                    clients.postgres,
                    project_id=project_id,
                    source_codigo=pair.source_codigo,
                    target_codigo=pair.target_codigo,
                    merged_by=user.user_id if user else None,
                    memo=payload.memo,
                )
                details.append({
                    "source": pair.source_codigo,
                    "target": pair.target_codigo,
                    "merged_count": count,
                })
            total_merged += count
        api_logger.info(
            "api.auto_merge_candidates.completed",
            project=payload.project,
            total_merged=total_merged,
            pairs_processed=len(payload.pairs),
        )

        resp: Dict[str, Any] = {
            "success": True if payload.dry_run else (total_merged > 0),
            "dry_run": bool(payload.dry_run),
            "total_merged": total_merged,
            "pairs_processed": len(payload.pairs),
            "details": details,
        }

        if payload.idempotency_key:
            store_idempotency_response(
                clients.postgres,
                project_id=project_id,
                endpoint=endpoint_key,
                idempotency_key=payload.idempotency_key,
                response=resp,
            )

        return resp
    except Exception as exc:
        api_logger.error(
            "api.auto_merge_candidates.error",
            error=str(exc),
            project=payload.project,
        )
        raise HTTPException(status_code=500, detail=f"Error en auto-fusión: {str(exc)}") from exc
    finally:
        clients.close()



# =============================================================================
# HIPÓTESIS Y MUESTREO TEÓRICO (Grounded Theory)
# =============================================================================

class ValidateHypothesisRequest(BaseModel):
    """Request para validar hipótesis con evidencia empírica."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="ID del proyecto")
    fragmento_id: str = Field(..., min_length=10, description="ID del fragmento que respalda la hipótesis")
    cita: str = Field(..., min_length=10, description="Cita textual del fragmento")


class RejectHypothesisRequest(BaseModel):
    """Request para rechazar hipótesis por falta de evidencia."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="ID del proyecto")
    razon: str = Field(..., min_length=5, description="Razón metodológica del rechazo")


@app.get("/api/codes/hypotheses")
async def api_list_hypotheses(
    project: str = Query(..., description="ID del proyecto"),
    limit: int = Query(50, ge=1, le=200),
    settings: AppSettings = Depends(get_settings),
    _user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Lista hipótesis pendientes de muestreo teórico.
    
    En Grounded Theory, las hipótesis son proposiciones emergentes del análisis
    estructural que requieren validación empírica antes de ser aceptadas.
    
    Según Charmaz (2014): "Theoretical sampling means sampling to develop 
    or refine emerging theoretical categories... keeps them grounded in data."
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    from app.postgres_block import list_hypothesis_codes, count_hypothesis_codes
    
    clients = build_clients_or_error(settings)
    try:
        hypotheses = list_hypothesis_codes(clients.postgres, project_id, limit)
        total = count_hypothesis_codes(clients.postgres, project_id)
    finally:
        clients.close()
    
    return {
        "hypotheses": hypotheses,
        "total": total,
        "project": project_id,
        "methodology_note": (
            "Estas hipótesis emergieron del análisis estructural (link prediction). "
            "Requieren validación empírica: busca fragmentos que respalden cada hipótesis, "
            "o rechaza aquellas sin evidencia para mantener el rigor metodológico."
        ),
    }


@app.get("/api/codes/hypotheses/count")
async def api_count_hypotheses(
    project: str = Query(..., description="ID del proyecto"),
    settings: AppSettings = Depends(get_settings),
    _user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Cuenta hipótesis pendientes de muestreo teórico."""
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    from app.postgres_block import count_hypothesis_codes
    
    clients = build_clients_or_error(settings)
    try:
        count = count_hypothesis_codes(clients.postgres, project_id)
    finally:
        clients.close()
    
    return {"count": count, "project": project_id}


@app.post("/api/codes/hypotheses/{hypothesis_id}/validate-with-evidence")
async def api_validate_hypothesis(
    hypothesis_id: int,
    payload: ValidateHypothesisRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Valida una hipótesis vinculándola a evidencia empírica.
    
    Este endpoint implementa el paso final del muestreo teórico en Grounded Theory:
    el investigador encontró un fragmento que respalda la hipótesis, por lo que
    esta se convierte en un código validado con su evidencia empírica correspondiente.
    
    Según Glaser & Strauss (1967): "Grounded theory is defined as the discovery 
    of theory from data."
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    from app.postgres_block import validate_hypothesis_with_evidence
    
    clients = build_clients_or_error(settings)
    try:
        result = validate_hypothesis_with_evidence(
            clients.postgres,
            hypothesis_id=hypothesis_id,
            project=project_id,
            fragmento_id=payload.fragmento_id,
            cita=payload.cita,
            validado_por=user.user_id if user else "investigador",
        )
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    finally:
        clients.close()
    
    return result


@app.post("/api/codes/hypotheses/{hypothesis_id}/reject")
async def api_reject_hypothesis(
    hypothesis_id: int,
    payload: RejectHypothesisRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Rechaza una hipótesis después de muestreo teórico fallido.
    
    En Grounded Theory, las hipótesis que no encuentran evidencia empírica
    deben ser rechazadas para mantener el rigor metodológico. El investigador
    debe documentar la razón del rechazo.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    from app.postgres_block import reject_hypothesis
    
    clients = build_clients_or_error(settings)
    try:
        result = reject_hypothesis(
            clients.postgres,
            hypothesis_id=hypothesis_id,
            project=project_id,
            razon=payload.razon,
            rechazado_por=user.user_id if user else "investigador",
        )
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    finally:
        clients.close()
    
    return result


@app.get("/api/codes/hypotheses/{hypothesis_id}/search-evidence")
async def api_search_evidence_for_hypothesis(
    hypothesis_id: int,
    project: str = Query(..., description="ID del proyecto"),
    limit: int = Query(10, ge=1, le=50),
    settings: AppSettings = Depends(get_settings),
    _user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Busca fragmentos que podrían servir como evidencia para una hipótesis.
    
    Realiza búsqueda semántica en los fragmentos existentes para encontrar
    evidencia empírica que respalde la hipótesis.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    from app.postgres_block import list_hypothesis_codes, search_evidence_for_hypothesis
    
    clients = build_clients_or_error(settings)
    try:
        # Get hypothesis code
        hypotheses = list_hypothesis_codes(clients.postgres, project_id, limit=100)
        hypothesis = next((h for h in hypotheses if h["id"] == hypothesis_id), None)
        
        if not hypothesis:
            raise HTTPException(404, f"Hipótesis {hypothesis_id} no encontrada")
        
        # Search for evidence
        fragments = search_evidence_for_hypothesis(
            clients.postgres,
            project_id,
            hypothesis["codigo"],
            limit=limit,
        )
    finally:
        clients.close()
    
    return {
        "hypothesis": hypothesis,
        "potential_evidence": fragments,
        "methodology_note": (
            "Estos fragmentos contienen términos relacionados con la hipótesis. "
            "Revisa si alguno respalda empíricamente la proposición."
        ),
    }


# =============================================================================
# DUPLICATE DETECTION - Limpieza Post-Hoc de códigos
# =============================================================================

class DetectDuplicatesRequest(BaseModel):
    """Request para detectar códigos duplicados."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto")
    threshold: float = Field(0.80, ge=0.5, le=1.0, description="Umbral de similitud")


def ensure_fuzzystrmatch(pg_conn) -> bool:
    """Habilita la extensión fuzzystrmatch si no existe."""
    try:
        with pg_conn.cursor() as cur:
            cur.execute("CREATE EXTENSION IF NOT EXISTS fuzzystrmatch;")
        pg_conn.commit()
        return True
    except Exception as exc:
        api_logger.warning("fuzzystrmatch.not_available", error=str(exc))
        # Try to rollback the failed transaction
        try:
            pg_conn.rollback()
        except Exception:
            pass
        return False


def ensure_unaccent(pg_conn) -> bool:
    """Habilita la extensión unaccent si no existe."""
    try:
        with pg_conn.cursor() as cur:
            cur.execute("CREATE EXTENSION IF NOT EXISTS unaccent;")
        pg_conn.commit()
        return True
    except Exception as exc:
        api_logger.warning("unaccent.not_available", error=str(exc))
        try:
            pg_conn.rollback()
        except Exception:
            pass
        return False


def _levenshtein_distance(s1: str, s2: str) -> int:
    """Calcula distancia de Levenshtein usando RapidFuzz (compatible con Azure)."""
    return Levenshtein.distance(s1, s2)


def _normalize_text(value: str) -> str:
    """Normaliza texto: trim, lower y remueve acentos (Unicode)."""
    if value is None:
        return ""
    trimmed = value.strip().lower()
    normalized = unicodedata.normalize("NFKD", trimmed)
    return "".join(ch for ch in normalized if not unicodedata.combining(ch))


def _tokenize_code(value: str) -> List[str]:
    """Tokeniza un código normalizado en tokens alfanuméricos."""
    if not value:
        return []
    normalized = _normalize_text(value)
    return [tok for tok in re.split(r"[^a-z0-9]+", normalized) if tok]


def _token_overlap_ok(a: str, b: str, min_jaccard: float = 0.5) -> bool:
    """Valida solapamiento de tokens para evitar falsos positivos semánticos."""
    ta = set(_tokenize_code(a))
    tb = set(_tokenize_code(b))
    if not ta or not tb:
        return False
    inter = ta.intersection(tb)
    union = ta.union(tb)
    if not union:
        return False
    jaccard = len(inter) / len(union)
    if jaccard < min_jaccard:
        return False

    # Si difieren solo en un token, exigir similitud alta entre esos tokens
    diff_a = list(ta - tb)
    diff_b = list(tb - ta)
    if len(diff_a) == 1 and len(diff_b) == 1:
        tok_a, tok_b = diff_a[0], diff_b[0]
        max_len = max(len(tok_a), len(tok_b))
        if max_len == 0:
            return False
        tok_sim = 1 - (_levenshtein_distance(tok_a, tok_b) / max_len)
        return tok_sim >= 0.8

    return True


def find_similar_codes_python_fallback(
    pg_conn,
    project_id: str,
    threshold: float = 0.80,
    limit: int = 50,
) -> List[Dict[str, Any]]:
    """
    Encuentra códigos similares usando RapidFuzz (Levenshtein).
    """
    # Get unique codes from database
    sql = """
    SELECT DISTINCT codigo
    FROM codigos_candidatos
    WHERE project_id = %s AND estado IN ('pendiente', 'validado', 'hipotesis')
    """
    with pg_conn.cursor() as cur:
        cur.execute(sql, (project_id,))
        rows = cur.fetchall()
    
    codes = [row[0] for row in rows if row[0]]
    normalized_codes = [_normalize_text(c) for c in codes]
    rep_by_norm: Dict[str, str] = {}
    for original, norm in zip(codes, normalized_codes):
        rep_by_norm.setdefault(norm, original)
    
    # Find duplicates using Python (RapidFuzz)
    duplicates = []
    
    # Check exact duplicates only when there are multiple distinct variants
    counts: Dict[str, int] = {}
    variants_by_norm: Dict[str, set[str]] = {}
    for original, norm in zip(codes, normalized_codes):
        counts[norm] = counts.get(norm, 0) + 1
        variants_by_norm.setdefault(norm, set()).add(original.strip())
    for norm, count in sorted(counts.items(), key=lambda x: x[1], reverse=True):
        if count <= 1:
            continue
        if len(variants_by_norm.get(norm, set())) <= 1:
            continue
        duplicates.append({
            "code1": rep_by_norm.get(norm, norm),
            "code2": rep_by_norm.get(norm, norm),
            "distance": 0,
            "similarity": 1.0,
            "is_exact_duplicate": True,
            "duplicate_count": count,
        })
        if len(duplicates) >= limit:
            break
    
    # Calculate Levenshtein similarity for unique codes
    unique_codes = list(set(_normalize_text(c) for c in codes))
    
    for i, c1 in enumerate(unique_codes):
        for c2 in unique_codes[i+1:]:
            max_len = max(len(c1), len(c2))
            if max_len == 0:
                continue
            # Pre-filter by length difference
            if abs(len(c1) - len(c2)) > int((1 - threshold) * max_len):
                continue
            
            distance = _levenshtein_distance(c1, c2)
            similarity = 1 - (distance / max_len)
            
            if similarity >= threshold and distance > 0:
                if not _token_overlap_ok(c1, c2):
                    continue
                duplicates.append({
                    "code1": rep_by_norm.get(c1, c1),
                    "code2": rep_by_norm.get(c2, c2),
                    "distance": distance,
                    "similarity": round(similarity, 3),
                    "is_exact_duplicate": False,
                })
                
                if len(duplicates) >= limit:
                    break
        
        if len(duplicates) >= limit:
            break
    
    # Sort by similarity
    duplicates.sort(key=lambda x: x["similarity"], reverse=True)
    return duplicates[:limit]


def find_similar_codes_posthoc(
    pg_conn,
    project_id: str,
    threshold: float = 0.80,
    limit: int = 50,
    include_exact: bool = True,
    use_unaccent: bool = False,
) -> List[Dict[str, Any]]:
    """
    Encuentra pares de códigos candidatos similares usando Levenshtein.
    
    Args:
        include_exact: Si True, incluye duplicados exactos (distancia=0)
    """
    max_distance = int((1 - threshold) * 15)
    normalize_expr = "unaccent(lower(trim(codigo)))" if use_unaccent else "lower(trim(codigo))"
    
    # Query para duplicados exactos (mismo código, diferentes entradas)
    exact_duplicates = []
    if include_exact:
        sql_exact = f"""
        SELECT 
            MIN(codigo) AS codigo,
            COUNT(*) as count,
            array_agg(DISTINCT fuente_origen) as sources,
            array_agg(DISTINCT archivo) FILTER (WHERE archivo IS NOT NULL) as files
        FROM codigos_candidatos
        WHERE project_id = %s AND estado IN ('pendiente', 'validado', 'hipotesis')
        GROUP BY {normalize_expr}
        HAVING COUNT(*) > 1
        ORDER BY count DESC
        LIMIT %s
        """
        
        with pg_conn.cursor() as cur:
            cur.execute(sql_exact, (project_id, limit))
            rows = cur.fetchall()
        
        for row in rows:
            exact_duplicates.append({
                "code1": row[0],
                "code2": row[0],  # Mismo código
                "distance": 0,
                "similarity": 1.0,
                "is_exact_duplicate": True,
                "duplicate_count": row[1],
                "sources": row[2] if row[2] else [],
                "files": row[3][:5] if row[3] else [],  # Max 5 archivos
            })
    
    # Query para similares (distancia > 0)
    # OPTIMIZACIÓN: Filtro de longitud para reducir O(N²) a O(N*k)
    # Solo comparamos códigos cuya diferencia de longitud sea <= max_distance
    # Esto evita comparar "Organización" (12 chars) con "Si" (2 chars)
        sql_similar = f"""
    WITH unique_codes AS (
                SELECT DISTINCT codigo, LENGTH({normalize_expr}) AS len, {normalize_expr} AS norm
        FROM codigos_candidatos
        WHERE project_id = %s AND estado IN ('pendiente', 'validado', 'hipotesis')
    )
    SELECT 
        c1.codigo AS code1,
        c2.codigo AS code2,
                levenshtein(c1.norm, c2.norm) AS distance,
        GREATEST(c1.len, c2.len) AS max_len,
                1.0 - (levenshtein(c1.norm, c2.norm)::float / 
               GREATEST(c1.len, c2.len)::float) AS similarity
    FROM unique_codes c1, unique_codes c2
        WHERE c1.norm < c2.norm
      -- OPTIMIZACIÓN: Pre-filtro por longitud (evita Levenshtein innecesario)
      AND ABS(c1.len - c2.len) <= %s
      -- Solo calcular Levenshtein si pasa el filtro de longitud
            AND levenshtein(c1.norm, c2.norm) <= %s
            AND levenshtein(c1.norm, c2.norm) > 0
    ORDER BY similarity DESC
    LIMIT %s
    """
    

    with pg_conn.cursor() as cur:
        # Parámetros: project_id, max_len_diff, max_distance, limit
        cur.execute(sql_similar, (project_id, max_distance, max_distance, limit))
        rows = cur.fetchall()

    
    similar_codes = [
        {
            "code1": row[0],
            "code2": row[1],
            "distance": row[2],
            "similarity": round(row[4], 3),
            "is_exact_duplicate": False,
        }
        for row in rows
        if row[4] >= threshold
    ]
    
    # Combinar: exactos primero, luego similares
    return exact_duplicates + similar_codes




@app.post("/api/codes/detect-duplicates")
async def api_detect_duplicates(
    payload: DetectDuplicatesRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Detecta códigos duplicados usando similitud de Levenshtein (Post-Hoc).
    
    Útil para limpiar datos históricos que no pasaron por normalización Pre-Hoc.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Asegurar que la tabla existe
        from app.postgres_block import ensure_candidate_codes_table
        ensure_candidate_codes_table(clients.postgres)
        
        # RapidFuzz no requiere extensiones en PostgreSQL (compatible con Azure)
        duplicates = find_similar_codes_python_fallback(
            clients.postgres,
            project_id,
            threshold=payload.threshold,
            limit=50,
        )

        api_logger.info(
            "api.detect_duplicates.completed",
            project=payload.project,
            threshold=payload.threshold,
            count=len(duplicates),
            method="rapidfuzz",
        )
        
        return {
            "success": True,
            "project": project_id,
            "threshold": payload.threshold,
            "duplicates": duplicates,
            "count": len(duplicates),
            "method": "rapidfuzz",
        }
    except HTTPException:
        raise
    except Exception as exc:
        api_logger.error("api.detect_duplicates.error", error=str(exc), project=payload.project)
        raise HTTPException(status_code=500, detail=f"Error detectando duplicados: {str(exc)}") from exc
        resp = {
            "success": True,
            "dry_run": bool(payload.dry_run),
            "total_merged": total_merged,
            "details": details,
        }

        if payload.idempotency_key:
            store_idempotency_response(
                clients.postgres,
                project_id=project_id,
                endpoint=endpoint_key,
                idempotency_key=payload.idempotency_key,
                response=resp,
            )

        return resp
    finally:
        clients.close()


@app.post("/api/coding/candidates/auto-merge")
async def api_auto_merge_candidates_legacy(
    payload: AutoMergeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Alias legacy (deprecado): usar /api/codes/candidates/auto-merge."""
    return await api_auto_merge_candidates(payload=payload, settings=settings, user=user)

# =============================================================================
# POST-HOC DUPLICATE DETECTION FOR NEO4J - Códigos en el grafo
# =============================================================================

class DetectDuplicatesNeo4jRequest(BaseModel):
    """Request para detectar códigos duplicados en Neo4j."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto")
    threshold: float = Field(0.80, ge=0.5, le=1.0, description="Umbral de similitud")


@app.post("/api/codes/detect-duplicates-neo4j")
async def api_detect_duplicates_neo4j(
    payload: DetectDuplicatesNeo4jRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Detecta códigos duplicados en Neo4j usando similitud de Levenshtein (Post-Hoc).
    
    Busca duplicados en los nodos :Codigo del grafo (códigos ya promocionados).
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Get all code names from Neo4j
        cypher = """
        MATCH (c:Codigo)
        WHERE c.project_id = $project_id
        RETURN c.nombre AS nombre
        ORDER BY nombre
        """
        with clients.neo4j.session() as session:
            result = session.run(cypher, project_id=project_id)
            codes = [record["nombre"] for record in result if record["nombre"]]
        
        if len(codes) < 2:
            return {
                "success": True,
                "project": project_id,
                "threshold": payload.threshold,
                "duplicates": [],
                "count": 0,
                "source": "neo4j",
                "total_codes": len(codes),
            }
        
        # Find duplicates using RapidFuzz Levenshtein
        duplicates = []
        unique_codes = list(set(codes))
        
        for i, c1 in enumerate(unique_codes):
            for c2 in unique_codes[i+1:]:
                c1_lower = c1.lower()
                c2_lower = c2.lower()
                max_len = max(len(c1_lower), len(c2_lower))
                
                if max_len == 0:
                    continue
                    
                # Pre-filter by length difference
                if abs(len(c1_lower) - len(c2_lower)) > int((1 - payload.threshold) * max_len):
                    continue
                
                distance = _levenshtein_distance(c1_lower, c2_lower)
                similarity = 1 - (distance / max_len)
                
                if similarity >= payload.threshold:
                    duplicates.append({
                        "code1": c1,
                        "code2": c2,
                        "distance": distance,
                        "similarity": round(similarity, 3),
                        "source": "neo4j",
                    })
                    
                    if len(duplicates) >= 50:
                        break
            
            if len(duplicates) >= 50:
                break
        
        # Sort by similarity
        duplicates.sort(key=lambda x: x["similarity"], reverse=True)
        
        return {
            "success": True,
            "project": project_id,
            "threshold": payload.threshold,
            "duplicates": duplicates,
            "count": len(duplicates),
            "source": "neo4j",
            "total_codes": len(unique_codes),
        }
    except HTTPException:
        raise
    except Exception as exc:
        api_logger.error("api.detect_duplicates_neo4j.error", error=str(exc), project=payload.project)
        raise HTTPException(status_code=500, detail=f"Error detectando duplicados en Neo4j: {str(exc)}") from exc
    finally:
        clients.close()


# =============================================================================
# QDRANT GROUPED SEARCH - Evitar sesgo de fuente
# =============================================================================

class GroupedSearchRequest(BaseModel):
    """Request para búsqueda semántica agrupada con filtros avanzados."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto")
    query: str = Field(..., min_length=3, description="Texto de búsqueda")
    limit: int = Field(10, ge=1, le=50, description="Máximo de grupos")
    group_by: str = Field("archivo", description="Campo para agrupar")
    group_size: int = Field(2, ge=1, le=5, description="Resultados por grupo")
    score_threshold: float = Field(0.3, ge=0.0, le=1.0, description="Umbral de score")
    
    # Filtros avanzados (Payload Filtering - Comparación Constante)
    genero: Optional[str] = Field(None, description="Filtrar por género (mujer/hombre)")
    actor_principal: Optional[str] = Field(None, description="Filtrar por rol/actor")
    area_tematica: Optional[str] = Field(None, description="Filtrar por área temática")
    periodo: Optional[str] = Field(None, description="Filtrar por periodo temporal")
    archivo: Optional[str] = Field(None, description="Filtrar por archivo específico")


@app.post("/api/qdrant/search-grouped")
async def api_qdrant_search_grouped(
    payload: GroupedSearchRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Búsqueda semántica agrupada para evitar sesgo de fuente.
    
    Garantiza diversidad en resultados: máximo N fragmentos por entrevista/speaker.
    Útil para muestreo teórico sin que una fuente domine.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Generar embedding del query
        from app.embeddings import embed_batch
        embeddings = embed_batch(clients.aoai, settings.azure.deployment_embed, [payload.query])
        if not embeddings or not embeddings[0]:
            raise HTTPException(status_code=500, detail="Error generando embedding")
        
        query_vector = embeddings[0]
        
        # Usar búsqueda agrupada con filtros
        from app.qdrant_block import search_similar_grouped
        groups = search_similar_grouped(
            clients.qdrant,
            settings.qdrant.collection,
            query_vector,
            limit=payload.limit,
            group_by=payload.group_by,
            group_size=payload.group_size,
            score_threshold=payload.score_threshold,
            project_id=project_id,
            exclude_interviewer=True,
            # Filtros avanzados
            genero=payload.genero or "",
            actor_principal=payload.actor_principal or "",
            area_tematica=payload.area_tematica or "",
            periodo=payload.periodo or "",
            archivo_filter=payload.archivo or "",
        )
        
        # Formatear resultados
        results = []
        for group in groups:
            group_key = group.id if hasattr(group, 'id') else str(group)
            hits = []
            for hit in (group.hits if hasattr(group, 'hits') else []):
                hits.append({
                    "id": hit.id,
                    "score": hit.score,
                    "fragmento": hit.payload.get("fragmento", "")[:200],
                    "archivo": hit.payload.get("archivo"),
                    "speaker": hit.payload.get("speaker"),
                    "actor_principal": hit.payload.get("actor_principal"),
                })
            results.append({
                "group_key": group_key,
                "hits": hits,
            })
        
        return {
            "success": True,
            "query": payload.query,
            "group_by": payload.group_by,
            "results": results,
            "total_groups": len(results),
        }
    except Exception as e:
        api_logger.error("qdrant.search_grouped.error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e)) from e
    finally:
        clients.close()


class HybridSearchRequest(BaseModel):
    """Request para búsqueda híbrida (semántica + keyword)."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto")
    query: str = Field(..., min_length=2, description="Texto de búsqueda")
    limit: int = Field(10, ge=1, le=50, description="Máximo de resultados")
    score_threshold: float = Field(0.3, ge=0.0, le=1.0, description="Umbral de score")
    keyword_boost: float = Field(0.3, ge=0.0, le=1.0, description="Boost para matches exactos")


@app.post("/api/qdrant/search-hybrid")
async def api_qdrant_search_hybrid(
    payload: HybridSearchRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Búsqueda híbrida: semántica + texto exacto.
    
    Útil para capturar acrónimos (CESFAM, PAC) que embeddings pierden.
    - Dense: encuentra conceptos relacionados
    - Keyword: garantiza palabras exactas
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Generar embedding
        from app.embeddings import embed_batch
        embeddings = embed_batch(clients.aoai, settings.azure.deployment_embed, [payload.query])
        if not embeddings or not embeddings[0]:
            raise HTTPException(status_code=500, detail="Error generando embedding")
        
        query_vector = embeddings[0]
        
        # Búsqueda híbrida
        from app.qdrant_block import search_hybrid
        results = search_hybrid(
            clients.qdrant,
            settings.qdrant.collection,
            payload.query,
            query_vector,
            limit=payload.limit,
            score_threshold=payload.score_threshold,
            project_id=project_id,
            keyword_boost=payload.keyword_boost,
        )
        
        # Formatear resultados
        formatted = [
            {
                "id": str(r.id),
                "score": r.score,
                "fragmento": r.payload.get("fragmento", "")[:300] if r.payload else "",
                "archivo": r.payload.get("archivo") if r.payload else None,
                "keyword_match": r.score > payload.score_threshold + payload.keyword_boost * 0.5,
            }
            for r in results
        ]
        
        return {
            "success": True,
            "query": payload.query,
            "results": formatted,
            "count": len(formatted),
        }
    except Exception as e:
        api_logger.error("qdrant.search_hybrid.error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e)) from e
    finally:
        clients.close()


@app.post("/api/codes/candidates/promote")
async def api_promote_candidates(
    payload: PromoteCandidatesRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Promueve códigos candidatos validados a la lista definitiva.
    
    Si SYNC_NEO4J_ON_PROMOTE=true (default), sincroniza las relaciones
    TIENE_CODIGO en Neo4j usando batch UNWIND.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    candidate_ids = payload.candidate_ids or []
    if not candidate_ids and not payload.promote_all_validated:
        raise HTTPException(
            status_code=400,
            detail="Debe enviar candidate_ids o establecer promote_all_validated=true",
        )
    
    clients = build_clients_or_error(settings)
    neo4j_result = {"neo4j_merged": 0, "neo4j_missing_fragments": 0}
    
    try:
        # 1. Promover en PostgreSQL
        result = promote_to_definitive(
            clients.postgres,
            project_id=project_id,
            candidate_ids=candidate_ids,
            promote_all_validated=payload.promote_all_validated,
            promoted_by=user.user_id if user else None,
        )
        
        # 2. Sincronizar Neo4j si está habilitado y hay datos para sync
        neo4j_sync_rows = result.pop("neo4j_sync_rows", [])
        
        if settings.sync_neo4j_on_promote and neo4j_sync_rows:
            try:
                from app.neo4j_block import merge_fragment_codes_bulk
                
                sync_result = merge_fragment_codes_bulk(
                    driver=clients.neo4j,
                    database=settings.neo4j.database,
                    rows=neo4j_sync_rows,
                    project_id=project_id,
                )
                neo4j_result = {
                    "neo4j_merged": sync_result.get("merged_count", 0),
                    "neo4j_missing_fragments": sync_result.get("missing_fragments", 0),
                }
                
                # Log para observabilidad
                if sync_result.get("missing_fragments", 0) > 0:
                    api_logger.warning(
                        "promote.neo4j_sync.missing_fragments",
                        project_id=project_id,
                        missing_count=sync_result["missing_fragments"],
                        missing_ids=sync_result.get("missing_fragment_ids", [])[:10],  # Limitar log
                    )
                
                api_logger.info(
                    "promote.neo4j_sync.success",
                    project_id=project_id,
                    merged=neo4j_result["neo4j_merged"],
                    missing=neo4j_result["neo4j_missing_fragments"],
                )
            except Exception as e:
                # Neo4j sync es best-effort, no falla la promoción
                api_logger.error(
                    "promote.neo4j_sync.error",
                    project_id=project_id,
                    error=str(e),
                )
                neo4j_result = {"neo4j_merged": 0, "neo4j_missing_fragments": 0, "neo4j_error": str(e)[:200]}
    finally:
        clients.close()

    return {
        "success": (result.get("promoted_count", 0) or 0) > 0,
        **result,
        **neo4j_result,
    }


# =============================================================================
# LINK PREDICTION - Guardar sugerencias en Bandeja de Candidatos
# =============================================================================

class LinkPredictionSaveRequest(BaseModel):
    """Request para guardar sugerencias de Link Prediction."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto")
    suggestions: List[Dict[str, Any]] = Field(
        ..., 
        description="Lista de sugerencias: [{source, target, score, algorithm, reason}]"
    )


@app.post("/api/link-prediction/save")
async def api_save_link_predictions(
    payload: LinkPredictionSaveRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Guarda sugerencias de Link Prediction en la tabla link_predictions.
    
    Las predicciones son hipótesis de relaciones axiales entre códigos.
    Se validan en el panel de Codificación Axial.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    if not payload.suggestions:
        raise HTTPException(status_code=400, detail="suggestions no puede estar vacío")
    
    from app.postgres_block import insert_link_predictions
    
    # Convertir sugerencias a formato de link_predictions
    predictions = []
    for i, suggestion in enumerate(payload.suggestions):
        source = suggestion.get("source", "")
        target = suggestion.get("target", "")
        raw_score = suggestion.get("score", 0.0)
        algorithm = suggestion.get("algorithm", "common_neighbors")
        reason = suggestion.get("reason", "")
        
        if not source or not target:
            continue

        try:
            score = float(raw_score)
        except Exception:
            score = 0.0
        score = max(score, 0.0)
            
        predictions.append({
            "project_id": project_id,
            "source_code": source,
            "target_code": target,
            "relation_type": "asociado_con",
            "algorithm": algorithm,
            "score": score,
            "rank": i + 1,
            "memo": reason if reason else f"Sugerido por {algorithm}",
        })
    
    if not predictions:
        raise HTTPException(status_code=400, detail="No se encontraron sugerencias válidas")
    
    clients = build_clients_or_error(settings)
    try:
        count = insert_link_predictions(clients.postgres, predictions)
    finally:
        clients.close()
    
    api_logger.info(
        "link_prediction.saved",
        project=project_id,
        count=count,
        user=user.user_id,
    )
    
    return {"success": True, "saved_count": count}


@app.get("/api/link-predictions")
async def api_get_link_predictions(
    project: str = Query(..., description="Proyecto"),
    estado: Optional[str] = Query(None, description="Filtrar por estado"),
    algorithm: Optional[str] = Query(None, description="Filtrar por algoritmo"),
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene predicciones de relaciones para el panel de Codificación Axial.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    from app.postgres_block import (
        count_link_predictions,
        get_link_predictions,
        get_link_predictions_stats,
    )
    
    clients = build_clients_or_error(settings)
    try:
        predictions = get_link_predictions(
            clients.postgres,
            project_id,
            estado=estado,
            algorithm=algorithm,
            limit=limit,
            offset=offset,
        )
        total = count_link_predictions(
            clients.postgres,
            project_id,
            estado=estado,
            algorithm=algorithm,
        )
        stats = get_link_predictions_stats(clients.postgres, project_id)
    finally:
        clients.close()
    
    return {
        "items": predictions,
        "total": total,
        "limit": limit,
        "offset": offset,
        "stats": stats,
        "project": project_id,
    }


class LinkPredictionUpdateRequest(BaseModel):
    """Request para actualizar estado de predicción."""
    model_config = ConfigDict(extra="forbid")
    
    estado: str = Field(..., description="Nuevo estado: pendiente, validado, rechazado")
    memo: Optional[str] = Field(None, description="Nota opcional")
    relation_type: Optional[str] = Field(None, description="Tipo de relación para crear en Neo4j")


@app.patch("/api/link-predictions/{prediction_id}")
async def api_update_link_prediction(
    prediction_id: int,
    payload: LinkPredictionUpdateRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Actualiza el estado de una predicción.
    
    Si estado='validado', crea la relación axial en Neo4j.
    """
    if payload.estado not in ("pendiente", "validado", "rechazado"):
        raise HTTPException(status_code=400, detail="estado debe ser: pendiente, validado, rechazado")
    
    from app.postgres_block import update_link_prediction_estado, get_link_prediction_by_id
    from app.neo4j_block import merge_axial_relationship
    
    clients = build_clients_or_error(settings)
    try:
        # Obtener predicción actual
        prediction = get_link_prediction_by_id(clients.postgres, prediction_id)
        if not prediction:
            raise HTTPException(status_code=404, detail="Predicción no encontrada")
        
        # Actualizar estado
        success = update_link_prediction_estado(
            clients.postgres,
            prediction_id,
            payload.estado,
            validado_por=user.user_id,
            memo=payload.memo,
            relation_type=payload.relation_type,
        )
        
        neo4j_synced = False
        # Si se valida, crear relación en Neo4j
        if success and payload.estado == "validado" and clients.neo4j:
            rel_type = payload.relation_type or prediction.get("relation_type", "asociado_con")
            try:
                merge_axial_relationship(
                    driver=clients.neo4j,
                    database=settings.neo4j.database,
                    project_id=prediction["project_id"],
                    source_code=prediction["source_code"],
                    target_code=prediction["target_code"],
                    relation_type=rel_type,
                )
                neo4j_synced = True
            except Exception as e:
                api_logger.warning(
                    "link_prediction.neo4j_sync_failed",
                    prediction_id=prediction_id,
                    error=str(e),
                )
    finally:
        clients.close()
    
    return {
        "success": success,
        "prediction_id": prediction_id,
        "estado": payload.estado,
        "neo4j_synced": neo4j_synced,
    }


class LinkPredictionBatchUpdateRequest(BaseModel):
    """Request para actualizar múltiples predicciones."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto")
    prediction_ids: List[int] = Field(..., description="IDs de predicciones")
    estado: str = Field(..., description="Nuevo estado")


@app.post("/api/link-predictions/batch-update")
async def api_batch_update_link_predictions(
    payload: LinkPredictionBatchUpdateRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Actualiza el estado de múltiples predicciones.
    """
    if payload.estado not in ("pendiente", "validado", "rechazado"):
        raise HTTPException(status_code=400, detail="estado debe ser: pendiente, validado, rechazado")
    
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    from app.postgres_block import batch_update_link_predictions_estado, get_link_predictions_by_ids
    from app.neo4j_block import merge_axial_relationships_bulk
    
    clients = build_clients_or_error(settings)
    try:
        count = batch_update_link_predictions_estado(
            clients.postgres,
            project_id,
            payload.prediction_ids,
            payload.estado,
            validado_por=user.user_id,
        )
        neo4j_synced = 0
        if payload.estado == "validado" and clients.neo4j:
            predictions = get_link_predictions_by_ids(clients.postgres, project_id, payload.prediction_ids)
            neo4j_synced = merge_axial_relationships_bulk(
                driver=clients.neo4j,
                database=settings.neo4j.database,
                project_id=project_id,
                rows=predictions,
            )
    finally:
        clients.close()
    
    return {
        "success": True,
        "updated_count": count,
        "estado": payload.estado,
        "neo4j_synced": neo4j_synced,
    }


# =============================================================================
# SYNC - Sincronización PostgreSQL → Neo4j (códigos abiertos)
# =============================================================================

class SyncNeo4jRequest(BaseModel):
    """Request para sincronizar códigos PostgreSQL → Neo4j."""
    model_config = ConfigDict(extra="forbid")
    project: str = Field(..., description="Proyecto a sincronizar")


@app.post("/api/sync/neo4j")
async def api_sync_postgres_to_neo4j(
    payload: SyncNeo4jRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Sincroniza códigos abiertos desde PostgreSQL a Neo4j.
    
    Crea nodos :Codigo y relaciones :TIENE_CODIGO para todos los códigos
    que existen en PostgreSQL pero no en Neo4j.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    log = api_logger.bind(endpoint="sync.neo4j", project=project_id)
    clients = build_clients_or_error(settings)
    try:
        # 1. Obtener códigos únicos de PostgreSQL con fragmento_id válido
        with clients.postgres.cursor() as cur:
            cur.execute("""
                SELECT DISTINCT codigo, fragmento_id, archivo 
                FROM analisis_codigos_abiertos 
                WHERE project_id = %s AND fragmento_id IS NOT NULL
            """, (project_id,))
            pg_codes = cur.fetchall()

        if not clients.neo4j:
            return {"success": False, "error": "Neo4j no está configurado"}

        db = settings.neo4j.database

        # 2. Obtener códigos existentes en Neo4j
        with clients.neo4j.session(database=db) as session:
            result = session.run(
                "MATCH (c:Codigo {project_id: $pid}) RETURN c.nombre AS name",
                pid=project_id
            )
            neo4j_codes = {r["name"] for r in result}

        neo4j_codes_before = len(neo4j_codes)
        synced_codes = 0
        synced_rels = 0
        missing_fragments = 0

        # 3. Sincronizar códigos y relaciones
        with clients.neo4j.session(database=db) as session:
            for row in pg_codes:
                codigo, frag_id, archivo = row
                
                # Crear código si no existe
                if codigo not in neo4j_codes:
                    session.run(
                        """
                        MERGE (c:Codigo {nombre: $codigo, project_id: $project_id})
                        SET c.status = 'active', c.synced_at = datetime()
                        """,
                        codigo=codigo, project_id=project_id
                    )
                    neo4j_codes.add(codigo)
                    synced_codes += 1

                # Crear relación TIENE_CODIGO
                result = session.run(
                    """
                    MATCH (f:Fragmento {id: $frag_id, project_id: $project_id})
                    MATCH (c:Codigo {nombre: $codigo, project_id: $project_id})
                    MERGE (f)-[rel:TIENE_CODIGO]->(c)
                    SET rel.archivo = $archivo, 
                        rel.project_id = $project_id, 
                        rel.synced_at = datetime()
                    RETURN count(rel) as cnt
                    """,
                    frag_id=frag_id, codigo=codigo, archivo=archivo, project_id=project_id
                )
                cnt = result.single()
                if cnt and cnt["cnt"] > 0:
                    synced_rels += 1
                else:
                    missing_fragments += 1

        log.info(
            "sync.neo4j.success",
            project_id=project_id,
            synced_codes=synced_codes,
            synced_rels=synced_rels,
            missing_fragments=missing_fragments,
        )

        return {
            "success": True,
            "pg_codes_total": len(pg_codes),
            "neo4j_codes_before": neo4j_codes_before,
            "synced_codes": synced_codes,
            "synced_relations": synced_rels,
            "missing_fragments": missing_fragments,
        }
    except Exception as exc:
        log.error("sync.neo4j.error", error=str(exc), project=payload.project)
        raise HTTPException(status_code=500, detail=f"Error en sincronización: {exc}") from exc
    finally:
        clients.close()


# =============================================================================
# ANALYSIS REPORTS - Guardar análisis IA en la base de datos
# =============================================================================

class SaveAnalysisReportRequest(BaseModel):
    """Request para guardar un informe de análisis IA."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto")
    report_type: str = Field(..., description="Tipo: link_prediction, discovery, axial, etc.")
    title: str = Field(..., description="Título del informe")
    content: str = Field(..., description="Contenido del análisis")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Metadatos adicionales")


def ensure_analysis_reports_table(pg_conn) -> None:
    """Crea tabla analysis_reports si no existe."""
    with pg_conn.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS analysis_reports (
                id SERIAL PRIMARY KEY,
                project_id VARCHAR(100) NOT NULL,
                report_type VARCHAR(50) NOT NULL,
                title VARCHAR(500) NOT NULL,
                content TEXT NOT NULL,
                metadata JSONB,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_analysis_reports_project 
            ON analysis_reports(project_id)
        """)
    pg_conn.commit()


@app.post("/api/analysis/save-report")
async def api_save_analysis_report(
    payload: SaveAnalysisReportRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Guarda un informe de análisis IA en la base de datos.
    
    Soporta múltiples tipos de análisis: link_prediction, discovery, etc.
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        ensure_analysis_reports_table(clients.postgres)
        
        with clients.postgres.cursor() as cur:
            cur.execute("""
                INSERT INTO analysis_reports (project_id, report_type, title, content, metadata)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING id
            """, (
                project_id,
                payload.report_type,
                payload.title,
                payload.content,
                json.dumps(payload.metadata, default=str),
            ))
            result = cur.fetchone()
            report_id = result[0] if result else 0
        
        clients.postgres.commit()
    finally:
        clients.close()
    
    api_logger.info(
        "analysis_report.saved",
        project=project_id,
        report_type=payload.report_type,
        report_id=report_id,
        user=user.user_id,
    )
    
    return {"success": True, "report_id": report_id}


@app.get("/api/analysis/reports")
async def api_list_analysis_reports(
    project: str = Query(..., description="Proyecto requerido"),
    report_type: Optional[str] = Query(None, description="Filtrar por tipo"),
    limit: int = Query(50, ge=1, le=200),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista informes de análisis IA guardados."""
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        ensure_analysis_reports_table(clients.postgres)
        
        with clients.postgres.cursor() as cur:
            if report_type:
                cur.execute("""
                    SELECT id, report_type, title, content, metadata, created_at
                    FROM analysis_reports
                    WHERE project_id = %s AND report_type = %s
                    ORDER BY created_at DESC
                    LIMIT %s
                """, (project_id, report_type, limit))
            else:
                cur.execute("""
                    SELECT id, report_type, title, content, metadata, created_at
                    FROM analysis_reports
                    WHERE project_id = %s
                    ORDER BY created_at DESC
                    LIMIT %s
                """, (project_id, limit))
            
            rows = cur.fetchall()
    finally:
        clients.close()
    
    reports = [
        {
            "id": row[0],
            "report_type": row[1],
            "title": row[2],
            "content": row[3],
            "metadata": row[4],
            "created_at": row[5].isoformat() if row[5] else None,
        }
        for row in rows
    ]
    
    return {"reports": reports, "count": len(reports)}


# =============================================================================
# INTERVIEW REPORTS - Informes por Entrevista
# =============================================================================

@app.get("/api/reports/interview")
async def api_list_interview_reports(
    project: str = Query(..., description="Proyecto requerido"),
    limit: int = Query(50, ge=1, le=200),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Lista informes de análisis por entrevista.
    
    Retorna los informes generados por el análisis LLM de cada entrevista.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        from app.reports import get_interview_reports
        reports = get_interview_reports(clients.postgres, project_id, limit)
        
        return {
            "reports": [r.to_dict() for r in reports],
            "count": len(reports),
        }
    except Exception as e:
        api_logger.error("interview_reports.list_error", error=str(e))
        raise HTTPException(status_code=500, detail=str(e)) from e
    finally:
        clients.close()


@app.get("/api/reports/interview/{archivo}")
async def api_get_interview_report(
    archivo: str,
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene el informe de análisis para una entrevista específica.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        from app.reports import get_interview_report
        report = get_interview_report(clients.postgres, project_id, archivo)
        
        if not report:
            raise HTTPException(
                status_code=404, 
                detail=f"No se encontró informe para: {archivo}"
            )
        
        return {"report": report.to_dict()}
    except HTTPException:
        raise
    except Exception as e:
        api_logger.error("interview_report.get_error", error=str(e), archivo=archivo)
        raise HTTPException(status_code=500, detail=str(e)) from e
    finally:
        clients.close()


@app.get("/api/codes/definitive")
async def api_definitive_codes(
    project: str = Query(..., description="Proyecto requerido"),
    limit: int = Query(100, ge=1, le=500),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista la 'lista definitiva' real (tabla analisis_codigos_abiertos).

    Nota: la bandeja de candidatos validados vive en `codigos_candidatos` y se expone por
    `/api/codes/candidates` con `estado=validado`.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        from app.coding import list_open_codes
        codes = list_open_codes(clients, project=project_id, limit=limit)
    finally:
        clients.close()

    return {"definitive_codes": codes, "count": len(codes)}


@app.get("/api/codes/catalog")
async def api_codes_catalog(
    project: str = Query(..., description="Proyecto requerido"),
    limit: int = Query(500, ge=1, le=2000),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista el catálogo ontológico de códigos (status + canonical)."""
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import ensure_codes_catalog_table

        ensure_codes_catalog_table(clients.postgres)
        with clients.postgres.cursor() as cur:
            cur.execute(
                """
                SELECT codigo, status, canonical_codigo, merged_at, merged_by, memo, created_at, updated_at
                  FROM catalogo_codigos
                 WHERE project_id = %s
                 ORDER BY updated_at DESC NULLS LAST, codigo ASC
                 LIMIT %s
                """,
                (project_id, limit),
            )
            rows = cur.fetchall() or []
        clients.postgres.commit()
    finally:
        clients.close()

    catalog = [
        {
            "codigo": r[0],
            "status": r[1],
            "canonical_codigo": r[2],
            "merged_at": r[3].isoformat().replace("+00:00", "Z") if r[3] else None,
            "merged_by": r[4],
            "memo": r[5],
            "created_at": r[6].isoformat().replace("+00:00", "Z") if r[6] else None,
            "updated_at": r[7].isoformat().replace("+00:00", "Z") if r[7] else None,
        }
        for r in rows
    ]
    return {"catalog": catalog, "count": len(catalog)}


class MergeDefinitiveCodesBody(BaseModel):
    source_codigo: str
    target_codigo: str
    memo: Optional[str] = None


@app.post("/api/codes/definitive/merge")
async def api_merge_definitive_codes(
    body: MergeDefinitiveCodesBody,
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Fusiona códigos definitivos (mueve evidencia y marca status=merged)."""
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import merge_definitive_codes_by_code
        from app.neo4j_block import mark_codigo_merged

        result = merge_definitive_codes_by_code(
            clients.postgres,
            project_id=project_id,
            source_codigo=body.source_codigo,
            target_codigo=body.target_codigo,
            merged_by=getattr(user, "email", None) or getattr(user, "username", None) or "user",
            memo=body.memo,
        )

        # Best-effort: also mark merged in Neo4j so GDS projections can filter.
        try:
            mark_codigo_merged(
                clients.neo4j,
                settings.neo4j.database,
                project_id=project_id,
                source_codigo=body.source_codigo,
                target_codigo=result.get("target") or body.target_codigo,
                merged_by=getattr(user, "email", None) or getattr(user, "username", None) or "user",
                memo=body.memo,
            )
        except Exception:
            pass
        return {"result": result}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as e:
        api_logger.error(
            "codes.definitive.merge_error",
            error=str(e),
            source=body.source_codigo,
            target=body.target_codigo,
        )
        raise HTTPException(status_code=500, detail=str(e)) from e
    finally:
        clients.close()


@app.get("/api/codes/stats/sources")
async def api_candidate_stats_by_source(
    project: str = Query(..., description="Proyecto requerido"),
    clients: PgOnlyClients = Depends(get_pg_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Obtiene estadísticas de códigos candidatos por origen."""
    try:
        # Run resolve_project in a thread to avoid blocking the event loop with sync DB I/O
        project_id = await asyncio.to_thread(resolve_project, project, allow_create=False, pg=clients.postgres)
        return await run_pg_query_with_timeout(get_candidate_stats_by_source, clients.postgres, project_id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc




@app.get("/api/codes/candidates/{candidate_id}/examples")
async def api_candidate_examples(
    candidate_id: int,
    project: str = Query(..., description="Proyecto requerido"),
    limit: int = Query(3, ge=1, le=10, description="Cantidad de ejemplos canónicos"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene ejemplos canónicos (citas validadas previas) del mismo código.
    
    Usado para comparación constante al validar nuevos candidatos.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Buscar el candidato para obtener su código
        all_candidates = list_candidate_codes(
            clients.postgres,
            project=project_id,
            limit=500,
        )
        candidate = next((c for c in all_candidates if c["id"] == candidate_id), None)
        
        if not candidate:
            raise HTTPException(status_code=404, detail="Candidato no encontrado")
        
        codigo = candidate["codigo"]
        
        # Obtener ejemplos canónicos
        examples = get_canonical_examples(
            clients.postgres,
            codigo=codigo,
            project=project_id,
            limit=limit,
        )
    finally:
        clients.close()
    
    return {
        "candidate_id": candidate_id,
        "codigo": codigo,
        "examples": examples,
        "examples_count": len(examples),
    }


@app.get("/api/codes/candidates/health")
async def api_candidates_health(
    project: str = Query(..., description="Proyecto requerido"),
    threshold_days: int = Query(3, ge=1, le=30, description="Días máximos sin resolver"),
    threshold_count: int = Query(50, ge=10, le=500, description="Máximos pendientes antes de alerta"),
    clients: PgOnlyClients = Depends(get_pg_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Verifica la salud del backlog de candidatos pendientes."""
    try:
        # Run resolve_project in a thread to avoid blocking the event loop with sync DB I/O
        project_id = await asyncio.to_thread(resolve_project, project, allow_create=False, pg=clients.postgres)
        return await run_pg_query_with_timeout(
            get_backlog_health,
            clients.postgres,
            project_id,
            threshold_days,
            threshold_count,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/codes/similar")
async def api_similar_codes(
    codigo: str = Query(..., description="Código a buscar similares"),
    project: str = Query(..., description="Proyecto requerido"),
    top_k: int = Query(5, ge=1, le=20, description="Máximo de códigos similares"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Encuentra códigos semánticamente similares para sugerir sinónimos o fusiones.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        similar = find_similar_codes(
            clients,
            settings,
            codigo=codigo,
            project=project_id,
            top_k=top_k,
        )
    finally:
        clients.close()
    
    return {
        "codigo": codigo,
        "similar_codes": similar,
        "count": len(similar),
    }


# Sprint 23: Pre-Hoc Check for Batch Codes
class CheckBatchCodesRequest(BaseModel):
    project: str = Field(..., description="Project ID")
    codigos: List[str] = Field(..., description="List of code names to check")
    threshold: float = Field(default=0.85, ge=0.5, le=1.0, description="Similarity threshold")


class AiPlanMergeRequest(BaseModel):
    """Request para generar propuestas de auto-merge con auditoría por run_id."""

    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="ID del proyecto")
    codigos: List[str] = Field(..., description="Lista de códigos (candidatos) a analizar")
    threshold: float = Field(default=0.92, ge=0.5, le=1.0, description="Umbral mínimo de similitud")
    limit: int = Field(default=200, ge=1, le=2000, description="Máximo de pares sugeridos")
    source: str = Field(default="ui", description="Origen del runner (ui|cli|script)")



@app.post("/api/codes/check-batch")
async def api_check_batch_codes(
    payload: CheckBatchCodesRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Sprint 23: Pre-hoc check for batch of codes.
    
    Checks each proposed code against existing codes and returns
    similarity matches for deduplication UI.
    """
    import os
    import time
    import statistics
    from collections import defaultdict

    from app.code_normalization import (
        find_similar_codes_with_stats,
        get_existing_codes_for_project,
        normalize_code,
    )
    
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    started = time.perf_counter()
    clients = build_clients_or_error(settings)
    try:
        api_logger.info(
            "api.codes.check_batch.start",
            project=project_id,
            threshold=float(payload.threshold),
            input_count=len(payload.codigos or []),
        )

        # Get existing codes from DB
        t0 = time.perf_counter()
        existing_codes = get_existing_codes_for_project(clients.postgres, project_id)
        t_fetch_ms = (time.perf_counter() - t0) * 1000.0

        # ---------------------------------------------------------------------
        # Pre-Hoc intra-batch dedup (evita "batch blindness")
        # Agrupamos por clave normalizada (tildes/underscores/espacios)
        # ---------------------------------------------------------------------
        t1 = time.perf_counter()

        groups: Dict[str, List[int]] = defaultdict(list)
        empty_indexes: List[int] = []
        for idx, codigo in enumerate(payload.codigos):
            if not (codigo or "").strip():
                empty_indexes.append(idx)
                continue
            norm_key = normalize_code(codigo)
            groups[norm_key].append(idx)

        empty_norm_key_count = len(groups.get("", [])) if "" in groups else 0

        duplicate_groups = {k: idxs for k, idxs in groups.items() if len(idxs) > 1 and k}
        batch_duplicates_total = sum(max(0, len(idxs) - 1) for idxs in duplicate_groups.values())

        # Pre-computar similares una vez por grupo (eficiencia)
        # Métricas agregadas del motor de similitud (para diagnóstico)
        scan_stats = {
            "groups_scanned": 0,
            "comparisons": 0,
            "skipped_len_prefilter": 0,
            "skipped_similarity": 0,
            "skipped_token_overlap": 0,
            "kept": 0,
            "slow_scans": 0,
            "scan_elapsed_ms_sum": 0.0,
        }

        similar_by_norm: Dict[str, List[Tuple[str, float]]] = {}
        for norm_key, idxs in groups.items():
            if not norm_key:
                similar_by_norm[norm_key] = []
                continue
            representative = payload.codigos[idxs[0]]

            sims, st = find_similar_codes_with_stats(
                representative,
                existing_codes,
                threshold=payload.threshold,
                limit=3,
            )
            similar_by_norm[norm_key] = sims

            scan_stats["groups_scanned"] += 1
            scan_stats["comparisons"] += int(st.get("comparisons", 0))
            scan_stats["skipped_len_prefilter"] += int(st.get("skipped_len_prefilter", 0))
            scan_stats["skipped_similarity"] += int(st.get("skipped_similarity", 0))
            scan_stats["skipped_token_overlap"] += int(st.get("skipped_token_overlap", 0))
            scan_stats["kept"] += int(st.get("kept", 0))
            scan_stats["scan_elapsed_ms_sum"] += float(st.get("elapsed_ms", 0.0) or 0.0)
            if float(st.get("elapsed_ms", 0.0) or 0.0) >= 200.0:
                scan_stats["slow_scans"] += 1

        t_scan_ms = (time.perf_counter() - t1) * 1000.0

        # Construir results en el mismo orden de entrada
        t2 = time.perf_counter()
        best_sims: List[float] = []
        results: List[Dict[str, Any]] = []
        for idx, codigo in enumerate(payload.codigos):
            if idx in empty_indexes:
                results.append(
                    {
                        "codigo": codigo,
                        "has_similar": False,
                        "similar": [],
                        "duplicate_in_batch": False,
                        "batch_group_size": 0,
                    }
                )
                continue

            norm_key = normalize_code(codigo)
            group_size = len(groups.get(norm_key, []))
            similar = similar_by_norm.get(norm_key, [])

            if similar:
                try:
                    best_sims.append(float(similar[0][1]))
                except Exception:
                    pass

            results.append(
                {
                    "codigo": codigo,
                    "has_similar": bool(similar),
                    "similar": [
                        {"existing": s[0], "similarity": round(s[1], 2)}
                        for s in similar
                    ],
                    "duplicate_in_batch": group_size > 1,
                    "batch_group_size": group_size,
                }
            )

        t_results_ms = (time.perf_counter() - t2) * 1000.0

        has_any_similar = any(r["has_similar"] for r in results)
        matched_count = sum(1 for r in results if r.get("has_similar"))

        # Distribución (sobre la mejor similitud por código) para calibración
        best_sims_sorted = sorted(best_sims)
        sim_stats: Dict[str, Any] = {}
        if best_sims_sorted:
            sim_stats = {
                "min": best_sims_sorted[0],
                "p50": round(float(statistics.median(best_sims_sorted)), 2),
                "max": best_sims_sorted[-1],
            }
            # percentil 90 aproximado (sin numpy)
            try:
                p90_idx = int(round(0.9 * (len(best_sims_sorted) - 1)))
                sim_stats["p90"] = best_sims_sorted[max(0, min(len(best_sims_sorted) - 1, p90_idx))]
            except Exception:
                pass

        elapsed_ms = (time.perf_counter() - started) * 1000.0

        # Logging de muestras (opcional) para depuración puntual
        should_log_samples = str(os.getenv("PREHOC_LOG_SAMPLES", "0")).strip() == "1"
        sample_matches = None
        if should_log_samples and results:
            sample_matches = []
            for r in results:
                if r.get("has_similar"):
                    sims = r.get("similar") or []
                    if sims:
                        sample_matches.append(
                            {
                                "codigo": str(r.get("codigo"))[:120],
                                "best_existing": str(sims[0].get("existing"))[:120],
                                "best_similarity": sims[0].get("similarity"),
                            }
                        )
                if len(sample_matches) >= 3:
                    break

        api_logger.info(
            "api.codes.check_batch.completed",
            project=project_id,
            threshold=float(payload.threshold),
            input_count=len(payload.codigos or []),
            empty_count=len(empty_indexes),
            existing_count=len(existing_codes),
            batch_unique_count=len(groups),
            empty_norm_key_count=int(empty_norm_key_count),
            batch_duplicate_groups=len(duplicate_groups),
            batch_duplicates_total=batch_duplicates_total,
            matched_count=matched_count,
            has_any_similar=has_any_similar,
            similarity_best_stats=sim_stats or None,
            phase_ms={
                "fetch_existing": round(t_fetch_ms, 2),
                "scan_groups_total": round(t_scan_ms, 2),
                "build_results": round(t_results_ms, 2),
                "total": round(elapsed_ms, 2),
            },
            similarity_engine=scan_stats,
            sample_matches=sample_matches,
        )

        return {
            "project": project_id,
            "threshold": payload.threshold,
            "results": results,
            "has_any_similar": has_any_similar,
            "checked_count": len(payload.codigos),
            "existing_count": len(existing_codes),
            # Campos adicionales (compatibles): métricas de intra-batch
            "batch_unique_count": len(groups),
            "batch_duplicate_groups": len(duplicate_groups),
            "batch_duplicates_total": batch_duplicates_total,
        }
    except Exception as exc:
        api_logger.exception(
            "api.codes.check_batch.failed",
            project=project_id,
            threshold=float(payload.threshold),
            input_count=len(payload.codigos or []),
            error=str(exc),
        )
        raise
    finally:
        clients.close()


@app.post("/api/codes/candidates/ai/plan-merges")
async def api_ai_plan_merges(
    payload: AiPlanMergeRequest,
    request: Request,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Runner IA (solo propuestas): sugiere pares source->target con alta confianza.

    Persistencia:
    - Guarda el plan completo (inputs + pares + meta) en Postgres con `run_id`.
    - No ejecuta merges.
    """
    from collections import defaultdict

    from app.code_normalization import find_similar_codes, get_existing_codes_for_project, normalize_code
    from app.postgres_block import insert_ai_merge_plan

    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    if not payload.codigos:
        raise HTTPException(status_code=400, detail="codigos no puede estar vacío")

    # Normalizar y deduplicar entrada conservando primera ocurrencia
    raw = [str(c or "").strip() for c in payload.codigos]
    raw = [c for c in raw if c]
    if not raw:
        raise HTTPException(status_code=400, detail="codigos no contiene valores válidos")

    seen_lower = set()
    deduped: List[str] = []
    for c in raw:
        k = c.lower()
        if k in seen_lower:
            continue
        seen_lower.add(k)
        deduped.append(c)

    clients = build_clients_or_error(settings)
    try:
        existing_codes = get_existing_codes_for_project(clients.postgres, project_id)

        # Agrupar por clave normalizada para evitar cómputo repetido
        groups: Dict[str, List[int]] = defaultdict(list)
        for idx, codigo in enumerate(deduped):
            norm_key = normalize_code(codigo)
            groups[norm_key].append(idx)

        # Calcular similitudes por grupo (solo el representante)
        similar_by_norm: Dict[str, List[Tuple[str, float]]] = {}
        for norm_key, idxs in groups.items():
            if not norm_key:
                similar_by_norm[norm_key] = []
                continue
            representative = deduped[idxs[0]]
            similar_by_norm[norm_key] = find_similar_codes(
                representative,
                existing_codes,
                threshold=float(payload.threshold),
                limit=3,
            )

        # Construir pares sugeridos (best match) y deduplicar por source
        pairs: List[Dict[str, Any]] = []
        by_source: Dict[str, Dict[str, Any]] = {}
        for codigo in deduped:
            norm_key = normalize_code(codigo)
            sims = similar_by_norm.get(norm_key, [])
            if not sims:
                continue
            best_existing, best_sim = sims[0]
            if not best_existing:
                continue
            if codigo.strip().lower() == str(best_existing).strip().lower():
                continue
            src_lower = codigo.strip().lower()
            candidate = {
                "source_codigo": codigo,
                "target_codigo": best_existing,
                "similarity": round(float(best_sim), 4),
                "reason": "best_existing",
            }
            prev = by_source.get(src_lower)
            if not prev or float(candidate["similarity"]) > float(prev.get("similarity", 0)):
                by_source[src_lower] = candidate

        pairs = list(by_source.values())
        pairs.sort(key=lambda x: float(x.get("similarity", 0)), reverse=True)
        if payload.limit and len(pairs) > int(payload.limit):
            pairs = pairs[: int(payload.limit)]

        run_id = str(uuid.uuid4())
        meta = {
            "request_id": getattr(getattr(request, "state", None), "request_id", None),
            "session_id": getattr(getattr(request, "state", None), "session_id", None),
            "existing_count": len(existing_codes),
            "deduped_input_count": len(deduped),
        }

        insert_ai_merge_plan(
            clients.postgres,
            run_id=run_id,
            project_id=project_id,
            created_by=user.user_id if user else None,
            source=str(payload.source or "ui"),
            threshold=float(payload.threshold),
            input_codigos=deduped,
            pairs=pairs,
            meta=meta,
        )

        api_logger.info(
            "api.ai.plan_merges.created",
            project=project_id,
            run_id=run_id,
            threshold=float(payload.threshold),
            input_count=len(deduped),
            pairs_count=len(pairs),
            user_id=user.user_id if user else None,
        )

        return {
            "project": project_id,
            "run_id": run_id,
            "threshold": float(payload.threshold),
            "input_count": len(deduped),
            "pairs_count": len(pairs),
            "pairs": pairs,
        }
    except Exception as exc:
        api_logger.error("api.ai.plan_merges.error", error=str(exc), project=payload.project)
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


@app.get("/api/codes/candidates/ai/plan-merges/{run_id}")
async def api_ai_get_plan_merges(
    run_id: str,
    project: str = Query(..., description="ID del proyecto"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Obtiene un plan IA por run_id (auditoría)."""
    from app.postgres_block import get_ai_merge_plan

    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        plan = get_ai_merge_plan(clients.postgres, project_id=project_id, run_id=run_id)
        if not plan:
            raise HTTPException(status_code=404, detail="run_id no encontrado")
        # Multi-tenant: require_auth ya limita org; aquí validamos proyecto.
        return plan
    finally:
        clients.close()


@app.get("/api/codes/candidates/ai/plan-merges")
async def api_ai_list_plan_merges(
    project: str = Query(..., description="ID del proyecto"),
    limit: int = Query(20, ge=1, le=200),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista planes recientes (resumen) para auditoría."""
    from app.postgres_block import list_ai_merge_plans

    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        rows = list_ai_merge_plans(clients.postgres, project_id=project_id, limit=limit)
        return {"project": project_id, "plans": rows, "count": len(rows)}
    finally:
        clients.close()

@app.post("/api/maintenance/delete_file")
async def api_maintenance_delete_file(
    payload: MaintenanceDeleteFileRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Elimina todos los datos asociados a un archivo dentro de un proyecto.

    Nota: Esta operación es destructiva. Borra:
    - PostgreSQL: fragmentos, códigos abiertos, relaciones axiales por archivo, reportes por entrevista
    - Qdrant: puntos (fragmentos) filtrados por project_id + archivo
    - Neo4j: (:Entrevista {nombre=archivo, project_id}) y sus (:Fragmento) relacionados del proyecto
    """
    if not payload.file.strip():
        raise HTTPException(status_code=400, detail="file es obligatorio")

    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    deleted: Dict[str, Any] = {"project_id": project_id, "file": payload.file}

    # 1) PostgreSQL deletions
    pg_counts: Dict[str, int] = {}
    try:
        with clients.postgres.cursor() as cur:
            # counts (best-effort)
            for table, where_sql in (
                (
                    "entrevista_fragmentos",
                    "project_id = %s AND archivo = %s",
                ),
                (
                    "analisis_codigos_abiertos",
                    "project_id = %s AND archivo = %s",
                ),
                (
                    "analisis_axial",
                    "project_id = %s AND archivo = %s",
                ),
                (
                    "interview_reports",
                    "project_id = %s AND archivo = %s",
                ),
            ):
                try:
                    cur.execute(f"SELECT COUNT(*) FROM {table} WHERE {where_sql}", (project_id, payload.file))
                    row = cur.fetchone()
                    pg_counts[table] = int(row[0] or 0) if row else 0
                except Exception:
                    # Table might not exist yet in fresh installs; treat as 0
                    pg_counts[table] = 0

            # deletes
            for table, where_sql in (
                ("analisis_codigos_abiertos", "project_id = %s AND archivo = %s"),
                ("analisis_axial", "project_id = %s AND archivo = %s"),
                ("entrevista_fragmentos", "project_id = %s AND archivo = %s"),
                ("interview_reports", "project_id = %s AND archivo = %s"),
            ):
                try:
                    cur.execute(f"DELETE FROM {table} WHERE {where_sql}", (project_id, payload.file))
                except Exception:
                    # Same reasoning: if table doesn't exist, ignore
                    continue
        clients.postgres.commit()
    except Exception as exc:
        clients.postgres.rollback()
        clients.close()
        raise HTTPException(status_code=500, detail=f"Error eliminando en PostgreSQL: {exc}") from exc
    deleted["postgres"] = {"counts": pg_counts}

    # 2) Qdrant deletions
    try:
        qdrant_filter = Filter(
            must=cast(
                List[Any],
                [
                    FieldCondition(key="project_id", match=MatchValue(value=project_id)),
                    FieldCondition(key="archivo", match=MatchValue(value=payload.file)),
                ],
            )
        )
        clients.qdrant.delete(
            collection_name=settings.qdrant.collection,
            points_selector=FilterSelector(filter=qdrant_filter),
        )
        deleted["qdrant"] = {"collection": settings.qdrant.collection, "filter": {"project_id": project_id, "archivo": payload.file}}
    except Exception as exc:
        clients.close()
        raise HTTPException(status_code=500, detail=f"Error eliminando en Qdrant: {exc}") from exc

    # 3) Neo4j deletions
    try:
        database = settings.neo4j.database
        cypher = """
        MATCH (e:Entrevista {nombre: $archivo})
        WHERE e.project_id = $project_id
        OPTIONAL MATCH (e)-[:TIENE_FRAGMENTO]->(f:Fragmento)
        WHERE f.project_id = $project_id
        DETACH DELETE f
        DETACH DELETE e
        """
        with clients.neo4j.session(database=database) as session:
            result = session.run(cypher, archivo=payload.file, project_id=project_id)
            summary = result.consume()
        deleted["neo4j"] = {
            "database": database,
            "nodes_deleted": summary.counters.nodes_deleted,
            "relationships_deleted": summary.counters.relationships_deleted,
        }
    except Exception as exc:
        clients.close()
        raise HTTPException(status_code=500, detail=f"Error eliminando en Neo4j: {exc}") from exc
    finally:
        clients.close()

    api_logger.info("maintenance.delete_file", **deleted)
    return {"status": "ok", "deleted": deleted}


@app.get("/api/coding/codes")
async def api_coding_codes(
    limit: int = 50,
    search: Optional[str] = None,
    archivo: Optional[str] = Query(default=None, description="Filtrar por archivo de entrevista"),
    project: str = Query(..., description="Proyecto requerido"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
        from app.coding import list_open_codes

        return {
            "codes": list_open_codes(
                clients,
                project_id,
                limit=limit,
                search=search,
                archivo=archivo,
            )
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc



@app.get("/api/coding/fragments")
async def api_coding_fragments(
    archivo: str,
    limit: int = 25,
    project: str = Query(..., description="Proyecto requerido"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
        return {"fragments": list_interview_fragments(clients, project_id, archivo=archivo, limit=limit)}
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.get("/api/coding/citations")
async def api_coding_citations(
    codigo: str,
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    try:
        rows = citations_for_code(clients, codigo, project_id)
    finally:
        clients.close()
    return {"citations": rows}


@app.get("/api/coding/saturation")
async def api_coding_saturation(
    project: str = Query(..., description="Proyecto requerido"),
    window: int = Query(default=3, ge=1, le=10, description="Ventana de entrevistas para detectar plateau"),
    threshold: int = Query(default=2, ge=0, le=10, description="Máximo de códigos nuevos para considerar plateau"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene datos de saturación teórica para el indicador visual.
    
    Retorna:
    - curve: Lista de entrevistas con códigos nuevos y acumulados
    - plateau: Si se alcanzó saturación (las últimas N entrevistas tienen ≤threshold códigos nuevos)
    - summary: Estadísticas de resumen
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    try:
        data = get_saturation_data(clients, project_id, window=window, threshold=threshold)
    finally:
        clients.close()
    return data


@app.get("/api/export/refi-qda")
async def api_export_refi_qda(
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> PlainTextResponse:
    """
    Exporta códigos en formato REFI-QDA XML (compatible con Atlas.ti 9+).
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    try:
        xml_content = export_codes_refi_qda(clients, project_id)
    finally:
        clients.close()
    return PlainTextResponse(
        content=xml_content,
        media_type="application/xml",
        headers={"Content-Disposition": f'attachment; filename="{project_id}_codes.qdpx"'}
    )


@app.get("/api/export/maxqda")
async def api_export_maxqda(
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> PlainTextResponse:
    """
    Exporta códigos en formato CSV compatible con MAXQDA.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    try:
        csv_content = export_codes_maxqda_csv(clients, project_id)
    finally:
        clients.close()
    return PlainTextResponse(
        content=csv_content,
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{project_id}_codes.csv"'}
    )


@app.get("/api/coding/codes/{codigo}/history")
async def api_code_history(
    codigo: str,
    project: str = Query(..., description="Proyecto requerido"),
    limit: int = Query(default=20, ge=1, le=100),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene el historial de versiones de un código.
    
    Retorna lista de cambios con versión, acción, memos y timestamp.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_clients_or_error(settings)
    try:
        history = fetch_code_history(clients, project_id, codigo, limit=limit)
    finally:
        clients.close()
    return {"codigo": codigo, "history": history, "total": len(history)}


@app.get("/api/fragments/sample")
async def api_fragments_sample(
    limit: int = 8,
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    clients = build_service_clients(settings)
    try:
        rows = sample_postgres(clients, project_id, limit=limit)
    finally:
        clients.close()
    return {"samples": rows}


class AnalyzeRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    project: str
    docx_path: str
    persist: bool = False
    sync: bool = True  # Default to sync execution (Celery worker not processing tasks)
    run_id: Optional[str] = None


@app.post("/api/analyze", status_code=202)
async def api_analyze(
    payload: AnalyzeRequest,
    request: Request,
    settings: AppSettings = Depends(get_settings),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    # Etapa 0 gate: requiere preparación o override aprobado.
    try:
        stage0_require_ready_or_override(
            clients.postgres,
            project=project_id,
            scope="analyze",
            user_id=user.user_id,
        )
    except PermissionError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc


    # Sprint 20: Gate de backlog - bloquear análisis si hay demasiados pendientes
    BACKLOG_THRESHOLD = int(os.getenv("CANDIDATE_BACKLOG_THRESHOLD", "100"))
    try:
        from app.postgres_block import count_pending_candidates
        pending_count = count_pending_candidates(clients.postgres, project_id)
        if pending_count > BACKLOG_THRESHOLD:
            raise HTTPException(
                status_code=429,
                detail={
                    "code": "BACKLOG_LIMIT_EXCEEDED",
                    "message": f"Hay {pending_count} códigos pendientes de validar. "
                               f"Valida al menos {pending_count - BACKLOG_THRESHOLD} "
                               "antes de ejecutar más análisis.",
                    "pending_count": pending_count,
                    "threshold": BACKLOG_THRESHOLD,
                }
            )
    except ImportError:
        pass  # Función no disponible, continuar sin validación
    except Exception as e:
        api_logger.warning("analyze.backlog_check_failed", error=str(e))

    request_id = getattr(getattr(request, "state", None), "request_id", None)
    run_id = payload.run_id or str(uuid.uuid4())

    docx_path = Path(payload.docx_path)
    downloaded_temp_path: Optional[Path] = None
    # Preserve original filename BEFORE any temp file download
    original_filename = docx_path.name
    # Security check: prevent traversing up
    if ".." in str(docx_path):
         raise HTTPException(status_code=400, detail="Ruta de archivo invalida.")

    if not docx_path.exists():
         # Try looking in project structure (unified paths)
         candidates = [
             Path(f"data/projects/{project_id}/interviews") / docx_path.name,
             Path(f"data/projects/{project_id}/audio/transcriptions") / docx_path.name,
             # Legacy paths
             Path("data/test_interviews/transcription_interviews") / docx_path.name,
             Path("data/interviews") / docx_path.name,
         ]
         found = False
         for candidate in candidates:
             if candidate.exists():
                 docx_path = candidate
                 found = True
                 break

         if not found:
             # Attempt to download from Azure Blob Storage (ingested files live there now)
             try:
                 from app.blob_storage import download_file, CONTAINER_INTERVIEWS

                 blob_path = f"{project_id}/{original_filename}"
                 file_bytes = download_file(CONTAINER_INTERVIEWS, blob_path)
                 with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                     tmp_file.write(file_bytes)
                     downloaded_temp_path = Path(tmp_file.name)
                 docx_path = downloaded_temp_path
                 found = True
                 api_logger.info("analyze.downloaded_blob", blob_path=blob_path, original_name=original_filename)
             except Exception as blob_error:
                 api_logger.warning(
                     "analyze.download_blob_failed",
                     file=docx_path.name,
                     project=project_id,
                     error=str(blob_error),
                 )

         if not found:
              raise HTTPException(status_code=404, detail=f"Archivo no encontrado: {payload.docx_path}")

    # Load fragments immediately (fast I/O)
    fragments = load_fragments(docx_path)

    # Clean up temporary download once fragments are loaded
    if downloaded_temp_path and downloaded_temp_path.exists():
        try:
            downloaded_temp_path.unlink()
        except OSError:
            pass
    
    # Synchronous execution (default - Celery worker not processing tasks)
    if payload.sync:
        from app.analysis import analyze_interview_text, persist_analysis
        from app.reports import generate_interview_report, save_interview_report
        from app.coding import get_all_codes_for_project
        
        api_logger.info("analyze.sync.start", file=docx_path.name, project=project_id)
        
        try:
            # Clear any aborted transaction state from previous operations
            try:
                clients.postgres.rollback()
            except Exception:
                pass  # Ignore if no transaction to rollback
            
            # Get existing codes for novelty calculation
            try:
                existing_codes = get_all_codes_for_project(clients.postgres, project_id)
            except Exception:
                clients.postgres.rollback()  # Clear failed transaction state
                existing_codes = []
            
            # Execute LLM analysis (stages 0-4)
            result = analyze_interview_text(
                clients,
                settings,
                fragments,
                fuente=original_filename,
                project_id=project_id,
                run_id=run_id,
                request_id=request_id,
            )
            
            # Persist if requested
            if payload.persist:
                persist_analysis(
                    clients,
                    settings,
                    original_filename,
                    result,
                    project_id=project_id,
                    run_id=run_id,
                    request_id=request_id,
                )
                api_logger.info("analyze.sync.persisted", file=original_filename)
            
            # Generate interview report
            try:
                report = generate_interview_report(
                    archivo=original_filename,
                    project_id=project_id,
                    analysis_result=result,
                    existing_codes=existing_codes,
                    fragments_total=len(fragments),
                    llm_model=settings.azure.deployment_chat,
                )
                save_interview_report(clients.postgres, report)
                api_logger.info("analyze.sync.report_saved", file=original_filename)
                result["interview_report"] = report.to_dict()
            except Exception as report_error:
                api_logger.warning("analyze.sync.report_error", error=str(report_error))
                result["interview_report"] = None
            
            api_logger.info("analyze.sync.complete", file=original_filename, project=project_id)
            return {"status": "success", "result": result}
            
        except Exception as e:
            api_logger.exception("analyze.sync.error", file=original_filename)
            raise HTTPException(status_code=500, detail=f"Error en análisis: {str(e)}")
    
    # Async execution (Celery - currently not working)
    task = cast(Any, task_analyze_interview).delay(
        project_id=project_id,
        docx_path=str(docx_path),
        fragments=fragments,
        persist=payload.persist,
        file_name=original_filename,
        run_id=run_id,
        request_id=request_id,
    )
    
    api_logger.info("analyze.queued", task_id=task.id, file=original_filename)
    
    return {"task_id": task.id, "status": "queued", "run_id": run_id, "request_id": request_id}


@app.get("/api/tasks/{task_id}")
async def get_task_status(
    task_id: str,
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        task_result = AsyncResult(task_id, app=celery_app)
        
        # Handle different task states
        if task_result.failed():
            # Task failed - get the exception info
            error_info = str(task_result.result) if task_result.result else "Unknown error"
            return {
                "task_id": task_id,
                "status": "FAILURE",
                "result": None,
                "error": error_info,
            }
        
        result = {
            "task_id": task_id,
            "status": task_result.status,
            "result": task_result.result if task_result.ready() and task_result.successful() else None,
        }
        return result
    except Exception as e:
        api_logger.error("task_status.error", task_id=task_id, error=str(e))
        return {
            "task_id": task_id,
            "status": "ERROR",
            "result": None,
            "error": str(e),
        }


class PersistAnalysisRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")
    project: str
    archivo: str
    analysis_result: Dict[str, Any]


@app.post("/api/analyze/persist")
async def api_analyze_persist(
    payload: PersistAnalysisRequest,
    request: Request,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    log = api_logger.bind(endpoint="analyze.persist", project=project_id, file=payload.archivo)

    try:
        # Etapa 0 gate: persistir análisis también requiere preparación o override aprobado.
        try:
            stage0_require_ready_or_override(
                clients.postgres,
                project=project_id,
                scope="analyze",
                user_id=user.user_id,
            )
        except PermissionError as exc:
            raise HTTPException(status_code=409, detail=str(exc)) from exc

        request_id = getattr(getattr(request, "state", None), "request_id", None)
        cm = payload.analysis_result.get("cognitive_metadata") if isinstance(payload.analysis_result, dict) else None
        run_id = cm.get("run_id") if isinstance(cm, dict) else None

        persist_analysis(
            clients,
            settings,
            payload.archivo,
            payload.analysis_result,
            project_id=project_id,
            run_id=run_id,
            request_id=request_id,
        )
        log.info("analyze.persisted_manual")
        return {"status": "ok", "message": "Analisis persistido correctamente."}
    except Exception as exc:
        log.error("api.analyze.persist.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


@app.post("/api/coding/suggestions")
async def api_coding_suggestions(
    payload: CodeSuggestionRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    clients = build_clients_or_error(settings)
    try:
        # Embed fragment
        vectors = embed_batch(clients.aoai, settings.azure.deployment_embed, [payload.fragment_text])
        vector = vectors[0] if vectors else []
        
        # Search Similar
        points = search_similar(
            clients.qdrant,
            settings.qdrant.collection,
            vector=vector,
            limit=payload.limit * 2, # Fetch more to aggregate
        )
        
        # Aggregate Codes
        # Assuming payload has 'codigos_ancla' or we can fetch from Neo4j given the ID.
        # For speed, let's rely on Qdrant payload 'codigos_ancla' if it exists.
        
        suggested_codes = Counter()
        evidence_map = {}
        
        for point in points:
            codes = point.payload.get("codigos_ancla", [])
            if isinstance(codes, list):
                for code in codes:
                    suggested_codes[code] += 1
                    if code not in evidence_map:
                         evidence_map[code] = point.payload.get("fragmento")
            elif isinstance(codes, str): # Legacy single code
                 suggested_codes[codes] += 1
                 if codes not in evidence_map:
                     evidence_map[codes] = point.payload.get("fragmento")

        top_suggestions = []
        for code, count in suggested_codes.most_common(5):
            top_suggestions.append({
                "code": code,
                "confidence": count / len(points), # Naive confidence
                "example": evidence_map.get(code)
            })
            
        return {"suggestions": top_suggestions}

    except Exception as exc:
        api_logger.error("api.suggestions.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


# =============================================================================
# Sprint 9: GraphRAG, Discovery API, Link Prediction
# =============================================================================

class GraphRAGRequest(BaseModel):
    """Request for GraphRAG query."""
    query: str = Field(..., description="Pregunta del usuario")
    project: str = Field(default="default", description="ID del proyecto")
    include_fragments: bool = Field(default=True, description="Incluir fragmentos en contexto")
    chain_of_thought: bool = Field(default=False, description="Usar razonamiento paso a paso")
    filters: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Filtros aplicados desde la UI (project_id, scope, date_range, relation_types, k_qdrant, include_discovery, etc.)",
    )


class DiscoverRequest(BaseModel):
    """Request for Discovery API search."""
    positive_texts: List[str] = Field(..., description="Textos/conceptos positivos")
    negative_texts: Optional[List[str]] = Field(default=None, description="Textos a evitar")
    target_text: Optional[str] = Field(default=None, description="Texto objetivo")
    top_k: int = Field(default=10, ge=1, le=50)
    project: str = Field(default="default")


class LinkPredictionRequest(BaseModel):
    """Request for Link Prediction."""
    source_type: str = Field(default="Categoria", description="Tipo de nodo fuente")
    target_type: str = Field(default="Codigo", description="Tipo de nodo destino")
    algorithm: str = Field(default="common_neighbors", description="Algoritmo a usar")
    top_k: int = Field(default=10, ge=1, le=50)
    project: str = Field(default="default")
    categoria: Optional[str] = Field(default=None, description="Filtrar por categoria")


@app.post("/api/graphrag/query")
async def api_graphrag_query(
    payload: GraphRAGRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Ejecuta una consulta GraphRAG con contexto de grafo.
    
    Combina busqueda semantica + estructura del grafo + LLM para respuestas
    contextualizadas sobre la investigacion cualitativa.
    """
    from app.graphrag import graphrag_query, graphrag_chain_of_thought
    
    clients = build_clients_or_error(settings)
    try:
        api_logger.info("api.graphrag.start", query=payload.query[:50], project=payload.project)
        
        if payload.chain_of_thought:
            result = graphrag_chain_of_thought(
                clients, settings,
                query=payload.query,
                project=payload.project,
            )
        else:
            result = graphrag_query(
                clients, settings,
                query=payload.query,
                project=payload.project,
                include_fragments=payload.include_fragments,
                filters_applied=payload.filters,
            )
        
        # safe answer length calculation: some responses use structured fields (graph_summary) instead of 'answer'
        answer_text = (
            result.get("answer")
            or result.get("graph_summary")
            or result.get("analysis")
            or ""
        )
        api_logger.info(
            "api.graphrag.complete",
            query=payload.query[:50],
            nodes=len(result.get("nodes") or []),
            answer_len=len(str(answer_text)),
        )
        
        return result
        
    except Exception as exc:
        api_logger.error("api.graphrag.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()

class GraphRAGSaveRequest(BaseModel):
    query: str
    answer: str
    context: Optional[str] = ""
    nodes: List[Dict[str, Any]] = []
    relationships: List[Dict[str, Any]] = []
    fragments: List[Dict[str, Any]] = []
    project: str = "default"


@app.post("/api/graphrag/save_report")
async def api_save_report(
    payload: GraphRAGSaveRequest,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, str]:
    """Guarda el reporte de GraphRAG como artefacto (Blob Storage, multi-tenant)."""
    try:
        # Multi-tenant: resolver project_id canónico + validar acceso.
        project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        
        # Safe slug from query (first 30 chars, alphanumeric only)
        safe_query = "".join(c if c.isalnum() else "_" for c in payload.query[:30]).strip("_")
        filename = f"{timestamp}_{safe_query}.md"
        logical_path = f"reports/{project_id}/{filename}"
        
        # Format Content
        lines = [
            f"# Reporte de Investigación: {payload.query}",
            f"**Fecha:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**Proyecto:** {project_id}",
            "\n## Respuesta",
            payload.answer,
            "\n## Contexto del Grafo",
            payload.context or "(Sin contexto)",
            f"\n## Evidencia ({len(payload.fragments)} fragmentos)",
        ]
        
        for idx, frag in enumerate(payload.fragments):
            source = frag.get("archivo", "Desconocido")
            text = frag.get("fragmento", "")
            lines.append(f"\n### Fragmento {idx+1} ({source})")
            lines.append(f"> {text}")
            
        content = "\n".join(lines)

        # Persist to Blob Storage (reports container, tenant-scoped prefix).
        try:
            from app.blob_storage import CONTAINER_REPORTS, logical_path_to_blob_name, upload_text

            blob_name = logical_path_to_blob_name(
                org_id=str(getattr(user, "organization_id", None) or ""),
                project_id=project_id,
                logical_path=logical_path,
            )
            blob_url = upload_text(
                container=CONTAINER_REPORTS,
                blob_name=blob_name,
                text=content,
                content_type="text/markdown; charset=utf-8",
            )
        except Exception as exc:
            api_logger.error("report.save.blob_failed", error=str(exc)[:200], project=project_id)
            raise HTTPException(status_code=503, detail="Blob Storage no disponible") from exc

        api_logger.info("report.saved", path=logical_path, blob=blob_name)
        return {"status": "ok", "path": logical_path, "filename": filename, "blob_url": blob_url}
        
    except HTTPException:
        raise
    except Exception as exc:
        api_logger.error("report.save.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc


class DoctoralReportRequest(BaseModel):
    """Request para generar informe doctoral."""
    stage: str  # "stage3" o "stage4"
    project: str = "default"


# -----------------------------------------------------------------------------
# Doctoral report jobs (async)
# -----------------------------------------------------------------------------

def _user_is_admin(user: User) -> bool:
    roles = user.roles or []
    return "admin" in {str(r).strip().lower() for r in roles if r}


def _assert_doctoral_task_access(*, task_id: str, task: Dict[str, Any], user: User) -> None:
    auth = task.get("auth") or {}
    owner_user_id = str(auth.get("user_id") or "").strip()
    if not owner_user_id:
        # Legacy/backward-compat: if task has no auth metadata, allow only admins.
        if _user_is_admin(user):
            return
        raise HTTPException(status_code=403, detail="Forbidden: task has no owner metadata")

    if owner_user_id != str(user.user_id):
        if _user_is_admin(user):
            return
        raise HTTPException(status_code=403, detail="Forbidden: task belongs to another user")


def _run_doctoral_report_task(*, task_id: str, payload: DoctoralReportRequest, settings: AppSettings, org_id: str) -> None:
    """Background worker for stage report generation."""
    from app.stage_reports import generate_stage3_report, generate_stage4_report
    from app.postgres_block import save_doctoral_report, update_report_job
    from app.blob_storage import CONTAINER_REPORTS, logical_path_to_blob_name, upload_text

    clients = build_clients_or_error(settings)
    try:
        # Mark running
        try:
            update_report_job(
                clients.postgres,
                task_id=task_id,
                status="running",
                message="Generando informe...",
                started_at=datetime.now().isoformat(),
            )
        except Exception:
            pass

        if payload.stage == "stage3":
            result = generate_stage3_report(clients, settings, payload.project, org_id=org_id)
        elif payload.stage == "stage4":
            result = generate_stage4_report(clients, settings, payload.project, org_id=org_id)
        else:
            update_report_job(
                clients.postgres,
                task_id=task_id,
                status="error",
                message=f"Etapa no válida: {payload.stage}. Use 'stage3' o 'stage4'.",
                errors=[f"Etapa no válida: {payload.stage}"],
                finished_at=datetime.now().isoformat(),
            )
            return

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        filename = f"{timestamp}_doctoral_{payload.stage}.md"
        logical_path = f"reports/{payload.project}/doctoral/{filename}"
        content = result.get("content") or ""

        blob_url: Optional[str] = None
        try:
            blob_name = logical_path_to_blob_name(org_id=org_id, project_id=payload.project, logical_path=logical_path)
            blob_url = upload_text(container=CONTAINER_REPORTS, blob_name=blob_name, text=content, content_type="text/markdown; charset=utf-8")
        except Exception as exc:
            api_logger.warning("stage_reports.report.blob_upload_failed", error=str(exc)[:200], task_id=task_id)

        # Persist DB record for history
        report_id = save_doctoral_report(
            clients.postgres,
            project=payload.project,
            stage=payload.stage,
            content=content,
            stats=result.get("stats"),
            file_path=logical_path,
        )

        final_result = {
            **(result or {}),
            "stage": payload.stage,
            "project": payload.project,
            "path": logical_path,
            "filename": filename,
            "report_id": report_id,
            "blob_url": blob_url,
            "blob_name": blob_name if blob_url else None,
        }

        update_report_job(
            clients.postgres,
            task_id=task_id,
            status="completed",
            message="Informe de avance generado",
            result=final_result,
            result_path=logical_path,
            finished_at=datetime.now().isoformat(),
        )
    except Exception as exc:
        try:
            update_report_job(
                clients.postgres,
                task_id=task_id,
                status="error",
                message=str(exc),
                errors=[str(exc)],
                finished_at=datetime.now().isoformat(),
            )
        except Exception:
            pass
        api_logger.error("stage_reports.report.job_error", error=str(exc), task_id=task_id)
    finally:
        clients.close()


class DoctoralReportJobStatusResponse(BaseModel):
    task_id: str
    status: str
    project: str
    stage: str
    message: Optional[str] = None
    started_at: Optional[str] = None
    finished_at: Optional[str] = None
    errors: Optional[List[str]] = None


class DoctoralReportJobResultResponse(BaseModel):
    task_id: str
    status: str
    result: Dict[str, Any]


@app.post("/api/reports/doctoral/execute")
async def api_execute_doctoral_report_job(
    payload: DoctoralReportRequest,
    background_tasks: BackgroundTasks,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Dict[str, Any]:
    """Start an async doctoral report generation job."""
    try:
        resolved_project = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    task_id = f"doctoral_{resolved_project}_{payload.stage}_{uuid.uuid4().hex[:12]}"
    auth = {
        "user_id": str(user.user_id),
        "org": str(user.organization_id),
        "roles": list(user.roles or []),
    }
    # Persist job row (durable)
    try:
        from app.postgres_block import create_report_job

        clients = build_clients_or_error(settings)
        try:
            # Multi-tenant: validate project access using the live pg connection.
            resolved_project = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
            create_report_job(
                clients.postgres,
                task_id=task_id,
                job_type="doctoral",
                project_id=resolved_project,
                payload={**payload.model_dump(), "project": resolved_project},
                auth=auth,
                message="Inicializando...",
            )
        finally:
            clients.close()
    except Exception as exc:
        api_logger.warning("doctoral.report.job_persist_failed", error=str(exc), task_id=task_id)

    api_logger.info("doctoral.report.job_started", task_id=task_id, project=resolved_project, stage=payload.stage)

    # FastAPI injects BackgroundTasks, but we keep it typed loosely to avoid extra imports.
    payload.project = resolved_project
    background_tasks.add_task(
        _run_doctoral_report_task,
        task_id=task_id,
        payload=payload,
        settings=settings,
        org_id=str(user.organization_id),
    )
    return {"task_id": task_id, "status": "started"}


@app.get("/api/reports/doctoral/status/{task_id}", response_model=DoctoralReportJobStatusResponse)
async def api_get_doctoral_report_job_status(
    task_id: str,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> DoctoralReportJobStatusResponse:
    # Load from durable store
    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import get_report_job

        task = get_report_job(clients.postgres, task_id)
    finally:
        clients.close()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task {task_id} not found")
    _assert_doctoral_task_access(task_id=task_id, task=task, user=user)
    return DoctoralReportJobStatusResponse(
        task_id=task_id,
        status=str(task.get("status") or "pending"),
        project=str(task.get("project_id") or "default"),
        stage=str(((task.get("payload") or {}) if isinstance(task.get("payload"), dict) else {}).get("stage") or ""),
        message=task.get("message"),
        started_at=task.get("started_at") or task.get("created_at"),
        finished_at=task.get("finished_at"),
        errors=task.get("errors") or None,
    )


@app.get("/api/reports/doctoral/result/{task_id}", response_model=DoctoralReportJobResultResponse)
async def api_get_doctoral_report_job_result(
    task_id: str,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> DoctoralReportJobResultResponse:
    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import get_report_job

        task = get_report_job(clients.postgres, task_id)
    finally:
        clients.close()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task {task_id} not found")
    _assert_doctoral_task_access(task_id=task_id, task=task, user=user)
    if task.get("status") != "completed":
        raise HTTPException(status_code=400, detail=f"Task not completed: {task.get('status')}")
    return DoctoralReportJobResultResponse(task_id=task_id, status="completed", result=task.get("result") or {})


@app.post("/api/reports/generate-doctoral")
async def api_generate_doctoral_report(
    payload: DoctoralReportRequest,
    settings: AppSettings = Depends(get_settings),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Genera informe de avance analítico para Etapa 3 o Etapa 4.
    
    - stage3: Codificación Abierta (códigos, saturación, memos)
    - stage4: Codificación Axial (categorías, comunidades, núcleo)
    
    El informe se guarda tanto en archivo como en la base de datos para
    trazabilidad y análisis histórico.
    """
    from app.stage_reports import generate_stage3_report, generate_stage4_report
    from app.postgres_block import save_doctoral_report
    
    try:
        project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
        if payload.stage == "stage3":
            result = generate_stage3_report(clients, settings, project_id, org_id=str(getattr(user, "organization_id", None) or ""))
        elif payload.stage == "stage4":
            result = generate_stage4_report(clients, settings, project_id, org_id=str(getattr(user, "organization_id", None) or ""))
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Etapa no válida: {payload.stage}. Use 'stage3' o 'stage4'."
            )
        
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        filename = f"{timestamp}_doctoral_{payload.stage}.md"

        logical_path = f"reports/{project_id}/doctoral/{filename}"
        content = result.get("content") or ""
        blob_url: Optional[str] = None
        try:
            from app.blob_storage import CONTAINER_REPORTS, logical_path_to_blob_name, upload_text

            blob_name = logical_path_to_blob_name(
                org_id=str(getattr(user, "organization_id", None) or ""),
                project_id=project_id,
                logical_path=logical_path,
            )
            blob_url = upload_text(
                container=CONTAINER_REPORTS,
                blob_name=blob_name,
                text=content,
                content_type="text/markdown; charset=utf-8",
            )
        except Exception as exc:
            api_logger.warning("stage_reports.report.blob_upload_failed", error=str(exc)[:200])

        result["path"] = logical_path
        result["filename"] = filename
        result["blob_url"] = blob_url
        
        # Sprint 26: Guardar en base de datos para persistencia
        report_id = save_doctoral_report(
            clients.postgres,
            project=project_id,
            stage=payload.stage,
            content=content,
            stats=result.get("stats"),
            file_path=logical_path,
        )
        result["report_id"] = report_id
        
        api_logger.info(
            "stage_reports.report.generated",
            stage=payload.stage,
            project=project_id,
            path=logical_path,
            report_id=report_id
        )
        
        return result
        
    except HTTPException:
        raise
    except Exception as exc:
        api_logger.error("stage_reports.report.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc


class DiscoverySaveMemoRequest(BaseModel):
    positive_texts: List[str]
    negative_texts: Optional[List[str]] = None
    target_text: Optional[str] = None
    fragments: List[Dict[str, Any]] = []
    project: str = "default"
    memo_title: Optional[str] = None
    ai_synthesis: Optional[str] = None  # Síntesis generada por IA



@app.post("/api/discovery/save_memo")
async def api_save_discovery_memo(
    payload: DiscoverySaveMemoRequest,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, str]:
    """Guarda los resultados de Discovery como memo (Blob Storage, multi-tenant)."""
    try:
        project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        title = payload.memo_title or "_".join(payload.positive_texts[:2])
        safe_title = "".join(c if c.isalnum() else "_" for c in title[:30]).strip("_")
        filename = f"{timestamp}_discovery_{safe_title}.md"
        logical_path = f"notes/{project_id}/{filename}"
        
        # Format Content
        lines = [
            f"# Memo de Exploración: {payload.memo_title or ', '.join(payload.positive_texts)}",
            f"**Fecha:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**Proyecto:** {project_id}",
            "",
            "## Criterios de Búsqueda",
            f"**Conceptos Positivos:** {', '.join(payload.positive_texts)}",
            f"**Conceptos Negativos:** {', '.join(payload.negative_texts or ['(ninguno)'])}",
            f"**Texto Objetivo:** {payload.target_text or '(ninguno)'}",
            "",
            f"## Fragmentos Encontrados ({len(payload.fragments)})",
        ]
        
        for idx, frag in enumerate(payload.fragments):
            source = frag.get("archivo", "Desconocido")
            text = frag.get("fragmento", "")
            score = frag.get("score", 0)
            lines.append(f"\n### [{idx+1}] {source} (relevancia: {score:.1%})")
            lines.append(f"> {text}")
        
        # Agregar síntesis IA si existe
        if payload.ai_synthesis:
            lines.append("")
            lines.append("## 🧠 Síntesis y Sugerencias (IA)")
            lines.append("")
            for line in payload.ai_synthesis.split('\n'):
                lines.append(line)
            
        lines.append("\n---")
        lines.append("*Generado automáticamente por Discovery Search*")
            
        content = "\n".join(lines)

        try:
            from app.blob_storage import CONTAINER_REPORTS, logical_path_to_blob_name, upload_text

            blob_name = logical_path_to_blob_name(
                org_id=str(getattr(user, "organization_id", None) or ""),
                project_id=project_id,
                logical_path=logical_path,
            )
            blob_url = upload_text(
                container=CONTAINER_REPORTS,
                blob_name=blob_name,
                text=content,
                content_type="text/markdown; charset=utf-8",
            )
        except Exception as exc:
            api_logger.error("discovery.memo.save.blob_failed", error=str(exc)[:200], project=project_id)
            raise HTTPException(status_code=503, detail="Blob Storage no disponible") from exc

        api_logger.info("discovery.memo.saved", path=logical_path, blob=blob_name, has_synthesis=bool(payload.ai_synthesis))
        return {"status": "ok", "path": logical_path, "filename": filename, "blob_url": blob_url}
        
    except HTTPException:
        raise
    except Exception as exc:
        api_logger.error("discovery.memo.save.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/search/discover")
async def api_discover_search(
    payload: DiscoverRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Busqueda exploratoria con triplete (positivo, negativo, target).
    
    Encuentra fragmentos similares a los conceptos positivos pero
    diferentes de los negativos.
    """
    from app.queries import discover_search
    
    clients = build_clients_or_error(settings)
    try:
        api_logger.info(
            "api.discover.start",
            positive_count=len(payload.positive_texts),
            negative_count=len(payload.negative_texts or []),
        )
        
        results = discover_search(
            clients, settings,
            positive_texts=payload.positive_texts,
            negative_texts=payload.negative_texts,
            target_text=payload.target_text,
            top_k=payload.top_k,
            project=payload.project,
        )
        
        api_logger.info("api.discover.complete", results=len(results))
        
        return {"fragments": results, "count": len(results)}
        
    except Exception as exc:
        api_logger.error("api.discover.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


class AnalyzePredictionsRequest(BaseModel):
    """Request para análisis IA de predicciones de enlaces."""
    model_config = ConfigDict(extra="forbid")
    project: str = Field(default="default", description="ID del proyecto")
    algorithm: str = Field(..., description="Algoritmo usado para la predicción")
    suggestions: List[Dict[str, Any]] = Field(..., description="Sugerencias de enlaces")
    persist: bool = Field(
        default=True,
        description="Si True, persiste el análisis como artefacto auditable (axial_ai_analyses).",
    )


@app.post("/api/axial/analyze-predictions")
async def api_analyze_predictions(
    payload: AnalyzePredictionsRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Analiza predicciones de enlaces con IA.
    
    Envía las sugerencias de link prediction al LLM para obtener
    una interpretación cualitativa y recomendaciones.
    """
    from app.clients import build_service_clients
    
    ALGORITHM_DESCRIPTIONS = {
        "common_neighbors": "Vecinos Comunes - cuenta nodos compartidos entre dos códigos",
        "jaccard": "Coeficiente Jaccard - mide similitud entre conjuntos de vecinos",
        "adamic_adar": "Adamic-Adar - pondera vecinos comunes por su exclusividad",
        "preferential_attachment": "Preferential Attachment - favorece códigos populares",
        "community_based": "Basado en Comunidades - detecta enlaces intra-comunidad faltantes",
    }
    
    clients = build_service_clients(settings)
    try:
        # Multi-tenant: resolver project_id canónico + validar acceso.
        try:
            project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        # Epistemic mode → prompts versionados (audit trail)
        from app.postgres_block import get_project_epistemic_mode
        from app.prompts.loader import get_system_prompt

        epistemic_mode = get_project_epistemic_mode(clients.postgres, project_id)
        system_prompt, prompt_version = get_system_prompt(epistemic_mode, "axial_coding")

        algorithm_desc = ALGORITHM_DESCRIPTIONS.get(
            payload.algorithm, 
            f"Algoritmo: {payload.algorithm}"
        )

        api_logger.info(
            "api.analyze_predictions.start",
            project=project_id,
            algorithm=payload.algorithm,
            epistemic_mode=epistemic_mode.value,
            prompt_version=prompt_version,
            suggestions_count=len(payload.suggestions or []),
        )
        
        MAX_SUGGESTIONS = 10
        POS_TOTAL = 24
        NEG_TOTAL = 12
        EXCERPT_CHARS = 240
        POS_IN_PROMPT = 3
        NEG_IN_PROMPT = 2

        suggestions_for_prompt = payload.suggestions[:MAX_SUGGESTIONS]

        # AX-AI-03: Evidence pack (positivo/negativo) para grounding + auditoría reproducible.
        from datetime import datetime, timezone
        evidence_pack: Dict[str, Any]
        try:
            from app.axial_evidence import build_link_prediction_evidence_pack

            evidence_pack = build_link_prediction_evidence_pack(
                clients,
                settings,
                project_id=project_id,
                suggestions=suggestions_for_prompt,
                positive_total=POS_TOTAL,
                negative_total=NEG_TOTAL,
                excerpt_chars=EXCERPT_CHARS,
            )
        except Exception as exc:  # noqa: BLE001
            api_logger.warning(
                "api.analyze_predictions.evidence_failed",
                project=project_id,
                error=str(exc)[:200],
            )
            now = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")
            evidence_pack = {
                "schema_version": 1,
                "generated_at": now,
                "limits": {
                    "positive_total": POS_TOTAL,
                    "negative_total": NEG_TOTAL,
                    "excerpt_chars": EXCERPT_CHARS,
                },
                "suggestions": [],
                "totals": {"positive": 0, "negative": 0, "by_method": {}},
                "notes": {"reason": "build_failed"},
            }

        evidence_by_id: Dict[int, Dict[str, Any]] = {}
        try:
            if isinstance(evidence_pack, dict) and isinstance(evidence_pack.get("suggestions"), list):
                for item in evidence_pack["suggestions"]:
                    if not isinstance(item, dict):
                        continue
                    try:
                        sid = int(item.get("id") or 0)
                    except Exception:
                        continue
                    if sid > 0:
                        evidence_by_id[sid] = item
        except Exception:
            evidence_by_id = {}

        def _one_line(value: Any) -> str:
            return str(value or "").replace("\r", " ").replace("\n", " ").strip()

        # Preparar contexto de sugerencias (+ evidencia resumida) para el prompt.
        suggestions_text: List[str] = []
        for i, sug in enumerate(suggestions_for_prompt, 1):  # Max 10 para el prompt
            source = sug.get("source", "?")
            target = sug.get("target", "?")
            score = sug.get("score", 0)
            try:
                score_f = float(score or 0.0)
            except Exception:
                score_f = 0.0

            block_lines = [f"{i}. {source} ↔ {target} (score: {score_f:.3f})"]
            ev = evidence_by_id.get(i) or {}
            pos_items = ev.get("positive") if isinstance(ev.get("positive"), list) else []
            neg_items = ev.get("negative") if isinstance(ev.get("negative"), list) else []

            if pos_items:
                block_lines.append("   EVIDENCIA_POSITIVA:")
                for e in pos_items[:POS_IN_PROMPT]:
                    if not isinstance(e, dict):
                        continue
                    fid = _one_line(e.get("fragmento_id"))
                    archivo = _one_line(e.get("archivo"))
                    par_idx = e.get("par_idx")
                    frag = _one_line(e.get("fragmento"))
                    method = _one_line(e.get("method"))
                    block_lines.append(
                        f"   - [{fid}] {archivo}:{par_idx} \"{frag}\" ({method})".strip()
                    )
            if neg_items:
                block_lines.append("   EVIDENCIA_NEGATIVA (tensión):")
                for e in neg_items[:NEG_IN_PROMPT]:
                    if not isinstance(e, dict):
                        continue
                    fid = _one_line(e.get("fragmento_id"))
                    archivo = _one_line(e.get("archivo"))
                    par_idx = e.get("par_idx")
                    frag = _one_line(e.get("fragmento"))
                    method = _one_line(e.get("method"))
                    ps = e.get("present_source")
                    pt = e.get("present_target")
                    present = ""
                    if ps is not None or pt is not None:
                        present = f" present(source={int(bool(ps))},target={int(bool(pt))})"
                    block_lines.append(
                        f"   - [{fid}] {archivo}:{par_idx}{present} \"{frag}\" ({method})".strip()
                    )

            suggestions_text.append("\n".join(block_lines))
        suggestions_block = "\n\n".join(suggestions_text)
        
        # Sprint 29+: Contrato JSON estructurado con etiquetas epistemológicas (compatible con legacy texto).
        prompt = f"""Analiza las siguientes predicciones de enlaces en un grafo de análisis cualitativo (Teoría Fundamentada).

Algoritmo utilizado: {algorithm_desc}

Para cada sugerencia se incluye evidencia en fragmentos (EVIDENCIA_POSITIVA / EVIDENCIA_NEGATIVA).
Cita la evidencia usando los `fragmento_id` proporcionados. No inventes IDs.

Sugerencias de relaciones axiales faltantes (IDs 1..{len(suggestions_for_prompt)}):
{suggestions_block}

Contexto: Estas son relaciones potenciales entre códigos/categorías detectadas algorítmicamente.

IMPORTANTE: Responde ÚNICAMENTE con un JSON válido (sin markdown, sin ```).
La síntesis debe distinguir explícitamente el estatus epistemológico de cada afirmación.
El JSON debe tener esta estructura exacta:

{{
    "memo_sintesis": [
        {{"type": "OBSERVATION", "text": "...", "evidence_ids": [1, 3], "evidence_fragment_ids": ["123", "456"]}},
        {{"type": "INTERPRETATION", "text": "...", "evidence_ids": [1, 3], "evidence_fragment_ids": ["123"]}},
        {{"type": "HYPOTHESIS", "text": "...", "evidence_ids": [2], "evidence_fragment_ids": ["789"]}},
        {{"type": "NORMATIVE_INFERENCE", "text": "...", "evidence_ids": [4]}}
    ]
}}

REGLAS:
1. memo_sintesis: lista de 3-6 statements.
2. type: uno de OBSERVATION | INTERPRETATION | HYPOTHESIS | NORMATIVE_INFERENCE.
3. text: una oración clara (español).
4. evidence_ids: lista de enteros referidos a la numeración de las sugerencias (1..{len(suggestions_for_prompt)}).
5. evidence_fragment_ids: lista de strings `fragmento_id` (de la evidencia provista). No inventes IDs.
6. PROHIBIDO: OBSERVATION sin evidence_fragment_ids no vacíos.
7. Sé conciso: evita listas largas o párrafos extensos."""

        # gpt-5.x models no soportan temperature != 1, omitir el parámetro
        response = clients.aoai.chat.completions.create(
            model=settings.azure.deployment_chat,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            max_completion_tokens=600,
            # temperature omitido: gpt-5.x usa default=1
        )
        
        choice = response.choices[0] if response.choices else None
        message = getattr(choice, "message", None)
        message_content = getattr(message, "content", None)
        raw_content = message_content or ""

        # Parsear JSON estructurado (con fallback al texto)
        import json
        import re

        parsed: Optional[Dict[str, Any]] = None
        try:
            clean_content = re.sub(r"^```json?\s*", "", raw_content.strip())
            clean_content = re.sub(r"\s*```$", "", clean_content)
            parsed = json.loads(clean_content)
        except Exception:
            parsed = None

        def _normalize_memo(value: Any) -> tuple[str, List[Dict[str, Any]]]:
            allowed = {"OBSERVATION", "INTERPRETATION", "HYPOTHESIS", "NORMATIVE_INFERENCE"}
            if isinstance(value, str):
                return value.strip(), []
            if not isinstance(value, list):
                return "", []
            out: List[Dict[str, Any]] = []
            lines: List[str] = []
            for item in value:
                if not isinstance(item, dict):
                    continue
                stype = str(item.get("type") or "").strip().upper()
                text = str(item.get("text") or "").strip()
                if not text:
                    continue
                if stype not in allowed:
                    stype = "INTERPRETATION"
                evidence_ids: List[int] = []
                raw_ids = item.get("evidence_ids")
                if isinstance(raw_ids, list):
                    for v in raw_ids:
                        try:
                            evidence_ids.append(int(v))
                        except Exception:
                            continue
                evidence_fragment_ids: List[str] = []
                raw_frag_ids = item.get("evidence_fragment_ids")
                if isinstance(raw_frag_ids, list):
                    for v in raw_frag_ids:
                        fid = str(v or "").strip()
                        if fid:
                            evidence_fragment_ids.append(fid)
                # Regla de seguridad: OBSERVATION requiere evidencia en fragmentos.
                if stype == "OBSERVATION" and not evidence_fragment_ids:
                    stype = "INTERPRETATION"
                entry: Dict[str, Any] = {"type": stype, "text": text}
                if evidence_ids:
                    entry["evidence_ids"] = evidence_ids
                if evidence_fragment_ids:
                    entry["evidence_fragment_ids"] = evidence_fragment_ids

                if evidence_fragment_ids:
                    preview = ", ".join(evidence_fragment_ids[:6])
                    suffix = "…" if len(evidence_fragment_ids) > 6 else ""
                    lines.append(f"[{stype}] {text} (frags: {preview}{suffix})")
                elif evidence_ids:
                    lines.append(f"[{stype}] {text} (links: {', '.join(str(i) for i in evidence_ids)})")
                else:
                    lines.append(f"[{stype}] {text}")
                out.append(entry)
            return "\n".join(lines).strip(), out
        
        api_logger.info(
            "api.analyze_predictions.complete",
            project=project_id,
            algorithm=payload.algorithm,
            suggestions_count=len(suggestions_for_prompt),
        )
        
        if isinstance(parsed, dict) and parsed.get("memo_sintesis") is not None:
            memo_text, memo_statements = _normalize_memo(parsed.get("memo_sintesis"))
            analysis = memo_text
            structured = True
        else:
            analysis = raw_content
            structured = False
            memo_statements = []

        analysis_id = 0
        persisted = False
        evidence_schema_version = int(evidence_pack.get("schema_version", 1) or 1) if isinstance(evidence_pack, dict) else 1
        if payload.persist:
            from app.postgres_block import insert_axial_ai_analysis

            analysis_id = insert_axial_ai_analysis(
                clients.postgres,
                project_id=project_id,
                source_type="analyze_predictions",
                algorithm=payload.algorithm,
                algorithm_description=algorithm_desc,
                suggestions_json=payload.suggestions[:10],
                analysis_text=analysis,
                memo_statements=memo_statements,
                structured=structured,
                estado="pendiente",
                created_by=user.user_id,
                epistemic_mode=epistemic_mode.value,
                prompt_version=prompt_version,
                llm_deployment=settings.azure.deployment_chat,
                llm_api_version=settings.azure.api_version,
                evidence_schema_version=evidence_schema_version,
                evidence_json=evidence_pack if isinstance(evidence_pack, dict) else None,
            )
            persisted = bool(analysis_id)
            api_logger.info(
                "api.analyze_predictions.persisted",
                project=project_id,
                analysis_id=analysis_id,
                estado="pendiente",
            )

        return {
            # Compatibilidad: se mantiene el texto plano.
            "analysis": analysis,
            # Nuevo: campos estructurados (opcionales para el frontend).
            "structured": structured,
            "memo_statements": memo_statements,
            "algorithm": payload.algorithm,
            "algorithm_description": algorithm_desc,
            "suggestions_analyzed": len(payload.suggestions[:10]),
            "project": project_id,
            "epistemic_mode": epistemic_mode.value,
            "prompt_version": prompt_version,
            "llm_deployment": settings.azure.deployment_chat,
            "llm_api_version": settings.azure.api_version,
            "analysis_id": analysis_id,
            "persisted": persisted,
            "evidence_schema_version": evidence_schema_version,
            "evidence_totals": evidence_pack.get("totals") if isinstance(evidence_pack, dict) else None,
        }
        
    except Exception as exc:
        api_logger.error("api.analyze_predictions.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


class AxialAiAnalysisUpdateRequest(BaseModel):
    """Request para actualizar estado de un artefacto IA axial."""
    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="Proyecto")
    estado: str = Field(..., description="Nuevo estado: pendiente, validado, rechazado")
    review_memo: Optional[str] = Field(None, description="Memo/nota humana opcional")


@app.get("/api/axial/ai-analyses")
async def api_list_axial_ai_analyses(
    project: str = Query(..., description="Proyecto requerido"),
    estado: Optional[str] = Query(None, description="Filtrar por estado"),
    source_type: Optional[str] = Query(None, description="Filtrar por origen (source_type)"),
    algorithm: Optional[str] = Query(None, description="Filtrar por algoritmo"),
    epistemic_mode: Optional[str] = Query(None, description="Filtrar por epistemic_mode"),
    created_from: Optional[str] = Query(None, description="Filtrar desde (ISO 8601)"),
    created_to: Optional[str] = Query(None, description="Filtrar hasta (ISO 8601)"),
    min_score: Optional[float] = Query(None, description="Filtrar por score mínimo (max score dentro del artefacto)"),
    has_evidence: Optional[bool] = Query(None, description="Filtrar por artefactos con/sin evidencia"),
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista artefactos IA axiales persistidos (audit trail)."""
    clients = build_pg_only(settings)
    try:
        try:
            project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        from app.postgres_block import list_axial_ai_analyses, count_axial_ai_analyses

        items = list_axial_ai_analyses(
            clients.postgres,
            project_id=project_id,
            estado=estado,
            source_type=source_type,
            algorithm=algorithm,
            epistemic_mode=epistemic_mode,
            created_from=created_from,
            created_to=created_to,
            min_score=min_score,
            has_evidence=has_evidence,
            limit=limit,
            offset=offset,
        )
        total = count_axial_ai_analyses(
            clients.postgres,
            project_id=project_id,
            estado=estado,
            source_type=source_type,
            algorithm=algorithm,
            epistemic_mode=epistemic_mode,
            created_from=created_from,
            created_to=created_to,
            min_score=min_score,
            has_evidence=has_evidence,
        )

        return {
            "project": project_id,
            "items": items,
            "total": total,
            "limit": limit,
            "offset": offset,
        }
    finally:
        clients.close()


@app.get("/api/axial/ai-analyses/{analysis_id}")
async def api_get_axial_ai_analysis(
    analysis_id: int,
    project: str = Query(..., description="Proyecto requerido"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Obtiene un artefacto IA axial por ID (scoped por proyecto)."""
    clients = build_pg_only(settings)
    try:
        try:
            project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        from app.postgres_block import get_axial_ai_analysis_by_id

        item = get_axial_ai_analysis_by_id(
            clients.postgres,
            project_id=project_id,
            analysis_id=int(analysis_id),
        )
        if not item:
            raise HTTPException(status_code=404, detail="Análisis IA axial no encontrado")
        return item
    finally:
        clients.close()


@app.patch("/api/axial/ai-analyses/{analysis_id}")
async def api_update_axial_ai_analysis(
    analysis_id: int,
    payload: AxialAiAnalysisUpdateRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Actualiza estado/memo de revisión de un artefacto IA axial (no crea relaciones)."""
    if payload.estado not in ("pendiente", "validado", "rechazado"):
        raise HTTPException(status_code=400, detail="estado debe ser: pendiente, validado, rechazado")

    clients = build_pg_only(settings)
    try:
        try:
            project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        from app.postgres_block import (
            get_axial_ai_analysis_by_id,
            log_project_action,
            update_axial_ai_analysis_estado,
        )

        current = get_axial_ai_analysis_by_id(
            clients.postgres,
            project_id=project_id,
            analysis_id=int(analysis_id),
        )
        if not current:
            raise HTTPException(status_code=404, detail="Análisis IA axial no encontrado")

        current_estado = str(current.get("estado") or "pendiente")
        new_estado = payload.estado
        effective_review_memo = (
            payload.review_memo if payload.review_memo is not None else current.get("review_memo")
        )
        is_admin = "admin" in (user.roles or [])

        if current_estado != new_estado:
            # Regla: cambios desde estados finales requieren rol admin + motivo.
            if current_estado in ("validado", "rechazado"):
                if not is_admin:
                    raise HTTPException(
                        status_code=403,
                        detail="No autorizado: solo admin puede cambiar un estado final (validado/rechazado).",
                    )
                if not (payload.review_memo or "").strip():
                    raise HTTPException(
                        status_code=400,
                        detail="Para cambiar un estado final se requiere review_memo (motivo).",
                    )
            # Regla: desde pendiente solo se permite cerrar (validado/rechazado).
            if current_estado == "pendiente" and new_estado not in ("validado", "rechazado"):
                raise HTTPException(
                    status_code=400,
                    detail="Transición inválida: pendiente solo puede cambiar a validado o rechazado.",
                )

        success = update_axial_ai_analysis_estado(
            clients.postgres,
            project_id=project_id,
            analysis_id=int(analysis_id),
            estado=payload.estado,
            reviewed_by=user.user_id,
            review_memo=payload.review_memo,
        )
        if not success:
            raise HTTPException(status_code=404, detail="Análisis IA axial no encontrado")

        # AX-AI-05: audit trail (antes/después).
        log_project_action(
            clients.postgres,
            project=project_id,
            user_id=user.user_id,
            action="axial_ai_analysis.update",
            entity_type="axial_ai_analysis",
            entity_id=str(int(analysis_id)),
            details={
                "from_estado": current_estado,
                "to_estado": new_estado,
                "from_review_memo": current.get("review_memo"),
                "to_review_memo": effective_review_memo,
                "changed_estado": bool(current_estado != new_estado),
                "is_admin": bool(is_admin),
            },
        )

        api_logger.info(
            "axial_ai_analysis.estado_updated",
            project=project_id,
            analysis_id=int(analysis_id),
            estado=payload.estado,
            user=user.user_id,
        )

        return {"success": True, "analysis_id": int(analysis_id), "estado": payload.estado}
    finally:
        clients.close()


class AxialAiSuggestionDecisionRequest(BaseModel):
    """Aplica una decisión humana sobre una sugerencia dentro de un artefacto IA.

    Nivel 1 (ya): validar/rechazar el artefacto (`axial_ai_analyses`).
    Nivel 2 (nuevo): aplicar/cerrar una sugerencia (crea/actualiza `link_predictions` y opcionalmente Neo4j).
    """
    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="Proyecto")
    suggestion_id: int = Field(..., ge=1, description="ID 1-based de la sugerencia dentro del artefacto")
    decision: str = Field(..., description="validate_apply | reject_close")
    relation_type: Optional[str] = Field(
        None,
        description="Tipo de relación a aplicar en Neo4j (solo para validate_apply).",
    )
    memo: Optional[str] = Field(None, description="Memo opcional para guardar en link_predictions")


@app.post("/api/axial/ai-analyses/{analysis_id}/suggestions/decision")
async def api_decide_axial_ai_suggestion(
    analysis_id: int,
    payload: AxialAiSuggestionDecisionRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Cierra/aplica una sugerencia del análisis IA axial.

    - `validate_apply`: marca la sugerencia como validada y crea la relación axial en Neo4j (best-effort).
    - `reject_close`: marca la sugerencia como rechazada (para que no reaparezca en sugerencias futuras).
    """
    if payload.decision not in ("validate_apply", "reject_close"):
        raise HTTPException(status_code=400, detail="decision debe ser: validate_apply | reject_close")

    from app.neo4j_block import merge_axial_relationship
    from app.postgres_block import (
        check_project_permission,
        ensure_link_predictions_table,
        ensure_project_members_table,
        get_axial_ai_analysis_by_id,
    )
    from psycopg2.extras import Json

    clients = build_clients_or_error(settings)
    try:
        try:
            project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        # Permisos: aplicar/cerrar sugerencias modifica el ledger (write).
        if not check_project_permission(clients.postgres, project_id, user.user_id, "codificador"):
            raise HTTPException(
                status_code=403,
                detail="Acceso denegado: se requiere rol 'codificador' (o 'admin') en el proyecto.",
            )

        analysis = get_axial_ai_analysis_by_id(
            clients.postgres,
            project_id=project_id,
            analysis_id=int(analysis_id),
        )
        if not analysis:
            raise HTTPException(status_code=404, detail="Análisis IA axial no encontrado")

        # Prefer evidence-pack suggestions (incluye id + metadata), fallback a suggestions_json.
        suggestions: List[Any] = []
        evidence_json = analysis.get("evidence_json")
        if isinstance(evidence_json, dict) and isinstance(evidence_json.get("suggestions"), list):
            suggestions = list(evidence_json.get("suggestions") or [])
        if not suggestions and isinstance(analysis.get("suggestions_json"), list):
            suggestions = list(analysis.get("suggestions_json") or [])

        total = len(suggestions)
        if total <= 0:
            raise HTTPException(status_code=400, detail="El artefacto no contiene sugerencias")

        sid = int(payload.suggestion_id)
        if sid < 1 or sid > total:
            raise HTTPException(status_code=400, detail=f"suggestion_id fuera de rango (1..{total})")

        item = suggestions[sid - 1]
        if not isinstance(item, dict):
            raise HTTPException(status_code=400, detail="Sugerencia inválida en el artefacto")

        source = str(item.get("source") or item.get("source_code") or "").strip()
        target = str(item.get("target") or item.get("target_code") or "").strip()
        if not source or not target:
            raise HTTPException(status_code=400, detail="Sugerencia sin source/target")
        if source == target:
            raise HTTPException(status_code=400, detail="Sugerencia inválida: source == target")

        # En este producto, el ledger link_predictions es Código↔Código.
        source_type = str(item.get("source_type") or "Codigo").strip() or "Codigo"
        target_type = str(item.get("target_type") or "Codigo").strip() or "Codigo"
        if source_type != "Codigo" or target_type != "Codigo":
            raise HTTPException(
                status_code=400,
                detail="Solo se soporta aplicar/cerrar sugerencias Codigo↔Codigo en link_predictions",
            )

        algorithm = str(item.get("algorithm") or analysis.get("algorithm") or "common_neighbors").strip()
        if not algorithm:
            algorithm = "common_neighbors"

        # Guardrails epistemológicos (server-side): aplicar requiere evidencia positiva
        # (o override admin + memo).
        is_admin = "admin" in (user.roles or [])
        positive_items = item.get("positive")
        negative_items = item.get("negative")
        pos_count = len(positive_items) if isinstance(positive_items, list) else 0
        neg_count = len(negative_items) if isinstance(negative_items, list) else 0
        notes = item.get("notes") if isinstance(item.get("notes"), dict) else {}
        positive_missing = bool(notes.get("positive_missing")) if isinstance(notes, dict) else False
        negative_missing = bool(notes.get("negative_missing")) if isinstance(notes, dict) else False

        raw_score = item.get("score", 0.0)
        try:
            score = float(raw_score)
        except Exception:
            score = 0.0
        score = max(score, 0.0)

        # Canonicalize (A,B) as undirected pair for storage.
        a, b = (source, target) if source <= target else (target, source)

        new_estado = "validado" if payload.decision == "validate_apply" else "rechazado"

        if payload.decision == "validate_apply" and pos_count <= 0:
            memo_clean = str(payload.memo or "").strip()
            if not (is_admin and len(memo_clean) >= 20):
                raise HTTPException(
                    status_code=409,
                    detail=(
                        "No se puede 'validar y aplicar' sin evidencia positiva. "
                        "Use evidencia (evidence pack) o solicite override admin con memo >= 20 chars."
                    ),
                )

        # Resolver predicción existente (si existe) para no pisar relation_type accidentalmente.
        existing_prediction = None
        with clients.postgres.cursor() as cur:
            cur.execute(
                """
                SELECT id, estado, memo, relation_type
                FROM link_predictions
                WHERE project_id = %s AND source_code = %s AND target_code = %s AND algorithm = %s
                """,
                (project_id, a, b, algorithm),
            )
            row = cur.fetchone()
            if row:
                existing_prediction = {
                    "id": row[0],
                    "estado": row[1],
                    "memo": row[2],
                    "relation_type": row[3],
                }
        relation_type_to_store = str(
            payload.relation_type
            or (existing_prediction.get("relation_type") if isinstance(existing_prediction, dict) else None)
            or "asociado_con"
        ).strip() or "asociado_con"

        # Asegurar tablas antes de transacción.
        ensure_link_predictions_table(clients.postgres)
        ensure_project_members_table(clients.postgres)

        # Transacción PG: upsert + update estado + audit log.
        try:
            with clients.postgres.cursor() as cur:
                # Snapshot previo (si existe).
                cur.execute(
                    """
                    SELECT id, estado, memo, relation_type
                    FROM link_predictions
                    WHERE project_id = %s AND source_code = %s AND target_code = %s AND algorithm = %s
                    FOR UPDATE
                    """,
                    (project_id, a, b, algorithm),
                )
                before_row = cur.fetchone()
                before_id = before_row[0] if before_row else None
                before_estado = str(before_row[1]) if before_row and before_row[1] else None
                before_memo = before_row[2] if before_row else None
                before_relation_type = before_row[3] if before_row else None

                # Upsert base.
                cur.execute(
                    """
                    INSERT INTO link_predictions (
                        project_id, source_code, target_code, relation_type,
                        algorithm, score, rank, memo
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT (project_id, source_code, target_code, algorithm) DO UPDATE SET
                        score = EXCLUDED.score,
                        rank = EXCLUDED.rank,
                        memo = COALESCE(EXCLUDED.memo, link_predictions.memo),
                        relation_type = EXCLUDED.relation_type,
                        updated_at = NOW()
                    RETURNING id
                    """,
                    (
                        project_id,
                        a,
                        b,
                        relation_type_to_store,
                        algorithm,
                        score,
                        sid,
                        None,
                    ),
                )
                prediction_id = cur.fetchone()[0]

                # Aplicar estado/memo/relation_type (si cambian).
                changed_estado = (before_estado or "pendiente") != new_estado
                memo_clean = str(payload.memo or "").strip()
                before_memo_clean = str(before_memo or "").strip()
                changed_memo = payload.memo is not None and memo_clean != before_memo_clean
                rel_type_clean = str(payload.relation_type or "").strip()
                before_rel_clean = str(before_relation_type or "").strip()
                changed_relation_type = payload.relation_type is not None and rel_type_clean != before_rel_clean

                if changed_estado or changed_memo or changed_relation_type:
                    cur.execute(
                        """
                        UPDATE link_predictions
                        SET estado = %s,
                            relation_type = COALESCE(%s, relation_type),
                            validado_por = COALESCE(%s, validado_por),
                            validado_en = CASE WHEN %s IN ('validado', 'rechazado') THEN NOW() ELSE validado_en END,
                            memo = COALESCE(%s, memo),
                            updated_at = NOW()
                        WHERE id = %s
                        RETURNING relation_type, estado, memo
                        """,
                        (
                            new_estado,
                            payload.relation_type,
                            user.user_id,
                            new_estado,
                            payload.memo,
                            prediction_id,
                        ),
                    )
                    updated_row = cur.fetchone()
                else:
                    updated_row = (relation_type_to_store, before_estado or "pendiente", before_memo)

                rel_type = str(updated_row[0] or "asociado_con").strip() or "asociado_con"

                # Audit log (mismo TX).
                cur.execute(
                    """
                    INSERT INTO project_audit_log (project_id, user_id, action, entity_type, entity_id, details)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (
                        project_id,
                        user.user_id,
                        "axial_ai_suggestion.decision",
                        "link_prediction",
                        str(prediction_id),
                        Json(
                            {
                                "analysis_id": int(analysis_id),
                                "suggestion_id": sid,
                                "decision": payload.decision,
                                "estado": new_estado,
                                "before_estado": before_estado,
                                "changed_estado": bool(changed_estado),
                                "changed_memo": bool(changed_memo),
                                "changed_relation_type": bool(changed_relation_type),
                                "idempotent": bool(not (changed_estado or changed_memo or changed_relation_type)),
                                "analysis_epistemic_mode": analysis.get("epistemic_mode"),
                                "analysis_prompt_version": analysis.get("prompt_version"),
                                "has_evidence_pack": bool(isinstance(analysis.get("evidence_json"), dict)),
                                "evidence_positive_count": int(pos_count),
                                "evidence_negative_count": int(neg_count),
                                "evidence_positive_missing": bool(positive_missing),
                                "evidence_negative_missing": bool(negative_missing),
                                "source_code": a,
                                "target_code": b,
                                "algorithm": algorithm,
                                "score": score,
                                "relation_type": rel_type,
                            }
                        ),
                    ),
                )

            clients.postgres.commit()
        except Exception:
            clients.postgres.rollback()
            raise

        neo4j_synced = False
        neo4j_error: Optional[str] = None
        if payload.decision == "validate_apply" and clients.neo4j:
            try:
                merge_axial_relationship(
                    driver=clients.neo4j,
                    database=settings.neo4j.database,
                    project_id=project_id,
                    source_code=prediction.get("source_code") or a,
                    target_code=prediction.get("target_code") or b,
                    relation_type=rel_type,
                )
                neo4j_synced = True
            except Exception as exc:  # noqa: BLE001
                neo4j_error = str(exc)[:200]
                api_logger.warning(
                    "axial_ai_suggestion.neo4j_sync_failed",
                    project=project_id,
                    analysis_id=int(analysis_id),
                    suggestion_id=sid,
                    prediction_id=prediction_id,
                    error=neo4j_error,
                )

        # Persistir estado de sync (best-effort).
        sync_status = "skipped"
        if payload.decision == "validate_apply":
            sync_status = "success" if neo4j_synced else "failed"
        try:
            with clients.postgres.cursor() as cur:
                cur.execute(
                    """
                    UPDATE link_predictions
                    SET neo4j_sync_status = %s,
                        neo4j_sync_error = %s,
                        neo4j_synced_at = CASE WHEN %s = 'success' THEN NOW() ELSE neo4j_synced_at END,
                        updated_at = NOW()
                    WHERE id = %s
                    """,
                    (sync_status, neo4j_error, sync_status, prediction_id),
                )
            clients.postgres.commit()
        except Exception as exc:  # noqa: BLE001
            clients.postgres.rollback()
            api_logger.warning(
                "axial_ai_suggestion.sync_status_update_failed",
                project=project_id,
                analysis_id=int(analysis_id),
                prediction_id=prediction_id,
                error=str(exc)[:200],
            )

        # Registrar sync Neo4j (best-effort) en audit log separado si falla.
        if neo4j_error:
            with clients.postgres.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO project_audit_log (project_id, user_id, action, entity_type, entity_id, details)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (
                        project_id,
                        user.user_id,
                        "axial_ai_suggestion.neo4j_sync_failed",
                        "link_prediction",
                        str(prediction_id),
                        Json(
                            {
                                "analysis_id": int(analysis_id),
                                "suggestion_id": sid,
                                "neo4j_error": neo4j_error,
                            }
                        ),
                    ),
                )
            clients.postgres.commit()

        return {
            "success": True,
            "analysis_id": int(analysis_id),
            "suggestion_id": sid,
            "prediction_id": prediction_id,
            "estado": new_estado,
            "neo4j_synced": neo4j_synced,
        }
    finally:
        clients.close()


class AnalyzeDiscoveryRequest(BaseModel):
    """Request para análisis IA de resultados de Discovery."""
    model_config = ConfigDict(extra="forbid")
    project: str = Field(default="default", description="ID del proyecto")
    positive_texts: List[str] = Field(..., description="Conceptos positivos usados")
    negative_texts: List[str] = Field(default=[], description="Conceptos negativos usados")
    target_text: Optional[str] = Field(default=None, description="Texto objetivo opcional")
    fragments: List[Dict[str, Any]] = Field(..., description="Fragmentos encontrados")


@app.post("/api/discovery/analyze")
async def api_discovery_analyze(
    payload: AnalyzeDiscoveryRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Analiza resultados de Discovery con IA.
    
    Genera síntesis temática, códigos sugeridos, memo de descubrimiento
    y sugerencias de conceptos para refinar la búsqueda.
    """
    from app.clients import build_service_clients
    
    clients = build_service_clients(settings)
    try:
        # Preparar fragmentos para el prompt
        fragment_sample = payload.fragments[:8] if payload.fragments else []
        fragments_text = []
        for i, frag in enumerate(fragment_sample, 1):  # Max 8 para el prompt
            fragmento = frag.get("fragmento", "")[:400]
            archivo = frag.get("archivo", "?")
            score = frag.get("score", 0)
            fragments_text.append(f"{i}. [{archivo}] (sim: {score:.1%}) {fragmento}")
        
        positives_str = ", ".join(payload.positive_texts)
        negatives_str = ", ".join(payload.negative_texts) if payload.negative_texts else "ninguno"
        target_str = payload.target_text or "no especificado"
        
        # Sprint 22+: Prompt para JSON estructurado con etiquetas epistemológicas
        prompt = f"""Analiza los resultados de una búsqueda exploratoria (Discovery) en un proyecto de análisis cualitativo.

**Parámetros de búsqueda:**
- Conceptos positivos: {positives_str}
- Conceptos negativos: {negatives_str}
- Texto objetivo: {target_str}

**Fragmentos encontrados:**
{chr(10).join(fragments_text)}

IMPORTANTE: Responde ÚNICAMENTE con un JSON válido (sin markdown, sin ```).
La síntesis debe distinguir explícitamente el estatus epistemológico de cada afirmación.
El JSON debe tener esta estructura exacta:

{{
    "memo_sintesis": [
        {{"type": "OBSERVATION", "text": "...", "evidence_ids": [1, 3]}},
        {{"type": "INTERPRETATION", "text": "...", "evidence_ids": [1, 3]}},
        {{"type": "HYPOTHESIS", "text": "...", "evidence_ids": [2]}}
    ],
  "codigos_sugeridos": ["codigo_uno", "codigo_dos", "codigo_tres"],
  "refinamiento_busqueda": {{
    "positivos": ["término1", "término2"],
    "negativos": ["término_excluir"],
    "target": "término_focalizar"
  }}
}}

REGLAS:
1. memo_sintesis: lista de 3-6 statements. Cada statement incluye:
    - type: uno de OBSERVATION | INTERPRETATION | HYPOTHESIS | NORMATIVE_INFERENCE
    - text: una oración clara (español)
    - evidence_ids: lista de enteros referidos a la numeración de "Fragmentos encontrados" (1..8)
2. PROHIBIDO: OBSERVATION sin evidence_ids no vacíos.
2. codigos_sugeridos: 3-5 códigos en snake_case (ej: identidad_cultural_amenazada)
3. refinamiento_busqueda: sugerencias para próxima búsqueda"""

        # gpt-5.x models no soportan temperature != 1, omitir el parámetro
        response = clients.aoai.chat.completions.create(
            model=settings.azure.deployment_chat,
            messages=[
                {"role": "system", "content": "Eres un experto en análisis cualitativo. Respondes SOLO con JSON válido, sin markdown ni explicaciones adicionales."},
                {"role": "user", "content": prompt}
            ],
            max_completion_tokens=700,
        )
        
        choice = response.choices[0] if response.choices else None
        message = getattr(choice, "message", None)
        message_content = getattr(message, "content", None)
        raw_content = message_content or ""
        
        # Sprint 22: Parsear JSON estructurado
        import json
        import re
        
        parsed_synthesis = None
        try:
            # Limpiar posibles bloques de código markdown
            clean_content = re.sub(r'^```json?\s*', '', raw_content.strip())
            clean_content = re.sub(r'\s*```$', '', clean_content)
            parsed_synthesis = json.loads(clean_content)
        except json.JSONDecodeError:
            api_logger.warning("api.discovery_analyze.json_parse_failed", raw=raw_content[:200])
            # Fallback: devolver texto sin parsear
            parsed_synthesis = None

        def _normalize_memo(value: Any) -> tuple[str, List[Dict[str, Any]]]:
            allowed = {"OBSERVATION", "INTERPRETATION", "HYPOTHESIS", "NORMATIVE_INFERENCE"}
            if isinstance(value, str):
                return value.strip(), []
            if not isinstance(value, list):
                return "", []
            out: List[Dict[str, Any]] = []
            lines: List[str] = []
            for item in value:
                if not isinstance(item, dict):
                    continue
                stype = str(item.get("type") or "").strip().upper()
                text = str(item.get("text") or "").strip()
                if not text:
                    continue
                if stype not in allowed:
                    stype = "INTERPRETATION"
                evidence_ids: List[int] = []
                raw_ids = item.get("evidence_ids")
                if isinstance(raw_ids, list):
                    for v in raw_ids:
                        try:
                            evidence_ids.append(int(v))
                        except Exception:
                            continue
                if stype == "OBSERVATION" and not evidence_ids:
                    stype = "INTERPRETATION"
                entry: Dict[str, Any] = {"type": stype, "text": text}
                if evidence_ids:
                    entry["evidence_ids"] = evidence_ids
                    lines.append(f"[{stype}] {text} (evid: {', '.join(str(i) for i in evidence_ids)})")
                else:
                    lines.append(f"[{stype}] {text}")
                out.append(entry)
            return "\n".join(lines).strip(), out
        
        api_logger.info(
            "api.discovery_analyze.complete",
            positive_count=len(payload.positive_texts),
            fragment_count=len(payload.fragments),
            parsed=parsed_synthesis is not None,
        )
        
        # Sprint 22: Respuesta estructurada con fallback
        if parsed_synthesis:
            memo_text, memo_statements = _normalize_memo(parsed_synthesis.get("memo_sintesis"))
            return {
                # Compatibilidad: el campo "analysis" se mantiene como texto.
                "analysis": memo_text,
                "structured": True,
                # Nuevo: lista de statements etiquetados.
                "memo_statements": memo_statements,
                "codigos_sugeridos": parsed_synthesis.get("codigos_sugeridos", []),
                "refinamiento_busqueda": parsed_synthesis.get("refinamiento_busqueda", {}),
                "positive_texts": payload.positive_texts,
                "negative_texts": payload.negative_texts,
                "target_text": payload.target_text,
                "fragments_analyzed": len(fragment_sample),
            }
        else:
            return {
                "analysis": raw_content,
                "structured": False,
                "memo_statements": [],
                "codigos_sugeridos": [],
                "refinamiento_busqueda": {},
                "positive_texts": payload.positive_texts,
                "negative_texts": payload.negative_texts,
                "target_text": payload.target_text,
                "fragments_analyzed": len(fragment_sample),
            }
        
    except Exception as exc:
        api_logger.error("api.discovery_analyze.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


# =============================================================================
# Sprint 24: Discovery Navigation Log - Muestreo Teórico
# E3-1.2: Extended for E3 actions traceability
# =============================================================================

class LogNavigationRequest(BaseModel):
    project: str = Field(..., description="Project ID")
    positivos: List[str] = Field(..., description="Positive concepts used")
    negativos: List[str] = Field(default=[], description="Negative concepts used")
    target_text: Optional[str] = Field(default=None, description="Target text")
    fragments_count: int = Field(..., description="Number of fragments found")
    codigos_sugeridos: Optional[List[str]] = Field(default=None, description="AI suggested codes")
    refinamientos_aplicados: Optional[Dict[str, Any]] = Field(default=None, description="Applied refinements")
    ai_synthesis: Optional[str] = Field(default=None, description="AI synthesis text")
    action_taken: str = Field(default="search", description="Action: search, refine, send_codes, e3_suggest, e3_send_candidates, e3_validate, e3_reject, e3_promote")
    busqueda_origen_id: Optional[str] = Field(default=None, description="Parent search UUID")
    # E3-1.2: Additional fields for E3 traceability
    seed_fragmento_id: Optional[str] = Field(default=None, description="Seed fragment ID for E3 suggestions")
    scope_archivo: Optional[str] = Field(default=None, description="Interview file scope filter")
    top_k: Optional[int] = Field(default=None, description="Number of similar fragments requested")
    include_coded: Optional[bool] = Field(default=None, description="Whether to include already coded fragments")


@app.post("/api/discovery/log-navigation")
async def api_log_discovery_navigation(
    payload: LogNavigationRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Sprint 24: Log a discovery navigation step for theoretical sampling traceability.
    E3-1.2: Extended with E3 action types and fields.
    """
    from app.postgres_block import log_discovery_navigation
    
    clients = build_clients_or_error(settings)
    try:
        busqueda_id = log_discovery_navigation(
            clients.postgres,
            project=payload.project,
            positivos=payload.positivos,
            negativos=payload.negativos,
            target_text=payload.target_text,
            fragments_count=payload.fragments_count,
            codigos_sugeridos=payload.codigos_sugeridos,
            refinamientos_aplicados=payload.refinamientos_aplicados,
            ai_synthesis=payload.ai_synthesis,
            action_taken=payload.action_taken,
            busqueda_origen_id=payload.busqueda_origen_id,
            # E3-1.2: Pass E3 fields
            seed_fragmento_id=payload.seed_fragmento_id,
            scope_archivo=payload.scope_archivo,
            top_k=payload.top_k,
            include_coded=payload.include_coded,
        )
        
        api_logger.info(
            "api.discovery.navigation_logged",
            project=payload.project,
            action=payload.action_taken,
            busqueda_id=busqueda_id,
            scope_archivo=payload.scope_archivo,
        )
        
        return {
            "success": True,
            "busqueda_id": busqueda_id,
            "action_taken": payload.action_taken,
        }
    except Exception as exc:
        api_logger.error("api.discovery.navigation_log_error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


@app.get("/api/discovery/navigation-history")
async def api_get_discovery_navigation_history(
    project: str = Query(..., description="Project ID"),
    limit: int = Query(default=50, ge=1, le=200, description="Max entries"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Sprint 24: Get discovery navigation history for a project.
    
    Returns chronological list of searches with their refinements.
    """
    from app.postgres_block import get_discovery_navigation_history
    
    clients = build_clients_or_error(settings)
    try:
        history = get_discovery_navigation_history(
            clients.postgres,
            project=project,
            limit=limit,
        )
        
        return {
            "project": project,
            "history": history,
            "count": len(history),
        }
    except Exception as exc:
        api_logger.error("api.discovery.navigation_history_error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


# =============================================================================
# Etapa 2: Familiarización - Review fragments before coding
# =============================================================================

@app.get("/api/familiarization/fragments")
async def api_get_familiarization_fragments(
    project: str = Query(default="default"),
    file_filter: Optional[str] = Query(default=None, description="Filter by archivo name"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene fragmentos ingestados para revisión/familiarización.
    
    Retorna fragmentos con información de hablante para que el investigador
    pueda revisar las transcripciones antes de codificar.
    """
    from app.clients import build_service_clients
    from qdrant_client.models import Filter, FieldCondition, MatchValue
    from typing import Any, List, cast
    
    clients = build_service_clients(settings)
    try:
        # Build filter for project (stored as project_id in Qdrant)
        filter_conditions = [
            FieldCondition(key="project_id", match=MatchValue(value=project))
        ]
        if file_filter:
            filter_conditions.append(
                FieldCondition(key="archivo", match=MatchValue(value=file_filter))
            )
        
        points_filter = Filter(must=cast(List[Any], filter_conditions))
        
        # Fetch from Qdrant
        points, _ = clients.qdrant.scroll(
            collection_name=settings.qdrant.collection,
            scroll_filter=points_filter,
            limit=500,  # Max fragments to return
            with_payload=True,
            with_vectors=False,
        )
        
        # Process fragments
        fragments = []
        files_set = set()
        
        for point in points:
            payload = point.payload or {}
            archivo = payload.get("archivo", "unknown")
            files_set.add(archivo)
            
            fragments.append({
                "id": str(point.id),
                "text": payload.get("fragmento", ""),
                "speaker": payload.get("speaker", "interviewee"),
                "archivo": archivo,
                "fragmento_idx": payload.get("fragmento_idx", 0),
                "char_count": len(payload.get("fragmento", "")),
                "interviewee_tokens": payload.get("interviewee_tokens", 0),
                "interviewer_tokens": payload.get("interviewer_tokens", 0),
            })
        
        # Sort by archivo and index
        fragments.sort(key=lambda x: (x["archivo"], x["fragmento_idx"]))
        
        api_logger.info(
            "api.familiarization.fragments",
            project=project,
            count=len(fragments),
        )
        
        return {
            "fragments": fragments,
            "total": len(fragments),
            "files": sorted(list(files_set)),
            "project": project,
        }
        
    except Exception as exc:
        api_logger.error("api.familiarization.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


class ConfirmRelationshipRequest(BaseModel):
    source: str
    target: str
    relation_type: str = "partede"
    project: str = "default"
    algorithm: Optional[str] = "hidden_relationships"
    score: Optional[float] = 0.0
    memo: Optional[str] = None


@app.get("/api/axial/hidden-relationships")
async def api_hidden_relationships(
    project: str = Query(default="default"),
    top_k: int = Query(default=20, ge=1, le=50),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Descubre relaciones ocultas/latentes entre códigos.
    
    Combina 3 métodos de descubrimiento:
    1. Co-ocurrencia en fragmentos
    2. Categoría compartida
    3. Misma comunidad (Louvain)
    
    Retorna sugerencias con score y razón.
    """
    from app.link_prediction import discover_hidden_relationships
    from app.project_state import resolve_project
    
    clients = build_clients_or_error(settings)
    try:
        project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
        api_logger.info("api.hidden_relationships.start", project=project_id)
        suggestions = discover_hidden_relationships(clients, settings, project_id, top_k)
        
        # Agrupar por método
        by_method = {}
        for s in suggestions:
            method = s.get("method", "unknown")
            if method not in by_method:
                by_method[method] = []
            by_method[method].append(s)
        
        return {
            "suggestions": suggestions,
            "by_method": by_method,
            "total": len(suggestions),
            "project": project_id,
        }
    except Exception as exc:
        api_logger.error("api.hidden_relationships.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


@app.post("/api/axial/confirm-relationship")
async def api_confirm_relationship(
    payload: ConfirmRelationshipRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Encola una relación oculta descubierta como hipótesis en PostgreSQL (ledger).

    Importante:
    - NO crea la relación en Neo4j.
    - Neo4j solo se actualiza cuando un humano valida/aplica la relación (workflow link_predictions).
    """
    from app.link_prediction import confirm_hidden_relationship
    from app.project_state import resolve_project
    from app.postgres_block import check_project_permission
    
    clients = build_clients_or_error(settings)
    try:
        project_id = resolve_project(payload.project, allow_create=False, pg=clients.postgres)

        # Permisos: encolar hipótesis es una operación de escritura (codificador/admin).
        is_admin = "admin" in {str(r).strip().lower() for r in (user.roles or [])}
        if not is_admin and not check_project_permission(clients.postgres, project_id, user.user_id, "codificador"):
            raise HTTPException(status_code=403, detail="Permiso insuficiente para confirmar relaciones")

        api_logger.info(
            "api.confirm_relationship.start",
            source=payload.source,
            target=payload.target,
            tipo=payload.relation_type,
            project=project_id,
        )
        result = confirm_hidden_relationship(
            clients, settings,
            source=payload.source,
            target=payload.target,
            relation_type=payload.relation_type,
            project=project_id,
            algorithm=payload.algorithm,
            score=payload.score,
            memo=payload.memo,
        )
        
        if result.get("status") == "error":
            raise HTTPException(status_code=400, detail=result.get("message"))
            
        return result
    except HTTPException:
        raise
    except Exception as exc:
        api_logger.error("api.confirm_relationship.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


# =============================================================================
# MULTI-USER COLLABORATION (E3)
# =============================================================================

from app.postgres_block import (
    add_project_member,
    assign_org_users_to_project,
    get_project_members,
    get_project_db,
    get_user_project_role,
    check_project_permission,
    get_user_projects,
    log_project_action,
    get_project_audit_log,
    PROJECT_ROLES,
)


class AddMemberRequest(BaseModel):
    """Request para agregar miembro a proyecto."""
    user_id: str = Field(..., description="ID del usuario a agregar")
    role: str = Field(default="lector", description="Rol: admin, codificador, lector")


class SyncOrgMembersRequest(BaseModel):
    """Request para sincronizar miembros por organización."""
    org_id: Optional[str] = Field(default=None, description="Organización a sincronizar")
    default_role: str = Field(default="codificador", description="Rol por defecto si no se usa rol del usuario")
    use_user_role: bool = Field(default=True, description="Usar rol del usuario para el rol de proyecto")
    include_inactive: bool = Field(default=False, description="Incluir usuarios inactivos")


@app.get("/api/projects/{project_id}/members")
async def api_get_project_members(
    project_id: str,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Obtiene los miembros de un proyecto."""
    try:
        project = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        members = get_project_members(clients.postgres, project)
    finally:
        clients.close()
    
    return {"project": project, "members": members, "roles_available": list(PROJECT_ROLES)}


@app.post("/api/projects/{project_id}/members")
async def api_add_project_member(
    project_id: str,
    payload: AddMemberRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Agrega un miembro al proyecto (requiere rol admin)."""
    try:
        project = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Verificar que el usuario actual es admin
        if not check_project_permission(clients.postgres, project, user.user_id, "admin"):
            # Si no hay miembros, permitir que el primero sea admin automáticamente
            existing = get_project_members(clients.postgres, project)
            if len(existing) == 0:
                # Auto-agregar al creador como admin
                add_project_member(clients.postgres, project, user.user_id, "admin", added_by="system")
            else:
                raise HTTPException(status_code=403, detail="Solo administradores pueden agregar miembros")
        
        # Agregar el nuevo miembro
        member = add_project_member(
            clients.postgres, project, payload.user_id, payload.role, added_by=user.user_id
        )
        
        # Registrar en audit log
        log_project_action(
            clients.postgres, project, user.user_id, "add_member",
            entity_type="member", entity_id=payload.user_id,
            details={"role": payload.role}
        )
    finally:
        clients.close()
    
    return {"status": "ok", "member": member}


@app.post("/api/projects/{project_id}/members/sync-org")
async def api_sync_project_org_members(
    project_id: str,
    payload: SyncOrgMembersRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Sincroniza miembros del proyecto según organización (admin)."""
    try:
        project = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    org_id = payload.org_id or user.organization_id
    if org_id != user.organization_id:
        raise HTTPException(status_code=403, detail="No autorizado para sincronizar otra organización")

    clients = build_clients_or_error(settings)
    try:
        project_entry = get_project_db(clients.postgres, project)
        if project_entry and project_entry.get("org_id") and project_entry.get("org_id") != org_id:
            raise HTTPException(status_code=403, detail="Proyecto fuera de la organización")

        if not check_project_permission(clients.postgres, project, user.user_id, "admin"):
            if "admin" not in {r.lower() for r in (user.roles or [])}:
                raise HTTPException(status_code=403, detail="Solo administradores pueden sincronizar miembros")

        result = assign_org_users_to_project(
            clients.postgres,
            project,
            org_id,
            default_role=payload.default_role,
            use_user_role=payload.use_user_role,
            include_inactive=payload.include_inactive,
            added_by=user.user_id,
        )

        log_project_action(
            clients.postgres,
            project,
            user.user_id,
            "sync_org_members",
            entity_type="organization",
            entity_id=org_id,
            details={
                "members_assigned": result.get("members_assigned"),
                "users_total": result.get("users_total"),
                "use_user_role": payload.use_user_role,
                "default_role": payload.default_role,
                "include_inactive": payload.include_inactive,
            },
        )
    finally:
        clients.close()

    return {"status": "ok", "result": result}


@app.delete("/api/projects/{project_id}/members/{member_user_id}")
async def api_remove_project_member(
    project_id: str,
    member_user_id: str,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Elimina un miembro del proyecto (requiere rol admin)."""
    try:
        project = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Verificar permisos
        if not check_project_permission(clients.postgres, project, user.user_id, "admin"):
            raise HTTPException(status_code=403, detail="Solo administradores pueden eliminar miembros")
        
        # No permitir eliminarse a sí mismo
        if member_user_id == user.user_id:
            raise HTTPException(status_code=400, detail="No puedes eliminarte a ti mismo del proyecto")
        
        # Eliminar
        sql = "DELETE FROM project_members WHERE project_id = %s AND user_id = %s"
        with clients.postgres.cursor() as cur:
            cur.execute(sql, (project, member_user_id))
        clients.postgres.commit()
        
        # Audit log
        log_project_action(
            clients.postgres, project, user.user_id, "remove_member",
            entity_type="member", entity_id=member_user_id
        )
    finally:
        clients.close()
    
    return {"status": "ok", "removed": member_user_id}


@app.get("/api/users/me/projects")
async def api_get_my_projects(
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Obtiene los proyectos a los que el usuario tiene acceso."""
    clients = build_clients_or_error(settings)
    try:
        projects = get_user_projects(clients.postgres, user.user_id)
    finally:
        clients.close()
    
    return {"user_id": user.user_id, "projects": projects}


@app.get("/api/projects/{project_id}/audit")
async def api_get_project_audit(
    project_id: str,
    limit: int = Query(default=50, ge=1, le=200),
    action: Optional[str] = Query(default=None),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Obtiene el historial de acciones del proyecto."""
    try:
        project = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Verificar acceso al proyecto
        if not check_project_permission(clients.postgres, project, user.user_id, "lector"):
            raise HTTPException(status_code=403, detail="No tienes acceso a este proyecto")
        
        log = get_project_audit_log(clients.postgres, project, limit=limit, action_filter=action)
    finally:
        clients.close()
    
    return {"project": project, "audit_log": log, "total": len(log)}


# =============================================================================
# ENDPOINTS DE INFORMES (Sprint B)
# =============================================================================

@app.get("/api/reports/interviews")
async def api_get_interview_reports(
    project: str = Query(default="default", description="ID del proyecto"),
    limit: int = Query(default=50, ge=1, le=100),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Lista informes de análisis por entrevista.
    
    Retorna métricas de codificación para cada entrevista analizada,
    incluyendo: códigos nuevos/reutilizados, categorías, saturación.
    """
    from app.reports import get_interview_reports
    
    clients = build_clients_or_error(settings)
    try:
        reports = get_interview_reports(clients.postgres, project, limit=limit)
        return {
            "project": project,
            "reports": [r.to_dict() for r in reports],
            "total": len(reports),
        }
    finally:
        clients.close()


@app.get("/api/reports/stage4-summary")
async def api_get_stage4_summary(
    project: str = Query(default="default", description="ID del proyecto"),
    clients: PgOnlyClients = Depends(get_pg_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Genera resumen consolidado de Etapa 4 para transición a Etapa 5.
    
    Incluye:
    - Total de entrevistas, códigos, categorías
    - Score de saturación
    - Informes por entrevista resumidos
    - Candidatos a núcleo selectivo (si disponibles)
    """
    from app.reports import generate_stage4_summary
    
    summary = await run_pg_query_with_timeout(generate_stage4_summary, clients.postgres, project)
    return summary.to_dict()


@app.post("/api/reports/stage4-final")
async def api_generate_stage4_final_report(
    project: str = Query(default="default", description="ID del proyecto"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Genera informe final de Etapa 4 con análisis IA.
    
    Este endpoint genera un informe completo para la transición a Etapa 5:
    - Métricas consolidadas
    - Candidatos a núcleo selectivo (basado en densidad de relaciones)
    - Análisis IA de patrones emergentes
    - Recomendaciones para selección del núcleo
    
    El informe se puede exportar como Markdown para documentación.
    """
    from app.reports import generate_stage4_final_report
    
    clients = build_clients_or_error(settings)
    try:
        report = generate_stage4_final_report(
            pg_conn=clients.postgres,
            neo4j_driver=clients.neo4j,
            database=settings.neo4j.database,
            aoai_client=clients.aoai,
            deployment_chat=settings.azure.deployment_chat,
            project_id=project,
            org_id=str(getattr(user, "organization_id", None) or ""),
        )
        return report
    except Exception as exc:
        api_logger.error("api.stage4_final.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()


# -----------------------------------------------------------------------------
# Stage4 final report jobs (async)
# -----------------------------------------------------------------------------

def _run_stage4_final_report_task(*, task_id: str, project: str, settings: AppSettings, org_id: str) -> None:
    from app.reports import generate_stage4_final_report
    from app.postgres_block import update_report_job
    from app.blob_storage import CONTAINER_REPORTS, logical_path_to_blob_name, upload_file

    clients = build_clients_or_error(settings)
    try:
        try:
            update_report_job(
                clients.postgres,
                task_id=task_id,
                status="running",
                message="Generando informe final Etapa 4...",
                started_at=datetime.now().isoformat(),
            )
        except Exception:
            pass

        report = generate_stage4_final_report(
            pg_conn=clients.postgres,
            neo4j_driver=clients.neo4j,
            database=settings.neo4j.database,
            aoai_client=clients.aoai,
            deployment_chat=settings.azure.deployment_chat,
            project_id=project,
            org_id=org_id,
        )

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        filename = f"{timestamp}_stage4_final.json"
        logical_path = f"reports/{project}/{filename}"
        payload_bytes = json.dumps(report, ensure_ascii=False, indent=2).encode("utf-8")

        blob_url: Optional[str] = None
        try:
            blob_name = logical_path_to_blob_name(org_id=org_id, project_id=project, logical_path=logical_path)
            blob_url = upload_file(
                container=CONTAINER_REPORTS,
                blob_name=blob_name,
                data=payload_bytes,
                content_type="application/json",
            )
        except Exception as exc:
            api_logger.warning("stage4_final.blob_upload_failed", error=str(exc)[:200], task_id=task_id)

        result = {
            "project": project,
            "path": logical_path,
            "filename": filename,
            "report": report,
            "blob_url": blob_url,
            "blob_name": blob_name if blob_url else None,
        }

        update_report_job(
            clients.postgres,
            task_id=task_id,
            status="completed",
            message="Informe final Etapa 4 generado",
            result=result,
            result_path=logical_path,
            finished_at=datetime.now().isoformat(),
        )
    except Exception as exc:
        try:
            update_report_job(
                clients.postgres,
                task_id=task_id,
                status="error",
                message=str(exc),
                errors=[str(exc)],
                finished_at=datetime.now().isoformat(),
            )
        except Exception:
            pass
        api_logger.error("stage4_final.job_error", error=str(exc), task_id=task_id)
    finally:
        clients.close()


@app.post("/api/reports/stage4-final/execute")
async def api_execute_stage4_final_report_job(
    background_tasks: BackgroundTasks,
    project: str = Query(default="default", description="ID del proyecto"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Dict[str, Any]:
    try:
        resolved_project = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    task_id = f"stage4final_{resolved_project}_{uuid.uuid4().hex[:12]}"
    auth = {
        "user_id": str(user.user_id),
        "org": str(user.organization_id),
        "roles": list(user.roles or []),
    }
    # Persist job row (durable)
    try:
        from app.postgres_block import create_report_job

        clients = build_clients_or_error(settings)
        try:
            resolved_project = resolve_project(project, allow_create=False, pg=clients.postgres)
            create_report_job(
                clients.postgres,
                task_id=task_id,
                job_type="stage4_final",
                project_id=resolved_project,
                payload={"project": resolved_project},
                auth=auth,
                message="Inicializando...",
            )
        finally:
            clients.close()
    except Exception as exc:
        api_logger.warning("stage4_final.job_persist_failed", error=str(exc), task_id=task_id)

    api_logger.info("stage4_final.job_started", task_id=task_id, project=resolved_project)
    background_tasks.add_task(
        _run_stage4_final_report_task,
        task_id=task_id,
        project=resolved_project,
        settings=settings,
        org_id=str(user.organization_id),
    )
    return {"task_id": task_id, "status": "started"}


@app.get("/api/reports/stage4-final/status/{task_id}")
async def api_get_stage4_final_report_job_status(
    task_id: str,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import get_report_job

        task = get_report_job(clients.postgres, task_id)
    finally:
        clients.close()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task {task_id} not found")
    _assert_doctoral_task_access(task_id=task_id, task=task, user=user)
    return {
        "task_id": task_id,
        "status": task.get("status") or "pending",
        "project": task.get("project_id") or "default",
        "message": task.get("message"),
        "started_at": task.get("started_at") or task.get("created_at"),
        "finished_at": task.get("finished_at"),
        "errors": task.get("errors") or None,
    }


@app.get("/api/reports/stage4-final/result/{task_id}")
async def api_get_stage4_final_report_job_result(
    task_id: str,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    clients = build_clients_or_error(settings)
    try:
        from app.postgres_block import get_report_job

        task = get_report_job(clients.postgres, task_id)
    finally:
        clients.close()
    if not task:
        raise HTTPException(status_code=404, detail=f"Task {task_id} not found")
    _assert_doctoral_task_access(task_id=task_id, task=task, user=user)
    if task.get("status") != "completed":
        raise HTTPException(status_code=400, detail=f"Task not completed: {task.get('status')}")
    return {"task_id": task_id, "status": "completed", "result": task.get("result") or {}}


@app.get("/api/reports/nucleus-candidates")
async def api_get_nucleus_candidates(
    project: str = Query(default="default", description="ID del proyecto"),
    top_k: int = Query(default=5, ge=1, le=20),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Lista candidatas a categoría núcleo selectivo.
    
    Identifica las categorías con mayor densidad de relaciones y centralidad
    como candidatas para la Etapa 5.
    """
    from app.reports import identify_nucleus_candidates
    
    clients = build_clients_or_error(settings)
    try:
        candidates = identify_nucleus_candidates(
            clients.neo4j,
            settings.neo4j.database,
            project,
            top_k=top_k,
        )
        return {
            "project": project,
            "candidates": candidates,
            "total": len(candidates),
        }
    finally:
        clients.close()


@app.get("/api/reports/artifacts")
async def api_list_report_artifacts(
    project: str = Query(..., description="ID del proyecto"),
    limit: int = Query(default=50, ge=1, le=200),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista artefactos recientes (FS + DB) para la vista Reportes.

    Incluye (best-effort): GraphRAG reports, discovery/runner memos, runner post-mortems,
    reportes por entrevista (DB) y reportes doctorales.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        from app.report_artifacts import list_recent_report_artifacts

        artifacts = list_recent_report_artifacts(
            clients.postgres,
            project_id,
            org_id=str(getattr(user, "organization_id", None) or ""),
            limit=limit,
        )
        return {
            "project": project_id,
            "artifacts": artifacts,
            "count": len(artifacts),
        }
    finally:
        clients.close()


@app.get("/api/reports/artifacts/download")
async def api_download_report_artifact(
    project: str = Query(..., description="ID del proyecto"),
    path: str = Query(..., description="Ruta relativa del artefacto (ej: reports/<project>/x.md)"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Any:
    """Descarga segura de un artefacto (previene path traversal).

    Permitido únicamente bajo roots conocidos del proyecto:
    - reports/<project>/
    - reports/runner/<project>/
    - logs/runner_reports/<project>/
    - logs/runner_checkpoints/<project>/
    """
    try:
        project_id = resolve_project(project, allow_create=False, pg=clients.postgres)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    if not path or not path.strip():
        raise HTTPException(status_code=400, detail="Missing required query param: path")

    rel_norm = path.replace("\\", "/").lstrip("/")
    target_rel = Path(rel_norm)

    # Block absolute paths / Windows drive traversal / dot-dot traversal.
    if target_rel.is_absolute() or ".." in target_rel.parts or ":" in (target_rel.parts[0] if target_rel.parts else ""):
        raise HTTPException(status_code=400, detail="Absolute paths are not allowed")

    # Validate logical roots (independent of filesystem existence).
    allowed_prefixes = [
        f"reports/{project_id}/",
        f"reports/runner/{project_id}/",
        f"logs/runner_reports/{project_id}/",
        f"logs/runner_checkpoints/{project_id}/",
        f"notes/{project_id}/",  # allow notes via this endpoint too (optional)
    ]
    if not any(rel_norm.startswith(p) for p in allowed_prefixes):
        raise HTTPException(status_code=400, detail="Invalid artifact path")

    # Prefer tenant-scoped Blob artifacts (cloud mode).
    try:
        from app.blob_storage import CONTAINER_REPORTS, download_file, logical_path_to_blob_name

        blob_name = logical_path_to_blob_name(
            org_id=str(getattr(user, "organization_id", None) or ""),
            project_id=project_id,
            logical_path=rel_norm,
        )
        data = download_file(CONTAINER_REPORTS, blob_name)
        suffix = target_rel.suffix.lower()
        if suffix in {".md", ".markdown"}:
            media_type = "text/markdown"
        elif suffix == ".json":
            media_type = "application/json"
        elif suffix == ".csv":
            media_type = "text/csv"
        else:
            media_type = "application/octet-stream"
        headers = {"Content-Disposition": f'attachment; filename="{target_rel.name}"'}
        return Response(content=data, media_type=media_type, headers=headers)
    except Exception as exc:
        api_logger.debug("reports.artifact.blob_fallback", error=str(exc)[:200], project=project_id)

    # Backward-compat: try legacy job blob_url by result_path (if any).
    try:
        from app.blob_storage import download_by_url

        with clients.postgres.cursor() as cur:
            cur.execute(
                """
                SELECT result->>'blob_url'
                  FROM report_jobs
                 WHERE project_id = %s
                   AND replace(COALESCE(result_path,''),'\\\\','/') = %s
                 ORDER BY updated_at DESC
                 LIMIT 1
                """,
                (project_id, rel_norm),
            )
            row = cur.fetchone()
        blob_url = (row[0] if row else None)
        if isinstance(blob_url, str) and blob_url.strip():
            data = download_by_url(blob_url)
            suffix = target_rel.suffix.lower()
            if suffix in {".md", ".markdown"}:
                media_type = "text/markdown"
            elif suffix == ".json":
                media_type = "application/json"
            elif suffix == ".csv":
                media_type = "text/csv"
            else:
                media_type = "application/octet-stream"
            headers = {"Content-Disposition": f'attachment; filename="{target_rel.name}"'}
            return Response(content=data, media_type=media_type, headers=headers)
    except Exception:
        pass

    # Final fallback: legacy local filesystem.
    allowed_roots = [
        (Path("reports") / project_id).resolve(),
        (Path("reports") / "runner" / project_id).resolve(),
        (Path("logs") / "runner_reports" / project_id).resolve(),
        (Path("logs") / "runner_checkpoints" / project_id).resolve(),
        (Path("notes") / project_id).resolve(),
    ]
    target = target_rel.resolve()

    def _is_under_allowed_root(candidate: Path) -> bool:
        for root in allowed_roots:
            if root == candidate or root in candidate.parents:
                return True
        return False

    if not _is_under_allowed_root(target):
        raise HTTPException(status_code=400, detail="Invalid artifact path")

    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail="Artifact not found")

    suffix = target.suffix.lower()
    if suffix in {".md", ".markdown"}:
        media_type = "text/markdown"
    elif suffix == ".json":
        media_type = "application/json"
    elif suffix == ".csv":
        media_type = "text/csv"
    else:
        media_type = "application/octet-stream"

    return FileResponse(path=str(target), media_type=media_type, filename=target.name)


@app.post("/api/reports/product/generate")
async def api_generate_product_artifacts(
    project: str = Query(..., description="ID del proyecto"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Genera artefactos 'producto' de alto nivel.

    Archivos generados bajo reports/<project>/:
    - executive_summary.md
    - top_10_insights.json
    - open_questions.md
    - product_manifest.json

    Nota: Los artefactos también aparecerán en /api/reports/artifacts.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        from app.product_artifacts import generate_and_write_product_artifacts
        try:
            return generate_and_write_product_artifacts(
                clients,
                settings,
                project_id,
                org_id=str(getattr(user, "organization_id", None) or ""),
                changed_by=getattr(user, "user_id", None),
            )
        except Exception as exc:
            # Storage failures should not be reported as "invalid payload".
            msg = str(exc)
            if "AZURE_STORAGE_CONNECTION_STRING" in msg or "Blob Storage" in msg or "azure-storage-blob" in msg:
                raise HTTPException(status_code=503, detail="Blob Storage no disponible") from exc

            # If the new product contracts fail hard validation, surface a user-friendly message.
            try:
                from pydantic import ValidationError

                if isinstance(exc, (ValidationError, ValueError)):
                    raise HTTPException(
                        status_code=400,
                        detail=f"Invalid product artifact payload: {str(exc)}",
                    ) from exc
            except Exception:
                # If pydantic isn't available for some reason, fall through.
                pass
            raise
    finally:
        clients.close()


@app.get("/api/reports/product/latest")
async def api_get_latest_product_artifacts(
    project: str = Query(..., description="ID del proyecto"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Devuelve contenido de los artefactos producto si existen (best-effort)."""
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    org_id = str(getattr(user, "organization_id", None) or "")
    exec_logical = f"reports/{project_id}/executive_summary.md"
    top_logical = f"reports/{project_id}/top_10_insights.json"
    open_logical = f"reports/{project_id}/open_questions.md"
    manifest_logical = f"reports/{project_id}/product_manifest.json"

    def _read_text(logical_path: str) -> Optional[str]:
        # Blob first (cloud mode)
        try:
            from app.blob_storage import CONTAINER_REPORTS, download_file, logical_path_to_blob_name

            blob_name = logical_path_to_blob_name(org_id=org_id, project_id=project_id, logical_path=logical_path)
            data = download_file(CONTAINER_REPORTS, blob_name)
            return data.decode("utf-8", errors="ignore")
        except Exception:
            pass

        # Legacy local fallback
        try:
            local_path = Path(logical_path)
            if not local_path.exists():
                return None
            return local_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return None

    def _read_json(logical_path: str) -> Optional[Any]:
        raw = _read_text(logical_path)
        if raw is None:
            return None
        try:
            return json.loads(raw)
        except Exception:
            return None

    return {
        "project": project_id,
        "executive_summary": _read_text(exec_logical),
        "top_10_insights": _read_json(top_logical),
        "open_questions": _read_text(open_logical),
        "manifest": _read_json(manifest_logical),
        "paths": {
            "executive_summary": exec_logical,
            "top_10_insights": top_logical,
            "open_questions": open_logical,
            "manifest": manifest_logical,
        },
    }


@app.get("/api/reports/jobs")
async def api_list_report_jobs(
    project: str = Query(..., description="ID del proyecto"),
    limit: int = Query(default=50, ge=1, le=200),
    offset: int = Query(default=0, ge=0, le=100000),
    status: Optional[str] = Query(default=None, description="Opcional: filtrar por status (pending|running|completed|error)"),
    job_type: Optional[str] = Query(default=None, description="Opcional: filtrar por tipo de job"),
    task_id: Optional[str] = Query(default=None, description="Opcional: filtrar por task_id exacto"),
    task_id_prefix: Optional[str] = Query(default=None, description="Opcional: filtrar por task_id (prefijo)"),
    q: Optional[str] = Query(default=None, description="Opcional: búsqueda de texto en message"),
    user_id: Optional[str] = Query(default=None, description="Opcional: filtrar por user_id (solo admin o mismo usuario)"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Dict[str, Any]:
    """List async report jobs (durable) for Reportes v2 history."""
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    # Analysts can list project jobs, but cannot explicitly filter to other users.
    if user_id and str(user_id) != str(user.user_id) and not _user_is_admin(user):
        raise HTTPException(status_code=403, detail="Forbidden: cannot query other users")

    svc = build_clients_or_error(settings)
    try:
        from app.postgres_block import list_report_jobs

        jobs = list_report_jobs(
            svc.postgres,
            project_id=project_id,
            limit=int(limit),
            offset=int(offset),
            user_id=str(user_id) if user_id else None,
            status=str(status) if status else None,
            job_type=str(job_type) if job_type else None,
            task_id=str(task_id) if task_id else None,
            task_id_prefix=str(task_id_prefix) if task_id_prefix else None,
            q=str(q) if q else None,
        )
        next_offset = int(offset) + len(jobs)
        has_more = len(jobs) == int(limit)
        return {
            "project": project_id,
            "jobs": jobs,
            "count": len(jobs),
            "limit": int(limit),
            "offset": int(offset),
            "next_offset": next_offset,
            "has_more": has_more,
            "filters": {
                "status": status,
                "job_type": job_type,
                "user_id": user_id,
                "task_id": task_id,
                "task_id_prefix": task_id_prefix,
                "q": q,
            },
        }
    finally:
        svc.close()


@app.get("/api/reports/blob/download")
async def api_download_report_blob_by_url(
    url: str = Query(..., description="Full Azure Blob URL (used to locate container/blob)."),
    filename: Optional[str] = Query(default=None, description="Optional filename override for Content-Disposition"),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Response:
    """Proxy-download a blob by its URL.

    Uses the server-side storage connection string, so it works for private containers.
    This is intended for Reportes v2 job history downloads when only blob_url is known.

    Strict multi-tenant: only allows blob URLs that are registered in `report_jobs`
    for a project the caller can access.
    """
    try:
        from urllib.parse import urlparse

        from app.blob_storage import download_by_url

        # Multi-tenant guard: URL must belong to an existing job the user can access.
        try:
            with clients.postgres.cursor() as cur:
                cur.execute(
                    """
                    SELECT project_id
                      FROM report_jobs
                     WHERE result->>'blob_url' = %s
                     ORDER BY updated_at DESC
                     LIMIT 1
                    """,
                    (url,),
                )
                row = cur.fetchone()
        except Exception:
            row = None

        if not row or not row[0]:
            raise HTTPException(status_code=404, detail="Unknown blob_url (not registered)")

        job_project_id = str(row[0])
        try:
            resolve_project(job_project_id, allow_create=False, pg=clients.postgres)
        except Exception:
            raise HTTPException(status_code=403, detail="Forbidden") from None

        # Optional extra hardening: only allow downloads from the reports container.
        parsed = urlparse(url)
        container = (parsed.path or "").lstrip("/").split("/", 1)[0]
        if container and container != "reports":
            raise HTTPException(status_code=400, detail="Only 'reports' container is allowed for this endpoint")

        data = download_by_url(url)

        inferred = Path(urlparse(url).path or "").name
        name = filename or inferred or "report_blob.bin"

        suffix = Path(name).suffix.lower()
        if suffix in {".md", ".markdown"}:
            media_type = "text/markdown"
        elif suffix == ".json":
            media_type = "application/json"
        elif suffix == ".csv":
            media_type = "text/csv"
        else:
            media_type = "application/octet-stream"

        headers = {"Content-Disposition": f'attachment; filename="{name}"'}
        return Response(content=data, media_type=media_type, headers=headers)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        api_logger.warning("reports.blob_download_failed", error=str(exc)[:200])
        raise HTTPException(status_code=502, detail="Failed to download blob") from exc


# ---------- Analytics ----------


class AnalyticsEvent(BaseModel):
    """Evento de analytics de frontend."""
    event_type: str = Field(..., pattern="^(click|navigation|action|error)$")
    element_id: str
    timestamp: str
    metadata: Optional[Dict[str, Any]] = None
    page: Optional[str] = None
    session_id: Optional[str] = None


class AnalyticsTrackRequest(BaseModel):
    """Batch de eventos de analytics."""
    model_config = ConfigDict(extra="forbid")
    events: List[AnalyticsEvent]


def _ensure_analytics_table(pg_conn) -> None:
    """Crea tabla de analytics si no existe."""
    with pg_conn.cursor() as cur:
        cur.execute("""
            CREATE TABLE IF NOT EXISTS ui_analytics (
                id SERIAL PRIMARY KEY,
                event_type VARCHAR(50) NOT NULL,
                element_id VARCHAR(255) NOT NULL,
                event_timestamp TIMESTAMPTZ NOT NULL,
                session_id VARCHAR(100),
                page VARCHAR(255),
                metadata JSONB,
                created_at TIMESTAMPTZ DEFAULT NOW()
            )
        """)
        # Index for querying by time and event type
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_analytics_event_type 
            ON ui_analytics (event_type, event_timestamp DESC)
        """)
        cur.execute("""
            CREATE INDEX IF NOT EXISTS idx_analytics_session 
            ON ui_analytics (session_id, event_timestamp DESC)
        """)
    pg_conn.commit()


@app.post("/api/analytics/track")
async def api_analytics_track(
    payload: AnalyticsTrackRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Recibe batch de eventos de analytics desde el frontend.
    
    Los eventos se almacenan en PostgreSQL para análisis posterior.
    """
    if not payload.events:
        return {"tracked": 0}
    
    clients = build_clients_or_error(settings)
    try:
        _ensure_analytics_table(clients.postgres)
        
        with clients.postgres.cursor() as cur:
            for event in payload.events:
                cur.execute(
                    """
                    INSERT INTO ui_analytics 
                    (event_type, element_id, event_timestamp, session_id, page, metadata)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (
                        event.event_type,
                        event.element_id,
                        event.timestamp,
                        event.session_id,
                        event.page,
                        json.dumps(event.metadata) if event.metadata else None,
                    )
                )
        clients.postgres.commit()
        
        api_logger.info(
            "analytics.track",
            events=len(payload.events),
            session_id=payload.events[0].session_id if payload.events else None,
        )
        
        return {"tracked": len(payload.events)}
    except Exception as exc:
        api_logger.warning("analytics.track.error", error=str(exc))
        # Don't fail the request - analytics should be fire-and-forget
        return {"tracked": 0, "error": "failed_to_persist"}
    finally:
        clients.close()


@app.get("/api/analytics/summary")
async def api_analytics_summary(
    project: Optional[str] = None,
    days: int = Query(7, ge=1, le=90),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Resumen de analytics de los últimos N días.
    """
    clients = build_clients_or_error(settings)
    try:
        _ensure_analytics_table(clients.postgres)
        
        with clients.postgres.cursor() as cur:
            # Event counts by type
            cur.execute("""
                SELECT event_type, COUNT(*) as count
                FROM ui_analytics
                WHERE event_timestamp > NOW() - INTERVAL '%s days'
                GROUP BY event_type
                ORDER BY count DESC
            """, (days,))
            by_type = [{"type": row[0], "count": row[1]} for row in cur.fetchall()]
            
            # Top clicked elements
            cur.execute("""
                SELECT element_id, COUNT(*) as count
                FROM ui_analytics
                WHERE event_type = 'click'
                  AND event_timestamp > NOW() - INTERVAL '%s days'
                GROUP BY element_id
                ORDER BY count DESC
                LIMIT 20
            """, (days,))
            top_clicks = [{"element": row[0], "count": row[1]} for row in cur.fetchall()]
            
            # Unique sessions
            cur.execute("""
                SELECT COUNT(DISTINCT session_id)
                FROM ui_analytics
                WHERE event_timestamp > NOW() - INTERVAL '%s days'
            """, (days,))
            session_row = cur.fetchone()
            unique_sessions = (session_row[0] if session_row else 0) or 0
        
        return {
            "period_days": days,
            "by_event_type": by_type,
            "top_clicks": top_clicks,
            "unique_sessions": unique_sessions,
        }
    finally:
        clients.close()


# NOTE: Candidate code endpoints are defined above (lines ~2492-2622)
# Removed duplicate definitions that were causing argument mismatch errors




# =============================================================================
# GRAPHRAG METRICS - Sprint 15 Anti-Alucinaciones
# =============================================================================

@app.get("/api/graphrag/metrics")
async def api_get_graphrag_metrics(
    project: str = Query(..., description="ID del proyecto"),
    days: int = Query(30, ge=1, le=365, description="Días a analizar"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene métricas de calidad de respuestas GraphRAG.
    
    Incluye:
    - Tasa de respuestas grounded vs rechazadas
    - Score promedio de evidencia
    - Distribución de niveles de confianza
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        from app.graphrag_metrics import ensure_metrics_table, get_metrics_summary, get_recent_rejections
        
        ensure_metrics_table(clients.postgres)
        
        summary = get_metrics_summary(clients.postgres, project_id, days=days)
        rejections = get_recent_rejections(clients.postgres, project_id, limit=10)
        
        return {
            "project": project_id,
            "period_days": days,
            "summary": summary,
            "recent_rejections": rejections,
        }
    finally:
        clients.close()


@app.get("/api/codes/similar")
async def api_get_similar_codes(
    codigo: str = Query(..., description="Código a buscar similares"),
    project: str = Query(..., description="ID del proyecto"),
    top_k: int = Query(5, ge=1, le=20, description="Máximo de resultados"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Busca códigos semánticamente similares (posibles sinónimos o duplicados).
    
    Usa distancia de Levenshtein para encontrar códigos con nombres parecidos.
    Útil para sugerir fusiones durante la validación.
    """
    from difflib import SequenceMatcher
    
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        # Obtener todos los códigos únicos del proyecto
        with clients.postgres.cursor() as cur:
            cur.execute("""
                SELECT DISTINCT codigo, COUNT(*) as occurrences
                FROM codigos_candidatos
                WHERE project_id = %s AND estado IN ('validado', 'pendiente')
                GROUP BY codigo
                ORDER BY occurrences DESC
            """, (project_id,))
            rows = cur.fetchall()
        
        codigo_lower = codigo.lower().strip()
        similar_codes = []
        
        for row in rows:
            other_codigo = row[0]
            occurrences = row[1]
            
            if other_codigo.lower().strip() == codigo_lower:
                continue  # Skip self
            
            # Calcular similitud
            score = SequenceMatcher(
                None, 
                codigo_lower, 
                other_codigo.lower().strip()
            ).ratio()
            
            if score >= 0.5:  # Umbral mínimo de similitud
                similar_codes.append({
                    "codigo": other_codigo,
                    "score": round(score, 3),
                    "occurrences": occurrences,
                })
        
        # Ordenar por score descendente y limitar
        similar_codes.sort(key=lambda x: x["score"], reverse=True)
        similar_codes = similar_codes[:top_k]
        
        return {
            "codigo": codigo,
            "similar_codes": similar_codes,
            "count": len(similar_codes),
        }
    finally:
        clients.close()


# =============================================================================
# CONFIGURACIÓN DE PROYECTO
# =============================================================================

@app.get("/api/projects/{project_id}/config")
async def api_get_project_config(
    project_id: str,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Obtiene la configuración de un proyecto.
    
    Incluye:
    - discovery_threshold: Umbral de similitud para Discovery (0.0-1.0)
    - analysis_temperature: Temperatura LLM para análisis
    - analysis_max_tokens: Max tokens para respuesta LLM
    """
    try:
        resolved = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    
    config = get_project_config(clients.postgres, resolved)
    return {
        "project_id": resolved,
        "config": config,
    }


class ProjectConfigUpdate(BaseModel):
    """Request para actualizar configuración de proyecto."""
    model_config = ConfigDict(extra="forbid")
    
    discovery_threshold: Optional[float] = Field(
        None, ge=0.0, le=1.0, description="Umbral de Discovery (0.0-1.0)"
    )
    analysis_temperature: Optional[float] = Field(
        None, ge=0.0, le=2.0, description="Temperatura LLM para análisis"
    )
    analysis_max_tokens: Optional[int] = Field(
        None, ge=100, le=8000, description="Max tokens para respuesta LLM"
    )


class ProjectDetailsUpdate(BaseModel):
    """Request para actualizar nombre/descripcion de proyecto."""
    model_config = ConfigDict(extra="forbid")

    name: Optional[str] = Field(None, min_length=1, max_length=120)
    description: Optional[str] = Field(None, max_length=500)


@app.put("/api/projects/{project_id}/config")
async def api_update_project_config(
    project_id: str,
    payload: ProjectConfigUpdate,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Actualiza la configuración de un proyecto.
    
    Solo actualiza los campos proporcionados en el payload.
    """
    try:
        resolved = resolve_project(project_id, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    
    # Construir dict con solo los campos proporcionados
    updates = {}
    if payload.discovery_threshold is not None:
        updates["discovery_threshold"] = payload.discovery_threshold
    if payload.analysis_temperature is not None:
        updates["analysis_temperature"] = payload.analysis_temperature
    if payload.analysis_max_tokens is not None:
        updates["analysis_max_tokens"] = payload.analysis_max_tokens
    
    if not updates:
        raise HTTPException(status_code=400, detail="No se proporcionaron campos para actualizar")
    
    try:
        config = update_project_config(clients.postgres, resolved, updates)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    
    api_logger.info(
        "project.config.updated",
        project_id=resolved,
        user=user.user_id,
        updates=list(updates.keys()),
    )
    
    return {
        "project_id": resolved,
        "config": config,
        "updated": list(updates.keys()),
    }


@app.patch("/api/projects/{project_id}")
async def api_update_project_details(
    project_id: str,
    payload: ProjectDetailsUpdate,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Actualiza nombre/descripcion de proyecto (requiere rol codificador/admin)."""
    try:
        resolved = resolve_project(project_id, allow_create=False, pg=clients.postgres)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc

    if payload.name is None and payload.description is None:
        raise HTTPException(status_code=400, detail="No se proporcionaron campos para actualizar")

    is_admin = "admin" in {str(r).strip().lower() for r in (user.roles or [])}
    if not is_admin and not check_project_permission(clients.postgres, resolved, user.user_id, "codificador"):
        raise HTTPException(status_code=403, detail="Permiso insuficiente para editar el proyecto")

    try:
        updated = update_project_details(
            clients.postgres,
            resolved,
            name=payload.name,
            description=payload.description,
        )
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc

    api_logger.info(
        "project.updated",
        project_id=resolved,
        user=user.user_id,
        updates=[k for k in ("name", "description") if getattr(payload, k) is not None],
    )

    return updated


class ProjectEpistemicModeUpdate(BaseModel):
    """Request para actualizar el modo epistémico de un proyecto (endpoint dedicado)."""
    model_config = ConfigDict(extra="forbid")

    epistemic_mode: str = Field(..., description="constructivist | post_positivist")


@app.put("/api/projects/{project_id}/epistemic-mode")
async def api_update_project_epistemic_mode(
    project_id: str,
    payload: ProjectEpistemicModeUpdate,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Actualiza `epistemic_mode` con guardrails de estado (no es PATCH genérico)."""
    from app.postgres_block import (
        check_project_permission,
        get_project_epistemic_mode,
        log_project_action,
        set_project_epistemic_mode,
    )
    from app.project_state import resolve_project
    from app.settings import EpistemicMode

    try:
        resolved = resolve_project(project_id, allow_create=False, pg=clients.postgres)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc

    requested_raw = str(payload.epistemic_mode or "").strip().lower()
    if requested_raw not in ("constructivist", "post_positivist"):
        raise HTTPException(
            status_code=400,
            detail="epistemic_mode inválido: use constructivist | post_positivist",
        )

    # Permisos: alinea con edición de proyecto (codificador/admin).
    is_admin = "admin" in {str(r).strip().lower() for r in (user.roles or [])}
    if not is_admin and not check_project_permission(clients.postgres, resolved, user.user_id, "codificador"):
        raise HTTPException(status_code=403, detail="Permiso insuficiente para cambiar epistemic_mode")

    current = get_project_epistemic_mode(clients.postgres, resolved)
    desired = EpistemicMode.from_string(requested_raw)
    if desired.value == current.value:
        return {
            "project_id": resolved,
            "epistemic_mode": current.value,
            "changed": False,
            "message": "epistemic_mode ya estaba configurado",
        }

    ok, message = set_project_epistemic_mode(clients.postgres, resolved, desired)
    if not ok:
        # Guardrails / lock de modo: conflicto de estado.
        raise HTTPException(status_code=409, detail=message)

    try:
        log_project_action(
            clients.postgres,
            project=resolved,
            user_id=user.user_id,
            action="project.epistemic_mode.updated",
            entity_type="project",
            entity_id=resolved,
            details={"from": current.value, "to": desired.value},
        )
    except Exception:
        # Best-effort audit trail: no bloquear el cambio ya aplicado.
        pass

    api_logger.info(
        "project.epistemic_mode.updated",
        project_id=resolved,
        user=user.user_id,
        from_mode=current.value,
        to_mode=desired.value,
    )
    return {
        "project_id": resolved,
        "epistemic_mode": desired.value,
        "changed": True,
        "message": message,
    }


# =============================================================================
# SPRINT 17: SUGERENCIA DE ACCIÓN CON IA
# =============================================================================

class FragmentSelection(BaseModel):
    """Fragmento seleccionado para batch."""
    fragmento_id: str
    archivo: str
    cita: str
    score: float = 0.0


class BatchCandidateRequest(BaseModel):
    """Request para enviar múltiples fragmentos a bandeja de candidatos."""
    project: str
    codigo: str
    memo: Optional[str] = None
    fragments: List[FragmentSelection]


@app.post("/api/codes/candidates/batch")
async def api_submit_candidates_batch(
    payload: BatchCandidateRequest,
    user: User = Depends(require_auth),
    settings: AppSettings = Depends(get_settings),
) -> Dict[str, Any]:
    """
    Envía múltiples fragmentos a la bandeja de candidatos.
    
    Sprint 17 - E1: Sugerencia de Acción con selección múltiple.
    
    Todos los fragmentos se agrupan bajo el mismo código y memo.
    Los códigos pasan a validación antes de ser definitivos.
    """
    if not payload.fragments:
        raise HTTPException(status_code=400, detail="No hay fragmentos para procesar")
    
    clients = build_clients_or_error(settings)
    try:
        ensure_candidate_codes_table(clients.postgres)
        
        candidates = [
            {
                "project_id": payload.project,
                "codigo": payload.codigo,
                "cita": frag.cita[:500] if frag.cita else "",
                "fragmento_id": frag.fragmento_id,
                "archivo": frag.archivo,
                "fuente_origen": "semantic_suggestion",
                "fuente_detalle": "Sugerencia de Acción IA",
                "memo": payload.memo,
                "score_confianza": frag.score,
            }
            for frag in payload.fragments
        ]
        
        count = insert_candidate_codes(clients.postgres, candidates)
        
        api_logger.info(
            "candidates.batch.submitted",
            project=payload.project,
            codigo=payload.codigo,
            count=count,
            user=user.user_id,
        )
        
        return {
            "submitted": count,
            "codigo": payload.codigo,
            "project": payload.project,
            "fragments_count": len(payload.fragments),
        }
    finally:
        clients.close()


class SuggestCodeRequest(BaseModel):
    """Request para sugerencia de código IA."""
    project: str
    fragments: List[Dict[str, Any]]  # Lista de fragmentos con texto
    llm_model: Optional[str] = None


@app.post("/api/coding/suggest-code")
async def api_suggest_code_from_fragments(
    payload: SuggestCodeRequest,
    user: User = Depends(require_auth),
    settings: AppSettings = Depends(get_settings),
) -> Dict[str, Any]:
    """
    Sugiere nombre de código y memo basado en fragmentos.
    
    Sprint 17 - E2: IA que propone código y justificación.
    
    Returns:
        - suggested_code: Nombre propuesto en snake_case
        - memo: Justificación del agrupamiento
        - confidence: alta/media/baja
    """
    if not payload.fragments:
        raise HTTPException(status_code=400, detail="No hay fragmentos para analizar")
    
    clients = build_clients_or_error(settings)
    try:
        from app.coding import suggest_code_from_fragments, list_open_codes
        
        # Obtener códigos existentes para evitar duplicados
        existing_codes = [c["codigo"] for c in list_open_codes(clients, payload.project, limit=100)]
        
        result = suggest_code_from_fragments(
            clients=clients,
            settings=settings,
            fragments=payload.fragments,
            existing_codes=existing_codes,
            llm_model=payload.llm_model,
            project=payload.project,
        )
        
        api_logger.info(
            "coding.suggest_code.completed",
            project=payload.project,
            suggested_code=result.get("suggested_code"),
            confidence=result.get("confidence"),
            user=user.user_id,
        )
        
        return result
    finally:
        clients.close()


# =============================================================================
# Sprint 27: Analysis Insights - Muestreo Teórico Automatizado
# =============================================================================

class InsightsListRequest(BaseModel):
    project: str
    status: Optional[str] = None  # 'pending', 'executed', 'dismissed'
    source_type: Optional[str] = None  # 'discovery', 'coding', 'link_prediction', 'report'
    limit: int = 20


@app.post("/api/insights/list")
async def api_list_insights(
    payload: InsightsListRequest,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Lista insights de un proyecto con filtros opcionales."""
    from app.postgres_block import list_insights, count_insights_by_status
    
    insights = list_insights(
        clients.postgres,
        payload.project,
        status=payload.status,
        source_type=payload.source_type,
        limit=payload.limit,
    )
    
    counts = count_insights_by_status(clients.postgres, payload.project)
    
    return {
        "insights": insights,
        "counts": counts,
        "total": len(insights),
    }


class InsightActionRequest(BaseModel):
    insight_id: int
    project: str


@app.post("/api/insights/dismiss")
async def api_dismiss_insight(
    payload: InsightActionRequest,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Descarta un insight (no ejecutar)."""
    from app.postgres_block import update_insight_status
    
    success = update_insight_status(
        clients.postgres,
        payload.insight_id,
        status="dismissed",
    )
    
    api_logger.info(
        "insights.dismissed",
        insight_id=payload.insight_id,
        project=payload.project,
        user=user.user_id,
    )
    
    return {"success": success, "insight_id": payload.insight_id}


@app.post("/api/insights/execute")
async def api_execute_insight(
    payload: InsightActionRequest,
    settings: AppSettings = Depends(get_settings),
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Ejecuta la sugerencia de un insight.
    
    Dependiendo del tipo de sugerencia, ejecuta:
    - search: Búsqueda Discovery
    - analyze: Análisis de codificación
    - link_prediction: Predicción de enlaces
    """
    from app.postgres_block import get_insight, update_insight_status
    
    insight = get_insight(clients.postgres, payload.insight_id)
    if not insight:
        raise HTTPException(status_code=404, detail="Insight no encontrado")
    
    suggested_query = insight.get("suggested_query", {})
    action = suggested_query.get("action", "search")
    
    result = {
        "insight_id": payload.insight_id,
        "action": action,
        "executed": False,
        "message": "",
    }
    
    try:
        if action == "search" and suggested_query.get("positivos"):
            # Ejecutar búsqueda semántica usando embeddings
            from app.qdrant_block import search_similar
            
            positivos = suggested_query.get("positivos", [])
            query_text = " ".join(positivos)
            
            api_logger.info(
                "insights.execute.search.start",
                insight_id=payload.insight_id,
                query_text=query_text,
                project=payload.project,
            )
            
            try:
                # Generar embedding del término de búsqueda
                embedding_response = clients.aoai.embeddings.create(
                    model=settings.azure.deployment_embed,
                    input=query_text,
                )
                vector = embedding_response.data[0].embedding
                
                api_logger.info(
                    "insights.execute.embedding.created",
                    insight_id=payload.insight_id,
                    vector_dim=len(vector),
                )
                
                # Usar el nombre de colección correcto desde settings
                collection_name = settings.qdrant.collection
                
                api_logger.info(
                    "insights.execute.search.query",
                    insight_id=payload.insight_id,
                    collection=collection_name,
                    project=payload.project,
                )
                
                search_results = search_similar(
                    clients.qdrant,
                    collection=collection_name,
                    vector=vector,
                    limit=10,
                    score_threshold=0.4,  # Umbral más permisivo
                    project_id=payload.project,
                )
                
                result["executed"] = True
                result["message"] = f"Búsqueda ejecutada: {len(search_results)} fragmentos encontrados para '{query_text}'"
                result["fragments_count"] = len(search_results)
                
                # Formatear resultados para preview
                result["fragments"] = [
                    {
                        "id": str(r.id), 
                        "score": round(r.score, 3), 
                        "fragmento": (r.payload.get("fragmento", "") or "")[:200] + "..."
                    }
                    for r in search_results[:5]
                ]
                
                api_logger.info(
                    "insights.execute.search.complete",
                    insight_id=payload.insight_id,
                    fragments_found=len(search_results),
                    top_score=round(search_results[0].score, 3) if search_results else 0,
                )
                
            except Exception as search_err:
                result["message"] = f"Error en búsqueda: {str(search_err)}"
                api_logger.error(
                    "insights.execute.search.error",
                    insight_id=payload.insight_id,
                    error=str(search_err),
                    error_type=type(search_err).__name__,
                )
            
        elif action == "compare":
            # Comparar códigos para posible merge
            codes = suggested_query.get("codes", [])
            result["executed"] = True
            result["message"] = f"Comparar códigos: {codes}. Revisar manualmente para confirmar fusión."
            result["codes"] = codes
            
            api_logger.info(
                "insights.execute.compare",
                insight_id=payload.insight_id,
                codes=codes,
            )
            
        else:
            result["message"] = f"Acción '{action}' requiere implementación manual."
            api_logger.warning(
                "insights.execute.unsupported_action",
                insight_id=payload.insight_id,
                action=action,
            )
        
        # Actualizar status del insight
        update_insight_status(
            clients.postgres,
            payload.insight_id,
            status="executed",
            execution_result=result,
        )
        
        api_logger.info(
            "insights.executed",
            insight_id=payload.insight_id,
            action=action,
            project=payload.project,
            user=user.user_id,
        )
        
    except Exception as e:
        result["message"] = f"Error ejecutando: {str(e)}"
        api_logger.error("insights.execute.error", error=str(e))
    
    return result


class GenerateInsightsRequest(BaseModel):
    project: str
    source: str = "coding"  # 'coding', 'report'


@app.post("/api/insights/generate")
async def api_generate_insights(
    payload: GenerateInsightsRequest,
    clients: ServiceClients = Depends(get_service_clients),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Genera insights manualmente desde una fuente.
    
    Útil para trigger manual de análisis de códigos poco frecuentes.
    """
    from app.insights import extract_insights_from_coding
    
    if payload.source == "coding":
        insight_ids = extract_insights_from_coding(
            clients.postgres,
            project=payload.project,
        )
        
        return {
            "source": "coding",
            "insights_created": len(insight_ids),
            "insight_ids": insight_ids,
        }
    
    return {"source": payload.source, "insights_created": 0, "message": "Fuente no soportada"}


# ═══════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════
# ADMIN PANEL ENDPOINTS - Gestión de usuarios, limpieza y análisis
# ═══════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════════

# ─────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────
# 1. USERS MANAGEMENT ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────

@app.get("/api/admin/users")
async def api_admin_list_users(
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """Lista todos los usuarios de la organización del admin autenticado."""
    clients = build_clients_or_error(settings)
    try:
        with clients.postgres.cursor() as cur:
            cur.execute(
                """
                SELECT id, email, full_name, role, organization_id, is_active, created_at, last_login
                  FROM app_users
                 WHERE organization_id = %s
                 ORDER BY created_at DESC
                """,
                (user.organization_id,),
            )
            rows = cur.fetchall()
        
        users = [
            {
                "id": row[0],
                "email": row[1],
                "full_name": row[2],
                "role": row[3],
                "organization_id": row[4],
                "is_active": row[5],
                "created_at": row[6].isoformat() if row[6] else None,
                "last_login": row[7].isoformat() if row[7] else None,
            }
            for row in rows
        ]
        
        return {"users": users, "total": len(users)}
    finally:
        clients.close()


@app.get("/api/admin/stats")
async def api_admin_get_stats(
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """Obtiene estadísticas generales de la organización."""
    clients = build_clients_or_error(settings)
    try:
        with clients.postgres.cursor() as cur:
            # Contar usuarios por rol
            cur.execute(
                """
                SELECT role, COUNT(*) as count
                  FROM app_users
                 WHERE organization_id = %s AND is_active = true
                 GROUP BY role
                """,
                (user.organization_id,),
            )
            role_counts = {row[0]: row[1] for row in cur.fetchall()}
            
            # Total de usuarios
            cur.execute(
                "SELECT COUNT(*) FROM app_users WHERE organization_id = %s",
                (user.organization_id,),
            )
            total_users = cur.fetchone()[0] or 0
            
            # Total de fragmentos
            cur.execute(
                """
                SELECT COUNT(DISTINCT ef.id) FROM entrevista_fragmentos ef
                 JOIN proyectos p ON ef.project_id = p.id
                 WHERE p.org_id = %s
                """,
                (user.organization_id,),
            )
            total_fragments = cur.fetchone()[0] or 0
            
            # Sesiones activas (últimos 30 minutos)
            cur.execute(
                """
                SELECT COUNT(DISTINCT user_id) FROM app_user_sessions
                 WHERE organization_id = %s AND last_activity > NOW() - INTERVAL '30 minutes'
                """,
                (user.organization_id,),
            )
            active_sessions = cur.fetchone()[0] or 0
        
        return {
            "organization_id": user.organization_id,
            "total_users": total_users,
            "users_by_role": role_counts,
            "total_fragments": total_fragments,
            "active_sessions": active_sessions,
        }
    finally:
        clients.close()


class UserUpdateRequest(BaseModel):
    role: Optional[str] = None
    is_active: Optional[bool] = None


@app.patch("/api/admin/users/{user_id}")
async def api_admin_update_user(
    user_id: str,
    payload: UserUpdateRequest,
    settings: AppSettings = Depends(get_settings),
    admin: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """Actualiza el rol o estado de un usuario."""
    clients = build_clients_or_error(settings)
    try:
        # Verificar que el usuario a actualizar pertenece a la misma organización
        with clients.postgres.cursor() as cur:
            cur.execute(
                "SELECT organization_id FROM app_users WHERE id = %s",
                (user_id,),
            )
            row = cur.fetchone()
            if not row or row[0] != admin.org_id:
                raise HTTPException(status_code=404, detail="User not found")
            
            # Actualizar
            updates = []
            values = []
            if payload.role:
                updates.append("role = %s")
                values.append(payload.role)
            if payload.is_active is not None:
                updates.append("is_active = %s")
                values.append(payload.is_active)
            
            if not updates:
                raise HTTPException(status_code=400, detail="No updates provided")
            
            values.append(user_id)
            cur.execute(
                f"UPDATE app_users SET {', '.join(updates)}, updated_at = NOW() WHERE id = %s",
                values,
            )
            clients.postgres.commit()
        
        api_logger.info("admin.user_updated", user_id=user_id, admin_id=admin.user_id, org_id=admin.org_id)
        return {"user_id": user_id, "status": "updated"}
    finally:
        clients.close()


@app.delete("/api/admin/users/{user_id}")
async def api_admin_delete_user(
    user_id: str,
    settings: AppSettings = Depends(get_settings),
    admin: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """Elimina un usuario (soft-delete o hard-delete según config)."""
    clients = build_clients_or_error(settings)
    try:
        # Verificar que el usuario a eliminar pertenece a la misma organización y no es el admin
        if user_id == admin.user_id:
            raise HTTPException(status_code=400, detail="No puedes eliminarte a ti mismo")
        
        with clients.postgres.cursor() as cur:
            cur.execute(
                "SELECT organization_id FROM app_users WHERE id = %s",
                (user_id,),
            )
            row = cur.fetchone()
            if not row or row[0] != admin.org_id:
                raise HTTPException(status_code=404, detail="User not found")
            
            # Soft delete (marcar como inactivo)
            cur.execute(
                "UPDATE app_users SET is_active = false, updated_at = NOW() WHERE id = %s",
                (user_id,),
            )
            clients.postgres.commit()
        
        api_logger.info("admin.user_deleted", user_id=user_id, admin_id=admin.user_id, org_id=admin.org_id)
        return {"user_id": user_id, "status": "deleted"}
    finally:
        clients.close()


# ─────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────
# 2. DATA CLEANUP ENDPOINTS (Destructive - Admin Only)
# ─────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────

class CleanupConfirmRequest(BaseModel):
    """Solicitud de confirmación para operaciones destructivas."""
    confirm: bool = Field(default=False, description="Debe ser true para ejecutar")
    reason: str = Field(default="", description="Razón o comentario (logging)")


@app.post("/api/admin/cleanup/all-data")
async def api_admin_cleanup_all_data(
    payload: CleanupConfirmRequest,
    project: str = Query(default="default", description="Proyecto a limpiar"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """
    DESTRUCTIVE: Elimina todos los datos de un proyecto (PostgreSQL, Qdrant, Neo4j).
    Requiere confirm=true.
    """
    if not payload.confirm:
        return {
            "status": "not_confirmed",
            "message": "Operación requiere confirm=true. Esta acción es irreversible.",
            "project": project,
        }
    
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    counts: Dict[str, int] = {"postgres": 0, "qdrant": 0, "neo4j": 0}
    
    try:
        # 1. PostgreSQL cleanup
        with clients.postgres.cursor() as cur:
            tables = [
                "entrevista_fragmentos",
                "analisis_codigos_abiertos",
                "analisis_relacion_axial",
                "reporte_entrevista_staging",
                "report_jobs",
            ]
            cur.execute(
                """
                SELECT table_name
                FROM information_schema.tables
                WHERE table_schema = 'public'
                  AND table_name = ANY(%s)
                """,
                (tables,),
            )
            existing_tables = {row[0] for row in cur.fetchall()}
            for table in tables:
                if table not in existing_tables:
                    api_logger.warning("admin.cleanup.table_missing", table=table, project=project_id)
                    continue
                cur.execute(f"DELETE FROM {table} WHERE project_id = %s", (project_id,))
                counts["postgres"] += cur.rowcount
            if "proyectos" in existing_tables:
                cur.execute("DELETE FROM proyectos WHERE id = %s", (project_id,))
                counts["postgres"] += cur.rowcount
            clients.postgres.commit()
        
        # 2. Qdrant cleanup - use global collection with project_id filter
        try:
            # Use settings.qdrant.collection (global collection) instead of per-project collection
            clients.qdrant.delete(
                collection_name=settings.qdrant.collection,
                points_selector=models.Filter(
                    must=[
                        models.FieldCondition(
                            key="project_id",
                            match=models.MatchValue(value=project_id),
                        )
                    ]
                ),
            )
            counts["qdrant"] = 1  # Marker
        except Exception as qd_err:
            api_logger.warning("admin.cleanup.qdrant_error", error=str(qd_err), project=project_id)
        
        # 3. Neo4j cleanup
        try:
            with clients.neo4j.session(database=settings.neo4j.database) as session:
                result = session.run(
                    "MATCH (n {project_id: $project}) DETACH DELETE n",
                    project=project_id,
                )
                counts["neo4j"] = result.consume().counters.nodes_deleted
        except Exception as neo_err:
            api_logger.warning("admin.cleanup.neo4j_error", error=str(neo_err), project=project_id)
        
        api_logger.info(
            "admin.cleanup_all_data",
            project=project_id,
            admin_id=user.user_id,
            counts=counts,
            reason=payload.reason,
        )
        
        return {
            "status": "completed",
            "project": project_id,
            "counts": counts,
            "message": f"Limpieza completada: {counts['postgres']} registros PostgreSQL, {counts['qdrant']} Qdrant, {counts['neo4j']} Neo4j",
        }
    except Exception as exc:
        api_logger.error("admin.cleanup_all_data.error", error=str(exc), project=project_id)
        raise HTTPException(status_code=500, detail=f"Error limpiando: {str(exc)}") from exc
    finally:
        clients.close()


@app.post("/api/admin/cleanup/projects")
async def api_admin_cleanup_projects(
    payload: CleanupConfirmRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """
    DESTRUCTIVE: Elimina datos de proyectos marcados como deleted.
    Requiere confirm=true.
    """
    if not payload.confirm:
        return {
            "status": "not_confirmed",
            "message": "Operación requiere confirm=true.",
        }
    
    clients = build_clients_or_error(settings)
    try:
        with clients.postgres.cursor() as cur:
            # Validar que la columna is_deleted exista (compatibilidad con esquemas antiguos)
            cur.execute(
                """
                SELECT 1
                FROM information_schema.columns
                WHERE table_name = 'proyectos'
                  AND column_name = 'is_deleted'
                LIMIT 1
                """
            )
            has_is_deleted = cur.fetchone() is not None
            if not has_is_deleted:
                api_logger.warning(
                    "admin.cleanup_projects.missing_column",
                    column="is_deleted",
                    table="proyectos",
                )
                return {
                    "status": "not_supported",
                    "message": "La columna is_deleted no existe en proyectos. No hay proyectos marcados como deleted.",
                    "cleaned_projects": [],
                }

            # Obtener proyectos deleted
            cur.execute(
                "SELECT id FROM proyectos WHERE org_id = %s AND is_deleted = true",
                (user.organization_id,),
            )
            deleted_projects = [row[0] for row in cur.fetchall()]
            
            if not deleted_projects:
                return {
                    "status": "no_action_needed",
                    "message": "No hay proyectos marcados como deleted",
                    "cleaned_projects": [],
                }
            
            cleaned = []
            for project_id in deleted_projects:
                # Limpiar datos del proyecto
                tables = [
                    "entrevista_fragmentos",
                    "analisis_codigos_abiertos",
                    "analisis_relacion_axial",
                ]
                cur.execute(
                    """
                    SELECT table_name
                    FROM information_schema.tables
                    WHERE table_schema = 'public'
                      AND table_name = ANY(%s)
                    """,
                    (tables,),
                )
                existing_tables = {row[0] for row in cur.fetchall()}
                for table in tables:
                    if table not in existing_tables:
                        api_logger.warning("admin.cleanup.table_missing", table=table, project=project_id)
                        continue
                    cur.execute(f"DELETE FROM {table} WHERE project_id = %s", (project_id,))
                cleaned.append({"project_id": project_id, "rows_deleted": cur.rowcount})
            
            # Eliminar registros de proyectos
            cur.execute(
                "DELETE FROM proyectos WHERE org_id = %s AND is_deleted = true",
                (user.organization_id,),
            )
            clients.postgres.commit()
        
        api_logger.info(
            "admin.cleanup_projects",
            org_id=user.organization_id,
            admin_id=user.user_id,
            projects_cleaned=len(deleted_projects),
        )
        
        return {
            "status": "completed",
            "cleaned_projects": cleaned,
            "message": f"Limpieza completada: {len(deleted_projects)} proyectos eliminados",
        }
    except Exception as exc:
        api_logger.error("admin.cleanup_projects.error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error limpiando proyectos: {str(exc)}") from exc
    finally:
        clients.close()


@app.post("/api/admin/cleanup/orphans")
async def api_admin_cleanup_orphans(
    payload: CleanupConfirmRequest,
    project: str = Query(default="default", description="Proyecto a limpiar"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """
    DESTRUCTIVE: Elimina archivos huérfanos del proyecto (PostgreSQL + Neo4j).
    Un archivo huérfano existe en DB pero no existe en filesystem local.
    Requiere confirm=true.
    """
    if not payload.confirm:
        return {
            "status": "not_confirmed",
            "message": "Operación requiere confirm=true.",
            "project": project,
        }

    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        with clients.postgres.cursor() as cur:
            cur.execute(
                """
                SELECT DISTINCT archivo FROM entrevista_fragmentos
                 WHERE project_id = %s
                 ORDER BY archivo
                """,
                (project_id,),
            )
            files_in_db = [row[0] for row in cur.fetchall() if row and row[0]]

        blob_enabled = False
        blob_check_failed = False
        try:
            from app import blob_storage

            conn_str = os.environ.get("AZURE_STORAGE_CONNECTION_STRING")
            blob_enabled = bool(conn_str) and bool(getattr(blob_storage, "_AZURE_BLOB_AVAILABLE", False))
        except Exception:
            blob_enabled = False

        orphans: list[str] = []
        for filename in files_in_db:
            file_path = Path("data") / project_id / filename
            exists_in_fs = file_path.exists()
            exists_in_blob = None
            if blob_enabled:
                try:
                    from app.blob_storage import CONTAINER_INTERVIEWS, file_exists

                    blob_name = f"{project_id}/{filename}"
                    exists_in_blob = bool(file_exists(CONTAINER_INTERVIEWS, blob_name))
                except Exception:
                    blob_check_failed = True
                    exists_in_blob = None

            if blob_enabled:
                is_orphan = (not exists_in_fs) and (exists_in_blob is False)
            else:
                is_orphan = not exists_in_fs

            if is_orphan:
                orphans.append(filename)

        if not orphans:
            return {
                "status": "no_action_needed",
                "project": project_id,
                "message": "No hay archivos huérfanos para limpiar",
                "cleaned": [],
                "deleted_codes": 0,
                "deleted_fragments": 0,
                "deleted_neo4j_nodes": 0,
            }

        deleted_codes = 0
        deleted_fragments = 0
        deleted_neo4j_nodes = 0
        cleaned: list[Dict[str, Any]] = []

        for filename in orphans:
            file_deleted_codes = 0
            file_deleted_fragments = 0
            file_deleted_neo4j = 0
            with clients.postgres.cursor() as cur:
                cur.execute(
                    """
                    DELETE FROM analisis_codigos_abiertos
                     WHERE project_id = %s AND archivo = %s
                    """,
                    (project_id, filename),
                )
                file_deleted_codes = cur.rowcount

                cur.execute(
                    """
                    DELETE FROM entrevista_fragmentos
                     WHERE project_id = %s AND archivo = %s
                    """,
                    (project_id, filename),
                )
                file_deleted_fragments = cur.rowcount
            clients.postgres.commit()

            try:
                with clients.neo4j.session(database=settings.neo4j.database) as session:
                    result = session.run(
                        """
                        MATCH (e:Entrevista {nombre: $archivo, project_id: $project})
                        OPTIONAL MATCH (e)-[:TIENE_FRAGMENTO]->(f:Fragmento)
                        DETACH DELETE e, f
                        RETURN count(e) AS deleted
                        """,
                        archivo=filename,
                        project=project_id,
                    )
                    record = result.single()
                    file_deleted_neo4j = int(record["deleted"]) if record and record.get("deleted") else 0
            except Exception as neo_err:
                api_logger.warning(
                    "admin.cleanup_orphans.neo4j_error",
                    error=str(neo_err),
                    project=project_id,
                    archivo=filename,
                )

            deleted_codes += file_deleted_codes
            deleted_fragments += file_deleted_fragments
            deleted_neo4j_nodes += file_deleted_neo4j
            cleaned.append(
                {
                    "archivo": filename,
                    "deleted_codes": file_deleted_codes,
                    "deleted_fragments": file_deleted_fragments,
                    "deleted_neo4j_nodes": file_deleted_neo4j,
                }
            )

        api_logger.info(
            "admin.cleanup_orphans",
            project=project_id,
            admin_id=user.user_id,
            orphans=len(orphans),
            deleted_codes=deleted_codes,
            deleted_fragments=deleted_fragments,
            deleted_neo4j_nodes=deleted_neo4j_nodes,
            reason=payload.reason,
            blob_enabled=blob_enabled,
            blob_check_failed=blob_check_failed,
        )

        return {
            "status": "completed",
            "project": project_id,
            "orphans_cleaned": len(orphans),
            "deleted_codes": deleted_codes,
            "deleted_fragments": deleted_fragments,
            "deleted_neo4j_nodes": deleted_neo4j_nodes,
            "cleaned": cleaned,
            "message": f"Huérfanos eliminados: {len(orphans)}",
            "blob_enabled": blob_enabled,
        }
    except Exception as exc:
        api_logger.error("admin.cleanup_orphans.error", error=str(exc), project=project_id)
        raise HTTPException(status_code=500, detail=f"Error limpiando huérfanos: {str(exc)}") from exc
    finally:
        clients.close()


@app.post("/api/admin/cleanup/neo4j-orphans")
async def api_admin_cleanup_neo4j_orphans(
    payload: CleanupConfirmRequest,
    project: str = Query(default="default", description="Proyecto a limpiar"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """
    DESTRUCTIVE: Elimina nodos huérfanos en Neo4j comparando con PostgreSQL.
    Huérfanos:
      - Fragmentos que no existen en entrevista_fragmentos
      - Entrevistas cuyo archivo no existe en entrevista_fragmentos
    Requiere confirm=true.
    """
    if not payload.confirm:
        return {
            "status": "not_confirmed",
            "message": "Operación requiere confirm=true. Esta acción es irreversible.",
            "project": project,
        }

    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    clients = build_clients_or_error(settings)
    try:
        # PostgreSQL: ids válidos
        with clients.postgres.cursor() as cur:
            cur.execute(
                "SELECT id FROM entrevista_fragmentos WHERE project_id = %s",
                (project_id,),
            )
            pg_fragments = {str(row[0]) for row in cur.fetchall() if row and row[0] is not None}

            cur.execute(
                "SELECT DISTINCT archivo FROM entrevista_fragmentos WHERE project_id = %s",
                (project_id,),
            )
            pg_archivos = {str(row[0]) for row in cur.fetchall() if row and row[0] is not None}

        # Neo4j: ids actuales
        neo_fragments: list[str] = []
        neo_archivos: list[str] = []
        with clients.neo4j.session(database=settings.neo4j.database) as session:
            result = session.run(
                "MATCH (f:Fragmento {project_id: $project}) RETURN f.id AS id",
                project=project_id,
            )
            for record in result:
                frag_id = record.get("id")
                if frag_id is not None:
                    neo_fragments.append(str(frag_id))

            result = session.run(
                "MATCH (e:Entrevista {project_id: $project}) RETURN e.nombre AS nombre",
                project=project_id,
            )
            for record in result:
                nombre = record.get("nombre")
                if nombre is not None:
                    neo_archivos.append(str(nombre))

        orphan_fragments = [fid for fid in neo_fragments if fid not in pg_fragments]
        orphan_archivos = [a for a in neo_archivos if a not in pg_archivos]

        deleted_fragments = 0
        deleted_entrevistas = 0
        batch = 500

        if orphan_fragments:
            for i in range(0, len(orphan_fragments), batch):
                batch_ids = orphan_fragments[i : i + batch]
                with clients.neo4j.session(database=settings.neo4j.database) as session:
                    result = session.run(
                        """
                        UNWIND $ids AS id
                        MATCH (f:Fragmento {project_id: $project, id: id})
                        DETACH DELETE f
                        RETURN count(f) AS deleted
                        """,
                        project=project_id,
                        ids=batch_ids,
                    )
                    record = result.single()
                    if record and record.get("deleted") is not None:
                        deleted_fragments += int(record["deleted"])

        if orphan_archivos:
            for i in range(0, len(orphan_archivos), batch):
                batch_names = orphan_archivos[i : i + batch]
                with clients.neo4j.session(database=settings.neo4j.database) as session:
                    result = session.run(
                        """
                        UNWIND $nombres AS nombre
                        MATCH (e:Entrevista {project_id: $project, nombre: nombre})
                        OPTIONAL MATCH (e)-[:TIENE_FRAGMENTO]->(f:Fragmento)
                        DETACH DELETE e, f
                        RETURN count(e) AS deleted
                        """,
                        project=project_id,
                        nombres=batch_names,
                    )
                    record = result.single()
                    if record and record.get("deleted") is not None:
                        deleted_entrevistas += int(record["deleted"])

        api_logger.info(
            "admin.cleanup_neo4j_orphans",
            project=project_id,
            admin_id=user.user_id,
            orphan_fragments=len(orphan_fragments),
            orphan_entrevistas=len(orphan_archivos),
            deleted_fragments=deleted_fragments,
            deleted_entrevistas=deleted_entrevistas,
            reason=payload.reason,
        )

        return {
            "status": "completed",
            "project": project_id,
            "orphan_fragments": len(orphan_fragments),
            "orphan_entrevistas": len(orphan_archivos),
            "deleted_fragments": deleted_fragments,
            "deleted_entrevistas": deleted_entrevistas,
            "message": (
                f"Neo4j huérfanos eliminados: {deleted_fragments} fragmentos, "
                f"{deleted_entrevistas} entrevistas."
            ),
        }
    except Exception as exc:
        api_logger.error("admin.cleanup_neo4j_orphans.error", error=str(exc), project=project_id)
        raise HTTPException(status_code=500, detail=f"Error limpiando Neo4j: {str(exc)}") from exc
    finally:
        clients.close()


@app.post("/api/admin/cleanup/neo4j-unscoped")
async def api_admin_cleanup_neo4j_unscoped(
    payload: CleanupConfirmRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """
    DESTRUCTIVE: Elimina nodos Neo4j que no tienen project_id.
    Útil para limpiar datos antiguos creados sin scope de proyecto.
    Requiere confirm=true.
    """
    if not payload.confirm:
        return {
            "status": "not_confirmed",
            "message": "Operación requiere confirm=true. Esta acción es irreversible.",
        }

    clients = build_clients_or_error(settings)
    labels = ["Fragmento", "Entrevista", "Codigo", "Categoria"]
    counts: Dict[str, int] = {"nodes_deleted": 0}

    try:
        with clients.neo4j.session(database=settings.neo4j.database) as session:
            for label in labels:
                result = session.run(
                    f"MATCH (n:{label}) WHERE n.project_id IS NULL DETACH DELETE n",
                )
                counts["nodes_deleted"] += result.consume().counters.nodes_deleted

        api_logger.info(
            "admin.cleanup_neo4j_unscoped",
            admin_id=user.user_id,
            labels=labels,
            counts=counts,
            reason=payload.reason,
        )

        return {
            "status": "completed",
            "counts": counts,
            "message": f"Neo4j sin project_id eliminado: {counts['nodes_deleted']} nodos.",
        }
    except Exception as exc:
        api_logger.error("admin.cleanup_neo4j_unscoped.error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error limpiando Neo4j: {str(exc)}") from exc
    finally:
        clients.close()


@app.post("/api/admin/cleanup/neo4j-unscoped-relationships")
async def api_admin_cleanup_neo4j_unscoped_relationships(
    payload: CleanupConfirmRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin"])),
) -> Dict[str, Any]:
    """
    DESTRUCTIVE: Elimina relaciones Neo4j cuyo origen o destino no tienen project_id
    (o la propia relación carece de project_id). Útil para limpiar datos antiguos.
    Requiere confirm=true.
    """
    if not payload.confirm:
        return {
            "status": "not_confirmed",
            "message": "Operación requiere confirm=true. Esta acción es irreversible.",
        }

    clients = build_clients_or_error(settings)
    counts: Dict[str, int] = {"relationships_deleted": 0}

    try:
        with clients.neo4j.session(database=settings.neo4j.database) as session:
            result = session.run(
                """
                MATCH (a)-[r]->(b)
                WHERE a.project_id IS NULL OR b.project_id IS NULL OR r.project_id IS NULL
                DELETE r
                RETURN count(r) AS deleted
                """
            )
            record = result.single()
            if record and record.get("deleted") is not None:
                counts["relationships_deleted"] = int(record["deleted"])

        api_logger.info(
            "admin.cleanup_neo4j_unscoped_relationships",
            admin_id=user.user_id,
            counts=counts,
            reason=payload.reason,
        )

        return {
            "status": "completed",
            "counts": counts,
            "message": f"Relaciones Neo4j sin project_id eliminadas: {counts['relationships_deleted']}.",
        }
    except Exception as exc:
        api_logger.error("admin.cleanup_neo4j_unscoped_relationships.error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error limpiando Neo4j: {str(exc)}") from exc
    finally:
        clients.close()


@app.post("/api/admin/cleanup/duplicate-codes")
async def api_admin_cleanup_duplicate_codes(
    project: str = Query(default="default", description="Proyecto a analizar"),
    threshold: float = Query(default=0.85, description="Umbral de similitud (0-1)"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Dict[str, Any]:
    """
    Detecta códigos posiblemente duplicados usando similitud de strings.
    No es destructiva, solo reporta candidatos.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        from difflib import SequenceMatcher
        
        with clients.postgres.cursor() as cur:
            cur.execute(
                "SELECT codigo, COUNT(*) FROM analisis_codigos_abiertos WHERE project_id = %s GROUP BY codigo",
                (project_id,),
            )
            codes = {row[0]: row[1] for row in cur.fetchall()}
        
        # Encontrar similares
        similar_groups = []
        checked = set()
        
        for code1 in codes.keys():
            if code1 in checked:
                continue
            group = [code1]
            for code2 in codes.keys():
                if code1 != code2 and code2 not in checked:
                    sim = SequenceMatcher(None, code1, code2).ratio()
                    if sim >= threshold:
                        group.append(code2)
                        checked.add(code2)
            if len(group) > 1:
                similar_groups.append(group)
                checked.add(code1)
        
        api_logger.info(
            "admin.duplicate_codes_detection",
            project=project_id,
            total_codes=len(codes),
            duplicate_groups=len(similar_groups),
            threshold=threshold,
        )
        
        return {
            "status": "completed",
            "project": project_id,
            "total_codes": len(codes),
            "duplicate_groups": similar_groups,
            "groups_count": len(similar_groups),
            "threshold": threshold,
        }
    except Exception as exc:
        api_logger.error("admin.duplicate_codes.error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error detectando duplicados: {str(exc)}") from exc
    finally:
        clients.close()


# ─────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────
# 3. INTEGRITY ANALYSIS ENDPOINTS (Non-destructive - Analyst+)
# ─────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────────

@app.get("/api/admin/analysis/orphan-files")
async def api_admin_find_orphan_files(
    project: str = Query(default="default", description="Proyecto a analizar"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Dict[str, Any]:
    """
    Detecta archivos en PostgreSQL que no existen en Blob Storage o filesystem.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        with clients.postgres.cursor() as cur:
            cur.execute(
                """
                SELECT DISTINCT archivo FROM entrevista_fragmentos
                 WHERE project_id = %s
                 ORDER BY archivo
                """,
                (project_id,),
            )
            files_in_db = {row[0] for row in cur.fetchall()}

        # Validar contra filesystem local y Blob Storage cuando esté disponible
        orphans = []
        blob_enabled = False
        blob_check_failed = False
        try:
            from app import blob_storage

            conn_str = os.environ.get("AZURE_STORAGE_CONNECTION_STRING")
            blob_enabled = bool(conn_str) and bool(getattr(blob_storage, "_AZURE_BLOB_AVAILABLE", False))
        except Exception:
            blob_enabled = False

        for filename in files_in_db:
            file_path = Path("data") / project_id / filename
            exists_in_fs = file_path.exists()
            exists_in_blob = None
            if blob_enabled:
                try:
                    from app.blob_storage import CONTAINER_INTERVIEWS, file_exists

                    blob_name = f"{project_id}/{filename}"
                    exists_in_blob = bool(file_exists(CONTAINER_INTERVIEWS, blob_name))
                except Exception:
                    blob_check_failed = True
                    exists_in_blob = None

            if blob_enabled:
                is_orphan = (not exists_in_fs) and (exists_in_blob is False)
            else:
                is_orphan = not exists_in_fs

            if is_orphan:
                orphans.append(
                    {
                        "filename": filename,
                        "exists_in_db": True,
                        "exists_in_fs": exists_in_fs,
                        "exists_in_blob": exists_in_blob,
                    }
                )
        
        api_logger.info(
            "admin.orphan_files_detection",
            project=project_id,
            total_files=len(files_in_db),
            orphans_found=len(orphans),
            blob_enabled=blob_enabled,
            blob_check_failed=blob_check_failed,
        )
        
        return {
            "status": "completed",
            "project": project_id,
            "total_files": len(files_in_db),
            "orphans": orphans,
            "orphans_count": len(orphans),
            "blob_enabled": blob_enabled,
        }
    except Exception as exc:
        api_logger.error("admin.orphan_files.error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error analizando archivos: {str(exc)}") from exc
    finally:
        clients.close()


@app.get("/api/admin/analysis/integrity")
async def api_admin_integrity_check(
    project: str = Query(default="default", description="Proyecto a verificar"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_role(["admin", "analyst"])),
) -> Dict[str, Any]:
    """
    Chequeo de integridad general: fragmentos sin códigos, códigos sin fragmentos, etc.
    """
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    clients = build_clients_or_error(settings)
    try:
        checks = {}
        
        with clients.postgres.cursor() as cur:
            # 1. Fragmentos sin códigos
            cur.execute(
                """
                SELECT COUNT(DISTINCT ef.id) FROM entrevista_fragmentos ef
                 LEFT JOIN analisis_codigos_abiertos ac ON ef.id = ac.fragmento_id
                 WHERE ef.project_id = %s AND ac.id IS NULL
                """,
                (project_id,),
            )
            checks["fragments_without_codes"] = cur.fetchone()[0] or 0
            
            # 2. Total de fragmentos
            cur.execute(
                "SELECT COUNT(*) FROM entrevista_fragmentos WHERE project_id = %s",
                (project_id,),
            )
            checks["total_fragments"] = cur.fetchone()[0] or 0
            
            # 3. Total de códigos
            cur.execute(
                "SELECT COUNT(DISTINCT codigo) FROM analisis_codigos_abiertos WHERE project_id = %s",
                (project_id,),
            )
            checks["unique_codes"] = cur.fetchone()[0] or 0
            
            # 4. Total de asignaciones de código
            cur.execute(
                "SELECT COUNT(*) FROM analisis_codigos_abiertos WHERE project_id = %s",
                (project_id,),
            )
            checks["total_code_assignments"] = cur.fetchone()[0] or 0
        
        api_logger.info(
            "admin.integrity_check",
            project=project_id,
            checks=checks,
        )
        
        return {
            "status": "completed",
            "project": project_id,
            "checks": checks,
        }
    except Exception as exc:
        api_logger.error("admin.integrity_check.error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error verificando integridad: {str(exc)}") from exc
    finally:
        clients.close()


# Force reload
