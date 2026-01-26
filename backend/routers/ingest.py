"""
Ingest router - Document upload, ingestion, and transcription endpoints.

Endpoints:
    POST /api/upload-and-ingest  - Upload DOCX file and ingest
    POST /api/ingest             - Ingest from server paths
    POST /api/transcribe         - Transcribe audio (sync)
    POST /api/transcribe/stream  - Transcribe audio (async with progress)
    POST /api/transcribe/batch   - Batch transcription
    POST /api/transcribe/merge   - Merge transcriptions to DOCX
    GET  /api/jobs/{task_id}/status  - Job status polling
"""
import base64
import json
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, cast

import structlog
from celery.result import AsyncResult
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel, ConfigDict, Field

from app.clients import ServiceClients, build_service_clients
from app.ingestion import ingest_documents
from app.project_state import resolve_project
from app.blob_storage import (
    upload_file,
    CONTAINER_INTERVIEWS,
    CONTAINER_AUDIO,
    logical_path_to_blob_name,
    upload_local_path,
    tenant_upload,
    TenantRequiredError,
    build_artifact_contract,
)
from app.settings import AppSettings, load_settings
from app.postgres_block import stage0_require_ready_or_override
from backend.auth import User, get_current_user
from backend.celery_worker import celery_app, task_transcribe_audio



# =============================================================================
# LOGGER
# =============================================================================

api_logger = structlog.get_logger("app.api")


# =============================================================================
# DEPENDENCIES
# =============================================================================

def get_settings() -> AppSettings:
    import os
    env_file = os.getenv("APP_ENV_FILE")
    return load_settings(env_file)


from typing import AsyncGenerator

async def get_service_clients(settings: AppSettings = Depends(get_settings)) -> AsyncGenerator[ServiceClients, None]:
    """
    CRITICAL: Uses yield + finally to ensure connections are returned to pool!
    """
    clients = build_service_clients(settings)
    try:
        yield clients
    finally:
        clients.close()


async def require_auth(user: User = Depends(get_current_user)) -> User:
    return user


# =============================================================================
# MODELS - INGESTION
# =============================================================================

class IngestRequest(BaseModel):
    model_config = ConfigDict(extra="forbid")

    project: str
    inputs: List[str]
    meta_json: Optional[str] = None
    batch_size: int = 64
    min_chars: int = 200
    max_chars: int = 1200
    run_id: Optional[str] = None


# =============================================================================
# MODELS - TRANSCRIPTION
# =============================================================================

class TranscribeRequest(BaseModel):
    """Request para transcripción de audio con diarización."""
    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="Proyecto donde se ingesta (si aplica)")
    audio_base64: str = Field(..., description="Archivo de audio codificado en base64")
    filename: str = Field(..., description="Nombre original del archivo (con extensión)")
    diarize: bool = Field(True, description="Usar diarización para separar speakers")
    language: str = Field("es", description="Código de idioma")
    ingest: bool = Field(False, description="Ingestar directamente al pipeline")
    save_to_folder: bool = Field(True, description="Guardar DOCX en carpeta del proyecto")
    min_chars: int = Field(200, description="Mínimo de caracteres por fragmento")
    max_chars: int = Field(1200, description="Máximo de caracteres por fragmento")


class TranscribeSegmentResponse(BaseModel):
    speaker: str
    text: str
    start: float
    end: float


class TranscribeResponse(BaseModel):
    """Respuesta de transcripción con segmentos y metadata."""
    text: str = Field(..., description="Transcripción completa")
    segments: List[TranscribeSegmentResponse] = Field(default_factory=list)
    speaker_count: int = Field(0, description="Número de speakers detectados")
    duration_seconds: float = Field(0.0, description="Duración del audio")
    fragments_ingested: Optional[int] = Field(None, description="Fragmentos creados si ingest=True")
    saved_path: Optional[str] = Field(None, description="Ruta donde se guardó el DOCX")


class StreamTranscribeRequest(BaseModel):
    """Request para transcripción asíncrona con entregas incrementales."""
    model_config = ConfigDict(extra="forbid")

    project: str = Field(..., description="Proyecto destino")
    audio_base64: str = Field(..., description="Audio codificado en base64")
    filename: str = Field(..., description="Nombre del archivo")
    diarize: bool = Field(True, description="Usar diarización")
    language: str = Field("es", description="Código de idioma")
    ingest: bool = Field(True, description="Ingestar al pipeline")
    min_chars: int = Field(200)
    max_chars: int = Field(1200)


class StreamTranscribeResponse(BaseModel):
    """Respuesta de transcripción stream: task_id para polling."""
    task_id: str
    filename: str
    message: str


class BatchTranscribeFileItem(BaseModel):
    """Un archivo en el batch de transcripción."""
    audio_base64: str = Field(..., description="Audio codificado en base64")
    filename: str = Field(..., description="Nombre del archivo")


class BatchTranscribeRequest(BaseModel):
    """Request para transcripción paralela de múltiples archivos."""
    model_config = ConfigDict(extra="forbid")
    
    project: str = Field(..., description="Proyecto destino")
    files: List[BatchTranscribeFileItem] = Field(..., description="Lista de archivos (max 20)")
    diarize: bool = Field(True, description="Usar diarización")
    language: str = Field("es", description="Código de idioma")
    ingest: bool = Field(True, description="Ingestar al pipeline")
    min_chars: int = Field(200)
    max_chars: int = Field(1200)


class BatchJobInfo(BaseModel):
    """Info de un job en el batch."""
    task_id: str
    filename: str


class BatchTranscribeResponse(BaseModel):
    """Respuesta de batch: IDs de jobs para tracking."""
    batch_id: str
    jobs: List[BatchJobInfo]
    message: str


class JobStatusResponse(BaseModel):
    """Estado de un job individual con soporte para progreso incremental."""
    task_id: str
    status: str  # PENDING, PROCESSING, SUCCESS, FAILURE
    filename: Optional[str] = None
    stage: Optional[str] = None  # transcribing, ingesting
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    chunk_index: Optional[int] = None
    total_chunks: Optional[int] = None
    chunks_completed: Optional[int] = None
    text_preview: Optional[str] = None
    segments_count: Optional[int] = None
    speakers_count: Optional[int] = None


# =============================================================================
# CONSTANTS
# =============================================================================

MAX_BATCH_SIZE = 20
SUPPORTED_DOCX_EXTENSIONS = {".docx"}


# =============================================================================
# HELPERS
# =============================================================================

def _expand_inputs(raw_inputs: List[str]) -> List[str]:
    """Expande patrones glob y normaliza rutas para compatibilidad Docker."""
    expanded: List[str] = []
    seen = set()
    for item in raw_inputs:
        if not item:
            continue
        item = item.replace("\\", "/")
        if any(ch in item for ch in ("*", "?", "[", "]")):
            for match in sorted(Path().glob(item)):
                if match.exists():
                    path = str(match)
                    if path not in seen:
                        expanded.append(path)
                        seen.add(path)
        else:
            if item not in seen:
                expanded.append(item)
                seen.add(item)
    return expanded


def _build_clients_or_error(settings: AppSettings) -> ServiceClients:
    clients = build_service_clients(settings)
    if clients.postgres is None:
        clients.close()  # CRITICAL: Close connections before raising!
        raise HTTPException(status_code=503, detail="PostgreSQL no disponible")
    return clients


# =============================================================================
# ROUTER
# =============================================================================

router = APIRouter(prefix="/api", tags=["Ingestion"])


# =============================================================================
# NEW ENDPOINT: UPLOAD AND INGEST
# =============================================================================

@router.post("/upload-and-ingest")
async def api_upload_and_ingest(
    file: UploadFile = File(..., description="Archivo DOCX a ingestar"),
    project: str = Form(..., description="ID del proyecto"),
    batch_size: int = Form(20, description="Tamaño del batch"),
    min_chars: int = Form(200, description="Mínimo de caracteres por fragmento"),
    max_chars: int = Form(1200, description="Máximo de caracteres por fragmento"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Sube un archivo DOCX desde el navegador y lo ingesta al proyecto.
    
    Este endpoint permite:
    1. Subir un archivo .docx desde cualquier ubicación
    2. Guardarlo en data/projects/{project}/interviews/
    3. Ingestarlo automáticamente al pipeline
    
    Returns:
        Dict con paths del archivo guardado y resultado de la ingesta
    """
    log = api_logger.bind(endpoint="upload_and_ingest", project=project)
    
    # Validar proyecto
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    # Sanitizar batch_size para evitar timeouts en Qdrant
    batch_size = max(1, min(batch_size, 32))
    
    # Validar extensión
    if not file.filename:
        raise HTTPException(status_code=400, detail="Archivo sin nombre")
    
    suffix = Path(file.filename).suffix.lower()
    if suffix not in SUPPORTED_DOCX_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado: {suffix}. Solo se permiten archivos .docx"
        )
    
    # Guardar archivo con nombre único
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = file.filename.replace(" ", "_")
    saved_filename = f"{Path(safe_filename).stem}_{timestamp}.docx"
    blob_name = f"{project_id}/{saved_filename}"
    
    # Ensure tenant present for blob writes
    org_id = getattr(user, "organization_id", None)
    if not org_id:
        raise HTTPException(status_code=409, detail="Missing organization_id (tenant). Contact admin or use a tenant-scoped API key.")

    tmp_path = None
    try:
        contents = await file.read()

        # Guardar en Azure Blob Storage using tenant-aware logical path (relative to project)
        logical = f"interviews/{saved_filename}"
        # Tenant-aware upload: central helper computes sha256 and returns canonical blob info
        try:
            blob_info = tenant_upload(
                container=CONTAINER_INTERVIEWS,
                org_id=org_id,
                project_id=project_id,
                logical_path=logical,
                data=contents,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except TenantRequiredError:
            raise HTTPException(status_code=409, detail="Missing organization_id (tenant). Contact admin or use a tenant-scoped API key.")

        blob_name_mapped = blob_info["name"]
        blob_url = blob_info["url"]

        log.info(
            "api.upload_and_ingest.saved",
            blob=blob_name,
            url=blob_url,
            size=len(contents),
        )

        # Crear archivo temporal con el mismo nombre para la ingesta
        tmp_path = Path(tempfile.gettempdir()) / saved_filename
        with open(tmp_path, "wb") as f:
            f.write(contents)

    except Exception as exc:
        log.error("api.upload_and_ingest.save_error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error guardando archivo: {exc}") from exc

    # Ingestar el archivo (lee desde el temporal generado a partir del blob)
    clients = _build_clients_or_error(settings)
    try:
        # Etapa 0 gate: requiere preparación o override aprobado.
        try:
            stage0_require_ready_or_override(
                clients.postgres,
                project=project_id,
                scope="ingest",
                user_id=str(getattr(user, "user_id", None) or getattr(user, "id", None) or "unknown"),
            )
        except PermissionError as exc:
            raise HTTPException(status_code=409, detail=str(exc)) from exc

        result = ingest_documents(
            clients,
            settings,
            files=[str(tmp_path)],
            batch_size=batch_size,
            min_chars=min_chars,
            max_chars=max_chars,
            logger=log,
            project=project_id,
        )

        log.info(
            "api.upload_and_ingest.complete",
            fragments=result.get("totals", {}).get("fragments", 0),
        )

        return {
            "project": project_id,
            "filename": saved_filename,
            "artifact_version": 1,
            "docx_logical_path": logical,
            "docx_blob": {"container": CONTAINER_INTERVIEWS, "name": blob_name_mapped, "url": blob_url},
            "result": result,
            "status": "success",
            "message": "Archivo subido a Azure e ingestión completada",
        }
    except Exception as exc:
        log.error("api.upload_and_ingest.ingest_error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()
        if tmp_path:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass


@router.post("/upload-and-transcribe", response_model=TranscribeResponse)
async def api_upload_and_transcribe(
    file: UploadFile = File(..., description="Archivo de audio a transcribir"),
    project: str = Form(..., description="ID del proyecto"),
    diarize: bool = Form(True, description="Usar diarización"),
    language: str = Form("es", description="Código de idioma"),
    ingest: bool = Form(True, description="Ingestar automáticamente"),
    batch_size: int = Form(64, description="Tamaño del batch ingest"),
    min_chars: int = Form(200, description="Mínimo chars ingest"),
    max_chars: int = Form(1200, description="Máximo chars ingest"),
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Sube un archivo de audio y lo transcribe (Síncrono/Bloqueante).
    
    Flujo:
    1. Sube archivo a data/projects/{project}/audio/raw/
    2. Ejecuta transcripción (puede tardar)
    3. Guarda DOCX en data/projects/{project}/audio/transcriptions/
    4. OPCIONAL: Ingesta el DOCX al pipeline
    """
    from app.transcription import (
        transcribe_audio_chunked,
        save_transcription_docx,
        SUPPORTED_AUDIO_FORMATS,
    )

    log = api_logger.bind(endpoint="upload_and_transcribe", project=project)

    # Validate tenant and project
    try:
        project_id = resolve_project(project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    org_id = getattr(user, "organization_id", None)
    if not org_id:
        raise HTTPException(status_code=409, detail="Missing organization_id (tenant). Contact admin or use a tenant-scoped API key.")

    # Validar extensión
    if not file.filename:
        raise HTTPException(status_code=400, detail="Archivo sin nombre")

    suffix = Path(file.filename).suffix.lower()
    if suffix not in SUPPORTED_AUDIO_FORMATS:
         raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado: {suffix}. Formatos válidos: {', '.join(SUPPORTED_AUDIO_FORMATS)}"
        )

    # Guardar audio original (to blob, using logical naming)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = file.filename.replace(" ", "_")
    saved_filename = f"{Path(safe_filename).stem}_{timestamp}{suffix}"
    logical_audio = f"audio/raw/{saved_filename}"
    try:
        contents = await file.read()

        # Map common extensions to reliable mime types
        ext = suffix.lower()
        ext_to_mime = {
            ".wav": "audio/wav",
            ".mp3": "audio/mpeg",
            ".m4a": "audio/mp4",
            ".mp4": "audio/mp4",
            ".flac": "audio/flac",
            ".ogg": "audio/ogg",
        }
        content_type = ext_to_mime.get(ext, f"audio/{ext.lstrip('.')}")

        # Tenant-aware upload for audio: helper returns sha256
        try:
            audio_blob = tenant_upload(
                container=CONTAINER_AUDIO,
                org_id=org_id,
                project_id=project_id,
                logical_path=logical_audio,
                data=contents,
                content_type=content_type,
            )
            audio_sha256 = audio_blob["sha256"]
            blob_name = audio_blob["name"]
            blob_url = audio_blob["url"]
        except TenantRequiredError:
            raise HTTPException(status_code=409, detail="Missing organization_id (tenant). Contact admin or use a tenant-scoped API key.")
        except Exception:
            # Non-tenant errors (upload failure) -> fallback to local hash
            blob_url = None
            import hashlib as _hashlib
            audio_sha256 = _hashlib.sha256(contents).hexdigest()
        log.info("api.upload_and_transcribe.saved_blob", blob=blob_name, url=blob_url, size=len(contents))

        # También guardar copia local temporal para transcripción (se eliminará)
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmpf:
            tmpf.write(contents)
            saved_path = Path(tmpf.name)

        trans_dir = Path(tempfile.gettempdir())

    except Exception as exc:
        log.error("api.upload_and_transcribe.save_error", error=str(exc))
        raise HTTPException(status_code=500, detail=f"Error guardando audio: {exc}") from exc


    # Transcribir
    try:
        result = transcribe_audio_chunked(
            saved_path,
            settings,
            diarize=diarize,
            language=language,
        )

        log.info(
            "api.upload_and_transcribe.complete",
            speakers=result.speaker_count,
            duration=result.duration_seconds,
        )
        
        # Construir respuesta base
        response_data = {
            "text": result.text,
            "segments": [
                {"speaker": seg.speaker, "text": seg.text, "start": seg.start, "end": seg.end}
                for seg in result.segments
            ],
            "speaker_count": result.speaker_count,
            "duration_seconds": result.duration_seconds,
            "fragments_ingested": None,
        }

        # Guardar DOCX
        docx_filename = f"{Path(safe_filename).stem}_{timestamp}.docx"
        docx_path = trans_dir / docx_filename
        save_transcription_docx(result, docx_path)

        # Upload final DOCX to interviews container using logical mapping (relative to project)
        logical_docx = f"interviews/audio/transcriptions/{docx_filename}"
        # Tenant-aware upload for DOCX: use helper to upload and compute sha256
        docx_blob = tenant_upload(
            container=CONTAINER_INTERVIEWS,
            org_id=org_id,
            project_id=project_id,
            logical_path=logical_docx,
            file_path=str(docx_path),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        docx_blob_name = docx_blob["name"]
        docx_url = docx_blob["url"]
        docx_hash = docx_blob["sha256"]

        # Clean local docx and audio temp files
        try:
            docx_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            saved_path.unlink(missing_ok=True)
        except Exception:
            pass

        # Ingestar (opcional)
        if ingest:
            clients = _build_clients_or_error(settings)
            try:
                ingest_result = ingest_documents(
                    clients,
                    settings,
                    files=[str(docx_path)],
                    batch_size=batch_size,
                    min_chars=min_chars,
                    max_chars=max_chars,
                    logger=log,
                    project=project_id,
                )
                totals = ingest_result.get("totals", {})
                response_data["fragments_ingested"] = totals.get("fragments_total", 0)
            finally:
                clients.close()
        # Return standardized artifact contract
        return {
            "artifact_version": 1,
            "docx_logical_path": logical_docx,
            "docx_blob": {"container": CONTAINER_INTERVIEWS, "name": docx_blob_name, "url": docx_url, "sha256": docx_hash},
            "audio_blob": {"container": CONTAINER_AUDIO, "name": blob_name, "url": blob_url, "sha256": audio_sha256},
            "text": result.text,
            "segments": response_data["segments"],
            "speaker_count": result.speaker_count,
            "duration_seconds": result.duration_seconds,
            "fragments_ingested": response_data["fragments_ingested"],
        }

    except Exception as exc:
        log.exception("api.upload_and_transcribe.error")
        raise HTTPException(status_code=500, detail=f"Error en transcripción: {exc}") from exc


# =============================================================================
# INGEST (PATH-BASED)
# =============================================================================

@router.post("/ingest")
async def api_ingest(
    payload: IngestRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Ingesta archivos desde rutas del servidor.
    
    Útil para CLI y procesamiento batch de archivos ya en el servidor.
    Para subir archivos desde el navegador, usa POST /api/upload-and-ingest
    """
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    inputs = _expand_inputs(payload.inputs)
    if not inputs:
        raise HTTPException(status_code=400, detail="Debe especificar al menos un archivo o patrón.")

    meta = None
    if payload.meta_json:
        meta_path = Path(payload.meta_json)
        if not meta_path.exists():
            raise HTTPException(status_code=400, detail="Archivo meta_json no encontrado.")
        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"No se pudo leer meta_json: {exc}") from exc

    clients = _build_clients_or_error(settings)
    log = api_logger.bind(endpoint="ingest", project=project_id)
    try:
        # Etapa 0 gate: requiere preparación o override aprobado.
        try:
            stage0_require_ready_or_override(
                clients.postgres,
                project=project_id,
                scope="ingest",
                user_id=str(getattr(user, "user_id", None) or getattr(user, "id", None) or "unknown"),
            )
        except PermissionError as exc:
            raise HTTPException(status_code=409, detail=str(exc)) from exc

        result = ingest_documents(
            clients,
            settings,
            inputs,
            batch_size=payload.batch_size,
            min_chars=payload.min_chars,
            max_chars=payload.max_chars,
            metadata=meta,
            run_id=payload.run_id,
            logger=log,
            project=project_id,
        )
    except Exception as exc:
        log.error("api.ingest.error", error=str(exc))
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    finally:
        clients.close()

    return {
        "project": project_id,
        "files": inputs,
        "result": result,
    }


# =============================================================================
# TRANSCRIPTION (SYNC)
# =============================================================================

@router.post("/transcribe", response_model=TranscribeResponse)
async def api_transcribe(
    payload: TranscribeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Transcribe archivo de audio con diarización automática.
    
    Soporta formatos: MP3, MP4, M4A, WAV, WebM, FLAC, OGG
    """
    from app.transcription import (
        transcribe_audio_chunked,
        save_transcription_docx,
        SUPPORTED_AUDIO_FORMATS,
    )
    
    # Validar proyecto
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    # Ensure tenant present for blob writes
    org_id = getattr(user, "organization_id", None)
    if not org_id:
        raise HTTPException(status_code=409, detail="Missing organization_id (tenant). Contact admin or use a tenant-scoped API key.")
    
    # Validar extensión
    suffix = Path(payload.filename).suffix.lower()
    if suffix not in SUPPORTED_AUDIO_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado: {suffix}. Formatos válidos: {', '.join(SUPPORTED_AUDIO_FORMATS)}"
        )
    
    log = api_logger.bind(endpoint="transcribe", project=project_id)
    
    # Decodificar audio
    try:
        audio_bytes = base64.b64decode(payload.audio_base64)
    except Exception as exc:
        raise HTTPException(status_code=400, detail="No se pudo decodificar audio_base64") from exc

    # Guardar audio a Blob (logical) y en temporal para transcripción
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = Path(payload.filename).stem
    saved_filename = f"{base_name}_{timestamp}{suffix}"
    logical_audio = f"audio/raw/{saved_filename}"
    # Tenant-aware upload for audio (central helper computes sha256)
    ext = suffix.lower()
    ext_to_mime = {
        ".wav": "audio/wav",
        ".mp3": "audio/mpeg",
        ".m4a": "audio/mp4",
        ".mp4": "audio/mp4",
        ".flac": "audio/flac",
        ".ogg": "audio/ogg",
    }
    content_type = ext_to_mime.get(ext, f"audio/{ext.lstrip('.')}")
    try:
        audio_blob = tenant_upload(
            container=CONTAINER_AUDIO,
            org_id=org_id,
            project_id=project_id,
            logical_path=logical_audio,
            data=audio_bytes,
            content_type=content_type,
        )
        blob_name = audio_blob["name"]
        blob_url = audio_blob["url"]
        audio_sha256 = audio_blob["sha256"]
    except Exception:
        blob_url = None
        import hashlib as _hashlib
        audio_sha256 = _hashlib.sha256(audio_bytes).hexdigest()

    # Guardar archivo temporal para transcripción
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(audio_bytes)
        tmp_path = Path(tmp.name)
    
    try:
        result = transcribe_audio_chunked(
            tmp_path,
            settings,
            diarize=payload.diarize,
            language=payload.language,
        )
        
        log.info(
            "api.transcribe.complete",
            speakers=result.speaker_count,
            segments=len(result.segments),
            duration=result.duration_seconds,
        )
        
        response_data = {
            "text": result.text,
            "segments": [
                {
                    "speaker": seg.speaker,
                    "text": seg.text,
                    "start": seg.start,
                    "end": seg.end,
                }
                for seg in result.segments
            ],
            "speaker_count": result.speaker_count,
            "duration_seconds": result.duration_seconds,
            "fragments_ingested": None,
            "saved_path": None,
        }
        
        # Guardar DOCX temporal y subir a Blob con mapping lógico
        trans_dir = Path(tempfile.gettempdir())
        docx_filename = f"{base_name}_{timestamp}.docx"
        docx_path = trans_dir / docx_filename
        save_transcription_docx(result, docx_path)

        logical_docx = f"interviews/audio/transcriptions/{docx_filename}"
        try:
            docx_blob = tenant_upload(
                container=CONTAINER_INTERVIEWS,
                org_id=org_id,
                project_id=project_id,
                logical_path=logical_docx,
                file_path=str(docx_path),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            docx_blob_name = docx_blob["name"]
            docx_url = docx_blob["url"]
            docx_hash = docx_blob["sha256"]
        except TenantRequiredError:
            raise HTTPException(status_code=409, detail="Missing organization_id (tenant). Contact admin or use a tenant-scoped API key.")
        except Exception:
            # Fallback to legacy upload and manual hash compute for non-tenant upload errors
            try:
                docx_blob_name = logical_path_to_blob_name(org_id=org_id, project_id=project_id, logical_path=logical_docx)
            except Exception:
                docx_blob_name = f"{docx_filename}"
            docx_url = upload_local_path(container=CONTAINER_INTERVIEWS, blob_name=docx_blob_name, file_path=str(docx_path), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            import hashlib
            h = hashlib.sha256()
            with open(docx_path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    h.update(chunk)
            docx_hash = h.hexdigest()

        # Clean local docx and audio temp files (will also be ensured in finally)
        try:
            docx_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

        # Ingestar si se solicita
        if payload.ingest:
            clients = _build_clients_or_error(settings)
            try:
                ingest_result = ingest_documents(
                    clients,
                    settings,
                    files=[str(docx_path)],
                    batch_size=20,
                    min_chars=payload.min_chars,
                    max_chars=payload.max_chars,
                    logger=log,
                    project=project_id,
                )
                totals = ingest_result.get("totals", {})
                response_data["fragments_ingested"] = totals.get("fragments_total", 0)
            finally:
                clients.close()
        # Return standardized artifact contract
        artifact = build_artifact_contract(audio={
            "container": CONTAINER_AUDIO,
            "name": blob_name,
            "url": blob_url,
            "sha256": audio_sha256,
            "size_bytes": None,
            "content_type": content_type,
            "logical_path": logical_audio,
        }, docx={
            "container": CONTAINER_INTERVIEWS,
            "name": docx_blob_name,
            "url": docx_url,
            "sha256": docx_hash,
            "size_bytes": None,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "logical_path": logical_docx,
        })

        return {
            **artifact,
            "text": result.text,
            "segments": response_data["segments"],
            "speaker_count": result.speaker_count,
            "duration_seconds": result.duration_seconds,
            "fragments_ingested": response_data["fragments_ingested"],
        }
        
    except ValueError as exc:
        log.error("api.transcribe.error", error=str(exc))
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        log.exception("api.transcribe.error")
        raise HTTPException(status_code=500, detail=f"Error en transcripción: {exc}") from exc
    finally:
        try:
            if 'docx_path' in locals():
                Path(docx_path).unlink(missing_ok=True)
        except Exception:
            pass
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


# =============================================================================
# TRANSCRIPTION (ASYNC STREAM)
# =============================================================================

@router.post("/transcribe/stream", response_model=StreamTranscribeResponse)
async def api_transcribe_stream(
    payload: StreamTranscribeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Inicia transcripción asíncrona con entregas incrementales.
    
    Retorna inmediatamente con un task_id para polling.
    Usa GET /api/jobs/{task_id}/status para consultar progreso.
    """
    from app.transcription import SUPPORTED_AUDIO_FORMATS
    
    log = api_logger.bind(endpoint="transcribe_stream", project=payload.project)
    
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    suffix = Path(payload.filename).suffix.lower()
    if suffix not in SUPPORTED_AUDIO_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Formato no soportado: {suffix}. Formatos válidos: {', '.join(SUPPORTED_AUDIO_FORMATS)}"
        )
    
    task = cast(Any, task_transcribe_audio).delay(
        org_id=getattr(user, "organization_id", None),
        project_id=project_id,
        audio_base64=payload.audio_base64,
        filename=payload.filename,
        diarize=payload.diarize,
        language=payload.language,
        ingest=payload.ingest,
        min_chars=payload.min_chars,
        max_chars=payload.max_chars,
        incremental=True,
    )
    
    log.info("api.transcribe_stream.started", task_id=task.id, filename=payload.filename)
    
    return {
        "task_id": task.id,
        "filename": payload.filename,
        "message": "Transcripción iniciada. Usa GET /api/jobs/{task_id}/status para consultar progreso.",
    }


# =============================================================================
# TRANSCRIPTION (BATCH)
# =============================================================================

@router.post("/transcribe/batch", response_model=BatchTranscribeResponse)
async def api_transcribe_batch(
    payload: BatchTranscribeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Inicia transcripción paralela de múltiples archivos.
    
    Los archivos se procesan en background por workers Celery.
    Límite: 20 archivos por batch.
    """
    log = api_logger.bind(endpoint="transcribe_batch", project=payload.project)
    
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    
    if len(payload.files) > MAX_BATCH_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"Máximo {MAX_BATCH_SIZE} archivos por batch. Recibidos: {len(payload.files)}"
        )
    
    if len(payload.files) == 0:
        raise HTTPException(status_code=400, detail="Debe incluir al menos un archivo")
    
    batch_id = f"batch_{uuid.uuid4().hex[:12]}"
    
    jobs = []
    for file_item in payload.files:
        task = cast(Any, task_transcribe_audio).delay(
            org_id=getattr(user, "organization_id", None),
            project_id=project_id,
            audio_base64=file_item.audio_base64,
            filename=file_item.filename,
            diarize=payload.diarize,
            language=payload.language,
            ingest=payload.ingest,
            min_chars=payload.min_chars,
            max_chars=payload.max_chars,
        )
        jobs.append({"task_id": task.id, "filename": file_item.filename})
    
    log.info("api.transcribe_batch.started", batch_id=batch_id, job_count=len(jobs))
    
    return {
        "batch_id": batch_id,
        "jobs": jobs,
        "message": f"Iniciados {len(jobs)} jobs de transcripción en paralelo",
    }


# =============================================================================
# JOB STATUS
# =============================================================================

@router.get("/jobs/{task_id}/status", response_model=JobStatusResponse)
async def api_job_status(
    task_id: str,
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Consulta el estado de un job de transcripción.
    
    Estados: PENDING, PROCESSING, SUCCESS, FAILURE
    """
    result = AsyncResult(task_id, app=celery_app)
    status = result.status
    
    if status == "STARTED":
        status = "PROCESSING"
    
    response = {
        "task_id": task_id,
        "status": status,
        "filename": None,
        "stage": None,
        "result": None,
        "error": None,
    }
    
    if status == "PROCESSING":
        meta = result.info or {}
        response["stage"] = meta.get("stage")
        response["filename"] = meta.get("filename")
        response["chunk_index"] = meta.get("chunk_index")
        response["total_chunks"] = meta.get("total_chunks")
        response["chunks_completed"] = meta.get("chunks_completed", 0)
        response["text_preview"] = meta.get("text_preview")
        response["segments_count"] = meta.get("segments_count")
        response["speakers_count"] = meta.get("speakers_count")
    
    elif status == "SUCCESS":
        task_result = result.result or {}
        response["result"] = task_result
        response["filename"] = task_result.get("filename")
    
    elif status == "FAILURE":
        response["error"] = str(result.result) if result.result else "Error desconocido"
    
    return response


@router.get("/jobs/batch/{batch_id}/status")
async def api_batch_status(
    batch_id: str,
    task_ids: str,  # Comma-separated list of task IDs
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """Consulta el estado de múltiples jobs de un batch."""
    ids = [tid.strip() for tid in task_ids.split(",") if tid.strip()]
    
    statuses = []
    completed = failed = processing = pending = 0
    
    for tid in ids:
        result = AsyncResult(tid, app=celery_app)
        status = result.status
        
        if status == "SUCCESS":
            completed += 1
        elif status == "FAILURE":
            failed += 1
        elif status == "PROCESSING":
            processing += 1
        else:
            pending += 1
        
        statuses.append({"task_id": tid, "status": status})
    
    return {
        "batch_id": batch_id,
        "total": len(ids),
        "completed": completed,
        "failed": failed,
        "processing": processing,
        "pending": pending,
        "all_done": (completed + failed) == len(ids),
        "jobs": statuses,
    }


# =============================================================================
# TRANSCRIPTION MERGE
# =============================================================================

class TranscribeMergeSegment(BaseModel):
    speaker: str
    text: str
    start: float
    end: float


class TranscribeMergeItem(BaseModel):
    filename: str
    text: str
    segments: List[TranscribeMergeSegment] = []


class TranscribeMergeRequest(BaseModel):
    """Request para combinar múltiples transcripciones en un DOCX."""
    model_config = ConfigDict(extra="forbid")
    
    project: str
    transcriptions: List[TranscribeMergeItem]


class TranscribeMergeResponse(BaseModel):
    docx_base64: str
    filename: str


@router.post("/transcribe/merge", response_model=TranscribeMergeResponse)
async def api_transcribe_merge(
    payload: TranscribeMergeRequest,
    settings: AppSettings = Depends(get_settings),
    user: User = Depends(require_auth),
) -> Dict[str, Any]:
    """
    Combina múltiples transcripciones en un único archivo DOCX.
    Retorna el archivo como base64 para descarga directa.
    """
    import tempfile
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Validar proyecto
    try:
        project_id = resolve_project(payload.project, allow_create=False)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    # Ensure tenant present for blob writes
    org_id = getattr(user, "organization_id", None)
    if not org_id:
        raise HTTPException(status_code=409, detail="Missing organization_id (tenant). Contact admin or use a tenant-scoped API key.")
    
    log = api_logger.bind(endpoint="transcribe_merge", project=project_id)
    
    if not payload.transcriptions:
        raise HTTPException(status_code=400, detail="No hay transcripciones para combinar")
    
    # Validar que al menos una transcripción tenga contenido real
    valid_count = sum(
        1 for t in payload.transcriptions
        if (t.text and t.text.strip()) or (t.segments and len(t.segments) > 0)
    )
    
    if valid_count == 0:
        log.warning(
            "api.transcribe_merge.empty_content",
            total=len(payload.transcriptions),
            filenames=[t.filename for t in payload.transcriptions],
        )
        raise HTTPException(
            status_code=400,
            detail="Las transcripciones no tienen contenido. Espera a que las transcripciones terminen de procesarse."
        )
    
    try:
        # Crear documento
        doc = Document()
        
        # Título
        title = doc.add_heading(f"Transcripción Combinada - {project_id}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Archivos: {len(payload.transcriptions)}")
        doc.add_paragraph("")
        
        # Agregar cada transcripción
        for item in payload.transcriptions:
            doc.add_heading(f"📁 {item.filename}", level=1)
            
            if item.segments:
                for seg in item.segments:
                    para = doc.add_paragraph()
                    speaker_run = para.add_run(f"[{seg.speaker}] ")
                    speaker_run.bold = True
                    speaker_run.font.color.rgb = RGBColor(0, 100, 180)
                    para.add_run(seg.text)
            else:
                doc.add_paragraph(item.text)
            
            doc.add_paragraph("")
        
        # Guardar a temp file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        
        doc.save(str(tmp_path))
        
        # Subir DOCX combinado a Blob (logical path) y devolver contrato de artefacto
        logical_docx = f"interviews/audio/transcriptions/combined_{timestamp}.docx"
        try:
            docx_blob = tenant_upload(
                container=CONTAINER_INTERVIEWS,
                org_id=org_id,
                project_id=project_id,
                logical_path=logical_docx,
                file_path=str(tmp_path),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            docx_blob_name = docx_blob["name"]
            docx_url = docx_blob["url"]
            docx_hash = docx_blob["sha256"]
        except Exception:
            try:
                docx_blob_name = logical_path_to_blob_name(org_id=org_id, project_id=project_id, logical_path=logical_docx)
            except Exception:
                docx_blob_name = f"combined_{timestamp}.docx"
            docx_url = upload_local_path(container=CONTAINER_INTERVIEWS, blob_name=docx_blob_name, file_path=str(tmp_path), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            import hashlib as _hashlib
            h = _hashlib.sha256()
            with open(tmp_path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    h.update(chunk)
            docx_hash = h.hexdigest()

        log.info("api.transcribe_merge.saved", blob=docx_blob_name, url=docx_url)

        return {
            "artifact_version": 1,
            "docx_logical_path": logical_docx,
            "docx_blob": {"container": CONTAINER_INTERVIEWS, "name": docx_blob_name, "url": docx_url, "sha256": docx_hash},
            "filename": f"transcripciones_{project_id}_{timestamp}.docx",
        }
    except Exception as exc:
        log.exception("api.transcribe_merge.error")
        raise HTTPException(status_code=500, detail=f"Error generando DOCX: {exc}") from exc
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass
